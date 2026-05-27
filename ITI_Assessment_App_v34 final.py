from flask import Flask, request, send_file, jsonify, session, redirect, url_for
import pandas as pd
import random, io, os, sqlite3, secrets, string
from datetime import datetime
from werkzeug.security import generate_password_hash, check_password_hash
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# ── ReportLab imports for PDF generation ──────────────────────────
try:
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                    Paragraph, Spacer, PageBreak)
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    _REPORTLAB_OK = True
except ImportError:
    _REPORTLAB_OK = False


def _find_libreoffice():
    import shutil, glob
    for cmd in ('libreoffice', 'soffice'):
        if shutil.which(cmd):
            return cmd
    for pattern in (
        r'C:\Program Files\LibreOffice\program\soffice.exe',
        r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
    ):
        m = glob.glob(pattern)
        if m:
            return m[0]
    return None

def _doc_to_docx_bytes(doc_bytes):
    """Convert .doc bytes to .docx bytes. Returns None on failure."""
    import tempfile, subprocess, io as _io
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_doc  = os.path.join(tmpdir, 'lh.doc')
        tmp_docx = os.path.join(tmpdir, 'lh.docx')
        with open(tmp_doc, 'wb') as f:
            f.write(doc_bytes)
        # Try MS Word via COM (Windows)
        try:
            import win32com.client
            word = win32com.client.Dispatch('Word.Application')
            word.Visible = False
            wb = word.Documents.Open(tmp_doc)
            wb.SaveAs2(tmp_docx, FileFormat=16)
            wb.Close(False)
            word.Quit()
        except Exception:
            # Fallback: LibreOffice
            lo = _find_libreoffice()
            if lo:
                subprocess.run([lo, '--headless', '--convert-to', 'docx',
                                '--outdir', tmpdir, tmp_doc],
                               capture_output=True, timeout=30)
        if os.path.exists(tmp_docx):
            return open(tmp_docx, 'rb').read()
    return None

app = Flask(__name__)
app.secret_key = secrets.token_hex(32)

# Server-side cache: stores generated marks per user so Annexure-II matches main report
_marks_cache_store  = {}  # user_id -> marks_cache
_trainee_data_store = {}  # user_id -> trainee_df (reused by progress card)

import hashlib, json as _json

def _file_hash(path):
    """Return MD5 hex digest of a file on disk, or '' if not found."""
    if not path or not os.path.exists(path):
        return ''
    h = hashlib.md5()
    with open(path, 'rb') as f:
        for chunk in iter(lambda: f.read(65536), b''):
            h.update(chunk)
    return h.hexdigest()

def _marks_cache_path(uid):
    return os.path.join(_user_upload_dir(uid), 'marks_cache.json')

def _save_marks_cache_disk(uid, marks_cache, trainee_hash, lo_hash):
    """Persist marks_cache + file hashes to disk as JSON."""
    try:
        payload = {
            'trainee_hash': trainee_hash,
            'lo_hash':      lo_hash,
            'marks_cache':  marks_cache,
        }
        path = _marks_cache_path(uid)
        with open(path, 'w', encoding='utf-8') as f:
            _json.dump(payload, f)
    except Exception:
        pass  # non-fatal — worst case we regenerate next time

def _load_marks_cache_disk(uid, trainee_hash, lo_hash):
    """
    Return persisted marks_cache if the stored file hashes match
    the current trainee + LO files.  Returns None otherwise.
    """
    try:
        path = _marks_cache_path(uid)
        if not os.path.exists(path):
            return None
        with open(path, 'r', encoding='utf-8') as f:
            payload = _json.load(f)
        if payload.get('trainee_hash') == trainee_hash and payload.get('lo_hash') == lo_hash:
            # Restore integer keys for lo_num (JSON serialises them as strings)
            raw = payload.get('marks_cache', {})
            restored = {}
            for roll_key, lo_dict in raw.items():
                try:
                    roll = int(roll_key)
                except (ValueError, TypeError):
                    roll = roll_key
                restored_lo = {}
                for lo_key, lo_data in lo_dict.items():
                    try:
                        lo_num = int(lo_key)
                    except (ValueError, TypeError):
                        lo_num = lo_key
                    restored_lo[lo_num] = lo_data
                restored[roll] = restored_lo
            return restored
        return None
    except Exception:
        return None

# Persistent file storage — survives logout/restart
UPLOAD_DIR = os.path.join(os.path.dirname(__file__), 'user_uploads')
os.makedirs(UPLOAD_DIR, exist_ok=True)

def _user_upload_dir(uid):
    d = os.path.join(UPLOAD_DIR, str(uid))
    os.makedirs(d, exist_ok=True)
    return d

def _save_user_file(uid, key, file_obj):
    """Save uploaded file to disk. key is 'trainee' or 'lo'."""
    path = os.path.join(_user_upload_dir(uid), f'{key}.xlsx')
    file_obj.seek(0)
    with open(path, 'wb') as f:
        f.write(file_obj.read())
    file_obj.seek(0)

def _save_user_file_ext(uid, key, file_obj, ext):
    """Save uploaded file to disk with given extension."""
    path = os.path.join(_user_upload_dir(uid), f'{key}.{ext}')
    file_obj.seek(0)
    with open(path, 'wb') as f:
        f.write(file_obj.read())
    file_obj.seek(0)
    return path

def _load_user_file_ext(uid, key, ext):
    """Return file path if saved file exists, else None."""
    path = os.path.join(_user_upload_dir(uid), f'{key}.{ext}')
    return path if os.path.exists(path) else None

def _load_user_file(uid, key):
    """Return file path if saved file exists, else None."""
    path = os.path.join(_user_upload_dir(uid), f'{key}.xlsx')
    return path if os.path.exists(path) else None

def _get_df(uid, key, request_file):
    """
    Return df. Prefers freshly uploaded file; falls back to disk.
    If a new file is uploaded, saves it to disk for future sessions.
    """
    if request_file and request_file.filename:
        _save_user_file(uid, key, request_file)
        df = pd.read_excel(request_file)
        df.columns = df.columns.str.strip()
        return df
    path = _load_user_file(uid, key)
    if path:
        df = pd.read_excel(path)
        df.columns = df.columns.str.strip()
        return df
    return None

DB = os.path.join(os.path.dirname(__file__), 'iti_users.db')

# ─── DATABASE SETUP ───────────────────────────────────────────────
def get_db():
    conn = sqlite3.connect(DB)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    with get_db() as db:
        db.execute('''CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            si_name TEXT NOT NULL,
            trade_name TEXT NOT NULL,
            iti_name TEXT NOT NULL,
            mobile TEXT NOT NULL,
            year_of_assessment TEXT,
            assessment_location TEXT,
            near_trade TEXT,
            trade_duration TEXT,
            semester TEXT,
            batch TEXT,
            status TEXT DEFAULT 'pending',
            user_id TEXT UNIQUE,
            password_hash TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )''')
        # Add new columns to existing DB if upgrading
        for col, coltype in [
            ('year_of_assessment','TEXT'), ('assessment_location','TEXT'),
            ('near_trade','TEXT'), ('trade_duration','TEXT'),
            ('semester','TEXT'), ('batch','TEXT'),
            ('approved_at','TEXT'), ('deregistered_at','TEXT'),
        ]:
            try:
                db.execute(f'ALTER TABLE users ADD COLUMN {col} {coltype}')
            except Exception:
                pass
        db.execute('''CREATE TABLE IF NOT EXISTS admin (
            id INTEGER PRIMARY KEY,
            password_hash TEXT NOT NULL
        )''')
        # Create default admin if not exists (password: admin123)
        admin = db.execute('SELECT * FROM admin WHERE id=1').fetchone()
        if not admin:
            db.execute('INSERT INTO admin VALUES (1, ?)',
                       (generate_password_hash('admin123'),))
        db.commit()

init_db()

# ─── HELPERS ──────────────────────────────────────────────────────
def gen_user_id(mobile):
    return mobile  # User ID is their mobile number

def gen_password(mobile):
    return mobile[-4:]  # Password is last 4 digits of mobile

def login_required(f):
    from functools import wraps
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'user_id' not in session:
            return jsonify({'error': 'Unauthorized', 'redirect': '/'}), 401
        return f(*args, **kwargs)
    return decorated

def admin_required(f):
    from functools import wraps
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('is_admin'):
            return jsonify({'error': 'Unauthorized'}), 401
        return f(*args, **kwargs)
    return decorated

# ─── BLANK TEMPLATE DOWNLOAD ROUTES ─────────────────────────────
@app.route('/download/irr-trainee-template')
def download_irr_trainee_template():
    wb = Workbook()
    ws = wb.active
    ws.title = 'Sheet1'
    headers = [
        'અટક ', 'તાલીમાર્થીનું નામ', 'પિતાનું નામ ',
        'એડ્રેસ-૧ ', 'એડ્રેસ-૨ ', 'પીનકોડ',
        'ટ્રેડ નું નામ અને બેચ', 'સુ.ઈ નું નામ)', 'આઈ.ટી.આઈ નું નામ '
    ]
    col_widths = [20.85, 21.85, 21.85, 38.57, 24.14, 9.85, 21.0, 27.0, 21.0]
    thin = Side(style='thin', color='000000')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for col_idx, (header, width) in enumerate(zip(headers, col_widths), start=1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = Font(name='Calibri', size=11, bold=True, color='000000')
        cell.fill = PatternFill('none')
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = border
        ws.column_dimensions[get_column_letter(col_idx)].width = width
    ws.row_dimensions[1].height = 45.0
    for row_idx in range(2, 11):
        ws.row_dimensions[row_idx].height = 17.25
        for col_idx in range(1, 10):
            cell = ws.cell(row=row_idx, column=col_idx, value='')
            cell.border = border
            cell.alignment = Alignment(horizontal='center', vertical='center')
    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    return send_file(out, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                     as_attachment=True, download_name='BLANK_TRAINEE_DATA_FOR_CHETAVANI_PATRA.xlsx')


@app.route('/download/trainee-template')
def download_trainee_template():
    wb = Workbook()
    ws = wb.active
    ws.title = "Trainee Details"
    headers = [
        'rollno', 'Firstname', 'Fathername', 'Lastname',
        'Birth_Date', 'EduQualification', 'DateofAdmission', 'DateofLeaving',
        'ES', 'WC_SC', 'ED', 'Totalof ES WCS ED',
        'All_LO_average_base_on_70_SEM_I',
        'Total', 'Working_days', 'Attendeate_days', 'Last Practical No'
    ]
    thin = Side(style='thin', color='000000')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = Font(name='Calibri', size=10, bold=True, color='000000')
        cell.fill = PatternFill('none')
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = border
        ws.column_dimensions[get_column_letter(col_idx)].width = 22
    ws.row_dimensions[1].height = 30
    # Add note row
    note_cell = ws.cell(row=2, column=13, value='← Fill marks out of 70 here')
    note_cell.font = Font(name='Calibri', size=9, bold=True, color='000000')
    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    return send_file(out, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                     as_attachment=True, download_name='Template_Trainee_Details.xlsx')

@app.route('/download/lo-template')
def download_lo_template():
    wb = Workbook()
    ws = wb.active
    ws.title = "LO Details"
    headers = ['lo', 'practfrom', 'practto', 'lo_name']
    thin = Side(style='thin', color='000000')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    hints = ['LO Number (e.g. 1)', 'First Practical No (e.g. 1)', 'Last Practical No (e.g. 5)', 'LO Name / Description']
    for col_idx, (header, hint) in enumerate(zip(headers, hints), start=1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = Font(name='Calibri', size=10, bold=True, color='000000')
        cell.fill = PatternFill('none')
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = border
        hint_cell = ws.cell(row=2, column=col_idx, value=hint)
        hint_cell.font = Font(name='Calibri', size=9, italic=True, color='000000')
        hint_cell.border = border
        ws.column_dimensions[get_column_letter(col_idx)].width = 30
    ws.row_dimensions[1].height = 28
    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    return send_file(out, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                     as_attachment=True, download_name='Template_LO_Details.xlsx')

# ─── MARK GENERATION (Reverse Distribution) ───────────────────────
# REVERSE FLOW:
#   All_LO_average (target, out of 70)
#     → scale to 100  → per-LO target
#       → each Practical in that LO gets the SAME grand-total target
#         → sub-marks for each category are distributed proportionally
#            to hit that exact grand total
#
# This guarantees:
#   avg(practical grand totals) = LO target
#   avg(all LO targets)         = All_LO_average target  ✓

MAX_MARKS = {
    'safety':     {'dress': 2, 'ppe': 5, 'safety': 8,  'total': 15},
    'hygiene':    {'personal': 3, 'scrap': 2, 'material': 5, 'total': 10},
    'attendance': {'initiative': 3, 'accountability': 3, 'participative': 4, 'total': 10},
    'manuals':    {'select': 1, 'search': 2, 'read': 2, 'total': 5},
    'knowledge':  {'plan': 4, 'tools': 3, 'review': 3, 'total': 10},
    'skills':     {'handle': 4, 'safety': 3, 'maintain': 3, 'total': 10},
    'speed':      {'sequence': 3, 'technique': 5, 'review': 2, 'total': 10},
    'quality':    {'accuracy': 7, 'conform': 3, 'satisfy': 5, 'total': 15},
    'viva':       {'clarity': 7, 'technical': 5, 'conscious': 3, 'total': 15},
}

# Category max totals (sum of sub-maxes), kept as a flat dict for distribution
_CAT_MAXES = {cat: v['total'] for cat, v in MAX_MARKS.items()}
# MAX_GRAND must equal sum of all category totals = 100
_MAX_GRAND = sum(_CAT_MAXES.values())  # 100
_MIN_GRAND = 27  # physical minimum: 9 categories × 3 sub-fields × min 1 each


def _distribute_to_subs(sub_maxes: dict, category_target: int) -> dict:
    """
    Distribute `category_target` across sub-fields.
    Each sub gets at least 1 mark, at most its declared maximum.
    Remainder is spread randomly so marks don't look uniform.
    """
    keys = list(sub_maxes.keys())
    maxes = [sub_maxes[k] for k in keys]
    total_max = sum(maxes)
    # Clamp target to achievable range [len(keys), total_max]
    target = max(len(keys), min(total_max, category_target))

    # Proportional floor allocation (minimum 1 each)
    result = [max(1, int(target * m / total_max)) for m in maxes]
    diff = target - sum(result)

    # Distribute leftover (or remove excess) in random order
    indices = list(range(len(keys)))
    random.shuffle(indices)
    for i in indices:
        if diff == 0:
            break
        if diff > 0:
            add = min(diff, maxes[i] - result[i])
            result[i] += add
            diff -= add
        else:
            sub = min(-diff, result[i] - 1)
            result[i] -= sub
            diff += sub

    return {keys[i]: result[i] for i in range(len(keys))}


def generate_row_marks(target_total=None):
    """
    Generate one practical's marks so that grand_total == target_total exactly.

    Algorithm (reverse distribution):
      1. Clamp target to [27, 100]  (27 is the physical minimum: 9 cats × min 3 subs × 1 each)
      2. Allocate the target proportionally across the 9 categories
      3. Adjust integer rounding so category allocations sum exactly to target
      4. For each category, distribute its allocated total across sub-fields
         proportionally (with random tie-breaking)

    The grand_total of the returned row is guaranteed to equal target_total
    for any target in [27, 100].
    """
    # Default target when none given (≈75 %)
    if target_total is None:
        target_total = 75

    # Clamp to physical bounds — max is 99 so no practical ever gets 100/100
    target = max(27, min(99, target_total))

    cat_names = list(_CAT_MAXES.keys())

    # Step 1 – proportional floor allocation per category
    cat_alloc = {
        cat: max(1, int(target * _CAT_MAXES[cat] / _MAX_GRAND))
        for cat in cat_names
    }

    # Step 2 – fix integer rounding so total == target
    diff = target - sum(cat_alloc.values())
    shuffled_cats = cat_names[:]
    random.shuffle(shuffled_cats)
    for cat in shuffled_cats:
        if diff == 0:
            break
        if diff > 0:
            add = min(diff, _CAT_MAXES[cat] - cat_alloc[cat])
            cat_alloc[cat] += add
            diff -= add
        else:
            sub = min(-diff, cat_alloc[cat] - 1)
            cat_alloc[cat] -= sub
            diff += sub

    # Step 3 – distribute each category allocation into its sub-fields
    r = {}

    s = _distribute_to_subs({k: v for k, v in MAX_MARKS['safety'].items() if k != 'total'}, cat_alloc['safety'])
    r['dress'], r['ppe'], r['apply_safety'] = s['dress'], s['ppe'], s['safety']
    r['safety_total'] = r['dress'] + r['ppe'] + r['apply_safety']

    h = _distribute_to_subs({k: v for k, v in MAX_MARKS['hygiene'].items() if k != 'total'}, cat_alloc['hygiene'])
    r['personal'], r['scrap'], r['material'] = h['personal'], h['scrap'], h['material']
    r['hygiene_total'] = r['personal'] + r['scrap'] + r['material']

    a = _distribute_to_subs({k: v for k, v in MAX_MARKS['attendance'].items() if k != 'total'}, cat_alloc['attendance'])
    r['initiative'], r['accountability'], r['participative'] = a['initiative'], a['accountability'], a['participative']
    r['attendance_total'] = r['initiative'] + r['accountability'] + r['participative']

    m = _distribute_to_subs({k: v for k, v in MAX_MARKS['manuals'].items() if k != 'total'}, cat_alloc['manuals'])
    r['select_manual'], r['search_topic'], r['read_manual'] = m['select'], m['search'], m['read']
    r['manuals_total'] = r['select_manual'] + r['search_topic'] + r['read_manual']

    k = _distribute_to_subs({k: v for k, v in MAX_MARKS['knowledge'].items() if k != 'total'}, cat_alloc['knowledge'])
    r['plan_work'], r['select_tools'], r['review_work'] = k['plan'], k['tools'], k['review']
    r['knowledge_total'] = r['plan_work'] + r['select_tools'] + r['review_work']

    sk = _distribute_to_subs({k: v for k, v in MAX_MARKS['skills'].items() if k != 'total'}, cat_alloc['skills'])
    r['handle_tools'], r['safety_handling'], r['care_maintain'] = sk['handle'], sk['safety'], sk['maintain']
    r['skills_total'] = r['handle_tools'] + r['safety_handling'] + r['care_maintain']

    sp = _distribute_to_subs({k: v for k, v in MAX_MARKS['speed'].items() if k != 'total'}, cat_alloc['speed'])
    r['sequence'], r['technique'], r['review_execution'] = sp['sequence'], sp['technique'], sp['review']
    r['speed_total'] = r['sequence'] + r['technique'] + r['review_execution']

    q = _distribute_to_subs({k: v for k, v in MAX_MARKS['quality'].items() if k != 'total'}, cat_alloc['quality'])
    r['accuracy'], r['conform'], r['satisfy'] = q['accuracy'], q['conform'], q['satisfy']
    r['quality_total'] = r['accuracy'] + r['conform'] + r['satisfy']

    v = _distribute_to_subs({k: v for k, v in MAX_MARKS['viva'].items() if k != 'total'}, cat_alloc['viva'])
    r['clarity'], r['technical'], r['conscious'] = v['clarity'], v['technical'], v['conscious']
    r['viva_total'] = r['clarity'] + r['technical'] + r['conscious']

    r['grand_total'] = (
        r['safety_total'] + r['hygiene_total'] + r['attendance_total'] +
        r['manuals_total'] + r['knowledge_total'] + r['skills_total'] +
        r['speed_total'] + r['quality_total'] + r['viva_total']
    )
    return r

def varied_targets_averaging_to(target, count, half_range=10, min_val=None):
    """
    Generate `count` integer mark values whose average equals `target` exactly.

    Each value is constrained within [target - half_range, target + half_range]
    and also within [_MIN_GRAND, 99] so no practical ever gets 100/100.
    Default half_range=10 gives ±10 variation; pass a smaller value for
    LO-level distribution to avoid compounding spread.

    min_val: optional hard lower bound for every generated value.
             Used to ensure all practicals stay >= 60 when target >= 60.
    """
    MAX_ALLOWED = 99

    target = max(_MIN_GRAND, min(MAX_ALLOWED, target))
    if count == 1:
        return [target]

    # Apply min_val floor (e.g. 60) on top of the normal _MIN_GRAND floor
    effective_min = _MIN_GRAND if min_val is None else max(_MIN_GRAND, min_val)

    lo_bound = max(effective_min, target - half_range)
    hi_bound = min(MAX_ALLOWED, target + half_range)

    values = [target] * count
    for _ in range(count * 10):
        i, j = random.sample(range(count), 2)
        max_add = hi_bound - values[i]
        max_sub = values[j] - lo_bound
        max_move = min(max_add, max_sub)
        if max_move > 0:
            delta = random.randint(1, max_move)
            values[i] += delta
            values[j] -= delta

    return values


def build_all_marks(trainee_df, lo_df):
    """
    Pre-generate ALL practical marks for every trainee × LO combination.
    Returns a nested dict:
        marks_cache[roll][lo_num] = {
            'pract_marks': [list of per-practical mark dicts],   # individual rows
            'cat_avgs':    {safety_total, ..., grand_total},     # LO-level averages
            'lo_avg':      int,                                   # avg of grand totals
        }
    Call this ONCE and pass the result to both create_excel and
    create_lo_summary_excel so both reports show identical numbers.
    """
    lo_rows_all = list(lo_df.iterrows())
    marks_cache = {}

    for _, trainee in trainee_df.iterrows():
        roll = trainee['rollno']
        raw_70 = trainee.get('All_LO_average_base_on_70_SEM_I')
        if raw_70 is not None and pd.notna(raw_70):
            target_avg = round((float(raw_70) / 70) * 100)
        else:
            target_avg = None

        last_pract_cap = trainee.get('Last Practical No')
        last_pract_cap = int(last_pract_cap) if last_pract_cap is not None and pd.notna(last_pract_cap) else None

        active_lo_rows = []
        for _, lo in lo_rows_all:
            p_from = int(lo['practfrom'])
            p_to   = int(lo['practto'])
            if last_pract_cap is not None:
                p_to = min(p_to, last_pract_cap)
            if p_from <= p_to:
                active_lo_rows.append(lo)

        num_active_los = len(active_lo_rows)
        # If trainee's mark (out of 70) is >= 60%, enforce all practicals >= 60
        min_pract = 60 if (target_avg is not None and target_avg >= 60) else None
        if target_avg is not None and num_active_los > 0:
            lo_targets = varied_targets_averaging_to(target_avg, num_active_los, half_range=6, min_val=min_pract)
        else:
            lo_targets = [target_avg] * num_active_los

        marks_cache[roll] = {}

        for lo_idx, lo in enumerate(active_lo_rows):
            lo_num     = lo['lo']
            pract_from = int(lo['practfrom'])
            pract_to   = int(lo['practto'])
            if last_pract_cap is not None:
                pract_to = min(pract_to, last_pract_cap)
            pract_count = pract_to - pract_from + 1
            lo_target   = lo_targets[lo_idx]

            if lo_target is not None and pract_count > 0:
                pract_targets = varied_targets_averaging_to(lo_target, pract_count, min_val=min_pract)
            else:
                pract_targets = [lo_target] * pract_count

            pract_marks = []
            cat_sums = {
                'safety_total': 0, 'hygiene_total': 0, 'attendance_total': 0,
                'manuals_total': 0, 'knowledge_total': 0, 'skills_total': 0,
                'speed_total': 0, 'quality_total': 0, 'viva_total': 0,
                'grand_total': 0,
            }
            for pract_offset in range(pract_count):
                m = generate_row_marks(target_total=pract_targets[pract_offset])
                pract_marks.append(m)
                for k in cat_sums:
                    cat_sums[k] += m[k]

            cat_avgs = {k: round(v / pract_count) for k, v in cat_sums.items()}
            lo_grand_totals = [m['grand_total'] for m in pract_marks]
            lo_avg = round(sum(lo_grand_totals) / len(lo_grand_totals))

            marks_cache[roll][lo_num] = {
                'pract_marks': pract_marks,
                'cat_avgs':    cat_avgs,
                'lo_avg':      lo_avg,
                'pract_from':  pract_from,
                'pract_to':    pract_to,
                'lo_name':     lo['lo_name'],
                'lo_target':   lo_target,
            }

    return marks_cache


def create_excel(trainee_df, lo_df, user_info=None, marks_cache=None):
    ui = user_info or {}
    u_si_name           = ui.get('si_name', 'N/A')
    u_trade_name        = ui.get('trade_name', 'N/A')
    u_iti_name          = ui.get('iti_name', 'N/A')
    u_year_of_assessment= ui.get('year_of_assessment', '')
    u_assessment_location=ui.get('assessment_location', '')
    u_trade_duration    = ui.get('trade_duration', '')
    u_semester          = ui.get('semester', '')
    u_batch             = ui.get('batch', '')

    wb = Workbook()
    wb.remove(wb.active)
    HEADER_BG="1B4F8A"; HEADER2_BG="2E75B6"; LO_SUMMARY_BG="D6E4F0"
    MAX_ROW_BG="FFF2CC"; ALT_ROW="EBF3FB"; WHITE="FFFFFF"
    thin = Side(style='thin', color='000000')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    def hdr_font(size=9, bold=True, color="000000"):
        return Font(name='Calibri', size=size, bold=bold, color=color)
    def cell_font(size=9, bold=False, color="000000"):
        return Font(name='Calibri', size=size, bold=bold, color=color)
    def center_align(wrap=True):
        return Alignment(horizontal='center', vertical='center', wrap_text=wrap)
    for _, trainee in trainee_df.iterrows():
        roll = trainee['rollno']
        full_name = f"{trainee['Firstname']} {trainee['Fathername']} {trainee['Lastname']}"
        # Take marks from All_LO_average_base_on_70_SEM_I (out of 70), scale to 100
        raw_70 = trainee.get('All_LO_average_base_on_70_SEM_I')
        if raw_70 is not None and pd.notna(raw_70):
            target_avg = round((float(raw_70) / 70) * 100)
        else:
            target_avg = None
        ws = wb.create_sheet(title=f"Roll_{roll}")

        # ── Row 1: Title — merged A1:AO1 ─────────────────────────────
        ws.merge_cells('A1:AO1')
        ws['A1'] = 'Internal Assessment'
        ws['A1'].font = Font(name='Calibri', size=14, bold=True, color='000000')
        ws['A1'].alignment = center_align()
        ws['A1'].fill = PatternFill('none')
        for col_idx in range(1, 42):
            ws.cell(row=1, column=col_idx).border = border

        # ── Rows 2-5: Info fields — 3 segments: A:K, L:S, T:AO ──────
        # Segment merges: cols 1-11, 12-19, 20-41
        info_rows = [
            (f"Name of Trainee: {full_name}", f"Roll NO: {roll}",                    f"Year of Assessment: {u_year_of_assessment}"),
            (f"Name of ITI: {u_iti_name}",    "Date of Assessment:",                  f"Batch: {u_batch}"),
            (f"Name of the Industry: {u_iti_name}", f"Assessment Location: {u_assessment_location}", ""),
            (f"Trade Name: {u_trade_name}",   f"Duration of the Trade: {u_trade_duration}", f"S.I.Name: {u_si_name}"),
        ]
        for i, (seg1, seg2, seg3) in enumerate(info_rows, start=2):
            ws.merge_cells(f'A{i}:K{i}')
            ws[f'A{i}'] = seg1; ws[f'A{i}'].font = cell_font(size=9, bold=True)
            ws.merge_cells(f'L{i}:S{i}')
            ws[f'L{i}'] = seg2; ws[f'L{i}'].font = cell_font(size=9, bold=True)
            ws.merge_cells(f'T{i}:AO{i}')
            ws[f'T{i}'] = seg3; ws[f'T{i}'].font = cell_font(size=9, bold=True)
            ws[f'T{i}'].alignment = Alignment(horizontal='center', vertical='center', wrap_text=False)
            for col_idx in range(1, 42):
                ws.cell(row=i, column=col_idx).border = border

        # ── Rows 6-7: Category and sub-column headers ─────────────────
        ROW_CAT=6; ROW_SUBCOL=7; ROW_MAX=8
        # Row 6: main category headers — A6:A7 merged, B6:B7 merged, then groups of 4 cols
        # A and B span rows 6-7 (merged vertically)
        ws.merge_cells(f'A{ROW_CAT}:A{ROW_SUBCOL}')
        ws.merge_cells(f'B{ROW_CAT}:B{ROW_SUBCOL}')
        # AM, AN, AO also span rows 6-7
        ws.merge_cells(f'AM{ROW_CAT}:AM{ROW_SUBCOL}')
        ws.merge_cells(f'AN{ROW_CAT}:AN{ROW_SUBCOL}')
        ws.merge_cells(f'AO{ROW_CAT}:AO{ROW_SUBCOL}')

        # (sc, ec, label, rotation) — 90deg for single-col headers, 0 for span headers
        categories = [
            (1,  1,  "LO No",                                             90),
            (2,  2,  "Practical No",                                       90),
            (3,  6,  "Safety\nconsciousness",                              0),
            (7,  10, "Workplace hygiene\n& Economical use\nof materials",  0),
            (11, 14, "Attendance/\nPunctuality",                           0),
            (15, 18, "Ability to follow\nManuals/\nWritten instructions",  0),
            (19, 22, "Application of\nKnowledge",                          0),
            (23, 26, "Skills to handle\ntools & equipment",                0),
            (27, 30, "Speed in doing work",                                0),
            (31, 34, "Quality in workmanship",                             0),
            (35, 38, "VIVA",                                               0),
            (39, 39, "Grand\nTotal",                                       90),
            (40, 40, "Signature\nTrainee",                                 90),
            (41, 41, "Signature\nSI",                                      90),
        ]
        for (sc, ec, label, rotation) in categories:
            if sc != ec:
                ws.merge_cells(f'{get_column_letter(sc)}{ROW_CAT}:{get_column_letter(ec)}{ROW_CAT}')
            cell = ws.cell(row=ROW_CAT, column=sc)
            cell.value = label
            cell.font = hdr_font(size=8)
            cell.fill = PatternFill('none')
            cell.alignment = Alignment(horizontal='center', vertical='center',
                                       wrap_text=(rotation == 0),
                                       text_rotation=rotation)
            cell.border = border

        # Apply borders to all cells in row 6 (including interior merged cells)
        for col_idx in range(1, 42):
            ws.cell(row=ROW_CAT, column=col_idx).border = border

        # Row 7: sub-column headers (cols 1-2 and 39-41 are blank — part of vertical merge)
        sub_cols = [
            (1,""), (2,""),
            (3,"Dress code"),(4,"Use PPE"),(5,"Apply/ practice safety"),(6,"Total"),
            (7,"Maintain personal & workplace cleanliness"),(8,"Dispose scrap"),(9,"Select material"),(10,"Total"),
            (11,"Initiative"),(12,"Account- ability"),(13,"Participative in work"),(14,"Total"),
            (15,"Select right manual"),(16,"Search topic"),(17,"Read & interpret"),(18,"Total"),
            (19,"Plan the work"),(20,"Select tools"),(21,"Review work"),(22,"Total"),
            (23,"Handle & use tools"),(24,"Maintain safety"),(25,"Care & maintain"),(26,"Total"),
            (27,"Properly sequence"),(28,"Use approp. technique"),(29,"Review execution"),(30,"Total"),
            (31,"Achieve high accuracy"),(32,"Conform to req."),(33,"Satisfy purpose"),(34,"Total"),
            (35,"Response with clarity"),(36,"Technical understand."),(37,"Conscious towards job role"),(38,"Total"),
            (39,""),(40,""),(41,""),
        ]
        for (col, label) in sub_cols:
            cell = ws.cell(row=ROW_SUBCOL, column=col)
            if label:  # only set value/font for actual sub-header cells
                cell.value = label
                cell.font = hdr_font(size=7)
                cell.fill = PatternFill('none')
                cell.alignment = Alignment(horizontal='center', vertical='bottom',
                                           wrap_text=False, text_rotation=90)
            cell.border = border

        # Row 7 borders for merged vertical cells (A7, B7, AM7, AN7, AO7)
        # These are bottom half of vertical merge — need L/R/B borders, no top
        for col_idx in [1, 2, 39, 40, 41]:
            c = ws.cell(row=ROW_SUBCOL, column=col_idx)
            c.border = Border(
                left=Side(style='thin', color='000000'),
                right=Side(style='thin', color='000000'),
                top=Side(style=None),
                bottom=Side(style='thin', color='000000')
            )

        # ── Row 8: Max marks ──────────────────────────────────────────
        max_row_vals = [None,None,2,5,8,15,3,2,5,10,3,3,4,10,1,2,2,5,4,3,3,10,4,3,3,10,3,5,2,10,7,3,5,15,7,5,3,15,100,None,None]
        for col, val in enumerate(max_row_vals, start=1):
            cell = ws.cell(row=ROW_MAX, column=col)
            if val is not None: cell.value = val
            cell.font = hdr_font(size=8, color="000000")
            cell.fill = PatternFill('none')
            cell.alignment = center_align()
            cell.border = border

        current_row = ROW_MAX + 1
        all_lo_totals = []

        # ── Use shared marks cache (generated once, reused in Annexure-II) ──
        trainee_cache = (marks_cache or {}).get(roll, {})

        for lo_num, lo_data in trainee_cache.items():
            pract_from  = lo_data['pract_from']
            pract_to    = lo_data['pract_to']
            lo_name     = lo_data['lo_name']
            pract_marks = lo_data['pract_marks']

            lo_grand_totals = []
            for pract_offset, pract_num in enumerate(range(pract_from, pract_to+1)):
                marks = pract_marks[pract_offset]
                lo_grand_totals.append(marks['grand_total'])
                row_vals = [
                    f"LO - {lo_num}", pract_num,
                    marks['dress'], marks['ppe'], marks['apply_safety'], marks['safety_total'],
                    marks['personal'], marks['scrap'], marks['material'], marks['hygiene_total'],
                    marks['initiative'], marks['accountability'], marks['participative'], marks['attendance_total'],
                    marks['select_manual'], marks['search_topic'], marks['read_manual'], marks['manuals_total'],
                    marks['plan_work'], marks['select_tools'], marks['review_work'], marks['knowledge_total'],
                    marks['handle_tools'], marks['safety_handling'], marks['care_maintain'], marks['skills_total'],
                    marks['sequence'], marks['technique'], marks['review_execution'], marks['speed_total'],
                    marks['accuracy'], marks['conform'], marks['satisfy'], marks['quality_total'],
                    marks['clarity'], marks['technical'], marks['conscious'], marks['viva_total'],
                    marks['grand_total'], "", "",
                ]
                for col, val in enumerate(row_vals, start=1):
                    cell = ws.cell(row=current_row, column=col)
                    cell.value = val
                    cell.font = cell_font(size=8)
                    cell.fill = PatternFill('none')
                    cell.alignment = center_align()
                    cell.border = border
                current_row += 1

            # LO summary row — merged A:AO, bold, centered, thin borders all sides
            lo_avg = round(sum(lo_grand_totals) / len(lo_grand_totals))
            all_lo_totals.append(lo_avg)
            ws.merge_cells(f'A{current_row}:AO{current_row}')
            cell = ws.cell(row=current_row, column=1)
            cell.value = f"{lo_name}     Average of LO{lo_num}  {lo_avg}"
            cell.font = Font(name='Calibri', size=8, bold=True, color='000000')
            cell.fill = PatternFill('none')
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = border
            ws.row_dimensions[current_row].height = 15
            for col_idx in range(1, 42):
                ws.cell(row=current_row, column=col_idx).border = border
            current_row += 1

        # Grand average row — merged A:AK, font 10pt bold, bottom border on all cells
        overall_avg = round(sum(all_lo_totals) / len(all_lo_totals)) if all_lo_totals else (target_avg or 0)
        ws.merge_cells(f'A{current_row}:AK{current_row}')
        cell = ws.cell(row=current_row, column=1)
        cell.value = f"Average of All LO     {overall_avg}"
        cell.font = Font(name='Calibri', size=10, bold=True, color='000000')
        cell.fill = PatternFill('none')
        cell.alignment = Alignment(horizontal='center', vertical='center')
        # First cell gets left+bottom border; interior cells get bottom only
        cell.border = Border(
            left=Side(style='thin', color='000000'),
            right=Side(style='thin', color='000000'),
            top=Side(style=None),
            bottom=Side(style='thin', color='000000')
        )
        for col_idx in range(2, 42):
            ws.cell(row=current_row, column=col_idx).border = Border(
                left=Side(style=None),
                right=Side(style='thin', color='000000') if col_idx == 41 else Side(style=None),
                top=Side(style=None),
                bottom=Side(style='thin', color='000000')
            )

        # ── Column widths (matching demo assessment) ──────────────────────
        ws.column_dimensions['A'].width = 5.43   # LO No
        ws.column_dimensions['B'].width = 2.71   # Practical No
        for col_idx in range(3, 42):
            ws.column_dimensions[get_column_letter(col_idx)].width = 3.5
        ws.column_dimensions['S'].width = 3.43   # Application of Knowledge col
        ws.column_dimensions['T'].width = 2.71
        ws.column_dimensions['AM'].width = 5.0   # Grand Total
        ws.column_dimensions['AN'].width = 7.29  # Signature Trainee

        # ── Row heights (matching demo assessment) ────────────────────────
        ws.row_dimensions[ROW_CAT].height = 71.25
        ws.row_dimensions[ROW_SUBCOL].height = 171.75
        ws.freeze_panes = 'A9'
    out=io.BytesIO(); wb.save(out); out.seek(0)
    return out

# ─── LO-WISE SUMMARY (ANNEXURE-II) GENERATION ────────────────────
# One sheet per LO, all trainees, category totals (max 100).
# Marks are derived from the SAME assessment logic used in create_excel —
# we re-run generate_row_marks with the same targets so the numbers are
# consistent. No existing logic is changed.

def create_lo_summary_excel(trainee_df, lo_df, user_info=None, marks_cache=None):
    """
    Build ANNEXURE-II style workbook: one sheet per LO.
    Each sheet lists all trainees with their 9 category totals and grand total.
    Layout mirrors reportsummary2.xlsx exactly.
    """
    ui = user_info or {}
    u_si_name            = ui.get('si_name', 'N/A')
    u_trade_name         = ui.get('trade_name', 'N/A')
    u_iti_name           = ui.get('iti_name', 'N/A')
    u_year_of_assessment = ui.get('year_of_assessment', '')
    u_assessment_location= ui.get('assessment_location', '')
    u_trade_duration     = ui.get('trade_duration', '')
    u_semester           = ui.get('semester', '')
    u_batch              = ui.get('batch', '')

    # ── Styles ────────────────────────────────────────────────────
    NAVY   = '0B1D3A'
    BLUE   = '1B4F8A'
    HDR2   = '2E75B6'
    YELLOW = 'FFF2CC'
    WHITE  = 'FFFFFF'
    LGRAY  = 'D9E1F2'
    ALT    = 'EBF3FB'
    thin   = Side(style='thin', color='000000')
    bdr    = Border(left=thin, right=thin, top=thin, bottom=thin)

    def _sf(size=9, bold=False, color='000000'):
        return Font(name='Calibri', size=size, bold=bold, color=color)

    def _hf(size=9, color='000000'):
        return Font(name='Calibri', size=size, bold=True, color=color)

    def _ca(h='center', v='center', wrap=True):
        return Alignment(horizontal=h, vertical=v, wrap_text=wrap)

    def _fill(hex_color):
        return PatternFill("none")

    def _s(ws, row, col, val, font=None, fill=None, align=None, border=None):
        c = ws.cell(row=row, column=col, value=val)
        if font:   c.font   = font
        if fill:   c.fill   = fill
        if align:  c.alignment = align
        if border: c.border = border
        return c

    def _ms(ws, row, c1, c2, val, font=None, fill=None, align=None, border=None):
        ws.merge_cells(f'{get_column_letter(c1)}{row}:{get_column_letter(c2)}{row}')
        _s(ws, row, c1, val, font, fill, align, border)
        for c in range(c1, c2 + 1):
            if border:
                ws.cell(row=row, column=c).border = border

    # ── Use shared marks cache ─────────────────────────────────────
    # Build trainee_lo_marks from cache (same marks as create_excel)
    lo_rows_all = list(lo_df.iterrows())
    trainee_lo_marks = {}
    for _, trainee in trainee_df.iterrows():
        roll = trainee['rollno']
        trainee_lo_marks[roll] = {}
        trainee_cache = (marks_cache or {}).get(roll, {})
        for lo_num, lo_data in trainee_cache.items():
            trainee_lo_marks[roll][lo_num] = lo_data['cat_avgs']

    # ── Build workbook ─────────────────────────────────────────────
    wb = Workbook()
    wb.remove(wb.active)

    # Collect all unique LOs
    all_lo_numbers = []
    seen = set()
    for _, lo in lo_rows_all:
        lo_num = lo['lo']
        if lo_num not in seen:
            seen.add(lo_num)
            all_lo_numbers.append(lo)

    for lo_row in all_lo_numbers:
        lo_num  = lo_row['lo']
        lo_name = lo_row['lo_name']

        ws = wb.create_sheet(title=f'LO_{lo_num}')

        # ── Row 1: ANNEXURE-II ────────────────────────────────────
        _ms(ws, 1, 1, 14, 'ANNEXURE-II',
            font=_hf(12, '000000'), fill=_fill('EBF3FB'), align=_ca(), border=bdr)

        # ── Row 2: Internal Assessment ───────────────────────────
        _ms(ws, 2, 1, 14, 'Internal Assessment',
            font=_hf(11, '000000'), fill=_fill('EBF3FB'), align=_ca(), border=bdr)

        # ── Rows 3-6: Header info (matching template merges) ─────
        lbl_font = _sf(9, bold=True, color='000000')
        val_font = _sf(9, bold=False, color='1A1A1A')
        lbl_fill = _fill(LGRAY)
        val_fill = _fill('FAFBFF')

        # Row 3: Assessor / Year of Enrolment
        _ms(ws, 3, 1, 3,  'Name & Address of the Assessor', font=lbl_font, fill=lbl_fill, align=_ca('left'), border=bdr)
        _ms(ws, 3, 4, 7,  u_si_name,                        font=val_font, fill=val_fill, align=_ca('left'), border=bdr)
        _ms(ws, 3, 8, 11, 'Year of Enrolment',              font=lbl_font, fill=lbl_fill, align=_ca('left'), border=bdr)
        _ms(ws, 3, 12, 14, u_year_of_assessment,            font=val_font, fill=val_fill, align=_ca(),       border=bdr)

        # Row 4: ITI / Date of Assessment
        _ms(ws, 4, 1, 3,  'Name & Address of ITI (Govt/Pvt)', font=lbl_font, fill=lbl_fill, align=_ca('left'), border=bdr)
        _ms(ws, 4, 4, 7,  u_iti_name,                          font=val_font, fill=val_fill, align=_ca('left'), border=bdr)
        _ms(ws, 4, 8, 11, 'Date of Assessment',                font=lbl_font, fill=lbl_fill, align=_ca('left'), border=bdr)
        _ms(ws, 4, 12, 14, '',                                  font=val_font, fill=val_fill, align=_ca(),       border=bdr)

        # Row 5: Industry / Assessment Location
        _ms(ws, 5, 1, 3,  'Name & Address of the Industry', font=lbl_font, fill=lbl_fill, align=_ca('left'), border=bdr)
        _ms(ws, 5, 4, 7,  u_iti_name,                        font=val_font, fill=val_fill, align=_ca('left'), border=bdr)
        _ms(ws, 5, 8, 11, 'Assessment Location',             font=lbl_font, fill=lbl_fill, align=_ca('left'), border=bdr)
        _ms(ws, 5, 12, 14, u_assessment_location,            font=val_font, fill=val_fill, align=_ca(),       border=bdr)

        # Row 6: Trade / Duration / Examination
        _ms(ws, 6, 1, 2,   'Trade Name',     font=lbl_font, fill=lbl_fill, align=_ca('left'), border=bdr)
        _ms(ws, 6, 3, 7,   u_trade_name,    font=val_font, fill=val_fill, align=_ca('left'), border=bdr)
        _ms(ws, 6, 8, 10,  'Duration Of Trade', font=lbl_font, fill=lbl_fill, align=_ca('left'), border=bdr)
        _ms(ws, 6, 11, 11, u_trade_duration, font=val_font, fill=val_fill, align=_ca(), border=bdr)
        _ms(ws, 6, 12, 13, 'Examination',   font=lbl_font, fill=lbl_fill, align=_ca('left'), border=bdr)
        _s(ws,  6, 14,     u_semester,      font=val_font, fill=val_fill, align=_ca(), border=bdr)

        # ── Row 7: Learning Outcome label ────────────────────────
        _ms(ws, 7, 1, 14, f'Learning Outcome :{lo_num}  —  {lo_name}',
            font=_hf(10, '000000'), fill=_fill(BLUE), align=_ca('left'), border=bdr)

        # ── Row 8: Max marks row ──────────────────────────────────
        mx_font = _hf(8, '000000')
        mx_fill = _fill(YELLOW)
        _ms(ws, 8, 1, 3, 'Maximum Marks (Total 100 Marks)', font=mx_font, fill=mx_fill, align=_ca(), border=bdr)
        for col, val in [(4,'15'),(5,'10'),(6,'10'),(7,'5'),(8,'10'),(9,'10'),(10,'10'),(11,'15'),(12,'15'),(13,''),(14,'')]:
            _s(ws, 8, col, val, font=mx_font, fill=mx_fill, align=_ca(), border=bdr)

        # ── Row 9: Column headers ─────────────────────────────────
        col_headers = [
            (1,  'Roll\nNo'),
            (2,  'Firstname'),
            (3,  'Fathername'),
            (4,  'Safety consciousness\n(Max 15)'),
            (5,  'Workplace hygiene & Economical use\nof materials\n(Max 10)'),
            (6,  'Attendance/\nPunctuality\n(Max 10)'),
            (7,  'Ability to follow Manuals/\nWritten instructions\n(Max 5)'),
            (8,  'Application of\nKnowledge\n(Max 10)'),
            (9,  'Skills to handle\ntools & equipment\n(Max 10)'),
            (10, 'Speed in\ndoing work\n(Max 10)'),
            (11, 'Quality in\nworkmanship\n(Max 15)'),
            (12, 'VIVA\n(Max 15)'),
            (13, 'Total Internal\nassessment\nMarks (Max 100)'),
            (14, 'Result\n(Y/N)'),
        ]
        hdr_fill = _fill(BLUE)
        for col, label in col_headers:
            _s(ws, 9, col, label,
               font=_hf(8, '000000'), fill=hdr_fill,
               align=_ca(), border=bdr)

        # ── Data rows: one row per trainee ────────────────────────
        for idx, (_, trainee) in enumerate(trainee_df.iterrows()):
            r        = 10 + idx
            row_fill = _fill(WHITE) if idx % 2 == 0 else _fill(ALT)
            roll     = trainee['rollno']
            fname    = str(trainee.get('Firstname', '') or '')
            fathname = str(trainee.get('Fathername', '') or '')

            # Get pre-computed LO averages for this trainee + LO
            lo_data = trainee_lo_marks.get(roll, {}).get(lo_num)

            if lo_data:
                s  = lo_data['safety_total']
                h  = lo_data['hygiene_total']
                a  = lo_data['attendance_total']
                mn = lo_data['manuals_total']
                k  = lo_data['knowledge_total']
                sk = lo_data['skills_total']
                sp = lo_data['speed_total']
                q  = lo_data['quality_total']
                v  = lo_data['viva_total']
                gt = lo_data['grand_total']
                result = 'Y' if gt >= 40 else 'N'
            else:
                s = h = a = mn = k = sk = sp = q = v = gt = ''
                result = ''

            row_vals = [roll, fname, fathname, s, h, a, mn, k, sk, sp, q, v, gt, result]
            for col, val in enumerate(row_vals, start=1):
                font = _sf(9, bold=(col == 13), color='000000')
                _s(ws, r, col, val, font=font, fill=row_fill, align=_ca(), border=bdr)

        # ── Column widths (matching template) ─────────────────────
        widths = {1:6, 2:16, 3:15, 4:7, 5:7, 6:7, 7:7, 8:7, 9:7, 10:7, 11:7, 12:7, 13:9, 14:7}
        for col, w in widths.items():
            ws.column_dimensions[get_column_letter(col)].width = w

        # Row heights
        ws.row_dimensions[1].height = 18
        ws.row_dimensions[2].height = 18
        ws.row_dimensions[7].height = 20
        ws.row_dimensions[9].height = 60
        ws.freeze_panes = 'A10'

    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    return out


# ─── PER-TRAINEE PER-LO REPORT (Proforma replication) ────────────
def create_trainee_lo_report(trainee_df, lo_df, user_info=None, marks_cache=None):
    """
    Build a workbook with one sheet per trainee per LO.
    Layout exactly replicates the demo proforma (demo_per_trainee_per_lo.xlsx):
      Row 1     : Title (merged A1:AN1)
      Rows 2-5  : Info rows - each split into 6 segments matching demo merges
      Row 6     : Category group headers (A6:A8 and B6:B8 merged vertically, LO/Pract rotated 90)
      Row 7     : Sub-column headers (all rotated 90, NO wrap_text)
      Row 8     : Max marks row
      Rows 9+   : Data rows (one per practical)
      Last row  : LO average summary (merged A:AN)
    No fill colors — plain white background like demo.
    """
    ui = user_info or {}
    u_si_name            = ui.get('si_name', '')
    u_trade_name         = ui.get('trade_name', '')
    u_iti_name           = ui.get('iti_name', '')
    u_year_of_assessment = ui.get('year_of_assessment', '')
    u_assessment_location= ui.get('assessment_location', '')
    u_trade_duration     = ui.get('trade_duration', '')
    u_semester           = ui.get('semester', '')
    u_batch              = ui.get('batch', '')

    thin = Side(style='thin', color='000000')
    bdr  = Border(left=thin, right=thin, top=thin, bottom=thin)
    no_fill = PatternFill('none')

    def _sf(size=8, bold=False):
        return Font(name='Calibri', size=size, bold=bold, color='000000')

    def _al(h='center', v='center', wrap=True, rot=0):
        return Alignment(horizontal=h, vertical=v, wrap_text=wrap, text_rotation=rot)

    def _s(ws, row, col, val, font=None, align=None):
        c = ws.cell(row=row, column=col, value=val)
        if font:  c.font  = font
        if align: c.alignment = align
        c.border = bdr
        c.fill   = no_fill
        return c

    def _ms(ws, row, c1, c2, val, font=None, align=None):
        """Merge cols c1..c2 in row, set value and apply border to all cells."""
        if c1 != c2:
            ws.merge_cells(f'{get_column_letter(c1)}{row}:{get_column_letter(c2)}{row}')
        ws.cell(row=row, column=c1).value = val
        if font:  ws.cell(row=row, column=c1).font  = font
        if align: ws.cell(row=row, column=c1).alignment = align
        ws.cell(row=row, column=c1).fill = no_fill
        for c in range(c1, c2 + 1):
            ws.cell(row=row, column=c).border = bdr
            ws.cell(row=row, column=c).fill   = no_fill

    def _ms_vt(ws, r1, r2, col, val, font=None, align=None):
        """Merge rows r1..r2 in single column, set value."""
        ws.merge_cells(f'{get_column_letter(col)}{r1}:{get_column_letter(col)}{r2}')
        ws.cell(row=r1, column=col).value = val
        if font:  ws.cell(row=r1, column=col).font  = font
        if align: ws.cell(row=r1, column=col).alignment = align
        ws.cell(row=r1, column=col).fill = no_fill
        for r in range(r1, r2 + 1):
            ws.cell(row=r, column=col).border = bdr
            ws.cell(row=r, column=col).fill   = no_fill

    wb = Workbook()
    wb.remove(wb.active)

    for _, trainee in trainee_df.iterrows():
        roll      = trainee['rollno']
        fname     = str(trainee.get('Firstname', '') or '')
        fathrname = str(trainee.get('Fathername', '') or '')
        lname     = str(trainee.get('Lastname', '') or '')
        full_name = f"{fname} {fathrname} {lname}".strip()

        trainee_cache = (marks_cache or {}).get(roll, {})

        for lo_num, lo_data in trainee_cache.items():
            pract_from  = lo_data['pract_from']
            pract_to    = lo_data['pract_to']
            lo_name     = lo_data['lo_name']
            pract_marks = lo_data['pract_marks']
            lo_avg      = lo_data['lo_avg']

            sheet_name = f"R{roll}_LO{lo_num}"[:31]
            ws = wb.create_sheet(title=sheet_name)

            # ── Column widths — exact from demo ───────────────────
            demo_widths = {
                1:8.42578125, 2:5.0, 3:3.42578125, 4:2.85546875, 5:3.140625,
                6:3.42578125, 7:4.5703125, 8:4.7109375, 9:4.0, 10:3.140625,
                11:3.42578125, 12:2.85546875, 13:3.42578125, 14:3.42578125,
                15:2.85546875, 16:3.28515625, 17:3.140625, 18:3.0,
                19:3.140625, 20:4.28515625, 21:3.28515625, 22:3.42578125,
                23:5.140625, 24:3.0, 25:2.85546875, 26:3.28515625,
                27:4.85546875, 28:4.5703125, 29:3.85546875, 30:4.28515625,
                31:3.0, 32:3.140625, 33:4.42578125, 34:3.0,
                35:3.140625, 36:4.140625, 37:4.42578125, 38:5.5703125,
                39:8.0, 40:9.0
            }
            for col, w in demo_widths.items():
                ws.column_dimensions[get_column_letter(col)].width = w

            # ── Row 1: Title merged A1:AN1 ────────────────────────
            _ms(ws, 1, 1, 40, 'Internal Assessment',
                font=_sf(12, bold=True),
                align=_al('center', 'center', wrap=True))

            # ── Rows 2-5: Info segments matching demo merges exactly ──
            # Demo merges per row:
            #   cols 1-3 (label), 4-17 (value), 18-22 (label), 23-27 (value), 28-34 (label), 35-40 (value)
            lbl_f = _sf(8, bold=True)
            val_f = _sf(8, bold=False)
            lbl_a = _al('left',   'center', wrap=True)
            val_l = _al('left',   'center', wrap=True)
            val_c = _al('center', 'center', wrap=True)

            # Row 2: Name of Trainee / Roll NO / Year of Enrollment / Sem
            _ms(ws, 2, 1,  3,  'Name of Trainee:',   font=lbl_f, align=lbl_a)
            _ms(ws, 2, 4,  17, full_name,             font=val_f, align=val_l)
            _ms(ws, 2, 18, 22, 'Roll NO:',            font=lbl_f, align=lbl_a)
            _ms(ws, 2, 23, 27, str(roll),             font=val_f, align=val_c)
            _ms(ws, 2, 28, 34, 'Year of Enrollment:', font=val_f, align=val_c)
            _ms(ws, 2, 35, 40, u_year_of_assessment,  font=val_f, align=val_c)

            # Row 3: Name of ITI / Date of Assessment / Batch
            _ms(ws, 3, 1,  3,  'Name of ITI:',        font=lbl_f, align=lbl_a)
            _ms(ws, 3, 4,  17, u_iti_name,             font=val_f, align=val_l)
            _ms(ws, 3, 18, 22, 'Date of Assessment:',  font=lbl_f, align=lbl_a)
            _ms(ws, 3, 23, 27, '',                     font=val_f, align=val_c)
            _ms(ws, 3, 28, 34, 'Batch:',               font=val_f, align=val_c)
            _ms(ws, 3, 35, 40, u_batch,                font=val_f, align=val_c)

            # Row 4: Name of Industry / Assessment Location / Sem
            _ms(ws, 4, 1,  3,  'Name of the Industry:', font=lbl_f, align=lbl_a)
            _ms(ws, 4, 4,  17, u_iti_name,               font=val_f, align=val_l)
            _ms(ws, 4, 18, 22, 'Assessment Location:',   font=lbl_f, align=lbl_a)
            _ms(ws, 4, 23, 27, u_assessment_location,    font=val_f, align=val_c)
            _ms(ws, 4, 28, 34, 'Sem:',                   font=val_f, align=val_c)
            _ms(ws, 4, 35, 40, u_semester,               font=val_f, align=val_c)

            # Row 5: Trade Name / Duration / SI Name
            _ms(ws, 5, 1,  3,  'Trade Name:',            font=lbl_f, align=lbl_a)
            _ms(ws, 5, 4,  17, u_trade_name,              font=val_f, align=val_l)
            _ms(ws, 5, 18, 22, 'Duration of the Trade:',  font=lbl_f, align=lbl_a)
            _ms(ws, 5, 23, 27, u_trade_duration,           font=val_f, align=val_c)
            _ms(ws, 5, 28, 34, 'S.I. Name:',              font=val_f, align=val_l)
            _ms(ws, 5, 35, 40, u_si_name,                  font=val_f, align=val_c)

            # ── Row 6: Category headers + vertical merges for A6:A8 and B6:B8 ──
            ws.row_dimensions[6].height = 22.5
            ws.row_dimensions[7].height = 150.75
            ws.row_dimensions[8].height = 15.0

            # A6:A8 merged — "Learning\nOutcome\nNumber" rotated 90
            _ms_vt(ws, 6, 8, 1,
                   'Learning\nOutcome\nNumber',
                   font=_sf(7, bold=False),
                   align=_al('center', None, wrap=None, rot=90))

            # B6:B8 merged — "Practical /\nProfessional\nSkill Number" rotated 90
            _ms_vt(ws, 6, 8, 2,
                   'Practical /\nProfessional\nSkill Number',
                   font=_sf(7, bold=False),
                   align=_al('center', None, wrap=None, rot=90))

            # Category span headers (row 6, cols 3-38) — wrap=True, NOT rotated
            cat_hf = _sf(8, bold=False)
            cat_al = _al('center', 'center', wrap=True, rot=0)
            _ms(ws, 6, 3,  6,  'Safety\nconsciousness',                           font=cat_hf, align=cat_al)
            _ms(ws, 6, 7,  10, 'Workplace hygiene &\nEconomical use of materials', font=cat_hf, align=cat_al)
            _ms(ws, 6, 11, 14, 'Attendance/\nPunctuality',                         font=cat_hf, align=cat_al)
            _ms(ws, 6, 15, 18, 'Ability to follow Manuals/\nWritten instructions', font=cat_hf, align=cat_al)
            _ms(ws, 6, 19, 22, 'Application of\nKnowledge',                        font=cat_hf, align=cat_al)
            _ms(ws, 6, 23, 26, 'Skills to handle\ntools & equipment',              font=cat_hf, align=cat_al)
            _ms(ws, 6, 27, 30, 'Speed in\ndoing work',                             font=cat_hf, align=cat_al)
            _ms(ws, 6, 31, 34, 'Quality in\nworkmanship',                          font=cat_hf, align=cat_al)
            _ms(ws, 6, 35, 38, 'VIVA',                                             font=cat_hf, align=cat_al)

            # AM6:AM7 merged — Grand Total rotated 90
            _ms_vt(ws, 6, 7, 39,
                   'Grand\nTotal',
                   font=_sf(8, bold=False),
                   align=_al('center', None, wrap=None, rot=90))

            # AN6:AN7 merged — Signature of Trainee rotated 90
            _ms_vt(ws, 6, 7, 40,
                   'Signature\nof Trainee',
                   font=_sf(8, bold=False),
                   align=_al('center', None, wrap=None, rot=90))

            # ── Row 7: Sub-column headers — ALL rotated 90, NO wrap_text ──
            sub_f  = _sf(7, bold=False)
            sub_al = _al('center', None, wrap=None, rot=90)
            sub_data = [
                (3,  'Dress\ncode'),
                (4,  'Use\nPPE'),
                (5,  'Apply/\npractice\nsafety'),
                (6,  'Total'),
                (7,  'Maintain\npersonal &\nworkplace\ncleanliness'),
                (8,  'Dispose\nscrap as per\nstandard\npractice'),
                (9,  'Select\nappropriate\nmaterial &\nminimize\nwastage'),
                (10, 'Total'),
                (11, 'Initiative'),
                (12, 'Account-\nability'),
                (13, 'Participative\nin work'),
                (14, 'Total'),
                (15, 'Select\nright\nmanual'),
                (16, 'Search for\nappropriate\ntopic'),
                (17, 'Read &\ninterpret\nthe manual'),
                (18, 'Total'),
                (19, 'Plan\nthe\nwork'),
                (20, 'Select\nappropriate\ntools &\nequipment'),
                (21, 'Review\nthe\nwork'),
                (22, 'Total'),
                (23, 'Handle &\nuse tools &\nequipment'),
                (24, 'Maintain\nsafety in\nhandling'),
                (25, 'Care &\nmaintain'),
                (26, 'Total'),
                (27, 'Properly\nsequence\nthe work'),
                (28, 'Use\nappropriate\ntechnique'),
                (29, 'Review the\nwork during\nexecution'),
                (30, 'Total'),
                (31, 'Achieve\nwork with\nhigh\naccuracy'),
                (32, 'Conform\nto\nrequirement'),
                (33, 'Satisfy\nthe\npurpose'),
                (34, 'Total'),
                (35, 'Response\nwith\nclarity'),
                (36, 'Technical\nunderstand.'),
                (37, 'Conscious\ntowards\njob role'),
                (38, 'Total'),
            ]
            for col, label in sub_data:
                _s(ws, 7, col, label, font=sub_f, align=sub_al)
            # Cols 1,2,39,40 in row 7 are already covered by vertical merges — just ensure borders
            for col in [1, 2, 39, 40]:
                ws.cell(row=7, column=col).border = bdr
                ws.cell(row=7, column=col).fill   = no_fill

            # ── Row 8: Max marks ───────────────────────────────────
            max_f  = _sf(8, bold=False)
            max_al = _al('center', 'center', wrap=True, rot=0)
            max_vals = [None, None, 2,5,8,15, 3,2,5,10, 3,3,4,10, 1,2,2,5, 4,3,3,10, 4,3,3,10, 3,5,2,10, 7,3,5,15, 7,5,3,15, 100, None]
            for col, val in enumerate(max_vals, start=1):
                _s(ws, 8, col, val, font=max_f, align=max_al)

            # ── Data rows ─────────────────────────────────────────
            data_f  = _sf(8, bold=False)
            data_al = _al('center', 'center', wrap=True, rot=0)
            for pract_offset, pract_num in enumerate(range(pract_from, pract_to + 1)):
                r = 9 + pract_offset
                m = pract_marks[pract_offset]
                row_vals = [
                    f'LO - {lo_num}', pract_num,
                    m['dress'], m['ppe'], m['apply_safety'], m['safety_total'],
                    m['personal'], m['scrap'], m['material'], m['hygiene_total'],
                    m['initiative'], m['accountability'], m['participative'], m['attendance_total'],
                    m['select_manual'], m['search_topic'], m['read_manual'], m['manuals_total'],
                    m['plan_work'], m['select_tools'], m['review_work'], m['knowledge_total'],
                    m['handle_tools'], m['safety_handling'], m['care_maintain'], m['skills_total'],
                    m['sequence'], m['technique'], m['review_execution'], m['speed_total'],
                    m['accuracy'], m['conform'], m['satisfy'], m['quality_total'],
                    m['clarity'], m['technical'], m['conscious'], m['viva_total'],
                    m['grand_total'], '',
                ]
                for col, val in enumerate(row_vals, start=1):
                    _s(ws, r, col, val, font=data_f, align=data_al)

            # ── Average of LO row (merged A:AN) ───────────────────
            avg_row = 9 + (pract_to - pract_from + 1)
            _ms(ws, avg_row, 1, 40,
                f'{lo_name}     Average of LO{lo_num}  {lo_avg}',
                font=_sf(8, bold=False),
                align=_al('center', 'center', wrap=True))
            ws.row_dimensions[avg_row].height = 15

            ws.freeze_panes = 'A9'

    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    return out



# ─── PROGRESS CARD GENERATION ────────────────────────────────────

def _is_2_year_trade(user_info):
    """Return True ONLY if the registered trade duration is explicitly 2 years.
    Defaults to False (1-year proforma) for any ambiguous or missing value."""
    dur = str((user_info or {}).get('trade_duration', '')).strip().lower()
    return dur in ('2', '2 year', '2 years', '2year', '2years', 'two year', 'two years')


def create_progress_card(trainee_df, user_info=None):
    """Dispatcher: 1-year → create_progress_card_1year, 2-year → create_progress_card_2year."""
    if _is_2_year_trade(user_info):
        return create_progress_card_2year(trainee_df, user_info)
    return create_progress_card_1year(trainee_df, user_info)


def _build_progress_card_sheet(ws, tr, ui, year2=False):
    """
    Write one progress-card sheet into ws, faithful to the uploaded proformas.
    year2=False → 1-year layout (18 rows, 11 cols)
    year2=True  → 2-year layout (29 rows, 11 cols)
    All borders are thin; all merges match the proforma exactly.
    """
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    from openpyxl.utils import get_column_letter as gcl

    THIN = Side(style='thin', color='000000')
    NO   = Side(style=None)

    def T(): return Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
    def bdr(l=0,r=0,t=0,b=0):
        return Border(
            left=THIN if l else NO, right=THIN if r else NO,
            top=THIN if t else NO,  bottom=THIN if b else NO)

    def fnt(sz=11, bold=False):
        return Font(name='Calibri', size=sz, bold=bold)

    def aln(h='center', v='center', wrap=False):
        return Alignment(horizontal=h, vertical=v, wrap_text=wrap)

    def sc(row, col, val=None, sz=11, bold=False, h='center', v='center', wrap=False, border=None):
        c = ws.cell(row=row, column=col, value=val)
        c.font = fnt(sz, bold)
        c.alignment = aln(h, v, wrap)
        if border: c.border = border
        return c

    def mr(r1,c1,r2,c2):
        ws.merge_cells(f'{gcl(c1)}{r1}:{gcl(c2)}{r2}')

    def fill_row_thin(row, c1, c2, top=False, bottom=False, left_col=None, right_col=None, v='center'):
        """Apply thin border to every cell in row from c1 to c2 with given sides."""
        for c in range(c1, c2+1):
            l = (c == c1 and left_col is None) or c == left_col
            r = (c == c2 and right_col is None) or c == right_col
            ws.cell(row=row, column=c).border = bdr(l=l, r=r, t=top, b=bottom)
            ws.cell(row=row, column=c).alignment = aln(v=v)

    def all_thin(row, c1, c2):
        for c in range(c1, c2+1):
            ws.cell(row=row, column=c).border = T()
            ws.cell(row=row, column=c).alignment = aln(v='center')

    # ── Helper values ────────────────────────────────────────────────
    def fmt_date(v):
        if v is None: return ''
        try:
            import math
            if isinstance(v, float) and math.isnan(v): return ''
        except Exception: pass
        if hasattr(v, 'strftime'): return v.strftime('%d-%m-%Y')
        return str(v)

    def to_int(v):
        try: return int(float(v))
        except Exception: return '' if (v is None or (isinstance(v,float) and __import__('math').isnan(v))) else v

    roll    = tr['rollno']
    nm      = ' '.join([str(tr.get(f,'') or '') for f in ['Firstname','Fathername','Lastname']]).strip()
    adm_dt  = fmt_date(tr.get('DateofAdmission'))
    dob_dt  = fmt_date(tr.get('Birth_Date'))
    leav_dt = fmt_date(tr.get('DateofLeaving'))
    edu_col = next((c for c in tr.index if 'Edu' in str(c)), None)
    edu     = str(tr.get(edu_col,'') or '') if edu_col else ''
    es      = to_int(tr.get('ES',''))
    wsc     = to_int(tr.get('WC_SC',''))
    ed      = to_int(tr.get('ED',''))
    lo70    = to_int(tr.get('All_LO_average_base_on_70_SEM_I',''))
    tot     = to_int(tr.get('Total',''))
    wday    = to_int(tr.get('Working_days',''))
    aday    = to_int(tr.get('Attendeate_days',''))
    try:    att_pct = round(float(aday)/float(wday)*100)
    except: att_pct = ''

    # Year-level attendance totals are ALWAYS blank — they equal H1+H2 which
    # the user fills manually once both halves are complete. FA Total (out of 200)
    # is also left blank for the same reason.
    yr_att_act = ''
    yr_att_pos = ''
    yr_att_pct = ''
    yr_fa_tot  = ''

    u_iti      = ui.get('iti_name','')
    u_trade    = ui.get('trade_name','')
    u_batch    = ui.get('batch','')
    u_semester = ui.get('semester','H1')   # which half the user is currently filing
    u_si       = ui.get('si_name','')

    # Route the current semester's data to the correct H1 or H2 row.
    # Proforma rows are always labelled H1 (row 14) and H2 (row 15) — fixed.
    sem_upper = u_semester.strip().upper()
    h1_act = aday    if sem_upper == 'H1' else ''
    h1_pos = wday    if sem_upper == 'H1' else ''
    h1_pct = att_pct if sem_upper == 'H1' else ''
    h1_lo  = lo70    if sem_upper == 'H1' else ''
    h1_es  = es      if sem_upper == 'H1' else ''
    h1_wsc = wsc     if sem_upper == 'H1' else ''
    h1_ed  = ed      if sem_upper == 'H1' else ''
    # GLO total = ES + WC_SC + ED (out of 30)
    try:    h1_glo = int(es) + int(wsc) + int(ed)  if sem_upper == 'H1' else ''
    except: h1_glo = ''                             if sem_upper != 'H1' else (es if es != '' else '')
    # Total out of 100 = SLO (A) + GLO (B)
    try:    h1_tot = int(lo70) + int(h1_glo)        if sem_upper == 'H1' else ''
    except: h1_tot = ''

    h2_act = aday    if sem_upper == 'H2' else ''
    h2_pos = wday    if sem_upper == 'H2' else ''
    h2_pct = att_pct if sem_upper == 'H2' else ''
    h2_lo  = lo70    if sem_upper == 'H2' else ''
    h2_es  = es      if sem_upper == 'H2' else ''
    h2_wsc = wsc     if sem_upper == 'H2' else ''
    h2_ed  = ed      if sem_upper == 'H2' else ''
    try:    h2_glo = int(es) + int(wsc) + int(ed)  if sem_upper == 'H2' else ''
    except: h2_glo = ''                             if sem_upper != 'H2' else (es if es != '' else '')
    try:    h2_tot = int(lo70) + int(h2_glo)        if sem_upper == 'H2' else ''
    except: h2_tot = ''

    # ══ ROW 1: Institute title — A1:N1 ══════════════════════════════
    mr(1,1,1,13)
    sc(1,1, f'INDUSTRIAL TRAINING INSTITUTE : {u_iti.upper()}', sz=14, bold=True, border=T())
    for c in range(2,14): ws.cell(1,c).border = bdr(t=1,b=1,r=(c==13))
    ws.row_dimensions[1].height = 18.75

    # ══ ROW 2: PROGRESS CARD — A2:N2 ════════════════════════════════
    mr(2,1,2,13)
    sc(2,1,'PROGRESS CARD', sz=14, bold=True, border=T())
    for c in range(2,14): ws.cell(2,c).border = bdr(t=1,b=1,r=(c==13))
    ws.row_dimensions[2].height = 18.75

    # ══ ROWS 3-4: ROLL NO / TRADE / (blank) / BATCH ═════════════════
    # A3:B4  ROLL NO.
    mr(3,1,4,2)
    sc(3,1,'ROLL NO.', border=T())
    ws.cell(3,2).border = bdr(r=1,t=1)
    ws.cell(4,1).border = bdr(l=1,b=1)
    ws.cell(4,2).border = bdr(r=1,b=1)

    # C3:C4  roll value
    mr(3,3,4,3)
    sc(3,3, str(roll), border=T())
    ws.cell(4,3).border = bdr(l=1,r=1,b=1)

    # D3:E4  TRADE
    mr(3,4,4,5)
    sc(3,4,'TRADE', border=T())
    ws.cell(3,5).border = bdr(r=1,t=1)
    ws.cell(4,4).border = bdr(l=1,b=1)
    ws.cell(4,5).border = bdr(r=1,b=1)

    # F3:K4  trade value
    mr(3,6,4,11)
    sc(3,6, u_trade, border=T())
    for c in range(7,12): ws.cell(3,c).border = bdr(t=1,r=(c==11))
    ws.cell(4,6).border = bdr(l=1,b=1)
    for c in range(7,12): ws.cell(4,c).border = bdr(b=1,r=(c==11))

    # L3:L4  BATCH
    mr(3,12,4,12)
    sc(3,12,'BATCH', border=T())
    ws.cell(4,12).border = bdr(l=1,r=1,b=1)

    # M3:M4  batch value
    mr(3,13,4,13)
    sc(3,13, str(u_batch), border=T())
    ws.cell(4,13).border = bdr(l=1,r=1,b=1)

    # ══ ROWS 5-6: NAME OF TRAINEE ═══════════════════════════════════
    # A5:B6  label
    mr(5,1,6,2)
    sc(5,1,'NAME OF TRAINEE', border=T())
    ws.cell(5,2).border = bdr(r=1,t=1)
    ws.cell(6,1).border = bdr(l=1,b=1)
    ws.cell(6,2).border = bdr(r=1,b=1)

    # C5:M6  name value
    mr(5,3,6,13)
    sc(5,3, nm, border=T())
    for c in range(4,14): ws.cell(5,c).border = bdr(t=1,r=(c==13))
    ws.cell(6,3).border = bdr(l=1,b=1)
    for c in range(4,14): ws.cell(6,c).border = bdr(b=1,r=(c==13))

    # ══ ROWS 7-8: DATE OF ADMISSION / DATE OF BIRTH ═════════════════
    # A7:C8  label
    mr(7,1,8,3)
    sc(7,1,'DATE OF ADMISSION', border=T())
    ws.cell(7,2).border = bdr(t=1); ws.cell(7,3).border = bdr(r=1,t=1)
    ws.cell(8,1).border = bdr(l=1,b=1); ws.cell(8,2).border = bdr(b=1); ws.cell(8,3).border = bdr(r=1,b=1)

    # D7:G8  admission date value (wider)
    mr(7,4,8,7)
    sc(7,4, adm_dt, border=T())
    ws.cell(7,5).border = bdr(t=1); ws.cell(7,6).border = bdr(t=1); ws.cell(7,7).border = bdr(r=1,t=1)
    ws.cell(8,4).border = bdr(l=1,b=1); ws.cell(8,5).border = bdr(b=1); ws.cell(8,6).border = bdr(b=1); ws.cell(8,7).border = bdr(r=1,b=1)

    # H7:J8  DATE OF BIRTH label
    mr(7,8,8,10)
    sc(7,8,'DATE OF BIRTH', border=T())
    ws.cell(7,9).border = bdr(t=1); ws.cell(7,10).border = bdr(r=1,t=1)
    ws.cell(8,8).border = bdr(l=1,b=1); ws.cell(8,9).border = bdr(b=1); ws.cell(8,10).border = bdr(r=1,b=1)

    # K7:M8  dob value
    mr(7,11,8,13)
    sc(7,11, dob_dt, border=T())
    ws.cell(7,12).border = bdr(t=1); ws.cell(7,13).border = bdr(r=1,t=1)
    ws.cell(8,11).border = bdr(l=1,b=1); ws.cell(8,12).border = bdr(b=1); ws.cell(8,13).border = bdr(r=1,b=1)

    # ══ ROWS 9-10: DATE OF LEAVING / EDU. QUA ═══════════════════════
    # A9:C10
    mr(9,1,10,3)
    sc(9,1,'DATE OF LEAVING', border=T())
    ws.cell(9,2).border = bdr(t=1); ws.cell(9,3).border = bdr(r=1,t=1)
    ws.cell(10,1).border = bdr(l=1,b=1); ws.cell(10,2).border = bdr(b=1); ws.cell(10,3).border = bdr(r=1,b=1)

    # D9:G10  leaving date value
    mr(9,4,10,7)
    sc(9,4, leav_dt, border=T())
    ws.cell(9,5).border = bdr(t=1); ws.cell(9,6).border = bdr(t=1); ws.cell(9,7).border = bdr(r=1,t=1)
    ws.cell(10,4).border = bdr(l=1,b=1); ws.cell(10,5).border = bdr(b=1); ws.cell(10,6).border = bdr(b=1); ws.cell(10,7).border = bdr(r=1,b=1)

    # H9:J10  EDU. QUA label
    mr(9,8,10,10)
    sc(9,8,'EDU. QUA', border=T())
    ws.cell(9,9).border = bdr(t=1); ws.cell(9,10).border = bdr(r=1,t=1)
    ws.cell(10,8).border = bdr(l=1,b=1); ws.cell(10,9).border = bdr(b=1); ws.cell(10,10).border = bdr(r=1,b=1)

    # K9:M10  edu value
    mr(9,11,10,13)
    sc(9,11, edu, border=T())
    ws.cell(9,12).border = bdr(t=1); ws.cell(9,13).border = bdr(r=1,t=1)
    ws.cell(10,11).border = bdr(l=1,b=1); ws.cell(10,12).border = bdr(b=1); ws.cell(10,13).border = bdr(r=1,b=1)

    # ══ ROW 11: HALF YEARLY ASSESSMENT | FORMATIVE ASSESSMENT ════════
    # A11:D11
    mr(11,1,11,4)
    sc(11,1,'HALF YEARLY ASSESSMENT', bold=True, border=T())
    ws.cell(11,2).border = bdr(t=1,b=1); ws.cell(11,3).border = bdr(t=1,b=1)
    ws.cell(11,4).border = bdr(r=1,t=1,b=1)
    # E11:M11
    mr(11,5,11,13)
    sc(11,5,'FORMATIVE ASSESSMENT', bold=True, border=T())
    for c in range(6,14): ws.cell(11,c).border = bdr(t=1,b=1,r=(c==13))

    # ══ ROWS 12-13: Column headers ═══════════════════════════════════
    # A12:A13  No. of Half/SEM
    mr(12,1,13,1)
    sc(12,1,'No.of \nHalf/\nSEM', sz=10, wrap=True, border=T())
    ws.cell(13,1).border = bdr(l=1,r=1,b=1)

    # B12:D12  Attendance During Semester (top span)
    mr(12,2,12,4)
    sc(12,2,'Attendance During Semester', sz=10, border=T())
    ws.cell(12,3).border = bdr(t=1,b=1); ws.cell(12,4).border = bdr(r=1,t=1,b=1)

    # B13  Actual, C13  Possible, D13  %
    sc(13,2,'Actual', sz=10, border=T())
    sc(13,3,'Possible', sz=10, border=T())
    sc(13,4,'%', sz=10, border=T())

    # E12:E13  Specific Learning Outcome (A)
    mr(12,5,13,5)
    sc(12,5,'Specific\nLearning Outcome\n(70 Marks)\nA', wrap=True, border=T())
    ws.cell(13,5).border = bdr(l=1,r=1,b=1)

    # F12:H12  Generic Learning Outcome (30 Marks) B — spans 3 sub-cols
    mr(12,6,12,8)
    sc(12,6,'Generic Learning Outcome (30 Marks) B', sz=10, wrap=True, border=T())
    for c in range(7,9): ws.cell(12,c).border = bdr(t=1,b=1,r=(c==8))

    # F13  ES, G13  WC_SC, H13  ED
    sc(13,6,'ES (10)', sz=9, wrap=True, border=T())
    sc(13,7,'WC_SC (10)', sz=9, wrap=True, border=T())
    sc(13,8,'ED (10)', sz=9, wrap=True, border=T())

    # I12:I13  Total Out of 100 (A+B)
    mr(12,9,13,9)
    sc(12,9,'Total Out \nof \n100\nA+B', wrap=True, border=T())
    ws.cell(13,9).border = bdr(l=1,r=1,b=1)

    # J12:J13  Sign of Trainee
    mr(12,10,13,10)
    sc(12,10,'Sign \nof Trainee', wrap=True, border=T())
    ws.cell(13,10).border = bdr(l=1,r=1,b=1)

    # K12:K13  Sign of S.I
    mr(12,11,13,11)
    sc(12,11,'Sign of  S.I', wrap=True, border=T())
    ws.cell(13,11).border = bdr(l=1,r=1,b=1)

    # L12:L13  Sign of F.I
    mr(12,12,13,12)
    sc(12,12,'Sign of\nF.I', wrap=True, border=T())
    ws.cell(13,12).border = bdr(l=1,r=1,b=1)

    # M12:M13  Sign of Principal
    mr(12,13,13,13)
    sc(12,13,'Sign \nof \nPrincipal', wrap=True, border=T())
    ws.cell(13,13).border = bdr(l=1,r=1,b=1)

    ws.row_dimensions[12].height = 25.5
    ws.row_dimensions[13].height = 33.0

    # ══ ROW 14: H1 data — always labelled "H1"; filled only when semester==H1 ══
    sc(14,1,'H1', border=T())
    for c in range(2,14): ws.cell(14,c).border = T(); ws.cell(14,c).alignment = aln(v='center')
    ws.cell(14,2).value = h1_act; ws.cell(14,2).alignment = aln(v='center')
    ws.cell(14,3).value = h1_pos; ws.cell(14,3).alignment = aln(v='center')
    ws.cell(14,4).value = h1_pct; ws.cell(14,4).alignment = aln(v='center')
    ws.cell(14,5).value = h1_lo;  ws.cell(14,5).alignment = aln(v='center')
    ws.cell(14,6).value = h1_es;  ws.cell(14,6).alignment = aln(v='center')
    ws.cell(14,7).value = h1_wsc; ws.cell(14,7).alignment = aln(v='center')
    ws.cell(14,8).value = h1_ed;  ws.cell(14,8).alignment = aln(v='center')
    ws.cell(14,9).value = h1_tot; ws.cell(14,9).alignment = aln(v='center')

    # ══ ROW 15: H2 data — always labelled "H2"; filled only when semester==H2 ═
    sc(15,1,'H2', border=T())
    for c in range(2,14): ws.cell(15,c).border = T(); ws.cell(15,c).alignment = aln(v='center')
    ws.cell(15,2).value = h2_act; ws.cell(15,2).alignment = aln(v='center')
    ws.cell(15,3).value = h2_pos; ws.cell(15,3).alignment = aln(v='center')
    ws.cell(15,4).value = h2_pct; ws.cell(15,4).alignment = aln(v='center')
    ws.cell(15,5).value = h2_lo;  ws.cell(15,5).alignment = aln(v='center')
    ws.cell(15,6).value = h2_es;  ws.cell(15,6).alignment = aln(v='center')
    ws.cell(15,7).value = h2_wsc; ws.cell(15,7).alignment = aln(v='center')
    ws.cell(15,8).value = h2_ed;  ws.cell(15,8).alignment = aln(v='center')
    ws.cell(15,9).value = h2_tot; ws.cell(15,9).alignment = aln(v='center')

    # ══ ROWS 16-18: NO. of Year / headers / YEAR-1 data (1-year layout)
    # If year2=False, also do:
    #   Row 16-17: A16:A17 (NO. of Year)  B16:B17(Actual) etc, E16:G17(Formative Total), H-K merged
    #   Row 18: data

    def write_year_summary_block(r_hdr, r_data, year_label, act, pos, pct, fa_tot, si):
        """
        Write the "NO. of Year" header block and the data row.
        r_hdr = first row of header (merged 2 rows), r_data = data row
        Matches exactly the proforma for rows 16-18 (1yr) / 17-19 and 27-29 (2yr).
        """
        # A(r_hdr):A(r_hdr+1)  NO. of Year
        mr(r_hdr,1,r_hdr+1,1)
        sc(r_hdr,1,'NO. of Year', wrap=True, border=T())
        ws.cell(r_hdr+1,1).border = bdr(l=1,r=1,b=1)

        # B(r_hdr):B(r_hdr+1)  Actual
        mr(r_hdr,2,r_hdr+1,2)
        sc(r_hdr,2,'Actual', sz=10, border=T())
        ws.cell(r_hdr+1,2).border = bdr(l=1,r=1,b=1)

        # C(r_hdr):C(r_hdr+1)  Possible
        mr(r_hdr,3,r_hdr+1,3)
        sc(r_hdr,3,'Possible', sz=10, border=T())
        ws.cell(r_hdr+1,3).border = bdr(l=1,r=1,b=1)

        # D(r_hdr):D(r_hdr+1)  %
        mr(r_hdr,4,r_hdr+1,4)
        sc(r_hdr,4,'%', sz=10, border=T())
        ws.cell(r_hdr+1,4).border = bdr(l=1,r=1,b=1)

        # E(r_hdr):I(r_hdr+1)  Formative Assessment Total (Out of 200)
        mr(r_hdr,5,r_hdr+1,9)
        sc(r_hdr,5,'Formative Assessment Total\n(Out of 200)', wrap=True, border=T())
        for c in range(6,10): ws.cell(r_hdr,c).border = bdr(t=1,r=(c==9))
        for c in range(5,10): ws.cell(r_hdr+1,c).border = bdr(b=1,l=(c==5),r=(c==9))

        # J(r_hdr):J(r_hdr+1)  Sign of Trainee
        mr(r_hdr,10,r_hdr+1,10)
        sc(r_hdr,10,'Sign \nof Trainee', wrap=True, border=T())
        ws.cell(r_hdr+1,10).border = bdr(l=1,r=1,b=1)

        # K(r_hdr):K(r_hdr+1)  Sign of S.I
        mr(r_hdr,11,r_hdr+1,11)
        sc(r_hdr,11,'Sign of S.I', border=T())
        ws.cell(r_hdr+1,11).border = bdr(l=1,r=1,b=1)

        # L(r_hdr):L(r_hdr+1)  Sign of F.I
        mr(r_hdr,12,r_hdr+1,12)
        sc(r_hdr,12,'Sign of\nF.I', wrap=True, border=T())
        ws.cell(r_hdr+1,12).border = bdr(l=1,r=1,b=1)

        # M(r_hdr):M(r_hdr+1)  Sign of Principal
        mr(r_hdr,13,r_hdr+1,13)
        sc(r_hdr,13,'Sign \nof \nPrincipal', wrap=True, border=T())
        ws.cell(r_hdr+1,13).border = bdr(l=1,r=1,b=1)

        ws.row_dimensions[r_hdr].height = 15.0

        # Data row r_data — cols 5-9 merged for FA Total (mirrors header merge E:I)
        sc(r_data,1, year_label, border=T())
        ws.cell(r_data,2).value = act;    ws.cell(r_data,2).border = T();  ws.cell(r_data,2).alignment = aln(v='center')
        ws.cell(r_data,3).value = pos;    ws.cell(r_data,3).border = T();  ws.cell(r_data,3).alignment = aln(v='center')
        ws.cell(r_data,4).value = pct;    ws.cell(r_data,4).border = T();  ws.cell(r_data,4).alignment = aln(v='center')
        mr(r_data,5,r_data,9)
        ws.cell(r_data,5).value = fa_tot; ws.cell(r_data,5).border = T();  ws.cell(r_data,5).alignment = aln(v='center')
        for c in range(6,10):
            ws.cell(r_data,c).border = bdr(t=1,b=1,r=(c==9))
        for c in range(10,14):
            ws.cell(r_data,c).value = ''; ws.cell(r_data,c).border = T();  ws.cell(r_data,c).alignment = aln(v='center')

    if not year2:
        # ── 1-YEAR LAYOUT ──────────────────────────────────────────
        # Row 16-17: header, Row 18: YEAR-1 data
        write_year_summary_block(16, 18, 'YEAR-1', yr_att_act, yr_att_pos, yr_att_pct, yr_fa_tot, u_si)
        ws.page_setup.fitToHeight = 1

    else:
        # ── 2-YEAR LAYOUT ──────────────────────────────────────────
        # Row 16: YEAR - 1 ASSESSMENT (full-width label)
        mr(16,1,16,13)
        sc(16,1,'YEAR - 1 ASSESSMENT', border=T())
        for c in range(2,14): ws.cell(16,c).border = bdr(t=1,b=1,r=(c==13))

        # Rows 17-18: NO. of Year headers, Row 19: YEAR-1 summary data
        write_year_summary_block(17, 19, 'YEAR-1', yr_att_act, yr_att_pos, yr_att_pct, yr_fa_tot, u_si)

        # Row 20: spacer row (left+right thin only, top+bottom thin)
        for c in range(1,14):
            ws.cell(20,c).border = bdr(t=1,b=1,l=(c==1),r=(c==13))

        # Row 21: HALF YEARLY ASSESSMENT | FORMATIVE ASSESSMENT (repeat for year 2)
        mr(21,1,21,4)
        sc(21,1,'HALF YEARLY ASSESSMENT', bold=True, border=T())
        ws.cell(21,2).border = bdr(t=1,b=1); ws.cell(21,3).border = bdr(t=1,b=1)
        ws.cell(21,4).border = bdr(r=1,t=1,b=1)
        mr(21,5,21,13)
        sc(21,5,'FORMATIVE ASSESSMENT', bold=True, border=T())
        for c in range(6,14): ws.cell(21,c).border = bdr(t=1,b=1,r=(c==13))

        # Rows 22-23: Column headers (same as rows 12-13 but for year 2)
        mr(22,1,23,1)
        sc(22,1,'No.of \nHalf/\nSEM', sz=10, wrap=True, border=T())
        ws.cell(23,1).border = bdr(l=1,r=1,b=1)
        mr(22,2,22,4)
        sc(22,2,'Attendance During Semester', sz=10, border=T())
        ws.cell(22,3).border = bdr(t=1,b=1); ws.cell(22,4).border = bdr(r=1,t=1,b=1)
        sc(23,2,'Actual', sz=10, border=T()); sc(23,3,'Possible', sz=10, border=T()); sc(23,4,'%', sz=10, border=T())
        mr(22,5,23,5); sc(22,5,'Specific\nLearning Outcome\n(70 Marks)\nA', wrap=True, border=T()); ws.cell(23,5).border = bdr(l=1,r=1,b=1)
        mr(22,6,22,8); sc(22,6,'Generic Learning Outcome (30 Marks) B', sz=10, wrap=True, border=T())
        for c in range(7,9): ws.cell(22,c).border = bdr(t=1,b=1,r=(c==8))
        sc(23,6,'ES (10)', sz=9, wrap=True, border=T())
        sc(23,7,'WC_SC (10)', sz=9, wrap=True, border=T())
        sc(23,8,'ED (10)', sz=9, wrap=True, border=T())
        mr(22,9,23,9); sc(22,9,'Total Out \nof \n100\nA+B', wrap=True, border=T()); ws.cell(23,9).border = bdr(l=1,r=1,b=1)
        mr(22,10,23,10); sc(22,10,'Sign \nof Trainee', wrap=True, border=T()); ws.cell(23,10).border = bdr(l=1,r=1,b=1)
        mr(22,11,23,11); sc(22,11,'Sign of  S.I', wrap=True, border=T()); ws.cell(23,11).border = bdr(l=1,r=1,b=1)
        mr(22,12,23,12); sc(22,12,'Sign of\nF.I', wrap=True, border=T()); ws.cell(23,12).border = bdr(l=1,r=1,b=1)
        mr(22,13,23,13); sc(22,13,'Sign \nof \nPrincipal', wrap=True, border=T()); ws.cell(23,13).border = bdr(l=1,r=1,b=1)
        ws.row_dimensions[22].height = 25.5
        ws.row_dimensions[23].height = 54.0

        # Row 24: Year-2 H1 row — always blank (Year 2 data not yet available)
        sc(24,1,'H1', border=T())
        for c in range(2,14): ws.cell(24,c).border = T(); ws.cell(24,c).alignment = aln(v='center')

        # Row 25: H2 blank
        sc(25,1,'H2', border=T())
        for c in range(2,14): ws.cell(25,c).border = T(); ws.cell(25,c).alignment = aln(v='center')

        # Row 26: YEAR - 2 ASSESSMENT label
        mr(26,1,26,13)
        sc(26,1,'YEAR - 2 ASSESSMENT', border=T())
        for c in range(2,14): ws.cell(26,c).border = bdr(t=1,b=1,r=(c==13))

        # Rows 27-28: headers, Row 29: YEAR-2 summary data
        write_year_summary_block(27, 29, 'YEAR-2', '', '', '', '', u_si)

    # ── Page setup ────────────────────────────────────────────────────
    ws.page_setup.fitToPage   = True
    ws.page_setup.fitToWidth  = 1
    ws.page_setup.fitToHeight = 1
    ws.page_setup.orientation = 'portrait'
    ws.print_area = 'A1:M' + ('18' if not year2 else '29')

    # ── Column widths (from proforma — no explicit widths set, use default) ──
    # Both proformas have no custom column widths; leave at default.


def create_progress_card_1year(trainee_df, user_info=None):
    """Generate 1-year ITI Progress Cards — one sheet per trainee."""
    ui = user_info or {}
    wb = Workbook()
    wb.remove(wb.active)
    for _, tr in trainee_df.iterrows():
        roll = tr['rollno']
        ws = wb.create_sheet(title=str(f'Roll_{roll}')[:31])
        _build_progress_card_sheet(ws, tr, ui, year2=False)
    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    return out


def create_progress_card_2year(trainee_df, user_info=None):
    """Generate 2-year ITI Progress Cards — one sheet per trainee."""
    ui = user_info or {}
    wb = Workbook()
    wb.remove(wb.active)
    for _, tr in trainee_df.iterrows():
        roll = tr['rollno']
        ws = wb.create_sheet(title=str(f'Roll_{roll}')[:31])
        _build_progress_card_sheet(ws, tr, ui, year2=True)
    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    return out

# ─── HTML PAGES ───────────────────────────────────────────────────

REGISTER_HTML = '''<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Register — ITI Assessment System</title>
<link href="https://fonts.googleapis.com/css2?family=Sora:wght@300;400;600;700;800&display=swap" rel="stylesheet">
<style>
*{margin:0;padding:0;box-sizing:border-box}
:root{--navy:#0B1D3A;--blue:#1A56DB;--accent:#F59E0B;--light:#EFF6FF;--muted:#64748B;--radius:16px}
body{font-family:'Sora',sans-serif;min-height:100vh;background:var(--navy);display:flex;align-items:center;justify-content:center;padding:20px;position:relative;overflow:hidden}
body::before{content:'';position:fixed;top:-200px;right:-200px;width:600px;height:600px;background:radial-gradient(circle,rgba(26,86,219,0.18) 0%,transparent 70%);pointer-events:none}
body::after{content:'';position:fixed;bottom:-200px;left:-200px;width:500px;height:500px;background:radial-gradient(circle,rgba(245,158,11,0.1) 0%,transparent 70%);pointer-events:none}
.dots{position:fixed;inset:0;background-image:radial-gradient(rgba(255,255,255,0.04) 1px,transparent 1px);background-size:32px 32px;pointer-events:none}
.wrap{display:grid;grid-template-columns:1fr 1fr;max-width:960px;width:100%;gap:0;border-radius:24px;overflow:hidden;box-shadow:0 40px 100px rgba(0,0,0,0.5);position:relative;z-index:1}
.left{background:linear-gradient(145deg,#1A56DB,#0B1D3A);padding:52px 44px;display:flex;flex-direction:column;justify-content:space-between}
.logo{display:flex;align-items:center;gap:12px;margin-bottom:48px}
.logo-icon{width:44px;height:44px;background:rgba(255,255,255,0.12);border-radius:12px;display:flex;align-items:center;justify-content:center;font-size:22px;border:1px solid rgba(255,255,255,0.15)}
.logo-text{font-size:15px;font-weight:700;color:#fff;letter-spacing:0.5px}
.logo-text span{display:block;font-size:11px;font-weight:400;color:rgba(255,255,255,0.5);letter-spacing:1px;text-transform:uppercase}
h2{font-size:36px;font-weight:800;color:#fff;line-height:1.2;margin-bottom:16px}
h2 em{font-style:normal;color:var(--accent)}
.tagline{color:rgba(255,255,255,0.55);font-size:14px;line-height:1.7;margin-bottom:40px}
.steps{display:flex;flex-direction:column;gap:16px}
.step{display:flex;align-items:center;gap:14px}
.step-num{width:32px;height:32px;border-radius:50%;background:rgba(255,255,255,0.1);border:1px solid rgba(255,255,255,0.2);color:#fff;font-size:13px;font-weight:700;display:flex;align-items:center;justify-content:center;flex-shrink:0}
.step-text{color:rgba(255,255,255,0.7);font-size:13px}
.step-text strong{color:#fff;display:block;font-size:13.5px;margin-bottom:2px}
.right{background:#fff;padding:48px 44px}
.right h3{font-size:22px;font-weight:700;color:var(--navy);margin-bottom:6px}
.right p{font-size:13.5px;color:var(--muted);margin-bottom:32px}
.form-group{margin-bottom:20px}
label{display:block;font-size:12.5px;font-weight:600;color:var(--navy);margin-bottom:7px;letter-spacing:0.3px;text-transform:uppercase}
input{width:100%;padding:13px 16px;border:2px solid #E2E8F0;border-radius:10px;font-family:'Sora',sans-serif;font-size:13.5px;color:var(--navy);transition:all 0.2s;outline:none;background:#FAFBFF}
input:focus{border-color:var(--blue);background:#fff;box-shadow:0 0 0 4px rgba(26,86,219,0.08)}
input::placeholder{color:#A0AEC0}
.btn-submit{width:100%;padding:15px;background:linear-gradient(135deg,var(--blue),#0B1D3A);color:#fff;border:none;border-radius:10px;font-family:'Sora',sans-serif;font-size:14px;font-weight:700;cursor:pointer;transition:all 0.25s;letter-spacing:0.5px;margin-top:8px;box-shadow:0 6px 20px rgba(26,86,219,0.35)}
.btn-submit:hover{transform:translateY(-2px);box-shadow:0 10px 28px rgba(26,86,219,0.45)}
.btn-submit:disabled{opacity:0.6;cursor:not-allowed;transform:none}
.login-link{text-align:center;margin-top:20px;font-size:13px;color:var(--muted)}
.login-link a{color:var(--blue);font-weight:600;text-decoration:none}
.login-link a:hover{text-decoration:underline}
#msg{margin-top:16px;padding:13px 16px;border-radius:10px;font-size:13px;font-weight:500;display:none;text-align:center}
#msg.success{background:#ECFDF5;color:#065F46;border:1px solid #A7F3D0;display:block}
#msg.error{background:#FEF2F2;color:#991B1B;border:1px solid #FECACA;display:block}
#msg.pending{background:#FFF7ED;color:#92400E;border:1px solid #FDE68A;display:block}
.spinner{display:inline-block;width:14px;height:14px;border:2px solid rgba(255,255,255,0.3);border-top-color:#fff;border-radius:50%;animation:spin 0.8s linear infinite;margin-right:8px;vertical-align:middle}
@keyframes spin{to{transform:rotate(360deg)}}
@media(max-width:700px){.wrap{grid-template-columns:1fr}.left{display:none}}
</style>
</head>
<body>
<div class="dots"></div>
<div class="wrap">
  <div class="left">
    <div>
      <div class="logo">
        <div class="logo-icon" style="overflow:hidden;padding:2px;"><img src="data:image/jpeg;base64,/9j/4AAQSkZJRgABAQEASABIAAD/2wBDAAYEBQYFBAYGBQYHBwYIChAKCgkJChQODwwQFxQYGBcUFhYaHSUfGhsjHBYWICwgIyYnKSopGR8tMC0oMCUoKSj/2wBDAQcHBwoIChMKChMoGhYaKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCj/wgARCAKbBLADASIAAhEBAxEB/8QAHAABAAMBAAMBAAAAAAAAAAAAAAUGBwEDBAgC/8QAGgEBAAMBAQEAAAAAAAAAAAAAAAEDBAIFBv/aAAwDAQACEAMQAAAB1QAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAH45TYzzPQ0VnTjvRWdDRWdDRWdDRWdDRWdDRuZ3w0XucjRWdDRWdDRWdDRWdDRe5xa7qZ7vO+hhBIAAAAAAAAAAAAAAAAA4jvPShors3KHC8Uat2Jlu9XRPQAAAAAAAAAAAAAAA4d568RFc+pEPxn07tGvPWjo6s5+f1SZXF88dt5+hnzyl9DPnkfQz554j6HfPJP0M+eR9DPnnkPod88j6GfPKX0M+eUPoZ88pfQz55H0M+eR9D+XFdqqkOZApMXKRfzH0IUXAAAAAAAAAAActlTte/HYB73iBIAAAAAAAAAAAAcO8567n2lfh+arv48ziuM2nw1Ec57HE+k5zDnNffz3hqsxDzGn6EOrAAAAAAAAAADg6/PrI9rtfh+Kbv+Mziuc2oQ9A7zns0THuM/6/POxUc6i2X+gX+/3As1KPeKP1GT953ZWAEAAAAkEAAAkEALHtmJ7Zm7CvoCkxcpF/MfQhRcAAAAAAAAAABy11S178dgHveIEgAAAAAAABwPFFxxMqbDc59K9DLvW4z6DEVVzmko7jmjnXI464AAAHO8TqsxDTOn6EOrAAAAAABw7z1YmK7Ao8PxRpsdl/g5zX+IrDnPIehxzRzvEcdcAAADvOlsv9Av9/uBZqUe8UfqMn7zuysIAAAAAAAAAABKx7Zie2Zewr6ApMXKRfzH0IUXAAAAAAAAAAActdUte/HYB73iBIAAA50OflH67Fw3Nds5nsPzn1CJzfnGa4Q8PznP5fF3nNDvCAAAAAAAAHO8TqsxDzGn6EOrDnQ4O8erHPtK7ERTefxmcVxn1CGoPec9niY5xn7+eoq4EAdcAAAAAADvOlsv9Av9/uBZqUe8UfqMn7zuysIAAAAAAAAAABKx7Zie2Zewr6ApMXKRfzH0IUXAAAAAAAAAAActdUte/HYB73iDku856cc+6rMPzTfvDmMZzm0uHpLnPOxPgcZ+d52K3BHeAAAAAAAAAAAA50nU5ivtHvWHlGiYq0yNzDw8579D1jvOf3/R/PeaDiOO8AAAAAAAAAAAAB3nS2X+gX+/3As1KPeKP1GT953ZWEAAAAAAAAAAAlY9sxPbMvYV9AUmLlIv5j6EKLgAAAAAAAAAAOWuqWvfjsA97xAlA5lpOa0eP3hXh64HeAAAAAAAAAAAAAAAB3g7wT1wgAAAAAAAAAAAAAAAAB3nS2X+gX+/3As1KPeKP1GT953ZWEAAAAAAAAAAAlY9sxPbMvYV9AUmLk4z5j6IKLQAAAAAAAAAAOWyp2vfjsA97xAlXs00vNKPHCvCAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA7zpbL/QL/f7gWalHvFH6jJ+87srCAAAAAAAAAAASse2YntmbsKugKv6F2YddKXVx3Se3UUpdRSl1FKXUUpdRSl1FKXUUpdRSl1FKXUUpdRSZ+X5bV3vO68wSr2Z6ZmdHj9FeEAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB3nS2X+gX+/wBwLNSrWnhjHdnWc4xzaEsXbQMXbQMXbQMX7s4xdtCGMc2hLF20DF20DF20DF20IYxzaOGZaccSESBxWvRxarn2luO7mphFzUwXNTBc1MFzUwXNTBc1MFzUwXNTBc1MFzUwXNS562uW6as4Sr2aaXmdHj9FeEAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB3nS2X+gX+/wBwLNQrBZ2N9s52Njg2Njg2Njg2Njg2Njg2Njg2Njg2Njg2Njg2Njg2Njg2Njg2NmOm8T0RIFJi5SL+Z+hDPcAAAAAAAAAABy11S178dgHveIEq9mml5pR44V4QAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAHedLZf6Bf7/cCzUo94o/UZP3ndlYQAAAAAAAAAAAn9ww/cM/YU9BKkxcpF/MfQhRcAAAAAAAAAABy11S178dgHveIEq9memZpR44V4QAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAHedLZf6Bf7/cCzUo94o/UZP3ndlYQAAAAAAAAAAAntxw7cc/YVdAUmLlIv5j6EKLgAAAAAAAAAAOWuqWvfjsA97xAlXs00vM6PH6K8IAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAADvOlsv9Av9/uBZqUe8UfqMn7zuysIAAAAAAAAAAAT244duOfsKugKTFykX8x9CFFwAAAAAAAAAAHLXVLXvx2Ae94gSrua6VmtHjhXhAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAd50tl/oF/v9wLNSj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vEHJV7NdXqdPmVRbHGWprYKmtiVTWwVNbEKmtiVTWxCprYKmtiVT7a0KolYmKQjgAAAAAAAAAAAAAAB3nQWWbK0tabamtiVTWxCprYlU1sQqa2CprYlVFrQqa2Cp9tY7fq1ZbvV6O9Cj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECXOgAAAAAAAAABRqZdKZn8Pg4ygAAAAAAAAAAAAAAO86NayXW7fRkBd6oAAAAAAAAAACj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECQAAAAAAAAAAFHpd0pefww4ygAAAAAAAAAAAAAAO86NbyTW7fRkBd6oAAAAAAAAAACj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECQAAAAAAAAAAFHpd0pefww4ygAAAAAAAAAAAAAAO86NbyTW7fRkBd6oAAAAAAAAAACj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECQAAAAAAAAAAFHpd0pefww4ygAAAAAAAAAAAAAAO86NbyTW7fRkBd6oAAAAAAAAAACj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECQAAAAAAAAAAFHpd0pefww4ygAAAAAAAAAAAAAAO86NbyTW7fRkBd6oAAAAAAAAAACj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECQAAAAAAAAAAFHpd0pefww4ygAAAAAAAAAAAAAAO86NbyTW7fRkBd6oAAAAAAAAAACj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECQAAAAAAAAAAFHpd0pmfw+DjKAAAAAAAAAAAAAAA7zo1rJdat9GRF3qgAAAAAAAAAAKPeKP1GT953ZWEAAAAAAAAAAAJ7ccO3HP2FXQFJi5yP+c9703uqbfSe7w9N7o9J7o9J7o9J7nT0nudPSe6PSe6PSe6PSe6PSe509G1wNj25JrvO+344SAAAAAAAAAAHCkUy91ejxozko4zRaUEWkxGJMRiTEYkxGJMRiTEYkxGJMRiTEYkxGJMRiTEYkxGJMRnZMRmtZxpVm/3hd6YAAAAAAAAAACj3eozGPdmO6+IZMiGTIhkyIZMiGTIhkyIZMiGTIhkyIZMiGTIhkyPLuORa5n66K+gOdIBIAAAAAAAAAABzqHOkgAAAAAAAAAAAOdAAAAAAAAAAAAAAAABzo46QCQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAHp+5FFC1D53+hz9vU9ck0bClsV6fP0A8EAWZVZ8zmxUu0l6eGMJkAhyYeHzBGQ5a1fmzy1Waxs1+Vza0lger6hKvW8R71f8AxNGYaxCTYeOIJtVPaLC8fkIbMtIyA30Dx+QZX4dWGUeLQMaLW1YU+6ej4SUeGALK8Ppkkq36LO9P3D847skAezLeLxnsorzHvqn7RYn4/ZTYa++uSxCk08EIWJVJw98ByIJhU/eJ54okm3OgELT53LzfSOJF6HrkuqciTbnQ/ECWFVZYlPF+64Rl2+cNuLH4YmaMb1qMmz802QxY+j/1XvOTTwxZNAEQS/58UWQNz+cdqLUgZw/TkMTSpyJNudAAAEVKxRhP0R87/RBR8v1Gpk9N20ZHEbf87n0TU7Dg57t7nrAVCtaoPnbUKhbyZx7YsePoADBd6wY1iLlP0ZdoudSRYXhuZH4/9AYqe7YPYupF5RteKHPb9/UijSllrhlG9YLvRWIGVqpZor1ulP3DCtUJTINfyA30AAEFi+0YufQoKHnWjZ4S8/ov7IqkXSgFg8UQKxolK/Zs/wA6fRfzsbvAz0EZBotZ3AqVW1b8mD7r897sVWsWesmxfO30T87lotc3NlGzb6Cz89615HdCgS0LuRUvSvQx6o77gx9Efv8AH7AKll+oZeb7lOrZSVi9+lp5T6LtUaZlr3zn9FGJ2j81o0Gv+tZSVrt69Y+dtXzH6IKdcPN+D503/APoEpmU/Q+Amn+5b/2QWL7Vix9CAYRu+Empen7s+fNeqZp9AlUtfvZ2VSx0q7Fhrn6Fa2jBNNLsAABFSseYJ9EYxs5SKtdoA0wH4+dvonFzXsd2UUbQsr9E2P08w9opWmQdqJLDvorOy/8AmxORNQ+fL/EmiS8fDFx8WUecteR2ufL1jO01Y9O7YjKGtYrqtDLDcK1ZRXLHCGQbTmetnzj9DUCINg/OZdKvolOvh7OQbNmxsAAAILF9vy42cFEql5rpp4MC0/y0M2Vlv7PcrUl5jT/nb6JxY1eCsMSZxtmWamPz+vyfPO65VrJUazdII0v53+iMXNSl42SFBv1PKZqtB0w+fNGulTLF6tRkSv0raMjPoD9/n9AFSy/W6Ea/lOrZ6etpdFvQj5D0j5/+isW2Q8ft5d6xpmZ+X1D3tH9GYPnP6IzD0DZPxmukHzt9A5BsR4sA+jMpNU8mOTZcsH+jKCXf2cSkTVPn68RhoM/ETZ84b/R68bTUq1pphe2VunGtMs/Joctm+jnlAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA/P6GVeDXOGW+HVemLaZYAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB//8QAMxAAAAUCBAQFBQEAAgMBAAAAAAIDBAUBBhUzNEASFCA1EBETFjEwMkFQYCEisCMkcJD/2gAIAQEAAQUC/wCg24qUHEUcRRxFHFQcVBxUHEUcRRxFHEUcRRxFHEUcRRxFHEUcRRxFHEUcRRxFHEUcRRxUHFQcZR5+f81WtKD1SD1SD1SD1SD1CD1CD1SD1CD1CD1SD1SD1SD1CD1CD1CD1CD1CD1CD1CD1CD1SD1SD1CD1SD1SD1SChy16pHz9fzqPOo86jzqPOo86jzqPOo86jzqPOo86jzqPOo86jzqPOo86jzqPOo86jzqPOo86jzqK1qInK/UHdIphSYblCs0pULyDlURla1afozrpECsq2JVWbCsm5UpAqHUV6LprWjTiOOI44zjjOOI44zjjOOI44jjiOOM44zjiOOI44jjiOOI44jjiOOI44jjjOOM44jjiOOI4gDG53pks/dRGT+jMcpQrINkwpNp0Csu4OFHKyg+fGoitHvjGoUKv26QUmkqBWZXMFHjhQVrWvRbub0XVpKbq39b0yWfuojK351kyBaVbJhWb/1STcnB1Dn+hURWj3Si6SYVlWyYVmwpJuTg6qin0bdzei6tJurf1vTJZ+6iMndVr5BV4gnRWZQKFZpWoVfOFBWta/UqIrR7WpqFCz9ulRWaSoFJlaoUeOFBWta/Vt3N6Lq0m6t/W9Mln7qIydsY5ShSSbECs2UKyzk4O4WUrsKiK0exOukQKyrYgVm6hSTcnB1VD7K3c3ourSbq39b0yWfuojJ2SjhJMKyzYgUmjVCsi5UBznPtaiK0f1KmoUKP26YVmkqBWZXMFHjhQV867a3c3ourSbq39b0yWfuojJ+rU1KBZ83SorNJUCkysYKO11BX/dzURWj+gdwkQKyzYgVm61Csm5ODqnPXdW7m9F1aTdW/remSz91EZP0TqkIFZNsSis3QKyrk9TrKn31RFaPxqalAo/bJhSaSoFZhcwVeLqiv+131u5vRdWk3Vv63pks/dRGT0qOkUwrMNyhWaUqFJByoDGqav6KPXTTaKyrYgUm6hWTcnB1FD/pLdzei6tJurf1vTJZ+6iMnomTmIz+f4+3c3ourSbq39b0yWfuojJ6JzRfx9u5vRdWk3Vv63pks/c1ERldE5ov4+3c3ourSbq39b0yWfuaiIyuic0X8fbub0XVpN1Aa3petVVFeRXHILjkFxyC45FcciuORXHIrjkVxyK45FcciuORXHIrjkVxyK45FcciuORXHIrjkVxyK45FcciuORXHILiPROin0Tmi/j7dzeieaKu2+BPhgT4YE+GBPhgT4YE+GBPhgT4YE+GBPhgT4YE+GBPhgT4YE+GBPhgT4YE+GBPhgT4YE+GBPhgT4YE+GBPhgT4RES6buf0s5ov4+3c3+FnNF/H27m/pXjtVJXEFxiC4xBcYguMQXGILjEFxiC4xBcYguMQXGILjEFxiC4xBcYguMQXGILjEFxiC4xBcYguMQXGILjEFxiC4j1jrJ9E5ov4+3c3onHirND3A+HuB8PcD4e4Hw9wPh7gfD3A+HuB8PcD4e4Hw9wPh7gfD3A+HuB8PcD4e4Hw9wPh7gfD3A+HuB8PcD4e4Hw9wPh7gfD3A+HuB8IyZduHXTJZ+5qInK6JzRfx9u5vRdOk3UFrumSz91E5PROaL+Pt3N6Lq0m6gtd0yWfuojJ6JzRfx9u5vRdWk3UFrumSz91EZPROaL+Pt3N6Lq0m6gtd0yWfuojJ6JzRfx9u5vRdWk3UFrumSz91EZPROaL+Pt3N6Lq0m6gtd0yWfuojJ6JzRfx9u5vRdWk3UFrumSz91EZPROaL+Pt3N6Lq0m6gtd0yWfuojJ6JzRfx9u5vRdWk3UFrumSz91EZPRIN6uUMEVGCKjBFRgiowRUYIqMEVGCKjBFRgiowRUYIqMEVGCKjBFRgiowRUYIqMEVGCKjBFQ9ZGafpUohRVPBFRgiowRUYIqMEVGCKjBFRgiowRUYIqMEVGCKjBFRgiowRUYKqMEVGCKjBFRgioi2B2h+i6tJuoLXdMln7qIyd1cH6T8sNLubq0m6gtd0yWfuojJ3VwfpPyw0u5urSbqC13TJZ+6iMndXB+k/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyt1cHz+j/LDS7m6tJuoLXdMln7qJyd1cH6T8sNLubp0m6gtd0v0FDrcqsOVWHKrDlVhyqw5VYcqsOVWHKrDlVhyqw5VYcqsOVWHKrDlVhyqw5VYcqsOVWHKrDlVhyqw5VYcqsOVWEamZNPdTSCqw5ByOQdDkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkXXmzLUrfc3Egq4bYW9GFvRhb0YW9GFvRhb0YW9GFvRhb0YW9GFvRhb0YW9GFvRhb0YW9GFvRhb0YW9GFvRhb0YW9GFvRhb0YW9GFvRDx7pJ3/+LztWqLdpcS6zun+0+pc75w2Wtpwq4Z/VkppFgpHvSPUeqacLNmkTLSC7zrllDJMoKRdLP/FSvCRa4XZFvcb0VuR5QUulxWvuN6Pcb0Qcou/U+jX4dzMkR1GKqLNOu4XjlolAST1062Uq7MzbRM+s8d/VUPRMhbibmcFr5lC1alSNMyfNNDmUbqqUSTSuNuo4p/tPoVr5UVuNumukeiif15PQxvcy/a4cJNytnzdzVy9bthj7AITDJatK0NTwVVIkVScYEqWeYGqg4RXpeGfaWgVUKkRGTaLn8V5JqgdJQipA4fNm4x9gEZhiqCKEPRy8QbC5nCThzb0g1QZ4wxDddNwRy/bNjILpuCOHSDeiMyzWWWUTTTbv2aygMYpaKyjNIY+wCUyxVBDlPQTfb7b7l0cBRwFEyUtGEJ3LgKOAoXWRalbP2zmqihEqHm2JDpqFUTWfNkQadYFqScYGqi5RXp4KSLEiiZimKuumgRs/bOTOHaDalZ9h5ozLFapTUNQOl0UCNHjVxUKSjNM9FSVTXmWKNaT7CtW7tBx0OJJo3HuBgEJZkuKnLQhJRmdTouTt9s9yDp4g1DZ2i6K5kWrYY+wDeUZuPExqFovMMURSeYVqi/arDzp5PJNn6SKhKSJJdjQhJVkofzp5VkWFFCmpUr6SZ+i0OUsgWXY8KUmzVOooVMiMm0WU8V5FqgdFUixHMi0SDo5TSDWVYlbllmRjUrStK/4HEmzQHuBgG8ozXFP9+jJ6GN7mT7bs0UWu5IdtblFBW32FSSdt+inByqzRyU1DFnJcrAjdB9NLNrZakClvMDlcwK7Oso+Wd1tLQTnb7e7j43J3KC7fcz9Ro2hIvFDUgGFKOLZaHETCKsnElFoyAnmKTFeFhGzxr7YZhgyTZJXdrGsuZvHNrfdvAyt1BstPdvtnuQmYysgVG1m9Ke3GPk6tdGpEl3cO7YOaO2032+2+5dU12+E7l4XbooWQLH1IxkJlRK1kaBRpSrJK1y1MS22RQrbLQ1JCPcxCtuylXqVfiS7pG6O59BGOF0VW1vVXHt9hwyVs0InESS7F0meihLu0NnamvxL/APGRZpvphNC22SYcW2zUo/ZOIlxb0nV8gsoVFOUmXD5ZhbVVCkt9gUri2mhw6SkIgkdXjkSfZ43J2+2e5C8/mIVdno3tpEKW8wOWXgDsy23LqFV/E/ILOHkZbrf0T2+wNRxa5RGMzNWr232taJIlM9JbLOpEbcaJKmp5JOO6NKebR/b7XhbIlUe0tlnWjW32rdWY7fBdz8bm7jAdvfwTZyo4RKm8b22zUQTtpmQ6qhGjd/KupJwwtknDgLDyf2yThjpVzHOEFSrJdcnoY3uZPtuzRWgWlXXgp9jjurX/ABrNqerJQiRUmHjdqBE3NpaCc7fbvcfG5O5QXb5aOJIIYdJxiqdwvUQ3uhucNXiDoou7V2tofC7tZaTEvD4T/b7Z7kJOdbsq45JOq81PAr+bTrMulXS9saCb7fbfcuqa7fCdy8Ls0VutCunpS0KUSEggxIpcq6xivJ0wq7naCYkHi7a0NXX4ku6x2jufQWmQpn/gb7ZKnlKRuiu7Q2dqa/Ev3KHLQjDwu+n/AKdnai5zGpGWoUhn/jIFKZoxpSkoT7PG5O32z3IXmLNJT0/CRp5s4/8AyTL9k1AGcKIrS8aE7oUKGlwM1wQ1DleaZDuiWWFPscd0Z6Z9pGPcyfYJnQQXc/G5u4wHb1Mt53NlpRd6pitoR6Vkp7nqPc9R7nqJdzzy1rKHMy65PQxvcyfbdmis/UeB/scd1aaa52R0HdtyqaiPg5cJNk5t/wA85tLQTBKnYRS1GsgkoVUgcLpt05RzzT2C7equkiC1oah0kzh/Cs102ZjtJROvmS7tXataVY+F3au1+2eE/wBvtnuT+tStGlCKSDciaaQrWhaXM6TcPLX0M32+2+5dU12+E7l4XborP1XhcZzmkbdRblZeF1u06NbP1dfiS7rHaO5tBaOu8DfbJ90jdFd2hs7U1+JfuUVofC79HZ2oetyum7ls5iXkdcbdYtH7WoXlmSJZKWXkSxtKlkCfZ43J2+2e5C8xZmR4P9Ix7oX7CuUTnr/oVZt1aXHENmyNnrqmDzTI18pNHKB/scd0Z6Z9pWdeGTT/ANIJROqjKPV5SRQVIsmFlSIpzDorp9AdvU+x9/kmxrSrQXEyq7ZwT0rF0lRBUnopj0UgdwxIqShaU635DKNWMO9Tfl+24mqrprbUc5aLeBv9KtDvayDYtSoOEE3Cb62TUOQs40oVxPKBCCdOlJuGUqLdaqtGdaedJa3SuDJNZljUrmeOEIV48UlIRfmIlI6DOXhCvjUjZdlXmp4gVTnHtIiAK1OJqJLIJpx0swP6s+qI0ixG1xxrp05gG6jZh4S6J12cFFPG76tPOkrblFjokm2Io/mz0LHyr6spAKENAtlWzSVSOszhIp43fdUokdZnFRD1F94XE1WdNbajXLVfwl4dKQoWMlo9QkhNFHDOuw9t41GttRrpq4r8Pod6o/ZEMm2nm6jlnbkY6auvA3w/h3qj9iQybW42izprbUc5aL1+JGHeqvo9MyTXwuRos7bW1HOWiwVImrRzbjJWvtVLzQtpmnWRM3j4+HTM5ky/4XxnG6jhnBRLxu+FzsHDsWwyXZpeDwlTt2kO+I/On6iLq21aKFJPNhzc7UYRJyCsZHpMEa086S1uqGVQxxrRoWaWc1p/41oZ8Z+2LUiBqedJa3VarIY62KwLMqufxL2+R0ZJlMMKldTxglDvnqkpBrevDonQZibgKuFG6U2yCOOLqp0rROUgkHopGy8ebnpwoNSddhjbX/NMlEyfz9ftfRcmRyWVl2o9yOKBSffqikZJyasVGJR6f/xGtPMcBBw0/wChr//EAC8RAAAEBAYABQMFAQAAAAAAAAABAgQDMDNREBESExQxFTJAQVAgYWIFIUJSkCL/2gAIAQMBAT8B/wADobDWnVmPDvyHh35Dw78h4d9x4d9x4d9x4d+Q8O/IeHfceHfceHfceHfkI7PaTqz9KqKhPZhT2GXQSeos/RKiJT2YU9hkIcTcTqx1DUMxqGoahqGoahqGoahqBYt6ZTX1L0BnkFOIafcKfpLogp8s+gqMtXZ4EIflKbnkFOIafcKfp9gp8s+gqOtXZjPBpTLA5ycW9Mpr6lMNaU9hTuGQU/8A6kFPIiga1K7+ohC8pSVREp7MKeQ0hT/+pBTuIr3BrUrv62lMsDnJxb0ymvqUhUZCezCnyC6Cn6j6IKcxFe4MzOUQh+QvoMyIKcw0+4U/SXRBT5Z9BUdauzGctpTLA5ycW9Mpr6liaiLsKdQ0+4U/L+JBT2IfQVFWrs55BERKUlmYU9hl0FPz/iQU7iK9wa1H36BpTLA5ycW9Mpr6li+P/v4hpTLA5ycW9Mpr6li+8/xDSmWBzk4oexElpIeIRBz4g58Qc+IOfEHPiDnxBz4g58Qc+IOfEHPiCK7XFLSrF95/iGlMschpGkaRpGQyGkaRpGkZDT9CWcRRZkOBFHAijgRRwIo4EUcCKOBFHAijgRRwIo4EUcCKIrVcMtSsX3n+IaUyxzGYzGYzGYzGYzGYzGYz+iBTKa+pYvvP8Q0plgc5OLemU19Sxfef4hpTLA5ycW9Mpr6li+8/xDSmWBzk4t6ZTX1LF95/iGlMsDnJxb0ymvqWLuCtas0kOLFsOLFsOLFsOLFsOLFsOLFsOLFsOLFsOLFsOLFsFpNB5H6JENUTyjixbDixbDixbDixbDixbDixbDixbDixbDixbDixbBuk0oyPA5ycW9Mpr6lOd1D9Ew7Oac5OLemU19SnO6h+iYdnNOcnFvTKa+pTndQ/RMOzmnOTi3plNfUpzuofomHZzTnJxb0ymvqU53UP0TDs5pzk4t6ZTX1Kc7qH6Jh2c05ycW9Mpr6lOd1D9F+n9nNOcnGBFQUMv3G8i43kXG8i43kXG8i43kXG8i43kXG6i43UXG8i43kXDyIlUPIjnOoajifsQ2l2G0uw2l2G0uw2l2G0uw2l2G0uw2l2G0uw2l2G0uw2l2G0uw2l2G0uw2l2DFBpM85pzi/3R//EACgRAAADBwQDAQEBAQAAAAAAAAABAwIQERITMDEEBRUzQFFhUCEgkP/aAAgBAgEBPwH/AIHKbhI1LAcn8HJ/Byfwcn8HJ/Byfwcn8HJ/Byfwcn8HJ/ByfwIa2s1LDxYkDULw4kJyBHFxFExRFEURRFEURRFEURRFEURRDTMr9T2Hd2/t8GYhUFQxE3li9MQNQGoYib2MOYzeVy/U9h3dv7bkYA2yFQG2Yj/ssWYkQNQhUBtnYTw5jN5XL9T2Hd2/tsTEKhCoYmO4WP8AM5A1BUMTGdxPDmM3lcv1PYd3b+3/AAbZCoDUMRPwI/wG2QqCc/BTw5jN5XL9T2Hd2/tepn8hPDmM3lcv1PYd3b+16mfyE8OYzeVy9vQJtnExxyQ45IcckOOSHHJDjkhxyQ45IcckOOSHHJDjkglo2EmpmXqZ/ITw4jgKpiqYqmKpiqYqmKpiqYqmKpiqYqmKphpqZ7WuTZOBjkEhyKQ5FIcikORSHIpDkUhyKQ5FIcikORSHIpBLVsKnKy9TP5CeHEURSMUjFJoUjFIxSMUmhSaFJoUmhSMUmhSMGzK/U9h3dv7XqZ/ITw5jN5XL9T2Hd2/tepn8hPDmM3lcv1PYd3b+16mfyE8OYzeVy/U9h3dv7XqZ/IYw5jN5XL9T2Hd2/te2RxEpiUxKYlMSmJTEhiUxKYlMQ8IiiJTEpiUxKYlMSmJTEpiUxKYZL+OYzeVy/U9h3dv7bzefCTusZvK5fqew7u39t5vPhJ3WM3lcv1PYd3b+283nwk7rGbyuX6nsO7t/bebz4Sd1jN5XL9T2Hd2/tvN58JO6xm8rl+p7Du7f23m8+EndYzeVy/U9h3dv7bzefCTusZvK5fqEmzUOBCip6FFT0KKnoUVPQoqehRU9Cip6FFT0KKnoUVPQoqehRU9DQptMqf0rzZHEQMQMQMQMQMQMQMQMQMQMQMQMQMQMQMQMQMJldYyJiExCYhMQmITEJiExCYhMQmITEJiCp/3/ALo//8QAPBAAAgEBBAULAwMDAwUAAAAAAQIAAxESMTIEITNxkRATICIjNEBBUFFhQmCBMFJyBRQkYoKwcJCSocH/2gAIAQEABj8C/wCA21mYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiZhMRMRNX21rIEzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeM1MOPSxMzGYmYmYmZjMxmYzEzEzEzEzEzEzEzMZmMzGYmYmYmYmYmYmZjMxmJ9K61RZ1bWM7NLN811LN0W3H0TruohFtp+J2dOWXru6Pzjlt/RFhImduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxgtZj+fsXrMBNdQQ3EJnVsTdOvUM19BfH6yBOtUE6ikzqqFnWqGayT0H6K+LH2D13Ama9unZ09XzM90fE67sf0V8X13Ama9unZ0+MzXd067sf0X6K+LHrmudaoJ2YLzs0A3zXUNntNZP6q+G1kCdaoPxOzUtOooE61VrJrP6r9FfFj1jrMBM9u6dkhJ+Z1bEnXqN4JfBdZ1E1G9unZU+MzXZ1nY/nwT9FfFj1Tr1FEsBLbp2dOyZ7N067E+GX9XWZrqA7p2SFp1AFmuqfxNZJ8M/RXxY9P1mdaoJ2alp1FCzr1TNevxS/o9Z1E1G9unZU+MzXN0tZyfz4t+ivix6X1mAme9unZU+M1EKJ16jHx69DWRNdQWzqKSZ1AFnWqH8S0m3x9Tor4sej9aosNy1jOzQLvm0I3S1iT6Gt9wJZbaZ2dOzfM93dLXcn0R+ivix6MShsMtOv7Pfor4sejH7QqdFfFj0Y/aFTor4sejH7QqdFfFjpWqNUwEwmEwmEwmEwmEwmEwmEwmEwmEwmEwmEwmEwmEwlj9E/aFTohaOMyjjMo4zKvGZV4zKvGZV4zKvGZV4zKvGZV4zKvGZV4zKsyrMomVeMyrxmVeMyrxmVeMyrxmUcZlWZVmUcZlHGX6oFnox+0H+xj9oVPRrFssnlPpn0z6Z9M+mfTPpn0z6Z9M+mfTPpn0z6Z9M+mfTPpn0z6Z9M+mfTPKWv0T9oVOiHo2W/MwpzCnMKcwpzCnMKcwpzCnMKcwpzCnMKcwpzCnMKcwpzCnMKcwpzCnMKcwpzCnMKcwpzCnAlQJd+PRj9oP0V8Wvox+0KnRXxa+jH7QqdFfFr6MftB+ivi19GP2g/RXxa+jH7QqdFfFr6MftB+ivi19GP2g/RXxa+jH7QqdFfFr6NcU2GbQTaCbQTaCbQTaCbQTaCbQTaCbQTaCbQTOJnEzibQTaCbQTaCZxOs1voocONc2gmcTaCbQTOJtBNoJtBNoJtBM4m0E2gm0E2gmcTOJtBNoJtBGLMDb0V8Wvrg9EETxS+LX1weiCJ4pfFr64PRBE8Uvi19cHogieKXxa+uD0QRPFL4tfXB6IInil8Wvrg9EETxS+LX1weiCJ4pfFr64PRBE8Uvi19cHogieKXxa+uD0QRPFL4tfXB6IInil8Wvrg9EETxS+LX1weiCJ4pfFr64PRBE8Uvi16VqoTMhmzMyGZDMhmQzIZkMyGZDMhmQzZmZDMhmQzIZszMhmQzZmZDMhmQzIZkM64s8WOaQtNk02LTZNNkZsjNk02TTZNNk02TTZNNk02TTZNNk02TTZNNk02TTZNNk02TTZNNk02TTZNNk02TTZNNk02TTZNNk02TTZGDsWiBhYfFAUULGd3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3abBoGqUWUf9l96gFtkFI0lstg/VUUahUQtWa8f1rlRWJ+JzlMED56d/R88VK2Q/6f0HambGirUqkr0CR5RlGjahO6zXoolgoLO6zusZa1G4B+kY6plB/bFetm/QB0bHdbLmk5P4+DNRVtMFJ6agH2/WLHATmrj222QHkYrjZCv02/tiM+YwucBOaCPbbZ+lbDSKPbbZA4wPgKu6Lvgl6s10SyhUDTtqqrNrLFrLb8y0G0cttRwo+ZZzwO6Wc9LaVRW3RYd8vVDYsuU6yluhcq1grS9TNo5O2rKs2s6tdfzLUYHdO3qBN8BouHEu1aoVp3hZeosGEu1qoUy/Sa8vvO2qKs5unUtaXqpAX5l2jUQt8cnWIG+daus2s6tdfzLUYEfHJU3RejlHCZRwlSxRF3zKOEyjhL1UqgllCqGMtqMFHzLvPAn4gdT1TO0rIPzNsJZzwnZVFbdy3HqJegKYS9Va6Jdo1QxnbVVWbWWLWFvzLVII5L1dgF+ZZo7qT8clx6yhpfvC77yxq4t+JtZ2NVW6HaVlBm1nUrrb8y9b1feXFrLe9ui0Xk7eoFltBwwna1lBm1nZ1lt5bWNgljV1t+Jtp2ddD+ZbbqlRDXW9ZhL5PVvYxf8hcIFWupJlvlLhqJelowlRDWW97QOx6t7GD/ACFlynWUtLzmxZcp1lLdC7VqhWl6mbVjJUrKG9ozKerexiA6QuEsFdbZaMOTtK62zazs662/pVd0XfByc3og67TnNOqs9Q+Us5v8w1NEcmz6TBRqkmmdVh8oCMDLqa6pwlrObvucJ2pZzLBTu7pz39PrHV9MC6QtjrDvlSJ0Ul2lqZvOGppFU2f+5ZzU7Mshl99ILJ7CDni2r2gSiTZ8znKpa34mapObpW2fMWLo2iC2sZzumVbtvvjFqc4xIjReQXapQztajsfiWWND/b1GD/Muljq8veLUHnKm6L06kXfyiO9285wnO1WKUzhbBfrMd0/t1YgWYwnSK7Numu8d86hZTL9Nzd8mEuVdoIY/8pT3QyzRR2jTnf6hWZnOuwSzm/zC+iOdX0mCm5JS2wgwMuB5Ghjn5ioW5rRlnXvOfmdnbTPxAQxs8mEu1dosLubAJzOjWqltmrzl/TKh1+QlnN275bTLUz8RlD87o5iH3MXd0Gi8lOHRtD6t7Fpe0qo1R5YEK/InO6OxZBxn9vXa8vlLZzAcqmFkWppB5xjNlZului1ip+ZzVSoahlSqS97Gc0cttkBvPxgdWe0Qj4h/nEHxKla172MFI5bbINb8YKiF7R7ypuib+ikaq5a9DSGW2yKxL2kQMC+r5lrHqqJzOjWqnsIH0xyzewmxhfQ3Ib2M5nSCSlthB8orpgf0Ku6Lvg5CbNfK26f7om6Nb7xLg6AdRraHfKkXopLpNjeRl6haR/plmk0LfxZO2RkltCoG5B0BDpDrafLleLyFB2lT2E/xdH1bpsl4S2pQvD2sltalzbQSpui9OpF38ogv4LLF1DktrNr9pd0TR/8A7LRSH/jNdIcJzelaNc+Y0Mb+Up7oZaeUw2e8p7uRoY++UwvKseNZOv0KgfCyCzC9F3dBovIkdvPlqbotn7oIa2inrftMu3GKD3Fssr6PLCxQ/MtU2iVN0/3Rd3I26H+Up7pU3Rf5QbuSpuib+ikbdG/lKe7kCDAwvzBqH4ndHndHndHnOJQZJY9ur3/Qq7ou+DkPK26H+UTdOdA6h84KFUhagw5S9ZgohZcgwhlQL7RTU1C3XAyG0HkL1WsAjOuW3VEg52oq2+5lqm0Trop/EY3BTPuJdovbY1mqKTjZBNR5Vi8rxZUK42Qc+dV7ziikAF+OS06hOyNtnnyVN0Xp1Iu/lEflYPbYMIrUQC3meXmbwLnyjQxv5Snuh6Bh3ynu5Ghjb5T3cojxqT4GXgDYDqMC6R2bzvFPjLWrqd0aloNNua82iBsbYu7oNF5Ej8tTdF/lBLi1VLe1s1ztKSH8TnqJuH9samxJQSpumv8AdF3cjbof5SnulTdFvfug3clQLjZFap5HXA9M2g8heo1iiF0yxI26Nb+6U7PbkNzMIU0hRdOrWMIGpqjA/E2a8Js14Tm2NIP7TqAWfH6FRUzGB2pWLbjBLtBbTC1endHKZfFLq3sYobGyFKqhll/Qn/BlihyvGWBCPxA/9Rrmz9sRdBodUS5XW601w1NGa658jOxD2SwIeED/ANSrG7+2D+0odnFSoLGl8VWV5/j1bw+DLGpE/iXKgKIfxOd0hr9Tk1G7UHnD/bhv9su3WHzZFGktbU84GoU7wgSsLG5WSkLWgerSurLDDU0Rrrexl1FZ0H5li6LYfeD+9rc3T8wIo0KmXHmZdrLY0dKYtYxXq0rqjpulMWsYHqUiFBx5btBbxjNXp3Ry25anvP8AFN4fEunRbx3SxuxUy8hatpBhavTujkZ1pdW3GIr4iFKItaX69O6vKYaiUrVtxiK+YS7QW8YzV6d0cjOlK1bcYiviOULQW8YzV6d0HkuVArfBlqg0901V2slrlqm+OFCpq1CAqPO2AdApRFrRXq0rq8icwl6yONIS6Tyuq4mB2pdW9bbLh1aoamiaQbfmWLa44ywUTbul7TXuj5l2nj5mWGGrohx+mXFRyIp0glKcs87JfFLq3rbYitjZCDDV0Q22+UuKr2fMVtJJWn7TXDU0c3KntOwDWfEsFM8IG/qVY3P2wDQ6PZiKlUWNyGtox63tLtJXu8YvOWokAbGX07Op7if4rlh8Sw0TbuljXkU/ic5ptW83sIFXAfcBsjVqFcvuMsr0C28TXohllHRrv4l7SbVH+qWLrfzb/olrmUcJgP8Aga//xAAtEAACAQIFAwQCAgMBAQAAAAAAAREhMUBBUWGhECDxMFBxkWCBsfBwweGwgP/aAAgBAQABPyH/AMDZ2Av2eaPNHmjzx5488eaPNHmjzR5o80eaPNHmjzR5o80eaPNHmjzR5o84eePJCS4n7anNq+0Su1bKjdweEHhB4QeEHjB4weEHjHTHhB4UeFHjR4weMd2MYxjwg8KPGjwg8IPChvDb27jpQj9nlDy55M8meVPKnkTyZ5M8meTPJnkzy55U8qeVPJnkzyZ5M8meTPKnlSD/ALhm3lt/PtDaV2TUZrKRQ6ZsJaf3DBNm2kcjttmxexScyjEIdWgY5X72ymhdg+/knbJOTkxRYnWMbxjet4xvWta1resY3rUjMISTn3LmKdi4P2NTKHdjmJmiIg+ybFkJvVD+r+HA6pZt6vrk9gqaVm7M6ew/a0diFpXk7i2H/ihzOOSPQfK9hK5inYu4/wCRfP7YsDGyXxmFbUGDlxZHfkxdbhVYllP8lsY2S7eTC9qKMfzGkLz6p5QrewK5inYuYtCVJfIzun0ciT/TDKh/sS8dshdB8v1MmFs9F0rN2SRj7DP9vYduDJxTGSLgP5eCfKFb2BXMU7FzD8+TJmVZZBFNCF5I2HBtfsu5dXgMnzg6tbjNxbVbYPaerD90o8kixa9QlFsM+UK3sCuYp2LmCb1ockDGjZNhPJO5sv8AH+g1lvuyMJk+fUr6KpUluxU3swU1Legwp/2Pm/isM1thuyMU+UK3sCuYp2LnrKZUluSj9eo3a3cmkm1F8MrSwzuP5OcTk+fRrcEzTNxI52WQyJLekWgZGpjPHnlCt7ArmKdi56SJtctyjpNBmhhynWSM9DKSMbk7Ta98seNPJkhS0eWweynjhnMm1b9kXKFb2BXMU7FzukojWUiH4NB/tcDiuxoGRs3fsTyP3UzGzR1aIdTU7Vuhl7I92X9ofKFb2BXMU7FztYp1kOa2Tzfs7repb2w8oVvYFcxTsXO4qy/FFyhW9gVzFWMu9vLFb8UXKFb3pFjLvbyxW/FFyBWxXO7oVmPNGw+zYfZsfs2ptfs2v2bX7Nr9m1+za/Ztfs2v2bX7Nr9m1+za/Ztfs2v2bX7Nr9m1+za/Ztfs2v2OC37HBaT7eWK34olZpq8xReopQhCEIQhCHmDzB5n0IQhCEJeYPOdZSR0e1+zcsVv8Kvlit+ah4/eYogAAAAAAAAAAAOAtUc7dvLFb8US+XdhehMzMzM+JPE4CZmZmZmZmZnxJOsr+wCwv9vLFb8UfMFbFcgWPHYvdvLFb8UHKFbFcgXbcxTsXO3lit+KLlCtiuQLtuYp2LnbyxW/FDyhWxXIF23MU7Fzt5Yrfij5QrYrkC7bmKdi528sVvxRcoVsVyBdtzFOxc7eQKy/FHyhWxXIF23MU7FzuCsvxQ8oVsVyBdtzFOxc7irfig5QrYrkC7bmKdi52v6uoxQYVERERERVVczERE4KKensmaFGd0YUTETERETVVWzER2U67eUK2K5Au25inYuYu45exrmOHiuUK2K5Au25inYuYu72SuY4eK5QrYrkC7bmKdi5i7jl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkCx47F338rmOHiuUK2K5Au24ZYl2Zexdxy9jXMcPFcwWK5Au2CIxTI3MzMzMyMzMjMjMzOFk6pr4ttQGgv+MeA9hZ3ZmZmZmZmZmZmZmZmZmZvDnzARdFVWKWObshf83Eqqqqqqqqqqqqqqqq+PEei7x3Ri4IxLI7II95j/7lmRFmBJK0JTqPK1XqsjjQY3lu/WUfQOSAMCJ+EiLm2Vs9CKbKjH506p9jUaWkwKjtkWZtfpi2VF+ybC9E2bH6ZtfpkICSnWvpUuV4FdNRVeLUh70j0JhZXYPsjhejBqfKyZnJYb1rRxIuJnAQos1PSyMkhZFOzGktRlS0aSKyJqg8D1r6M22Q2jSRlNVMrBRvHRCzrsfmXYTV7Rsc8fxIs2igWEtRdd2aNBMzgpUXyiEp3dXj8qq7ZtFJdmxxGLSr810vJaSNLt+iwN6BSR3qIerDX3BG13mNY7spQmaLsUUtASj5ViJNrUN1QjY2adJUt3QS1fomPR/RSEGF9xZt1Lk9sr6LzDkxoQc6T3Ch6L27S4gc1Nkh51sYVBhcZooxlDURyT4I0V8k4+fdGMqNXDTQxpN7NEC+qzZK6J5tLYnCv9EYMFdhZp9NyLBn2GDoVmRECksmsiQzoqIJT+hPOxp9W0lLcIcNLmScn9SHiWwNA2OsKwZKanATlSu+uGqOo6Kq7RLJUymT+5DxKUydBNNSqrowoIzYwlNFQuKC/RGpy8hlDcIIhYBAg1gonZFsMSMyskvs8RAivWeRanVr5Efmt9hHTRaZWQzKqrJswpdm3vGICrs0UklEieK8gtxaRNC8/awtslrMZJLaS3GrSKZJyf1IbpSWTcDJJTTXpxvDQsoOeS50M0ctGRDF6HUjc6swYhSoDk5Cky2aD29nVlAnTzXoTANqxJndcriRdG30zhCp/N2JUIlDoIDeIJnquqmoVau7ZKvL0OhFQLOY2ZB6g1BYfuRw/wDuCGxiVE4OzEwMExcioJLk9i09hPv6JNS+mcByEqtPWRoZSVpDNTVnQXZ0qtOhcnv4/TlboqZMlkEcvhsxshyzpBvpkWZCaJNBZnzh+0dWgqZlBRzr6nGEUX916AkzlQWkmZutMsbn3VGTrqF6WyKxaOJQgRTnGHUNaZbrqmYoo3UetlyC8aU0hMDUepHeUjKxiBXhbWurOJAn1Yk8lZULXD1rBOKirBxnfWikz6HbU7lq60JgTUGO/DQIrDbndE/yDskXQdBlgRmJoMKUurCwdlTF2yfaaqErGKppakOWWypYlRZOhEBWQMswqq4fi5E+r6DeK1IVQBF8IVuqhYJG61XDGGuBVBtXJFNpFC7PFXAcLNJArVuLaBFxqlE9wJiUrlelG8NdNtUhZ9eUESZKiLyNC2t0hUVKFKl9iyyuR0zgHJ7oPEbBvEqZ1T+i/ms5B0k1+z4lKa9/68jajrwDldJyU/B9isFqrE1BgJzkK0+rJoeegXJ7+P05W666N0wISoLJdKD2RbsdZzyzCODT1FXbWxlkXacY43ZVNBTLVuvHFf0tepHBOMVodimCFHVTY1VDs6y4dyVRlKk9inE9QUNqg4j0K2l19+qFGKOOY0FgU98wW7EQB6m92qDdPMpi77NPoFs8d05Q/u79cj+7ucV2Fit3yieP6MEi4NY3+nYZnWz50uOS0yw9KN4a7d8x1WqglWBm5wZirbotqdXchw0484SpLFIchLIQcWlNdFPFsxJocEfIxsoCSgyzQlf7UQ0FTagTnrhmQy+CdPTI2n2bn9eAcgnhWgIatq/YhB2nQ8tSLtkWQl6FnoXJ7+P05W7kGeyrhvCZFZ6xZ9KdFxO+q5wey6OCcTo/Bdr9mRFpRITZjCdV5MRSoPifAlPIlmzauWykGcR6EbyeyP7u5xR1St0qESQiaIkIYfnO44UlScugIneEWTaq0dOc6ziNxKoQ9R/2Mm1VQ6VP5CwBRhb1aaPotEpLbLSpwnr0XKGii5WTUej4rrk0dSzBZyU0h4WeBii/NCKJDaPQS1KKIdPRIJCHeBO2mNsO16rDV2mNzrrBC0JBjAwyY4R85oj4ItEEdXZxHxl8HI9QZ0ZjpUFhDFUmLPVjtMakwujlH3bxMyMxZH3RzRGbLohfFk6oq6VZSfyKsL8OkgBcQyFkhKFQVmqy1Duha1JkBIrsI1eUH7F6Ee/dus3gshq5F2xTkSnRiSrqu0L9EFh+hUE2Pyoz0NHVijtohcCiHhmqt9+Q4CKhlDCt0U9ojOjtetao7KPTduZX0IvlLDZbl+sP1yzbGPH6l4bXbTQWVCqolk2RJzvdUlSvAxA2QStCKo+BxHQnrMqYcRbNBBcKhrrBTch9i0no3pj3BianrH/EYWdlUCM44Iux9qATYFHZJ5shiZN2+lTedA0g5RdUcyqENOpJP2FMe5QbQtAdzXHIhp0Q4jVAiGMyX+hCVOdzFORKdxR9yltQhUjZNSL2B1ToMldQHF1/QCwISoTbDoxA5WmdHIyT5bJExJJXahFtQsqy7WY20PVKLyvxFZF7V3F2t5hB4MuiMaa75in4o6DPYVlJDBZjUfPgRm2BzlFzuorogDErk0yf7ERQqF+QK2hobVGPiLcqgyAQVmFJX/DEE9nnISZ7m1v0Q1muf4SRYTG65kmyv1/4Nf8A/9oADAMBAAIAAwAAABDzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzxTz3X3zxSzz3zxv3zzzzzzzzzzzzzzzzjDUVZLzzzzzzzzzzzzzzDDRUfbzzIDDXDDbLCLCCCLzx8MMMMMMMMMMMOxbzzzzzzzzzzzzzR0jlAGZ4PzzzzzzzzzzjwyqdmYpoPx033/8A/wD/AH//AP8Afy6/PPwwwwwwwwwwww/FvPPPPPPPOIOfWsQJ39//APt/zzzzjzSRHbXRC8/f/wDrD899/wD/AP8A/wD/AP8A/wD/AOmvzz8MMMMMMMMMMMPxbzzzjjmZDIae/wD/AP8A/wD/AP8A+sPzjxO0N4or/wDP/wD/AP8A/wD6w/Pff/8A/wD/AP8A/wD/AP8A+mvzz8MMMMMMMMMMMPxbhv6BCDp9/wD/AP8A/wD/AP8A/wD/AP8ArJtugh9//wD/AP8A/wD/AP8A/wD/AP6w/Pff/wD/AP8A/wD/AP8A/wD+mvzz8MMMMMMMMMMMPxagt8+//wD/AP8A/wD/AP8A/wD/AP8A/wD/AL3j/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wCsPxT3/wD/AP8A/wD/AP8A/wD/AKa/PCiwwwwwwwwwww7VuAv/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AOsPzXX/AP8A/wD/AP8A/wD/AP8A4YfLGTrjDDDDDDDDDP8A6hb/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/APrD8GHc889t/c889/HH84xuzzzzzzzzzzz/ANqBv/8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP6w/Fu8/wDPPP8A/wD/APPPN3ywXPPPPPPPPPPPP1agL/8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8ArD8V9/8A/wD/AP8A/wD/AP8A/wCoH7z8MMMMMMMMMMMPxahL/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wDrD899/wD/AP8A/wD/AP8A/wD/AOqnzz8MMMMMMMMMMMPxagb/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/APrD8U9//wD/AP8A/wD/AP8A/wD6qfPPwwwwwwwwwwww/FqKv/8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP6g/FPf/wD/AP8A/wD/AP8A/wD+qnzz8MMMMMMMMMMMPxaz/LLDDLDLLCIz/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP6zwsMsMssAsol/Pff/AP8A/wD/AP8A/wD/AP8Aqp88/DDDDDDDDDDDD8Wss888888888sy/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A7Yc88888888sM899/wD/AP8A/wD/AP8A/wD/AOqnzz8MMMMMMMMMMMPxbzzzzzzzzzzzxf8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AOqjzzzzzzzzzzzz33//AP8A/wD/AP8A/wD/APqp88/DDDDDDDDDDDD8W888888888888X//AP8A/wD/AP8A/wD/AP8A/wD/AP8A+qnzzzzzzzzzzzxT3/8A/wD/AP8A/wD/AP8A/qp88/DDDDDDDDDDDD8W888888888888X/8A/wD/AP8A/wD/AP8A/wD/AP8A/wD+qnzzzzzzzzzzzxT3/wD/AP8A/wD/AP8A/wD/AKqfPPwwwwwwwwwwww/FvPPPPPPPPPPPF/8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AKqfPPPPPPPPPPPFPf8A/wD/AP8A/wD/AP8A/wDqp88/DDDDDDDDDDDD8W888888888888X/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wDqo88888888888899//wD/AP8A/wD/AP8A/wD6qfPNwwwwwwwwwwww/FvPPPPPPPPPPPE//wD/AP8A/wD/AP8A/wD/AP8A/wD/APqU88888888888819//AP8A/wD/AP8A/wD/AP6qfODgxAAACCQwAACfdvPPPPPPPPPPPCE8/wD/AP8A/wD/AP8A/wD/AP8A/wD+PPzzzzzzzzzzzx/7rPPPPPPPPPJKfzy4xzzzzyxzzzzz7zzzzzzzzzzzyywwwwwwwwwwwwwwwwxz3zzzzzzzzzzzwzzzzzzzzzzzzwxzzzzzzzzzzzjzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzTzzzzzzzzzzzyzTDzjDDizDTzTjgBTTTRjjDQDzhAiTjTDjSRjjjDjDjDDzDDzxxTjDjzCASAQjTzTDxzjDDTzzzzxAjhxxgDBSgABzRzwBhhQiwDzzyhSxxiRBBiQhAgAyixyxxzzyxyggzTwBgiTyBTiCTjwRDzzzzwjwixijgTiBxjRxCDyjRiwQDzzyBSSxDCgChSiChBygBwBjzzDjywCxQxASCDBzyhByyTgTzzyzzxzwzwxzyzyzxzyyzyzzzyzzzzzzzzzxzzzzzzwzzzywiTTzyyzzyzzzzzyyyxwwzwyzxxzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzz//EACYRAAEDAgUFAQEBAAAAAAAAAAABETEQMGFxkaGxIUBBUPFRIJD/2gAIAQMBAT8Q/wADlEPLgP8Agf8AA/4H/I/5H/I/4H/A/wCR/wAj/kf8Hk4Ts1ViFhQbqGz9dlDSEGriTBKKrdamW2ACnrsxbnN2CMlYnR5qPylLFVVkmba6qZKxOhLFyMMLYKpZX+NC9CuzFuclyYIh596Z0pFbInD/ANTQ2VmGEIxXFQ/OZCw6ns6F6FdmLc5P7VUSaOftjxUSQlFtTQ2n8SikwPPRFEShiqVeq3dC9CuzFuclUl1MTQTh+aPKwqqs3podGjofsqY6V0EyfstC9CuzFuclVEY/aOq9toXoV2d3mrx+4yFUGwyGToZZlmWZZlmWZZlmWZZlnQI1eP1WVHGjBgwYNGjBgwYNGCI1U/EpkamRqZGpkamRqZGpkamRqZGpkamRqZGp4eV4/VZRo0aNGjRo0aNGjRoivXb3eavH6rQudKQrsxbnJXj9VoXoV2Ytzkrx+q0L0K7MW5yet9C9CuzFucleqxbttttsZiy9kuMh7dlstlibT60hehXZi3OT1u0L0K7MW5yet2hehXZi3OT1u0L0K7MW5yet2hehXZi3OT1u0L0K7MW5yet2hehXZ3eT0G2V2F6FUAVMfpg9TB6mD1MHqYPUwepg9TB6mH1MPqYPUweo1Qt5bKopjDEmJMSYkxJiTEmJMSYkxJiTEmJMSY3QYo12AyjKMoyjKMoyjKMoyjKMoyiVb2Df7Q//xAAoEQACAAQEBgMBAQAAAAAAAAAAAREwMWEQQEFxIIGRofDxIVBRkLH/2gAIAQIBAT8Q/gcl2jchMPuQ+5D7kPuQ+5CYfch9yH3Ifch9zSTKtVWICcVHJNFWNRCxw+II/wBN5uNxuNxuNxuN5uNxuIyGPdzf8BT4wGurGuiGyg3a40praVRr1EqiGKDdVlyBTwo4R4Y8EcNeD7udKW0qYmP8oZGzrxMpSXUMRH+UM6jbdeJlHCjkfu50uNoqxqoPSQ3alakJLKXA2lUadROiGyg6hy2UcKM3Xg+7nSwbSqI6jXRDFBuq57EqSLEx/lDbqNt1NpzKOFGbrwfdzpcJDIwybKOFGbrwfdzf8hfUjKOFHI7C5/JefnIvPzkXn5yLz85F5+ci8/ORefnIvPzkXn5yLz85F5+ci8/ORGQ4iwofTso4MaKnf/8A/wDsaLxd40UbxuG4bhuG4bhuG4bhuG4a9BYUPp2UcIiCnff1Xbmg8e7m0BYUPp2UcKORu7nSwofTso4Ucj93OlhQ+nZRwo5H7udLCl9OynhRyP3c6WD0SLZbLZbLZbLJbLZbGycHkbjKC2Wy2Wy2Wy2Wy2WxygeFHI/dzpTVyOxW5tHI/dzpZ7ZW5tHI/dzpZ7ZW5tHI/dzpZ7ZW5tHI/dzpZ7ZW5tHI/dzpZ7ZW5tHI/dzpZ/rc2jNfBtVtfwvejL3oy96MvejL3oy96Mvehe9C96MvejL3oy96MXW5LacxHAtFotFotFotFotFotFotFotFotDlGM1oLEuF0ul0ul0ul0ul0ul0uCH8OCBDggQIcMP7g//xAAuEAACAQIEBAYDAQEAAwAAAAAAAREhMRBBUfFAYXGhIIGRscHwMFDh0WCAoLD/2gAIAQEAAT8Q/wDeenCSeLn9XOE/pp8M8XJOE4sdQ25ojahtc2sbcNuG3DahtY2sbWNrYU2gbQNqG0DaptQ2obWNrG1jaxto28UPiC6UnJki/TMyG2J00G2nJcMk5MTm1kL9G3zG2TNTIlToTWMZOfmCav3ufheee2+/fJ9e+T798m2/A8UV96+cdvtvm2/9Nt/6bbNv/wC4DfSvnCr6d84HfWvk5QWQ34XZj4kmigUklCHnesN+G/cIMhvk34b8N+G/Dfhvw3r4IQhvw34b8N+G/DfmMISkW9Q7Xs3YVBfoW4JJJJEkoXVkYucyz6CxWjJKC9SMP9WCCdMWjL61GTOCOMbWZI2Tg0VW0lzYxhjyUc/Dgn5iR6nl8YbKovE/UY4UqSOOhZ4wkxo5qRF+9Yb98D323ebvNb1xv/C27zd+Ft//AOm/sP7/ADT9eL+/xt/d5u/Czjnvh21RG9phOfZisvA8ErLib53OBcczIyw5nliJW3WmVNkITGdUvcR5jY5a/wDgHbdQmX6szqkN2E7kdhhnxLZYmg6rkO+qSENknNLS+w7F6EJjF0suXehAGWr+A1lnet7iCskcoMx/R+Bi8A2FfLh3h3IsungdsBkuJvndjuLing62JhSyUasjY9WozhxySRHnFXzQnVhJKrzLXuj49CCvD6/6J5WRaxdWL3x9xHYfC4d2Jb6jbyEMaEubOsfA6ScPRPuQyTOPgWnQTIXQlvDoEhUr6hOOgqOplTF+Yjt/hDFhpxg3ciy6eB2wGS4m+d4O4uEbUwVJHUTgbmk5khayr0hvRChNcl8xhkNlUQyWmLbCgZzK1cjFCsOo+RPJUHW9/BGNXmXufR0FwEkjYtRIypDNCGevohp9hoSTKgJh3opUtE3KpNCRmJ1qP3EkqJIVLkuKHNl340dqMWGh3Hi7uRZdPA7YDJcTfO8HcXATh5jIzp5ls2MtEPySJlmXREoTH8Z1SSOC7ulFEKB5jZ1SSwgXkdDMrkLqZ+Oz6XPo6C/CycFhPMkcXbgU3Lk1n0Ojfgxh7CilSacKE9GSGnDBNhai5wJ2ukWs/DXX8CH9H4GLDQ7jxd3IsungdsBkuJvneDuL8slCXkKSllzNiSs8kQyNqifmL3TlO7C2FKsll6i3z/WNJsp8CmSdV+ez6XOw+ELxsuPozmFLQ+o6raKDCQrslixUWUsE51VAgomjYSjQYDFXUEidIROs+B6wTr+ZHajFhodx4u7kWXTwO2AyXE3zvB3F48ybid5JJkb1NdtAjaBDyNJpIoQmcWzNtof3u5g/gXXud7hJLJE0KZ3K+COBs+lzsMGOzE1kTQT1RM0EJVpdStsSWmhohcXE0dQ1OxNl2IU/cucKlo9CJ1akwS6SirFC8EYR+dE/L+Biw0O48XdyLLp4HbAZLib53g7iwkpqSSkUumMqFQcoOTWtw1tivmrePYytiXbzZIpE0tqGQmsvq5FRQ0oG3RYSeYrcPb1Ir6JMGVD7Y7eIjnK4+9u5agRN0QrZtbu+5UrWhvgNjVuwyFS0YN5458QjtBiw04wbuRZdPA7YDJcTfO8HcWDvXMbhVaXUgF0IZElE2SCfmJuqN/gPjSPkoGdbdsIU6Yt4THGa8uor3aVUlP0MlH1J+ZX8lJy9BLVFdoQzntxg0mrLfNlJuKi1wnCeMR2oxYaHcOLu5Fl08DtgMlxN87wdxYOxLGEO4PUFU2NuSaQoET0FQeiJ/QdBncac3MCihe51qTyJepXw5mfGIn5fwMWD4yLuRZdPA7YDJcTfO8HcQx2O+8C6/wDFo7QYsHxg3ciy6eLZcUu6HdYMdsIt9P8AjkXgxYO/GDdyLLpihl+BUQ+JXd4sdj7HIt9P+OR2gxDFxaz+MKWXQzx/wf3hzgSf8BRr6bEyt94WUpSlKUpSlKUpSlKUpJKzAYRLmjNbCsMdj7HIt9P+OR2gxYyi+NNAnUl9Bto24bcNuG3Dbhtw24bcNuG3DZJsk1/Rm3Dbhtw24bcNuGzjZJqehNtE/wDiFui6tLFbw0giEJRqVKlSpUqVKlSpUqVKlSpUqVKlSpUdoFztghjsfY5Fvp/xyO1GIY62kVcqi4h8i9hF/A3jXCCCCCCCCCCCCCCCCCCCCCCCCajnkIY7H2ORb6f8cjTZfAxYLBcT5YZeBlHODue0+xM+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5MTaMuhCxQ5wIY7H2ORb6f8cjtBiHhE4g0tBtT2mbMzYmbMzZmbMzZmP8AjM3YbsNmZszNmZszNmZszNmZszNmZszNmZsTNiZszF/GZd94KTGRBDFafDdFNI4mE+hOvqPBjsfY5Fr/AI5HajFxu2p2T4LF4GXhWrfib53g7rBjsfY5Fvp/xyNNl8DFxo2uGWLwO2AyXE3zvB3EMdj7HIt9P+OR2gxYacWNrgli6CxdsBkuJvneDuIY7H2ORa/45E/L+Biw0O4cVa4JYugsXbAZLib53g7iGOx9jkWv+OQ/o/AxYPi4tcEsXTwO2AyXE3zvB3EMdsIt9P8AjkdoMWD4sbXBLF0Fi7YDJcTfO8HcQx2/5BsjtRiwfFxa4JYungdsBkuJvneDuIY7Hf8A/HpXwbFg+Li1wSxdBYu2AyXE3zvB3EMdjuy2v3EDXI8/z0KeBSdt8FrCw04sbXBLF0Fi7YDJcTfO8HcWDl0kd+TmAXg5ejNlZsrNlZsrNlZshshsrNlZsLNhZsLNjZsxszNkZsrNkNkNmFMS0hWlfolUScCUttITIUSMnRGyM2c2Q2Q2dmys2VmyGyD/AIjNmZsLNhZsLNhZtxszNkZsjNkYu0UpJEeDQ7jxVrgli6eB2wGS4m+d4O4sGqjVcxIggjwQR+Sso9qIeX6GaQoguUOlxwQR4Y8MeFoqLV46HceKtcEsXQWLtgMlxN87wdxcT7QVqYZ/oMheKV6HceKtcEsXQWLtgMlxN87wdxcSvaF+iZcXV6HceKtcEsXTwO2AyXE3zvB3FxPsBfomR9LnxR6cWNrgli6CxdsBkuJvneDuLifYC/RMhhdn4l8XFrgli6CxdsBkuJvneDuLifYC/RMsK7PxL4uLXBLF08DtgMlxN87wdxcT7AX6JkILs/EvixtcEsXQWLtgMlxN87wdxcT7AX6JkfS5nZ+JfFxa4JYungdsBkuJvneDuLifYC/RMsK7PxL4uLXBLF08DtgMlxN87wdxcT7AX6JkILs/EvixtcEsXQWLtgMlxN87wdxcT7AX6JkfS58Ueh3DirXBLF0Fi7YDJcTfO8HcXE+wF+iZcXR6HceKtcEsXTwO2AyXE3zvB3XFewF+iZGT0cU+nFja4JYugsXYumXE3zux3XFewFb9FlxjP3DirJnZCxeBjcg+uZpTiFLS5j4r2QhywkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkVqlzyFVXi+W8RnmZObDKk8pYvA7DNqmhdPwzazY8P282828282828282z+lf4zbjbjbjbDazazbTbzbzbl/ptw+pT9CeuZSZNxcQ3DIzSE3QQc4MBsxt5t5txtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxCs8RObM1LDG0GzIy4dxmK/InuIaZgVSQNsNsNs/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/pT+AoOZuoRXJLWLUVEsYIwJPMggj8cEfiZzD4kkigJakEYkEEEEEEEEEEEEEEEEEEEEEEEEEEEEDWjIpcjqLiFCxGnE+eEc8I/8M5/LXQf6CXp+NfqHhP58/H0/FPApYQpqSPXumIKYJxUND8c+JvBqqH3IXhcUi1xYz+H2G5zMQocvHJD8LMh1lDaUW+hCaZl11geM4yN8M2hWPMmeT9vBIYlNTQ3+lG5GpvRH0H4GXONtIl2ERt2YMdP6PkXPs+Qq1BFQ5KozI/A7UKGa6wOHZXcS5wNsJ0P2PDOMNawUvIV7VUyV8GkFraZAmWkzG1+ZZbcolcnb8EVEk6UJPPBXEvUJqQIqOvdrAghqSsRDFCttMkRDr3SibVEr6JO7wz4F2IjYjiNVKJnLsIEaBDvXgO8mT1+591oKGXjaCuPq0WJpSzD9EJtUWsxGfDhXA4lFRsr1GeYzINS0Q5LSjg2hImW4liXqJT4tC36XM4qydi+RFwZskKjacKaXkJZmZkLjVcVpLLh7TH1J1INZzfoNimjNMQKYHEkMSyajS/YcKwtlcRTEzZVCuLE3ppCKX67EBl58oaxFro6v3CbQhVtK/oJisxChjkiKtEoXGI4SKc+h1HpWVbT7h0i6UchiLjvB0mU4SaBDLpSDT9DM757eHmpVajmG9u7gP8AihbDmRlUFQLHWknMaT2xsom0FP4C6gaUXTaloJDouHE1KepI5I1xQaE9dOT7DS7qzSRPY9VaRFs8mfS+DJJt2IAXAynaLXEcqlGjQ6Ic26MZrjxAub9BighmmckxucJUFkOCTT9MOSlmUJp+6iYySXZDyu5HyZF3UHQo6sbEvRtoFiGWbsLRQnd+lz3wekkZtwXZYTHag0OK+chdazhTA8rVJBqNTJlQmZFJZKaTXn4O0Z2xlQqn7dZ0h/kZXFlm9EVv9h3e6E+X+C01Ms05QhGTpbYS8xrRjiSRVgM2GVgig+5VzEUtHKH6oXZkUmTyESK82kNNukIeDqEO5I0DuORGXyjM2iwzIdQS0aixtIpuVAXmW7NZKoSjpoobcqhZMwrqxZ4ZPZIRHi2lNLOmFRTdlK0lkujrjVe2ZlUIyKV1KRr54NymMnvCS6iFyMozQ0JDNoSKKYHMDl/2HpDoU5PsJLws05Q/wd5O7e591oKafMX2ngqqjzFaywbTeTZPDFAFYAWyXJoXJC7zlyFNkE09UxtWjqOnOysUp5cskKfxUNpekcOSWaG64mtNHSj9BgJqKG41R2L5FTl1jAaqkivMVsVLnP5QokpAYhRiNGlVQqMKRNk+3oMQrpReMkPNiGnkSKDbD6zQt2EpSXyCSkSkL0iFDaojmVQpm0skCVYabKQkCzanNcy5z57dNeiKmfFAk2JWWTpeghsaqqPXB1S61cHRCk85SSE1axrFlsm2Jb/BorL7NRFOkXrDvnt+DvsdT6vMtdMJBVm8xgdWIaJjPySjNI2hcbsdNwSTYkHodxzB4BSSc+rF4KzWVWSQbdzkWo5NWJRsmGuVZs7x7FMKNJCpLVknsKhpoRtWPlqcxYqkiXZSOn2oAhg0mTK5NCsQH5quJQ/qVgU0bumUUrEdjvHsIKrqFrDIGETNWqcuYhzq7oENIlyZc0ILFGbR9URyNaXvRIhkzfLIkzcaKkVZdku5Tq2KzarzGyNJLQPIZ2QbOaI9c0RSVBMpaIUPol4O0Z2wigHUL7StQVUnzFkxyGhm+tRpjPJMqduRuPNRS3XWs06DVOlkhYlLJm3CbFzJQ0Z9yaMWrENFOMw8pdGqj+lcNtpKMpEJaHBXhMU3ydYkaCFs7Koj+AKJ7D2IFkpUh7jRWi7DPNE1WtiV1lmnWE4H0nQ/QO+dKSoSj0TrsPTLQ5VLHRYqVK186iIeiE9edDXoOpF0bvEwRiwaIqWMqlEToHVSIUrxkTin+Q1mJb0IKUoRLk27mucXmooiO0dKfJO6K/Zgbuum8hMdYC8feTu3ufdaFjzHzi2mZUxRNLSj4BBiyK6itFH2JoQjrQ2JqqoM3zFhHIqqGtHYvk7/AOx9fmK2Nzr8o7dDkKutD5j0jSnwVahWnqVs550HTK4bVKIKGKqkOqPU9wW+q9se1fsM4ijrKWkGZB3bCm0lLdBUaCts+YiSnvMX1/ogNQxNV/0R0fWDpyL89nFI5Btu0HfPb8HfY6n1eZa6Yd4ylRAz5ipqcZBFrIjtYs43kNVMo7T8xKxeKhJSINHZJ/0V1I5cqTMs3Yd49iXk/c7admJCr3Kw7rCqTX7CFSFH1kO3K6xxI7x7Cq6EHzqJluVtLVoWh6icDZRxYgKbiESJRlScDah5XP5eBpCvKSgh5SFDRM+o08HaM7IVi/rQ++jmXIsbYrb2GbhSSXqVpzLPoVycORLk9SpJKv8AlPQQqOzDb0ZSttCwvURuuZZM76dw9z6rTD6TR4C7T7He/Y7oPpNMO6HcvcsdMfv6HZo7l7HZ/c+g0weCamDFr0hNOwcsLvFf/QUynuC0Pms3uJKuhWKn4O8ndvc+60LHR+C/faM+jzK1uIiNzbLgUTmalMyrOkl5DJJaU8FiZnO5Esx78YkeaKOj+RmVbtJLkX5POYKr2YSmsGk+NtonoP6baDcpMpo9iiV5VSE5MlmUxUSd42Og/EjXqyFCqtHE3K4q2Gj6tEqiGRIBKTmDMzO2+B25sXdvbBHWIvDusPdtubUehyLLJaKhrBQ0JYhJDw87SFH55nszvnt+DvsdT6PMtdMO4FTmSssHalx5PARKJTWBM+04zXqvjPsVlWgnpHfPY7H7nbTtz2Me++x2f3OzYkd89jtvufYaLHuDtl7IQjLCT0epHRvlNOzaHe44Yq3UXWrWJyimUd+SFX5NAadcj3BOTqfUaeDt/hnbCsXdaO5e+PfvY7mGSY3Cj7FMNqRp5HJgTUoc1V3cZ9RdVTq1HSh6zjVpXQI6WsVULxvKoxxI6y6YfaaMp6X3O0ihY5aeQ2CKdOaRIaKmQNdMFBKSS1Fqt8MlNxBiiOTBkzDGCXxu6JN0dmhG1JS37BJhWiVuaUkvg0JzkMQ7Dgs2lUeSSBDqRNRJCpbFHoffPgzvudBiconPPoIzR0qBJ+n4Kfe9RjUV2ZqGm5GWsVPyQzovIIxFSROD7CJeg/tI5KImSPYyNHFhFFWmkXprhMcjDtT4pSRpdamRtm6InmJfJp3WfJqyGbVNozzMfKYrTMQEkIaeaHVkzS5aFvLVb6IlF7dGyIHPSOTrnk1ZEqCim92QyvAMVRn1KOTDhegY5GSys0w1CSHrFRa51Td65kAkSSougpusxU6h2IrqzqoYm4Kj9wyKGDkFsSabZUGf65jjC/qDlR1QiklA80MOXYmfJ5EjhUZKDlNYIWoIo/6JZ4dXDLRJf6Lj5XRkHOeXIYHerUK9GCJ8UYbdS5MdIwoicihJ0waLS5aXKWulJPB3kkdlaS/UvpbeFzbMXmrERdky6FiUoLqqjcsS7ieUie9QtYkrV2mhL76jpCcyKUKVoOCr3kBxgS+Jl0GS6wJszRGrJlI9egyZ9bVDqU4oJVsK0ldpoh7aDpEyR7dxDTjktHUVmJKScF7MhqTehVac6inoyyuVoyQgjMg/Qbg4rSb7sePMzPRSSldPoXgiUiELnWQsJUosNUdQA2OuMm6UeQ1D5hqICyMpzDVIKNoljJ8k+R+fn9EjzFlRBqlLnJdEoNjOWoYX6ShHmivuSJJ8mPHSkjiuTdSVSpdI10VZH33U9TgYcqpKlRBQZJoxHUrfQKksGh5GUwOQKlpnA5udCkfkhK0ioh8yuvFCuIVtmU/yYz0lLUTZdcPJCKakourdKshK+ocMXwqaeXJjmzVwlJnSSdx6OKR5CYilT6seZOrHqIpkboonRkx8uzYRaQm4SLq1UVQlVOXzuHvoAbli8a5YPkZlTPw10FhDM/BD1Jeh0KiM8ameEY1w6PwIz8HTwsX4o8GpUzwyph1xqVMvHHjnwYtB6jUxEsQtIdBBRRbLb81Il0W9WW97TPfKJQ6JhLpkeiC8knSVb5chTm/Ejl4M8eRkVIeuNcPMfLDr4ankVyxrPirOCGlyHJ6/9GihJzUl0PVCvvuSrD1/+DT/AP/Z" style="width:100%;height:100%;object-fit:contain;border-radius:8px;"></div>
        <div class="logo-text">ITI Assessment<span>Management System</span></div>
      </div>
      <h2>Register for <em>Access</em></h2>
      <p class="tagline">Submit your details to request access. Once approved by the administrator, you'll receive your login credentials.</p>
      <div class="steps">
        <div class="step"><div class="step-num">1</div><div class="step-text"><strong>Fill Registration Form</strong>Enter your SI name, trade, ITI and mobile</div></div>
        <div class="step"><div class="step-num">2</div><div class="step-text"><strong>Wait for Approval</strong>Admin reviews and approves your request</div></div>
        <div class="step"><div class="step-num">3</div><div class="step-text"><strong>Receive Credentials</strong>Get your User ID & Password</div></div>
        <div class="step"><div class="step-num">4</div><div class="step-text"><strong>Login & Generate</strong>Upload files and generate reports instantly</div></div>
      </div>
    </div>
  </div>
  <div class="right">
    <h3>Create Account</h3>
    <p>All fields are required. Your request will be reviewed before access is granted.</p>
    <form id="regForm">
      <div class="form-group">
        <label>S.I. Name</label>
        <input type="text" id="si_name" placeholder="R S Thakkar" required>
      </div>
      <div class="form-group">
        <label>Trade Name</label>
        <input type="text" id="trade_name" placeholder="Trade Name" required>
      </div>
      <div class="form-group">
        <label>Name of ITI</label>
        <input type="text" id="iti_name" placeholder="Name of ITI" required>
      </div>
      <div class="form-group">
        <label>Mobile Number</label>
        <input type="tel" id="mobile" placeholder="Mobile Number" maxlength="10" required>
      </div>
      <div class="form-group">
        <label>Year of Assessment</label>
        <input type="text" id="year_of_assessment" placeholder="Year of Assessment" required>
      </div>
      <div class="form-group">
        <label>Assessment Location</label>
        <input type="text" id="assessment_location" placeholder="Assessment Location" required>
      </div>
      <div class="form-group">
        <label>Duration of Trade</label>
        <input type="text" id="trade_duration" placeholder="Duration of Trade" required>
      </div>
      <div class="form-group">
        <label>Batch</label>
        <input type="text" id="batch" placeholder="Batch" required>
      </div>
      <button type="submit" class="btn-submit" id="regBtn">Submit Registration Request</button>
    </form>
    <div id="msg"></div>
    <div class="login-link">Already have an account? <a href="/login">Login here</a></div>
  </div>
</div>
<script>
document.getElementById('regForm').addEventListener('submit', async(e)=>{
  e.preventDefault();
  const btn=document.getElementById('regBtn');
  const msg=document.getElementById('msg');
  btn.disabled=true;
  btn.innerHTML='<span class="spinner"></span>Submitting...';
  msg.className=''; msg.style.display='none';
  const res = await fetch('/api/register',{
    method:'POST',headers:{'Content-Type':'application/json'},
    body:JSON.stringify({
      si_name:document.getElementById('si_name').value.trim(),
      trade_name:document.getElementById('trade_name').value.trim(),
      iti_name:document.getElementById('iti_name').value.trim(),
      mobile:document.getElementById('mobile').value.trim(),
      year_of_assessment:document.getElementById('year_of_assessment').value.trim(),
      assessment_location:document.getElementById('assessment_location').value.trim(),
      trade_duration:document.getElementById('trade_duration').value.trim(),
      batch:document.getElementById('batch').value.trim()
    })
  });
  const data = await res.json();
  btn.disabled=false; btn.innerHTML='Submit Registration Request';
  if(res.ok){
    msg.className='pending';
    msg.textContent='✅ Registration submitted! Please wait for admin approval. You will receive your credentials after approval.';
    document.getElementById('regForm').reset();
  } else {
    msg.className='error';
    msg.textContent='❌ '+data.error;
  }
});
</script>
</body></html>'''

LOGIN_HTML = '''<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Login — ITI Assessment System</title>
<link href="https://fonts.googleapis.com/css2?family=Sora:wght@300;400;600;700;800&display=swap" rel="stylesheet">
<style>
*{margin:0;padding:0;box-sizing:border-box}
:root{--navy:#0B1D3A;--blue:#1A56DB;--accent:#F59E0B}
body{font-family:'Sora',sans-serif;min-height:100vh;background:var(--navy);display:flex;align-items:center;justify-content:center;padding:20px;position:relative;overflow:hidden}
body::before{content:'';position:fixed;top:-200px;right:-200px;width:600px;height:600px;background:radial-gradient(circle,rgba(26,86,219,0.2) 0%,transparent 70%);pointer-events:none}
.dots{position:fixed;inset:0;background-image:radial-gradient(rgba(255,255,255,0.04) 1px,transparent 1px);background-size:32px 32px;pointer-events:none}
.card{background:#fff;border-radius:24px;padding:52px 48px;width:100%;max-width:440px;box-shadow:0 40px 100px rgba(0,0,0,0.4);position:relative;z-index:1;animation:up 0.5s cubic-bezier(.34,1.56,.64,1)}
@keyframes up{from{opacity:0;transform:translateY(24px)}to{opacity:1;transform:translateY(0)}}
.icon{width:60px;height:60px;background:linear-gradient(135deg,var(--blue),#0B1D3A);border-radius:16px;display:flex;align-items:center;justify-content:center;font-size:28px;margin:0 auto 24px;box-shadow:0 8px 24px rgba(26,86,219,0.35)}
h2{font-size:26px;font-weight:800;color:var(--navy);text-align:center;margin-bottom:6px}
.sub{text-align:center;color:#64748B;font-size:13.5px;margin-bottom:36px}
label{display:block;font-size:12px;font-weight:600;color:var(--navy);margin-bottom:7px;letter-spacing:0.5px;text-transform:uppercase}
input{width:100%;padding:14px 16px;border:2px solid #E2E8F0;border-radius:10px;font-family:'Sora',sans-serif;font-size:13.5px;color:var(--navy);transition:all 0.2s;outline:none;background:#FAFBFF;margin-bottom:18px}
input:focus{border-color:var(--blue);background:#fff;box-shadow:0 0 0 4px rgba(26,86,219,0.08)}
.btn{width:100%;padding:15px;background:linear-gradient(135deg,var(--blue),#0B1D3A);color:#fff;border:none;border-radius:10px;font-family:'Sora',sans-serif;font-size:14px;font-weight:700;cursor:pointer;transition:all 0.25s;box-shadow:0 6px 20px rgba(26,86,219,0.35)}
.btn:hover{transform:translateY(-2px);box-shadow:0 10px 28px rgba(26,86,219,0.45)}
.btn:disabled{opacity:0.6;cursor:not-allowed;transform:none}
#msg{margin-top:16px;padding:12px 16px;border-radius:10px;font-size:13px;font-weight:500;display:none;text-align:center}
#msg.error{background:#FEF2F2;color:#991B1B;border:1px solid #FECACA;display:block}
.links{text-align:center;margin-top:20px;font-size:13px;color:#64748B}
.links a{color:var(--blue);font-weight:600;text-decoration:none}
.admin-link{text-align:center;margin-top:12px;font-size:12px}
.admin-link a{color:#94A3B8;text-decoration:none}
.admin-link a:hover{color:var(--blue)}
.spinner{display:inline-block;width:14px;height:14px;border:2px solid rgba(255,255,255,0.3);border-top-color:#fff;border-radius:50%;animation:spin 0.8s linear infinite;margin-right:8px;vertical-align:middle}
@keyframes spin{to{transform:rotate(360deg)}}
</style>
</head>
<body>
<div class="dots"></div>
<div class="card">
  <div class="icon" style="background:#fff;border:2px solid #E2E8F0;overflow:hidden;padding:4px;"><img src="data:image/jpeg;base64,/9j/4AAQSkZJRgABAQEASABIAAD/2wBDAAYEBQYFBAYGBQYHBwYIChAKCgkJChQODwwQFxQYGBcUFhYaHSUfGhsjHBYWICwgIyYnKSopGR8tMC0oMCUoKSj/2wBDAQcHBwoIChMKChMoGhYaKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCj/wgARCAKbBLADASIAAhEBAxEB/8QAHAABAAMBAAMBAAAAAAAAAAAAAAUGBwEDBAgC/8QAGgEBAAMBAQEAAAAAAAAAAAAAAAEDBAIFBv/aAAwDAQACEAMQAAAB1QAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAH45TYzzPQ0VnTjvRWdDRWdDRWdDRWdDRWdDRuZ3w0XucjRWdDRWdDRWdDRWdDRe5xa7qZ7vO+hhBIAAAAAAAAAAAAAAAAA4jvPShors3KHC8Uat2Jlu9XRPQAAAAAAAAAAAAAAA4d568RFc+pEPxn07tGvPWjo6s5+f1SZXF88dt5+hnzyl9DPnkfQz554j6HfPJP0M+eR9DPnnkPod88j6GfPKX0M+eUPoZ88pfQz55H0M+eR9D+XFdqqkOZApMXKRfzH0IUXAAAAAAAAAAActlTte/HYB73iBIAAAAAAAAAAAAcO8567n2lfh+arv48ziuM2nw1Ec57HE+k5zDnNffz3hqsxDzGn6EOrAAAAAAAAAADg6/PrI9rtfh+Kbv+Mziuc2oQ9A7zns0THuM/6/POxUc6i2X+gX+/3As1KPeKP1GT953ZWAEAAAAkEAAAkEALHtmJ7Zm7CvoCkxcpF/MfQhRcAAAAAAAAAABy11S178dgHveIEgAAAAAAABwPFFxxMqbDc59K9DLvW4z6DEVVzmko7jmjnXI464AAAHO8TqsxDTOn6EOrAAAAAABw7z1YmK7Ao8PxRpsdl/g5zX+IrDnPIehxzRzvEcdcAAADvOlsv9Av9/uBZqUe8UfqMn7zuysIAAAAAAAAAABKx7Zie2Zewr6ApMXKRfzH0IUXAAAAAAAAAAActdUte/HYB73iBIAAA50OflH67Fw3Nds5nsPzn1CJzfnGa4Q8PznP5fF3nNDvCAAAAAAAAHO8TqsxDzGn6EOrDnQ4O8erHPtK7ERTefxmcVxn1CGoPec9niY5xn7+eoq4EAdcAAAAAADvOlsv9Av9/uBZqUe8UfqMn7zuysIAAAAAAAAAABKx7Zie2Zewr6ApMXKRfzH0IUXAAAAAAAAAAActdUte/HYB73iDku856cc+6rMPzTfvDmMZzm0uHpLnPOxPgcZ+d52K3BHeAAAAAAAAAAAA50nU5ivtHvWHlGiYq0yNzDw8579D1jvOf3/R/PeaDiOO8AAAAAAAAAAAAB3nS2X+gX+/3As1KPeKP1GT953ZWEAAAAAAAAAAAlY9sxPbMvYV9AUmLlIv5j6EKLgAAAAAAAAAAOWuqWvfjsA97xAlA5lpOa0eP3hXh64HeAAAAAAAAAAAAAAAB3g7wT1wgAAAAAAAAAAAAAAAAB3nS2X+gX+/3As1KPeKP1GT953ZWEAAAAAAAAAAAlY9sxPbMvYV9AUmLk4z5j6IKLQAAAAAAAAAAOWyp2vfjsA97xAlXs00vNKPHCvCAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA7zpbL/QL/f7gWalHvFH6jJ+87srCAAAAAAAAAAASse2YntmbsKugKv6F2YddKXVx3Se3UUpdRSl1FKXUUpdRSl1FKXUUpdRSl1FKXUUpdRSZ+X5bV3vO68wSr2Z6ZmdHj9FeEAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB3nS2X+gX+/wBwLNSrWnhjHdnWc4xzaEsXbQMXbQMXbQMX7s4xdtCGMc2hLF20DF20DF20DF20IYxzaOGZaccSESBxWvRxarn2luO7mphFzUwXNTBc1MFzUwXNTBc1MFzUwXNTBc1MFzUwXNS562uW6as4Sr2aaXmdHj9FeEAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB3nS2X+gX+/wBwLNQrBZ2N9s52Njg2Njg2Njg2Njg2Njg2Njg2Njg2Njg2Njg2Njg2Njg2Njg2NmOm8T0RIFJi5SL+Z+hDPcAAAAAAAAAABy11S178dgHveIEq9mml5pR44V4QAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAHedLZf6Bf7/cCzUo94o/UZP3ndlYQAAAAAAAAAAAn9ww/cM/YU9BKkxcpF/MfQhRcAAAAAAAAAABy11S178dgHveIEq9memZpR44V4QAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAHedLZf6Bf7/cCzUo94o/UZP3ndlYQAAAAAAAAAAAntxw7cc/YVdAUmLlIv5j6EKLgAAAAAAAAAAOWuqWvfjsA97xAlXs00vM6PH6K8IAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAADvOlsv9Av9/uBZqUe8UfqMn7zuysIAAAAAAAAAAAT244duOfsKugKTFykX8x9CFFwAAAAAAAAAAHLXVLXvx2Ae94gSrua6VmtHjhXhAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAd50tl/oF/v9wLNSj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vEHJV7NdXqdPmVRbHGWprYKmtiVTWwVNbEKmtiVTWxCprYKmtiVT7a0KolYmKQjgAAAAAAAAAAAAAAB3nQWWbK0tabamtiVTWxCprYlU1sQqa2CprYlVFrQqa2Cp9tY7fq1ZbvV6O9Cj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECXOgAAAAAAAAABRqZdKZn8Pg4ygAAAAAAAAAAAAAAO86NayXW7fRkBd6oAAAAAAAAAACj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECQAAAAAAAAAAFHpd0pefww4ygAAAAAAAAAAAAAAO86NbyTW7fRkBd6oAAAAAAAAAACj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECQAAAAAAAAAAFHpd0pefww4ygAAAAAAAAAAAAAAO86NbyTW7fRkBd6oAAAAAAAAAACj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECQAAAAAAAAAAFHpd0pefww4ygAAAAAAAAAAAAAAO86NbyTW7fRkBd6oAAAAAAAAAACj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECQAAAAAAAAAAFHpd0pefww4ygAAAAAAAAAAAAAAO86NbyTW7fRkBd6oAAAAAAAAAACj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECQAAAAAAAAAAFHpd0pefww4ygAAAAAAAAAAAAAAO86NbyTW7fRkBd6oAAAAAAAAAACj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECQAAAAAAAAAAFHpd0pmfw+DjKAAAAAAAAAAAAAAA7zo1rJdat9GRF3qgAAAAAAAAAAKPeKP1GT953ZWEAAAAAAAAAAAJ7ccO3HP2FXQFJi5yP+c9703uqbfSe7w9N7o9J7o9J7o9J7nT0nudPSe6PSe6PSe6PSe6PSe509G1wNj25JrvO+344SAAAAAAAAAAHCkUy91ejxozko4zRaUEWkxGJMRiTEYkxGJMRiTEYkxGJMRiTEYkxGJMRiTEYkxGJMRnZMRmtZxpVm/3hd6YAAAAAAAAAACj3eozGPdmO6+IZMiGTIhkyIZMiGTIhkyIZMiGTIhkyIZMiGTIhkyPLuORa5n66K+gOdIBIAAAAAAAAAABzqHOkgAAAAAAAAAAAOdAAAAAAAAAAAAAAAABzo46QCQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAHp+5FFC1D53+hz9vU9ck0bClsV6fP0A8EAWZVZ8zmxUu0l6eGMJkAhyYeHzBGQ5a1fmzy1Waxs1+Vza0lger6hKvW8R71f8AxNGYaxCTYeOIJtVPaLC8fkIbMtIyA30Dx+QZX4dWGUeLQMaLW1YU+6ej4SUeGALK8Ppkkq36LO9P3D847skAezLeLxnsorzHvqn7RYn4/ZTYa++uSxCk08EIWJVJw98ByIJhU/eJ54okm3OgELT53LzfSOJF6HrkuqciTbnQ/ECWFVZYlPF+64Rl2+cNuLH4YmaMb1qMmz802QxY+j/1XvOTTwxZNAEQS/58UWQNz+cdqLUgZw/TkMTSpyJNudAAAEVKxRhP0R87/RBR8v1Gpk9N20ZHEbf87n0TU7Dg57t7nrAVCtaoPnbUKhbyZx7YsePoADBd6wY1iLlP0ZdoudSRYXhuZH4/9AYqe7YPYupF5RteKHPb9/UijSllrhlG9YLvRWIGVqpZor1ulP3DCtUJTINfyA30AAEFi+0YufQoKHnWjZ4S8/ov7IqkXSgFg8UQKxolK/Zs/wA6fRfzsbvAz0EZBotZ3AqVW1b8mD7r897sVWsWesmxfO30T87lotc3NlGzb6Cz89615HdCgS0LuRUvSvQx6o77gx9Efv8AH7AKll+oZeb7lOrZSVi9+lp5T6LtUaZlr3zn9FGJ2j81o0Gv+tZSVrt69Y+dtXzH6IKdcPN+D503/APoEpmU/Q+Amn+5b/2QWL7Vix9CAYRu+Empen7s+fNeqZp9AlUtfvZ2VSx0q7Fhrn6Fa2jBNNLsAABFSseYJ9EYxs5SKtdoA0wH4+dvonFzXsd2UUbQsr9E2P08w9opWmQdqJLDvorOy/8AmxORNQ+fL/EmiS8fDFx8WUecteR2ufL1jO01Y9O7YjKGtYrqtDLDcK1ZRXLHCGQbTmetnzj9DUCINg/OZdKvolOvh7OQbNmxsAAAILF9vy42cFEql5rpp4MC0/y0M2Vlv7PcrUl5jT/nb6JxY1eCsMSZxtmWamPz+vyfPO65VrJUazdII0v53+iMXNSl42SFBv1PKZqtB0w+fNGulTLF6tRkSv0raMjPoD9/n9AFSy/W6Ea/lOrZ6etpdFvQj5D0j5/+isW2Q8ft5d6xpmZ+X1D3tH9GYPnP6IzD0DZPxmukHzt9A5BsR4sA+jMpNU8mOTZcsH+jKCXf2cSkTVPn68RhoM/ETZ84b/R68bTUq1pphe2VunGtMs/Joctm+jnlAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA/P6GVeDXOGW+HVemLaZYAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB//8QAMxAAAAUCBAQFBQEAAgMBAAAAAAIDBAUBBhUzNEASFCA1EBETFjEwMkFQYCEisCMkcJD/2gAIAQEAAQUC/wCg24qUHEUcRRxFHFQcVBxUHEUcRRxFHEUcRRxFHEUcRRxFHEUcRRxFHEUcRRxFHEUcRRxUHFQcZR5+f81WtKD1SD1SD1SD1SD1CD1CD1SD1CD1CD1SD1SD1SD1CD1CD1CD1CD1CD1CD1CD1CD1SD1SD1CD1SD1SD1SChy16pHz9fzqPOo86jzqPOo86jzqPOo86jzqPOo86jzqPOo86jzqPOo86jzqPOo86jzqPOo86jzqK1qInK/UHdIphSYblCs0pULyDlURla1afozrpECsq2JVWbCsm5UpAqHUV6LprWjTiOOI44zjjOOI44zjjOOI44jjiOOM44zjiOOI44jjiOOI44jjiOOI44jjjOOM44jjiOOI4gDG53pks/dRGT+jMcpQrINkwpNp0Csu4OFHKyg+fGoitHvjGoUKv26QUmkqBWZXMFHjhQVrWvRbub0XVpKbq39b0yWfuojK351kyBaVbJhWb/1STcnB1Dn+hURWj3Si6SYVlWyYVmwpJuTg6qin0bdzei6tJurf1vTJZ+6iMndVr5BV4gnRWZQKFZpWoVfOFBWta/UqIrR7WpqFCz9ulRWaSoFJlaoUeOFBWta/Vt3N6Lq0m6t/W9Mln7qIydsY5ShSSbECs2UKyzk4O4WUrsKiK0exOukQKyrYgVm6hSTcnB1VD7K3c3ourSbq39b0yWfuojJ2SjhJMKyzYgUmjVCsi5UBznPtaiK0f1KmoUKP26YVmkqBWZXMFHjhQV867a3c3ourSbq39b0yWfuojJ+rU1KBZ83SorNJUCkysYKO11BX/dzURWj+gdwkQKyzYgVm61Csm5ODqnPXdW7m9F1aTdW/remSz91EZP0TqkIFZNsSis3QKyrk9TrKn31RFaPxqalAo/bJhSaSoFZhcwVeLqiv+131u5vRdWk3Vv63pks/dRGT0qOkUwrMNyhWaUqFJByoDGqav6KPXTTaKyrYgUm6hWTcnB1FD/pLdzei6tJurf1vTJZ+6iMnomTmIz+f4+3c3ourSbq39b0yWfuojJ6JzRfx9u5vRdWk3Vv63pks/c1ERldE5ov4+3c3ourSbq39b0yWfuaiIyuic0X8fbub0XVpN1Aa3petVVFeRXHILjkFxyC45FcciuORXHIrjkVxyK45FcciuORXHIrjkVxyK45FcciuORXHIrjkVxyK45FcciuORXHILiPROin0Tmi/j7dzeieaKu2+BPhgT4YE+GBPhgT4YE+GBPhgT4YE+GBPhgT4YE+GBPhgT4YE+GBPhgT4YE+GBPhgT4YE+GBPhgT4YE+GBPhgT4RES6buf0s5ov4+3c3+FnNF/H27m/pXjtVJXEFxiC4xBcYguMQXGILjEFxiC4xBcYguMQXGILjEFxiC4xBcYguMQXGILjEFxiC4xBcYguMQXGILjEFxiC4j1jrJ9E5ov4+3c3onHirND3A+HuB8PcD4e4Hw9wPh7gfD3A+HuB8PcD4e4Hw9wPh7gfD3A+HuB8PcD4e4Hw9wPh7gfD3A+HuB8PcD4e4Hw9wPh7gfD3A+HuB8IyZduHXTJZ+5qInK6JzRfx9u5vRdOk3UFrumSz91E5PROaL+Pt3N6Lq0m6gtd0yWfuojJ6JzRfx9u5vRdWk3UFrumSz91EZPROaL+Pt3N6Lq0m6gtd0yWfuojJ6JzRfx9u5vRdWk3UFrumSz91EZPROaL+Pt3N6Lq0m6gtd0yWfuojJ6JzRfx9u5vRdWk3UFrumSz91EZPROaL+Pt3N6Lq0m6gtd0yWfuojJ6JzRfx9u5vRdWk3UFrumSz91EZPRIN6uUMEVGCKjBFRgiowRUYIqMEVGCKjBFRgiowRUYIqMEVGCKjBFRgiowRUYIqMEVGCKjBFQ9ZGafpUohRVPBFRgiowRUYIqMEVGCKjBFRgiowRUYIqMEVGCKjBFRgiowRUYKqMEVGCKjBFRgioi2B2h+i6tJuoLXdMln7qIyd1cH6T8sNLubq0m6gtd0yWfuojJ3VwfpPyw0u5urSbqC13TJZ+6iMndXB+k/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyt1cHz+j/LDS7m6tJuoLXdMln7qJyd1cH6T8sNLubp0m6gtd0v0FDrcqsOVWHKrDlVhyqw5VYcqsOVWHKrDlVhyqw5VYcqsOVWHKrDlVhyqw5VYcqsOVWHKrDlVhyqw5VYcqsOVWEamZNPdTSCqw5ByOQdDkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkXXmzLUrfc3Egq4bYW9GFvRhb0YW9GFvRhb0YW9GFvRhb0YW9GFvRhb0YW9GFvRhb0YW9GFvRhb0YW9GFvRhb0YW9GFvRhb0YW9GFvRDx7pJ3/+LztWqLdpcS6zun+0+pc75w2Wtpwq4Z/VkppFgpHvSPUeqacLNmkTLSC7zrllDJMoKRdLP/FSvCRa4XZFvcb0VuR5QUulxWvuN6Pcb0Qcou/U+jX4dzMkR1GKqLNOu4XjlolAST1062Uq7MzbRM+s8d/VUPRMhbibmcFr5lC1alSNMyfNNDmUbqqUSTSuNuo4p/tPoVr5UVuNumukeiif15PQxvcy/a4cJNytnzdzVy9bthj7AITDJatK0NTwVVIkVScYEqWeYGqg4RXpeGfaWgVUKkRGTaLn8V5JqgdJQipA4fNm4x9gEZhiqCKEPRy8QbC5nCThzb0g1QZ4wxDddNwRy/bNjILpuCOHSDeiMyzWWWUTTTbv2aygMYpaKyjNIY+wCUyxVBDlPQTfb7b7l0cBRwFEyUtGEJ3LgKOAoXWRalbP2zmqihEqHm2JDpqFUTWfNkQadYFqScYGqi5RXp4KSLEiiZimKuumgRs/bOTOHaDalZ9h5ozLFapTUNQOl0UCNHjVxUKSjNM9FSVTXmWKNaT7CtW7tBx0OJJo3HuBgEJZkuKnLQhJRmdTouTt9s9yDp4g1DZ2i6K5kWrYY+wDeUZuPExqFovMMURSeYVqi/arDzp5PJNn6SKhKSJJdjQhJVkofzp5VkWFFCmpUr6SZ+i0OUsgWXY8KUmzVOooVMiMm0WU8V5FqgdFUixHMi0SDo5TSDWVYlbllmRjUrStK/4HEmzQHuBgG8ozXFP9+jJ6GN7mT7bs0UWu5IdtblFBW32FSSdt+inByqzRyU1DFnJcrAjdB9NLNrZakClvMDlcwK7Oso+Wd1tLQTnb7e7j43J3KC7fcz9Ro2hIvFDUgGFKOLZaHETCKsnElFoyAnmKTFeFhGzxr7YZhgyTZJXdrGsuZvHNrfdvAyt1BstPdvtnuQmYysgVG1m9Ke3GPk6tdGpEl3cO7YOaO2032+2+5dU12+E7l4XbooWQLH1IxkJlRK1kaBRpSrJK1y1MS22RQrbLQ1JCPcxCtuylXqVfiS7pG6O59BGOF0VW1vVXHt9hwyVs0InESS7F0meihLu0NnamvxL/APGRZpvphNC22SYcW2zUo/ZOIlxb0nV8gsoVFOUmXD5ZhbVVCkt9gUri2mhw6SkIgkdXjkSfZ43J2+2e5C8/mIVdno3tpEKW8wOWXgDsy23LqFV/E/ILOHkZbrf0T2+wNRxa5RGMzNWr232taJIlM9JbLOpEbcaJKmp5JOO6NKebR/b7XhbIlUe0tlnWjW32rdWY7fBdz8bm7jAdvfwTZyo4RKm8b22zUQTtpmQ6qhGjd/KupJwwtknDgLDyf2yThjpVzHOEFSrJdcnoY3uZPtuzRWgWlXXgp9jjurX/ABrNqerJQiRUmHjdqBE3NpaCc7fbvcfG5O5QXb5aOJIIYdJxiqdwvUQ3uhucNXiDoou7V2tofC7tZaTEvD4T/b7Z7kJOdbsq45JOq81PAr+bTrMulXS9saCb7fbfcuqa7fCdy8Ls0VutCunpS0KUSEggxIpcq6xivJ0wq7naCYkHi7a0NXX4ku6x2jufQWmQpn/gb7ZKnlKRuiu7Q2dqa/Ev3KHLQjDwu+n/AKdnai5zGpGWoUhn/jIFKZoxpSkoT7PG5O32z3IXmLNJT0/CRp5s4/8AyTL9k1AGcKIrS8aE7oUKGlwM1wQ1DleaZDuiWWFPscd0Z6Z9pGPcyfYJnQQXc/G5u4wHb1Mt53NlpRd6pitoR6Vkp7nqPc9R7nqJdzzy1rKHMy65PQxvcyfbdmis/UeB/scd1aaa52R0HdtyqaiPg5cJNk5t/wA85tLQTBKnYRS1GsgkoVUgcLpt05RzzT2C7equkiC1oah0kzh/Cs102ZjtJROvmS7tXataVY+F3au1+2eE/wBvtnuT+tStGlCKSDciaaQrWhaXM6TcPLX0M32+2+5dU12+E7l4XborP1XhcZzmkbdRblZeF1u06NbP1dfiS7rHaO5tBaOu8DfbJ90jdFd2hs7U1+JfuUVofC79HZ2oetyum7ls5iXkdcbdYtH7WoXlmSJZKWXkSxtKlkCfZ43J2+2e5C8xZmR4P9Ix7oX7CuUTnr/oVZt1aXHENmyNnrqmDzTI18pNHKB/scd0Z6Z9pWdeGTT/ANIJROqjKPV5SRQVIsmFlSIpzDorp9AdvU+x9/kmxrSrQXEyq7ZwT0rF0lRBUnopj0UgdwxIqShaU635DKNWMO9Tfl+24mqrprbUc5aLeBv9KtDvayDYtSoOEE3Cb62TUOQs40oVxPKBCCdOlJuGUqLdaqtGdaedJa3SuDJNZljUrmeOEIV48UlIRfmIlI6DOXhCvjUjZdlXmp4gVTnHtIiAK1OJqJLIJpx0swP6s+qI0ixG1xxrp05gG6jZh4S6J12cFFPG76tPOkrblFjokm2Io/mz0LHyr6spAKENAtlWzSVSOszhIp43fdUokdZnFRD1F94XE1WdNbajXLVfwl4dKQoWMlo9QkhNFHDOuw9t41GttRrpq4r8Pod6o/ZEMm2nm6jlnbkY6auvA3w/h3qj9iQybW42izprbUc5aL1+JGHeqvo9MyTXwuRos7bW1HOWiwVImrRzbjJWvtVLzQtpmnWRM3j4+HTM5ky/4XxnG6jhnBRLxu+FzsHDsWwyXZpeDwlTt2kO+I/On6iLq21aKFJPNhzc7UYRJyCsZHpMEa086S1uqGVQxxrRoWaWc1p/41oZ8Z+2LUiBqedJa3VarIY62KwLMqufxL2+R0ZJlMMKldTxglDvnqkpBrevDonQZibgKuFG6U2yCOOLqp0rROUgkHopGy8ebnpwoNSddhjbX/NMlEyfz9ftfRcmRyWVl2o9yOKBSffqikZJyasVGJR6f/xGtPMcBBw0/wChr//EAC8RAAAEBAYABQMFAQAAAAAAAAABAgQDMDNREBESExQxFTJAQVAgYWIFIUJSkCL/2gAIAQMBAT8B/wADobDWnVmPDvyHh35Dw78h4d9x4d9x4d9x4d+Q8O/IeHfceHfceHfceHfkI7PaTqz9KqKhPZhT2GXQSeos/RKiJT2YU9hkIcTcTqx1DUMxqGoahqGoahqGoahqBYt6ZTX1L0BnkFOIafcKfpLogp8s+gqMtXZ4EIflKbnkFOIafcKfp9gp8s+gqOtXZjPBpTLA5ycW9Mpr6lMNaU9hTuGQU/8A6kFPIiga1K7+ohC8pSVREp7MKeQ0hT/+pBTuIr3BrUrv62lMsDnJxb0ymvqUhUZCezCnyC6Cn6j6IKcxFe4MzOUQh+QvoMyIKcw0+4U/SXRBT5Z9BUdauzGctpTLA5ycW9Mpr6liaiLsKdQ0+4U/L+JBT2IfQVFWrs55BERKUlmYU9hl0FPz/iQU7iK9wa1H36BpTLA5ycW9Mpr6li+P/v4hpTLA5ycW9Mpr6li+8/xDSmWBzk4oexElpIeIRBz4g58Qc+IOfEHPiDnxBz4g58Qc+IOfEHPiCK7XFLSrF95/iGlMschpGkaRpGQyGkaRpGkZDT9CWcRRZkOBFHAijgRRwIo4EUcCKOBFHAijgRRwIo4EUcCKIrVcMtSsX3n+IaUyxzGYzGYzGYzGYzGYzGYz+iBTKa+pYvvP8Q0plgc5OLemU19Sxfef4hpTLA5ycW9Mpr6li+8/xDSmWBzk4t6ZTX1LF95/iGlMsDnJxb0ymvqWLuCtas0kOLFsOLFsOLFsOLFsOLFsOLFsOLFsOLFsOLFsOLFsFpNB5H6JENUTyjixbDixbDixbDixbDixbDixbDixbDixbDixbDixbBuk0oyPA5ycW9Mpr6lOd1D9Ew7Oac5OLemU19SnO6h+iYdnNOcnFvTKa+pTndQ/RMOzmnOTi3plNfUpzuofomHZzTnJxb0ymvqU53UP0TDs5pzk4t6ZTX1Kc7qH6Jh2c05ycW9Mpr6lOd1D9F+n9nNOcnGBFQUMv3G8i43kXG8i43kXG8i43kXG8i43kXG6i43UXG8i43kXDyIlUPIjnOoajifsQ2l2G0uw2l2G0uw2l2G0uw2l2G0uw2l2G0uw2l2G0uw2l2G0uw2l2G0uw2l2DFBpM85pzi/3R//EACgRAAADBwQDAQEBAQAAAAAAAAABAwIQERITMDEEBRUzQFFhUCEgkP/aAAgBAgEBPwH/AIHKbhI1LAcn8HJ/Byfwcn8HJ/Byfwcn8HJ/Byfwcn8HJ/ByfwIa2s1LDxYkDULw4kJyBHFxFExRFEURRFEURRFEURRFEURRDTMr9T2Hd2/t8GYhUFQxE3li9MQNQGoYib2MOYzeVy/U9h3dv7bkYA2yFQG2Yj/ssWYkQNQhUBtnYTw5jN5XL9T2Hd2/tsTEKhCoYmO4WP8AM5A1BUMTGdxPDmM3lcv1PYd3b+3/AAbZCoDUMRPwI/wG2QqCc/BTw5jN5XL9T2Hd2/tepn8hPDmM3lcv1PYd3b+16mfyE8OYzeVy9vQJtnExxyQ45IcckOOSHHJDjkhxyQ45IcckOOSHHJDjkglo2EmpmXqZ/ITw4jgKpiqYqmKpiqYqmKpiqYqmKpiqYqmKphpqZ7WuTZOBjkEhyKQ5FIcikORSHIpDkUhyKQ5FIcikORSHIpBLVsKnKy9TP5CeHEURSMUjFJoUjFIxSMUmhSaFJoUmhSMUmhSMGzK/U9h3dv7XqZ/ITw5jN5XL9T2Hd2/tepn8hPDmM3lcv1PYd3b+16mfyE8OYzeVy/U9h3dv7XqZ/IYw5jN5XL9T2Hd2/te2RxEpiUxKYlMSmJTEhiUxKYlMQ8IiiJTEpiUxKYlMSmJTEpiUxKYZL+OYzeVy/U9h3dv7bzefCTusZvK5fqew7u39t5vPhJ3WM3lcv1PYd3b+283nwk7rGbyuX6nsO7t/bebz4Sd1jN5XL9T2Hd2/tvN58JO6xm8rl+p7Du7f23m8+EndYzeVy/U9h3dv7bzefCTusZvK5fqEmzUOBCip6FFT0KKnoUVPQoqehRU9Cip6FFT0KKnoUVPQoqehRU9DQptMqf0rzZHEQMQMQMQMQMQMQMQMQMQMQMQMQMQMQMQMQMJldYyJiExCYhMQmITEJiExCYhMQmITEJiCp/3/ALo//8QAPBAAAgEBBAULAwMDAwUAAAAAAQIAAxESMTIEITNxkRATICIjNEBBUFFhQmCBMFJyBRQkYoKwcJCSocH/2gAIAQEABj8C/wCA21mYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiZhMRMRNX21rIEzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeM1MOPSxMzGYmYmYmZjMxmYzEzEzEzEzEzEzEzMZmMzGYmYmYmYmYmYmZjMxmJ9K61RZ1bWM7NLN811LN0W3H0TruohFtp+J2dOWXru6Pzjlt/RFhImduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxgtZj+fsXrMBNdQQ3EJnVsTdOvUM19BfH6yBOtUE6ikzqqFnWqGayT0H6K+LH2D13Ama9unZ09XzM90fE67sf0V8X13Ama9unZ0+MzXd067sf0X6K+LHrmudaoJ2YLzs0A3zXUNntNZP6q+G1kCdaoPxOzUtOooE61VrJrP6r9FfFj1jrMBM9u6dkhJ+Z1bEnXqN4JfBdZ1E1G9unZU+MzXZ1nY/nwT9FfFj1Tr1FEsBLbp2dOyZ7N067E+GX9XWZrqA7p2SFp1AFmuqfxNZJ8M/RXxY9P1mdaoJ2alp1FCzr1TNevxS/o9Z1E1G9unZU+MzXN0tZyfz4t+ivix6X1mAme9unZU+M1EKJ16jHx69DWRNdQWzqKSZ1AFnWqH8S0m3x9Tor4sej9aosNy1jOzQLvm0I3S1iT6Gt9wJZbaZ2dOzfM93dLXcn0R+ivix6MShsMtOv7Pfor4sejH7QqdFfFj0Y/aFTor4sejH7QqdFfFjpWqNUwEwmEwmEwmEwmEwmEwmEwmEwmEwmEwmEwmEwmEwlj9E/aFTohaOMyjjMo4zKvGZV4zKvGZV4zKvGZV4zKvGZV4zKvGZV4zKsyrMomVeMyrxmVeMyrxmVeMyrxmUcZlWZVmUcZlHGX6oFnox+0H+xj9oVPRrFssnlPpn0z6Z9M+mfTPpn0z6Z9M+mfTPpn0z6Z9M+mfTPpn0z6Z9M+mfTPKWv0T9oVOiHo2W/MwpzCnMKcwpzCnMKcwpzCnMKcwpzCnMKcwpzCnMKcwpzCnMKcwpzCnMKcwpzCnMKcwpzCnAlQJd+PRj9oP0V8Wvox+0KnRXxa+jH7QqdFfFr6MftB+ivi19GP2g/RXxa+jH7QqdFfFr6MftB+ivi19GP2g/RXxa+jH7QqdFfFr6NcU2GbQTaCbQTaCbQTaCbQTaCbQTaCbQTaCbQTOJnEzibQTaCbQTaCZxOs1voocONc2gmcTaCbQTOJtBNoJtBNoJtBM4m0E2gm0E2gmcTOJtBNoJtBGLMDb0V8Wvrg9EETxS+LX1weiCJ4pfFr64PRBE8Uvi19cHogieKXxa+uD0QRPFL4tfXB6IInil8Wvrg9EETxS+LX1weiCJ4pfFr64PRBE8Uvi19cHogieKXxa+uD0QRPFL4tfXB6IInil8Wvrg9EETxS+LX1weiCJ4pfFr64PRBE8Uvi16VqoTMhmzMyGZDMhmQzIZkMyGZDMhmQzZmZDMhmQzIZszMhmQzZmZDMhmQzIZkM64s8WOaQtNk02LTZNNkZsjNk02TTZNNk02TTZNNk02TTZNNk02TTZNNk02TTZNNk02TTZNNk02TTZNNk02TTZNNk02TTZNNk02TTZGDsWiBhYfFAUULGd3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3abBoGqUWUf9l96gFtkFI0lstg/VUUahUQtWa8f1rlRWJ+JzlMED56d/R88VK2Q/6f0HambGirUqkr0CR5RlGjahO6zXoolgoLO6zusZa1G4B+kY6plB/bFetm/QB0bHdbLmk5P4+DNRVtMFJ6agH2/WLHATmrj222QHkYrjZCv02/tiM+YwucBOaCPbbZ+lbDSKPbbZA4wPgKu6Lvgl6s10SyhUDTtqqrNrLFrLb8y0G0cttRwo+ZZzwO6Wc9LaVRW3RYd8vVDYsuU6yluhcq1grS9TNo5O2rKs2s6tdfzLUYHdO3qBN8BouHEu1aoVp3hZeosGEu1qoUy/Sa8vvO2qKs5unUtaXqpAX5l2jUQt8cnWIG+daus2s6tdfzLUYEfHJU3RejlHCZRwlSxRF3zKOEyjhL1UqgllCqGMtqMFHzLvPAn4gdT1TO0rIPzNsJZzwnZVFbdy3HqJegKYS9Va6Jdo1QxnbVVWbWWLWFvzLVII5L1dgF+ZZo7qT8clx6yhpfvC77yxq4t+JtZ2NVW6HaVlBm1nUrrb8y9b1feXFrLe9ui0Xk7eoFltBwwna1lBm1nZ1lt5bWNgljV1t+Jtp2ddD+ZbbqlRDXW9ZhL5PVvYxf8hcIFWupJlvlLhqJelowlRDWW97QOx6t7GD/ACFlynWUtLzmxZcp1lLdC7VqhWl6mbVjJUrKG9ozKerexiA6QuEsFdbZaMOTtK62zazs662/pVd0XfByc3og67TnNOqs9Q+Us5v8w1NEcmz6TBRqkmmdVh8oCMDLqa6pwlrObvucJ2pZzLBTu7pz39PrHV9MC6QtjrDvlSJ0Ul2lqZvOGppFU2f+5ZzU7Mshl99ILJ7CDni2r2gSiTZ8znKpa34mapObpW2fMWLo2iC2sZzumVbtvvjFqc4xIjReQXapQztajsfiWWND/b1GD/Muljq8veLUHnKm6L06kXfyiO9285wnO1WKUzhbBfrMd0/t1YgWYwnSK7Numu8d86hZTL9Nzd8mEuVdoIY/8pT3QyzRR2jTnf6hWZnOuwSzm/zC+iOdX0mCm5JS2wgwMuB5Ghjn5ioW5rRlnXvOfmdnbTPxAQxs8mEu1dosLubAJzOjWqltmrzl/TKh1+QlnN275bTLUz8RlD87o5iH3MXd0Gi8lOHRtD6t7Fpe0qo1R5YEK/InO6OxZBxn9vXa8vlLZzAcqmFkWppB5xjNlZului1ip+ZzVSoahlSqS97Gc0cttkBvPxgdWe0Qj4h/nEHxKla172MFI5bbINb8YKiF7R7ypuib+ikaq5a9DSGW2yKxL2kQMC+r5lrHqqJzOjWqnsIH0xyzewmxhfQ3Ib2M5nSCSlthB8orpgf0Ku6Lvg5CbNfK26f7om6Nb7xLg6AdRraHfKkXopLpNjeRl6haR/plmk0LfxZO2RkltCoG5B0BDpDrafLleLyFB2lT2E/xdH1bpsl4S2pQvD2sltalzbQSpui9OpF38ogv4LLF1DktrNr9pd0TR/8A7LRSH/jNdIcJzelaNc+Y0Mb+Up7oZaeUw2e8p7uRoY++UwvKseNZOv0KgfCyCzC9F3dBovIkdvPlqbotn7oIa2inrftMu3GKD3Fssr6PLCxQ/MtU2iVN0/3Rd3I26H+Up7pU3Rf5QbuSpuib+ikbdG/lKe7kCDAwvzBqH4ndHndHndHnOJQZJY9ur3/Qq7ou+DkPK26H+UTdOdA6h84KFUhagw5S9ZgohZcgwhlQL7RTU1C3XAyG0HkL1WsAjOuW3VEg52oq2+5lqm0Trop/EY3BTPuJdovbY1mqKTjZBNR5Vi8rxZUK42Qc+dV7ziikAF+OS06hOyNtnnyVN0Xp1Iu/lEflYPbYMIrUQC3meXmbwLnyjQxv5Snuh6Bh3ynu5Ghjb5T3cojxqT4GXgDYDqMC6R2bzvFPjLWrqd0aloNNua82iBsbYu7oNF5Ej8tTdF/lBLi1VLe1s1ztKSH8TnqJuH9samxJQSpumv8AdF3cjbof5SnulTdFvfug3clQLjZFap5HXA9M2g8heo1iiF0yxI26Nb+6U7PbkNzMIU0hRdOrWMIGpqjA/E2a8Js14Tm2NIP7TqAWfH6FRUzGB2pWLbjBLtBbTC1endHKZfFLq3sYobGyFKqhll/Qn/BlihyvGWBCPxA/9Rrmz9sRdBodUS5XW601w1NGa658jOxD2SwIeED/ANSrG7+2D+0odnFSoLGl8VWV5/j1bw+DLGpE/iXKgKIfxOd0hr9Tk1G7UHnD/bhv9su3WHzZFGktbU84GoU7wgSsLG5WSkLWgerSurLDDU0Rrrexl1FZ0H5li6LYfeD+9rc3T8wIo0KmXHmZdrLY0dKYtYxXq0rqjpulMWsYHqUiFBx5btBbxjNXp3Ry25anvP8AFN4fEunRbx3SxuxUy8hatpBhavTujkZ1pdW3GIr4iFKItaX69O6vKYaiUrVtxiK+YS7QW8YzV6d0cjOlK1bcYiviOULQW8YzV6d0HkuVArfBlqg0901V2slrlqm+OFCpq1CAqPO2AdApRFrRXq0rq8icwl6yONIS6Tyuq4mB2pdW9bbLh1aoamiaQbfmWLa44ywUTbul7TXuj5l2nj5mWGGrohx+mXFRyIp0glKcs87JfFLq3rbYitjZCDDV0Q22+UuKr2fMVtJJWn7TXDU0c3KntOwDWfEsFM8IG/qVY3P2wDQ6PZiKlUWNyGtox63tLtJXu8YvOWokAbGX07Op7if4rlh8Sw0TbuljXkU/ic5ptW83sIFXAfcBsjVqFcvuMsr0C28TXohllHRrv4l7SbVH+qWLrfzb/olrmUcJgP8Aga//xAAtEAACAQIFAwQCAgMBAQAAAAAAAREhMUBBUWGhECDxMFBxkWCBsfBwweGwgP/aAAgBAQABPyH/AMDZ2Av2eaPNHmjzx5488eaPNHmjzR5o80eaPNHmjzR5o80eaPNHmjzR5o84eePJCS4n7anNq+0Su1bKjdweEHhB4QeEHjB4weEHjHTHhB4UeFHjR4weMd2MYxjwg8KPGjwg8IPChvDb27jpQj9nlDy55M8meVPKnkTyZ5M8meTPJnkzy55U8qeVPJnkzyZ5M8meTPKnlSD/ALhm3lt/PtDaV2TUZrKRQ6ZsJaf3DBNm2kcjttmxexScyjEIdWgY5X72ymhdg+/knbJOTkxRYnWMbxjet4xvWta1resY3rUjMISTn3LmKdi4P2NTKHdjmJmiIg+ybFkJvVD+r+HA6pZt6vrk9gqaVm7M6ew/a0diFpXk7i2H/ihzOOSPQfK9hK5inYu4/wCRfP7YsDGyXxmFbUGDlxZHfkxdbhVYllP8lsY2S7eTC9qKMfzGkLz6p5QrewK5inYuYtCVJfIzun0ciT/TDKh/sS8dshdB8v1MmFs9F0rN2SRj7DP9vYduDJxTGSLgP5eCfKFb2BXMU7FzD8+TJmVZZBFNCF5I2HBtfsu5dXgMnzg6tbjNxbVbYPaerD90o8kixa9QlFsM+UK3sCuYp2LmCb1ockDGjZNhPJO5sv8AH+g1lvuyMJk+fUr6KpUluxU3swU1Legwp/2Pm/isM1thuyMU+UK3sCuYp2LnrKZUluSj9eo3a3cmkm1F8MrSwzuP5OcTk+fRrcEzTNxI52WQyJLekWgZGpjPHnlCt7ArmKdi56SJtctyjpNBmhhynWSM9DKSMbk7Ta98seNPJkhS0eWweynjhnMm1b9kXKFb2BXMU7FzukojWUiH4NB/tcDiuxoGRs3fsTyP3UzGzR1aIdTU7Vuhl7I92X9ofKFb2BXMU7FztYp1kOa2Tzfs7repb2w8oVvYFcxTsXO4qy/FFyhW9gVzFWMu9vLFb8UXKFb3pFjLvbyxW/FFyBWxXO7oVmPNGw+zYfZsfs2ptfs2v2bX7Nr9m1+za/Ztfs2v2bX7Nr9m1+za/Ztfs2v2bX7Nr9m1+za/Ztfs2v2OC37HBaT7eWK34olZpq8xReopQhCEIQhCHmDzB5n0IQhCEJeYPOdZSR0e1+zcsVv8Kvlit+ah4/eYogAAAAAAAAAAAOAtUc7dvLFb8US+XdhehMzMzM+JPE4CZmZmZmZmZnxJOsr+wCwv9vLFb8UfMFbFcgWPHYvdvLFb8UHKFbFcgXbcxTsXO3lit+KLlCtiuQLtuYp2LnbyxW/FDyhWxXIF23MU7Fzt5Yrfij5QrYrkC7bmKdi528sVvxRcoVsVyBdtzFOxc7eQKy/FHyhWxXIF23MU7FzuCsvxQ8oVsVyBdtzFOxc7irfig5QrYrkC7bmKdi52v6uoxQYVERERERVVczERE4KKensmaFGd0YUTETERETVVWzER2U67eUK2K5Au25inYuYu45exrmOHiuUK2K5Au25inYuYu72SuY4eK5QrYrkC7bmKdi5i7jl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkCx47F338rmOHiuUK2K5Au24ZYl2Zexdxy9jXMcPFcwWK5Au2CIxTI3MzMzMyMzMjMjMzOFk6pr4ttQGgv+MeA9hZ3ZmZmZmZmZmZmZmZmZmZvDnzARdFVWKWObshf83Eqqqqqqqqqqqqqqqq+PEei7x3Ri4IxLI7II95j/7lmRFmBJK0JTqPK1XqsjjQY3lu/WUfQOSAMCJ+EiLm2Vs9CKbKjH506p9jUaWkwKjtkWZtfpi2VF+ybC9E2bH6ZtfpkICSnWvpUuV4FdNRVeLUh70j0JhZXYPsjhejBqfKyZnJYb1rRxIuJnAQos1PSyMkhZFOzGktRlS0aSKyJqg8D1r6M22Q2jSRlNVMrBRvHRCzrsfmXYTV7Rsc8fxIs2igWEtRdd2aNBMzgpUXyiEp3dXj8qq7ZtFJdmxxGLSr810vJaSNLt+iwN6BSR3qIerDX3BG13mNY7spQmaLsUUtASj5ViJNrUN1QjY2adJUt3QS1fomPR/RSEGF9xZt1Lk9sr6LzDkxoQc6T3Ch6L27S4gc1Nkh51sYVBhcZooxlDURyT4I0V8k4+fdGMqNXDTQxpN7NEC+qzZK6J5tLYnCv9EYMFdhZp9NyLBn2GDoVmRECksmsiQzoqIJT+hPOxp9W0lLcIcNLmScn9SHiWwNA2OsKwZKanATlSu+uGqOo6Kq7RLJUymT+5DxKUydBNNSqrowoIzYwlNFQuKC/RGpy8hlDcIIhYBAg1gonZFsMSMyskvs8RAivWeRanVr5Efmt9hHTRaZWQzKqrJswpdm3vGICrs0UklEieK8gtxaRNC8/awtslrMZJLaS3GrSKZJyf1IbpSWTcDJJTTXpxvDQsoOeS50M0ctGRDF6HUjc6swYhSoDk5Cky2aD29nVlAnTzXoTANqxJndcriRdG30zhCp/N2JUIlDoIDeIJnquqmoVau7ZKvL0OhFQLOY2ZB6g1BYfuRw/wDuCGxiVE4OzEwMExcioJLk9i09hPv6JNS+mcByEqtPWRoZSVpDNTVnQXZ0qtOhcnv4/TlboqZMlkEcvhsxshyzpBvpkWZCaJNBZnzh+0dWgqZlBRzr6nGEUX916AkzlQWkmZutMsbn3VGTrqF6WyKxaOJQgRTnGHUNaZbrqmYoo3UetlyC8aU0hMDUepHeUjKxiBXhbWurOJAn1Yk8lZULXD1rBOKirBxnfWikz6HbU7lq60JgTUGO/DQIrDbndE/yDskXQdBlgRmJoMKUurCwdlTF2yfaaqErGKppakOWWypYlRZOhEBWQMswqq4fi5E+r6DeK1IVQBF8IVuqhYJG61XDGGuBVBtXJFNpFC7PFXAcLNJArVuLaBFxqlE9wJiUrlelG8NdNtUhZ9eUESZKiLyNC2t0hUVKFKl9iyyuR0zgHJ7oPEbBvEqZ1T+i/ms5B0k1+z4lKa9/68jajrwDldJyU/B9isFqrE1BgJzkK0+rJoeegXJ7+P05W666N0wISoLJdKD2RbsdZzyzCODT1FXbWxlkXacY43ZVNBTLVuvHFf0tepHBOMVodimCFHVTY1VDs6y4dyVRlKk9inE9QUNqg4j0K2l19+qFGKOOY0FgU98wW7EQB6m92qDdPMpi77NPoFs8d05Q/u79cj+7ucV2Fit3yieP6MEi4NY3+nYZnWz50uOS0yw9KN4a7d8x1WqglWBm5wZirbotqdXchw0484SpLFIchLIQcWlNdFPFsxJocEfIxsoCSgyzQlf7UQ0FTagTnrhmQy+CdPTI2n2bn9eAcgnhWgIatq/YhB2nQ8tSLtkWQl6FnoXJ7+P05W7kGeyrhvCZFZ6xZ9KdFxO+q5wey6OCcTo/Bdr9mRFpRITZjCdV5MRSoPifAlPIlmzauWykGcR6EbyeyP7u5xR1St0qESQiaIkIYfnO44UlScugIneEWTaq0dOc6ziNxKoQ9R/2Mm1VQ6VP5CwBRhb1aaPotEpLbLSpwnr0XKGii5WTUej4rrk0dSzBZyU0h4WeBii/NCKJDaPQS1KKIdPRIJCHeBO2mNsO16rDV2mNzrrBC0JBjAwyY4R85oj4ItEEdXZxHxl8HI9QZ0ZjpUFhDFUmLPVjtMakwujlH3bxMyMxZH3RzRGbLohfFk6oq6VZSfyKsL8OkgBcQyFkhKFQVmqy1Duha1JkBIrsI1eUH7F6Ee/dus3gshq5F2xTkSnRiSrqu0L9EFh+hUE2Pyoz0NHVijtohcCiHhmqt9+Q4CKhlDCt0U9ojOjtetao7KPTduZX0IvlLDZbl+sP1yzbGPH6l4bXbTQWVCqolk2RJzvdUlSvAxA2QStCKo+BxHQnrMqYcRbNBBcKhrrBTch9i0no3pj3BianrH/EYWdlUCM44Iux9qATYFHZJ5shiZN2+lTedA0g5RdUcyqENOpJP2FMe5QbQtAdzXHIhp0Q4jVAiGMyX+hCVOdzFORKdxR9yltQhUjZNSL2B1ToMldQHF1/QCwISoTbDoxA5WmdHIyT5bJExJJXahFtQsqy7WY20PVKLyvxFZF7V3F2t5hB4MuiMaa75in4o6DPYVlJDBZjUfPgRm2BzlFzuorogDErk0yf7ERQqF+QK2hobVGPiLcqgyAQVmFJX/DEE9nnISZ7m1v0Q1muf4SRYTG65kmyv1/4Nf8A/9oADAMBAAIAAwAAABDzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzxTz3X3zxSzz3zxv3zzzzzzzzzzzzzzzzjDUVZLzzzzzzzzzzzzzzDDRUfbzzIDDXDDbLCLCCCLzx8MMMMMMMMMMMOxbzzzzzzzzzzzzzR0jlAGZ4PzzzzzzzzzzjwyqdmYpoPx033/8A/wD/AH//AP8Afy6/PPwwwwwwwwwwww/FvPPPPPPPOIOfWsQJ39//APt/zzzzjzSRHbXRC8/f/wDrD899/wD/AP8A/wD/AP8A/wD/AOmvzz8MMMMMMMMMMMPxbzzzjjmZDIae/wD/AP8A/wD/AP8A+sPzjxO0N4or/wDP/wD/AP8A/wD6w/Pff/8A/wD/AP8A/wD/AP8A+mvzz8MMMMMMMMMMMPxbhv6BCDp9/wD/AP8A/wD/AP8A/wD/AP8ArJtugh9//wD/AP8A/wD/AP8A/wD/AP6w/Pff/wD/AP8A/wD/AP8A/wD+mvzz8MMMMMMMMMMMPxagt8+//wD/AP8A/wD/AP8A/wD/AP8A/wD/AL3j/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wCsPxT3/wD/AP8A/wD/AP8A/wD/AKa/PCiwwwwwwwwwww7VuAv/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AOsPzXX/AP8A/wD/AP8A/wD/AP8A4YfLGTrjDDDDDDDDDP8A6hb/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/APrD8GHc889t/c889/HH84xuzzzzzzzzzzz/ANqBv/8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP6w/Fu8/wDPPP8A/wD/APPPN3ywXPPPPPPPPPPPP1agL/8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8ArD8V9/8A/wD/AP8A/wD/AP8A/wCoH7z8MMMMMMMMMMMPxahL/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wDrD899/wD/AP8A/wD/AP8A/wD/AOqnzz8MMMMMMMMMMMPxagb/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/APrD8U9//wD/AP8A/wD/AP8A/wD6qfPPwwwwwwwwwwww/FqKv/8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP6g/FPf/wD/AP8A/wD/AP8A/wD+qnzz8MMMMMMMMMMMPxaz/LLDDLDLLCIz/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP6zwsMsMssAsol/Pff/AP8A/wD/AP8A/wD/AP8Aqp88/DDDDDDDDDDDD8Wss888888888sy/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A7Yc88888888sM899/wD/AP8A/wD/AP8A/wD/AOqnzz8MMMMMMMMMMMPxbzzzzzzzzzzzxf8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AOqjzzzzzzzzzzzz33//AP8A/wD/AP8A/wD/APqp88/DDDDDDDDDDDD8W888888888888X//AP8A/wD/AP8A/wD/AP8A/wD/AP8A+qnzzzzzzzzzzzxT3/8A/wD/AP8A/wD/AP8A/qp88/DDDDDDDDDDDD8W888888888888X/8A/wD/AP8A/wD/AP8A/wD/AP8A/wD+qnzzzzzzzzzzzxT3/wD/AP8A/wD/AP8A/wD/AKqfPPwwwwwwwwwwww/FvPPPPPPPPPPPF/8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AKqfPPPPPPPPPPPFPf8A/wD/AP8A/wD/AP8A/wDqp88/DDDDDDDDDDDD8W888888888888X/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wDqo88888888888899//wD/AP8A/wD/AP8A/wD6qfPNwwwwwwwwwwww/FvPPPPPPPPPPPE//wD/AP8A/wD/AP8A/wD/AP8A/wD/APqU88888888888819//AP8A/wD/AP8A/wD/AP6qfODgxAAACCQwAACfdvPPPPPPPPPPPCE8/wD/AP8A/wD/AP8A/wD/AP8A/wD+PPzzzzzzzzzzzx/7rPPPPPPPPPJKfzy4xzzzzyxzzzzz7zzzzzzzzzzzyywwwwwwwwwwwwwwwwxz3zzzzzzzzzzzwzzzzzzzzzzzzwxzzzzzzzzzzzjzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzTzzzzzzzzzzzyzTDzjDDizDTzTjgBTTTRjjDQDzhAiTjTDjSRjjjDjDjDDzDDzxxTjDjzCASAQjTzTDxzjDDTzzzzxAjhxxgDBSgABzRzwBhhQiwDzzyhSxxiRBBiQhAgAyixyxxzzyxyggzTwBgiTyBTiCTjwRDzzzzwjwixijgTiBxjRxCDyjRiwQDzzyBSSxDCgChSiChBygBwBjzzDjywCxQxASCDBzyhByyTgTzzyzzxzwzwxzyzyzxzyyzyzzzyzzzzzzzzzxzzzzzzwzzzywiTTzyyzzyzzzzzyyyxwwzwyzxxzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzz//EACYRAAEDAgUFAQEBAAAAAAAAAAABETEQMGFxkaGxIUBBUPFRIJD/2gAIAQMBAT8Q/wADlEPLgP8Agf8AA/4H/I/5H/I/4H/A/wCR/wAj/kf8Hk4Ts1ViFhQbqGz9dlDSEGriTBKKrdamW2ACnrsxbnN2CMlYnR5qPylLFVVkmba6qZKxOhLFyMMLYKpZX+NC9CuzFuclyYIh596Z0pFbInD/ANTQ2VmGEIxXFQ/OZCw6ns6F6FdmLc5P7VUSaOftjxUSQlFtTQ2n8SikwPPRFEShiqVeq3dC9CuzFuclUl1MTQTh+aPKwqqs3podGjofsqY6V0EyfstC9CuzFuclVEY/aOq9toXoV2d3mrx+4yFUGwyGToZZlmWZZlmWZZlmWZZlnQI1eP1WVHGjBgwYNGjBgwYNGCI1U/EpkamRqZGpkamRqZGpkamRqZGpkamRqZGp4eV4/VZRo0aNGjRo0aNGjRoivXb3eavH6rQudKQrsxbnJXj9VoXoV2Ytzkrx+q0L0K7MW5yet9C9CuzFucleqxbttttsZiy9kuMh7dlstlibT60hehXZi3OT1u0L0K7MW5yet2hehXZi3OT1u0L0K7MW5yet2hehXZi3OT1u0L0K7MW5yet2hehXZ3eT0G2V2F6FUAVMfpg9TB6mD1MHqYPUwepg9TB6mH1MPqYPUweo1Qt5bKopjDEmJMSYkxJiTEmJMSYkxJiTEmJMSY3QYo12AyjKMoyjKMoyjKMoyjKMoyiVb2Df7Q//xAAoEQACAAQEBgMBAQAAAAAAAAAAAREwMWEQQEFxIIGRofDxIVBRkLH/2gAIAQIBAT8Q/gcl2jchMPuQ+5D7kPuQ+5CYfch9yH3Ifch9zSTKtVWICcVHJNFWNRCxw+II/wBN5uNxuNxuNxuN5uNxuIyGPdzf8BT4wGurGuiGyg3a40praVRr1EqiGKDdVlyBTwo4R4Y8EcNeD7udKW0qYmP8oZGzrxMpSXUMRH+UM6jbdeJlHCjkfu50uNoqxqoPSQ3alakJLKXA2lUadROiGyg6hy2UcKM3Xg+7nSwbSqI6jXRDFBuq57EqSLEx/lDbqNt1NpzKOFGbrwfdzpcJDIwybKOFGbrwfdzf8hfUjKOFHI7C5/JefnIvPzkXn5yLz85F5+ci8/ORefnIvPzkXn5yLz85F5+ci8/ORGQ4iwofTso4MaKnf/8A/wDsaLxd40UbxuG4bhuG4bhuG4bhuG4a9BYUPp2UcIiCnff1Xbmg8e7m0BYUPp2UcKORu7nSwofTso4Ucj93OlhQ+nZRwo5H7udLCl9OynhRyP3c6WD0SLZbLZbLZbLJbLZbGycHkbjKC2Wy2Wy2Wy2Wy2WxygeFHI/dzpTVyOxW5tHI/dzpZ7ZW5tHI/dzpZ7ZW5tHI/dzpZ7ZW5tHI/dzpZ7ZW5tHI/dzpZ7ZW5tHI/dzpZ/rc2jNfBtVtfwvejL3oy96MvejL3oy96Mvehe9C96MvejL3oy96MXW5LacxHAtFotFotFotFotFotFotFotFotDlGM1oLEuF0ul0ul0ul0ul0ul0uCH8OCBDggQIcMP7g//xAAuEAACAQIEBAYDAQEAAwAAAAAAAREhMRBBUfFAYXGhIIGRscHwMFDh0WCAoLD/2gAIAQEAAT8Q/wDeenCSeLn9XOE/pp8M8XJOE4sdQ25ojahtc2sbcNuG3DahtY2sbWNrYU2gbQNqG0DaptQ2obWNrG1jaxto28UPiC6UnJki/TMyG2J00G2nJcMk5MTm1kL9G3zG2TNTIlToTWMZOfmCav3ufheee2+/fJ9e+T798m2/A8UV96+cdvtvm2/9Nt/6bbNv/wC4DfSvnCr6d84HfWvk5QWQ34XZj4kmigUklCHnesN+G/cIMhvk34b8N+G/Dfhvw3r4IQhvw34b8N+G/DfmMISkW9Q7Xs3YVBfoW4JJJJEkoXVkYucyz6CxWjJKC9SMP9WCCdMWjL61GTOCOMbWZI2Tg0VW0lzYxhjyUc/Dgn5iR6nl8YbKovE/UY4UqSOOhZ4wkxo5qRF+9Yb98D323ebvNb1xv/C27zd+Ft//AOm/sP7/ADT9eL+/xt/d5u/Czjnvh21RG9phOfZisvA8ErLib53OBcczIyw5nliJW3WmVNkITGdUvcR5jY5a/wDgHbdQmX6szqkN2E7kdhhnxLZYmg6rkO+qSENknNLS+w7F6EJjF0suXehAGWr+A1lnet7iCskcoMx/R+Bi8A2FfLh3h3IsungdsBkuJvndjuLing62JhSyUasjY9WozhxySRHnFXzQnVhJKrzLXuj49CCvD6/6J5WRaxdWL3x9xHYfC4d2Jb6jbyEMaEubOsfA6ScPRPuQyTOPgWnQTIXQlvDoEhUr6hOOgqOplTF+Yjt/hDFhpxg3ciy6eB2wGS4m+d4O4uEbUwVJHUTgbmk5khayr0hvRChNcl8xhkNlUQyWmLbCgZzK1cjFCsOo+RPJUHW9/BGNXmXufR0FwEkjYtRIypDNCGevohp9hoSTKgJh3opUtE3KpNCRmJ1qP3EkqJIVLkuKHNl340dqMWGh3Hi7uRZdPA7YDJcTfO8HcXATh5jIzp5ls2MtEPySJlmXREoTH8Z1SSOC7ulFEKB5jZ1SSwgXkdDMrkLqZ+Oz6XPo6C/CycFhPMkcXbgU3Lk1n0Ojfgxh7CilSacKE9GSGnDBNhai5wJ2ukWs/DXX8CH9H4GLDQ7jxd3IsungdsBkuJvneDuL8slCXkKSllzNiSs8kQyNqifmL3TlO7C2FKsll6i3z/WNJsp8CmSdV+ez6XOw+ELxsuPozmFLQ+o6raKDCQrslixUWUsE51VAgomjYSjQYDFXUEidIROs+B6wTr+ZHajFhodx4u7kWXTwO2AyXE3zvB3F48ybid5JJkb1NdtAjaBDyNJpIoQmcWzNtof3u5g/gXXud7hJLJE0KZ3K+COBs+lzsMGOzE1kTQT1RM0EJVpdStsSWmhohcXE0dQ1OxNl2IU/cucKlo9CJ1akwS6SirFC8EYR+dE/L+Biw0O48XdyLLp4HbAZLib53g7iwkpqSSkUumMqFQcoOTWtw1tivmrePYytiXbzZIpE0tqGQmsvq5FRQ0oG3RYSeYrcPb1Ir6JMGVD7Y7eIjnK4+9u5agRN0QrZtbu+5UrWhvgNjVuwyFS0YN5458QjtBiw04wbuRZdPA7YDJcTfO8HcWDvXMbhVaXUgF0IZElE2SCfmJuqN/gPjSPkoGdbdsIU6Yt4THGa8uor3aVUlP0MlH1J+ZX8lJy9BLVFdoQzntxg0mrLfNlJuKi1wnCeMR2oxYaHcOLu5Fl08DtgMlxN87wdxYOxLGEO4PUFU2NuSaQoET0FQeiJ/QdBncac3MCihe51qTyJepXw5mfGIn5fwMWD4yLuRZdPA7YDJcTfO8HcQx2O+8C6/wDFo7QYsHxg3ciy6eLZcUu6HdYMdsIt9P8AjkXgxYO/GDdyLLpihl+BUQ+JXd4sdj7HIt9P+OR2gxDFxaz+MKWXQzx/wf3hzgSf8BRr6bEyt94WUpSlKUpSlKUpSlKUpJKzAYRLmjNbCsMdj7HIt9P+OR2gxYyi+NNAnUl9Bto24bcNuG3Dbhtw24bcNuG3DZJsk1/Rm3Dbhtw24bcNuGzjZJqehNtE/wDiFui6tLFbw0giEJRqVKlSpUqVKlSpUqVKlSpUqVKlSpUdoFztghjsfY5Fvp/xyO1GIY62kVcqi4h8i9hF/A3jXCCCCCCCCCCCCCCCCCCCCCCCCajnkIY7H2ORb6f8cjTZfAxYLBcT5YZeBlHODue0+xM+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5MTaMuhCxQ5wIY7H2ORb6f8cjtBiHhE4g0tBtT2mbMzYmbMzZmbMzZmP8AjM3YbsNmZszNmZszNmZszNmZszNmZszNmZsTNiZszF/GZd94KTGRBDFafDdFNI4mE+hOvqPBjsfY5Fr/AI5HajFxu2p2T4LF4GXhWrfib53g7rBjsfY5Fvp/xyNNl8DFxo2uGWLwO2AyXE3zvB3EMdj7HIt9P+OR2gxYacWNrgli6CxdsBkuJvneDuIY7H2ORa/45E/L+Biw0O4cVa4JYugsXbAZLib53g7iGOx9jkWv+OQ/o/AxYPi4tcEsXTwO2AyXE3zvB3EMdsIt9P8AjkdoMWD4sbXBLF0Fi7YDJcTfO8HcQx2/5BsjtRiwfFxa4JYungdsBkuJvneDuIY7Hf8A/HpXwbFg+Li1wSxdBYu2AyXE3zvB3EMdjuy2v3EDXI8/z0KeBSdt8FrCw04sbXBLF0Fi7YDJcTfO8HcWDl0kd+TmAXg5ejNlZsrNlZsrNlZshshsrNlZsLNhZsLNjZsxszNkZsrNkNkNmFMS0hWlfolUScCUttITIUSMnRGyM2c2Q2Q2dmys2VmyGyD/AIjNmZsLNhZsLNhZtxszNkZsjNkYu0UpJEeDQ7jxVrgli6eB2wGS4m+d4O4sGqjVcxIggjwQR+Sso9qIeX6GaQoguUOlxwQR4Y8MeFoqLV46HceKtcEsXQWLtgMlxN87wdxcT7QVqYZ/oMheKV6HceKtcEsXQWLtgMlxN87wdxcSvaF+iZcXV6HceKtcEsXTwO2AyXE3zvB3FxPsBfomR9LnxR6cWNrgli6CxdsBkuJvneDuLifYC/RMhhdn4l8XFrgli6CxdsBkuJvneDuLifYC/RMsK7PxL4uLXBLF08DtgMlxN87wdxcT7AX6JkILs/EvixtcEsXQWLtgMlxN87wdxcT7AX6JkfS5nZ+JfFxa4JYungdsBkuJvneDuLifYC/RMsK7PxL4uLXBLF08DtgMlxN87wdxcT7AX6JkILs/EvixtcEsXQWLtgMlxN87wdxcT7AX6JkfS58Ueh3DirXBLF0Fi7YDJcTfO8HcXE+wF+iZcXR6HceKtcEsXTwO2AyXE3zvB3XFewF+iZGT0cU+nFja4JYugsXYumXE3zux3XFewFb9FlxjP3DirJnZCxeBjcg+uZpTiFLS5j4r2QhywkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkVqlzyFVXi+W8RnmZObDKk8pYvA7DNqmhdPwzazY8P282828282828282z+lf4zbjbjbjbDazazbTbzbzbl/ptw+pT9CeuZSZNxcQ3DIzSE3QQc4MBsxt5t5txtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxCs8RObM1LDG0GzIy4dxmK/InuIaZgVSQNsNsNs/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/pT+AoOZuoRXJLWLUVEsYIwJPMggj8cEfiZzD4kkigJakEYkEEEEEEEEEEEEEEEEEEEEEEEEEEEEDWjIpcjqLiFCxGnE+eEc8I/8M5/LXQf6CXp+NfqHhP58/H0/FPApYQpqSPXumIKYJxUND8c+JvBqqH3IXhcUi1xYz+H2G5zMQocvHJD8LMh1lDaUW+hCaZl11geM4yN8M2hWPMmeT9vBIYlNTQ3+lG5GpvRH0H4GXONtIl2ERt2YMdP6PkXPs+Qq1BFQ5KozI/A7UKGa6wOHZXcS5wNsJ0P2PDOMNawUvIV7VUyV8GkFraZAmWkzG1+ZZbcolcnb8EVEk6UJPPBXEvUJqQIqOvdrAghqSsRDFCttMkRDr3SibVEr6JO7wz4F2IjYjiNVKJnLsIEaBDvXgO8mT1+591oKGXjaCuPq0WJpSzD9EJtUWsxGfDhXA4lFRsr1GeYzINS0Q5LSjg2hImW4liXqJT4tC36XM4qydi+RFwZskKjacKaXkJZmZkLjVcVpLLh7TH1J1INZzfoNimjNMQKYHEkMSyajS/YcKwtlcRTEzZVCuLE3ppCKX67EBl58oaxFro6v3CbQhVtK/oJisxChjkiKtEoXGI4SKc+h1HpWVbT7h0i6UchiLjvB0mU4SaBDLpSDT9DM757eHmpVajmG9u7gP8AihbDmRlUFQLHWknMaT2xsom0FP4C6gaUXTaloJDouHE1KepI5I1xQaE9dOT7DS7qzSRPY9VaRFs8mfS+DJJt2IAXAynaLXEcqlGjQ6Ic26MZrjxAub9BighmmckxucJUFkOCTT9MOSlmUJp+6iYySXZDyu5HyZF3UHQo6sbEvRtoFiGWbsLRQnd+lz3wekkZtwXZYTHag0OK+chdazhTA8rVJBqNTJlQmZFJZKaTXn4O0Z2xlQqn7dZ0h/kZXFlm9EVv9h3e6E+X+C01Ms05QhGTpbYS8xrRjiSRVgM2GVgig+5VzEUtHKH6oXZkUmTyESK82kNNukIeDqEO5I0DuORGXyjM2iwzIdQS0aixtIpuVAXmW7NZKoSjpoobcqhZMwrqxZ4ZPZIRHi2lNLOmFRTdlK0lkujrjVe2ZlUIyKV1KRr54NymMnvCS6iFyMozQ0JDNoSKKYHMDl/2HpDoU5PsJLws05Q/wd5O7e591oKafMX2ngqqjzFaywbTeTZPDFAFYAWyXJoXJC7zlyFNkE09UxtWjqOnOysUp5cskKfxUNpekcOSWaG64mtNHSj9BgJqKG41R2L5FTl1jAaqkivMVsVLnP5QokpAYhRiNGlVQqMKRNk+3oMQrpReMkPNiGnkSKDbD6zQt2EpSXyCSkSkL0iFDaojmVQpm0skCVYabKQkCzanNcy5z57dNeiKmfFAk2JWWTpeghsaqqPXB1S61cHRCk85SSE1axrFlsm2Jb/BorL7NRFOkXrDvnt+DvsdT6vMtdMJBVm8xgdWIaJjPySjNI2hcbsdNwSTYkHodxzB4BSSc+rF4KzWVWSQbdzkWo5NWJRsmGuVZs7x7FMKNJCpLVknsKhpoRtWPlqcxYqkiXZSOn2oAhg0mTK5NCsQH5quJQ/qVgU0bumUUrEdjvHsIKrqFrDIGETNWqcuYhzq7oENIlyZc0ILFGbR9URyNaXvRIhkzfLIkzcaKkVZdku5Tq2KzarzGyNJLQPIZ2QbOaI9c0RSVBMpaIUPol4O0Z2wigHUL7StQVUnzFkxyGhm+tRpjPJMqduRuPNRS3XWs06DVOlkhYlLJm3CbFzJQ0Z9yaMWrENFOMw8pdGqj+lcNtpKMpEJaHBXhMU3ydYkaCFs7Koj+AKJ7D2IFkpUh7jRWi7DPNE1WtiV1lmnWE4H0nQ/QO+dKSoSj0TrsPTLQ5VLHRYqVK186iIeiE9edDXoOpF0bvEwRiwaIqWMqlEToHVSIUrxkTin+Q1mJb0IKUoRLk27mucXmooiO0dKfJO6K/Zgbuum8hMdYC8feTu3ufdaFjzHzi2mZUxRNLSj4BBiyK6itFH2JoQjrQ2JqqoM3zFhHIqqGtHYvk7/AOx9fmK2Nzr8o7dDkKutD5j0jSnwVahWnqVs550HTK4bVKIKGKqkOqPU9wW+q9se1fsM4ijrKWkGZB3bCm0lLdBUaCts+YiSnvMX1/ogNQxNV/0R0fWDpyL89nFI5Btu0HfPb8HfY6n1eZa6Yd4ylRAz5ipqcZBFrIjtYs43kNVMo7T8xKxeKhJSINHZJ/0V1I5cqTMs3Yd49iXk/c7admJCr3Kw7rCqTX7CFSFH1kO3K6xxI7x7Cq6EHzqJluVtLVoWh6icDZRxYgKbiESJRlScDah5XP5eBpCvKSgh5SFDRM+o08HaM7IVi/rQ++jmXIsbYrb2GbhSSXqVpzLPoVycORLk9SpJKv8AlPQQqOzDb0ZSttCwvURuuZZM76dw9z6rTD6TR4C7T7He/Y7oPpNMO6HcvcsdMfv6HZo7l7HZ/c+g0weCamDFr0hNOwcsLvFf/QUynuC0Pms3uJKuhWKn4O8ndvc+60LHR+C/faM+jzK1uIiNzbLgUTmalMyrOkl5DJJaU8FiZnO5Esx78YkeaKOj+RmVbtJLkX5POYKr2YSmsGk+NtonoP6baDcpMpo9iiV5VSE5MlmUxUSd42Og/EjXqyFCqtHE3K4q2Gj6tEqiGRIBKTmDMzO2+B25sXdvbBHWIvDusPdtubUehyLLJaKhrBQ0JYhJDw87SFH55nszvnt+DvsdT6PMtdMO4FTmSssHalx5PARKJTWBM+04zXqvjPsVlWgnpHfPY7H7nbTtz2Me++x2f3OzYkd89jtvufYaLHuDtl7IQjLCT0epHRvlNOzaHe44Yq3UXWrWJyimUd+SFX5NAadcj3BOTqfUaeDt/hnbCsXdaO5e+PfvY7mGSY3Cj7FMNqRp5HJgTUoc1V3cZ9RdVTq1HSh6zjVpXQI6WsVULxvKoxxI6y6YfaaMp6X3O0ihY5aeQ2CKdOaRIaKmQNdMFBKSS1Fqt8MlNxBiiOTBkzDGCXxu6JN0dmhG1JS37BJhWiVuaUkvg0JzkMQ7Dgs2lUeSSBDqRNRJCpbFHoffPgzvudBiconPPoIzR0qBJ+n4Kfe9RjUV2ZqGm5GWsVPyQzovIIxFSROD7CJeg/tI5KImSPYyNHFhFFWmkXprhMcjDtT4pSRpdamRtm6InmJfJp3WfJqyGbVNozzMfKYrTMQEkIaeaHVkzS5aFvLVb6IlF7dGyIHPSOTrnk1ZEqCim92QyvAMVRn1KOTDhegY5GSys0w1CSHrFRa51Td65kAkSSougpusxU6h2IrqzqoYm4Kj9wyKGDkFsSabZUGf65jjC/qDlR1QiklA80MOXYmfJ5EjhUZKDlNYIWoIo/6JZ4dXDLRJf6Lj5XRkHOeXIYHerUK9GCJ8UYbdS5MdIwoicihJ0waLS5aXKWulJPB3kkdlaS/UvpbeFzbMXmrERdky6FiUoLqqjcsS7ieUie9QtYkrV2mhL76jpCcyKUKVoOCr3kBxgS+Jl0GS6wJszRGrJlI9egyZ9bVDqU4oJVsK0ldpoh7aDpEyR7dxDTjktHUVmJKScF7MhqTehVac6inoyyuVoyQgjMg/Qbg4rSb7sePMzPRSSldPoXgiUiELnWQsJUosNUdQA2OuMm6UeQ1D5hqICyMpzDVIKNoljJ8k+R+fn9EjzFlRBqlLnJdEoNjOWoYX6ShHmivuSJJ8mPHSkjiuTdSVSpdI10VZH33U9TgYcqpKlRBQZJoxHUrfQKksGh5GUwOQKlpnA5udCkfkhK0ioh8yuvFCuIVtmU/yYz0lLUTZdcPJCKakourdKshK+ocMXwqaeXJjmzVwlJnSSdx6OKR5CYilT6seZOrHqIpkboonRkx8uzYRaQm4SLq1UVQlVOXzuHvoAbli8a5YPkZlTPw10FhDM/BD1Jeh0KiM8ameEY1w6PwIz8HTwsX4o8GpUzwyph1xqVMvHHjnwYtB6jUxEsQtIdBBRRbLb81Il0W9WW97TPfKJQ6JhLpkeiC8knSVb5chTm/Ejl4M8eRkVIeuNcPMfLDr4ankVyxrPirOCGlyHJ6/9GihJzUl0PVCvvuSrD1/+DT/AP/Z" style="width:100%;height:100%;object-fit:contain;border-radius:8px;"></div>
  <h2>Welcome Back</h2>
  <p class="sub">Enter your credentials to access the ITI Assessment Generator</p>
  <form id="loginForm">
    <label>User ID</label>
    <input type="text" id="uid" placeholder="e.g. 9999999999" required autocomplete="username">
    <label>Password</label>
    <input type="password" id="pwd" placeholder="Enter your password" required autocomplete="current-password">
    <label>Semester</label>
    <select id="semester" required style="width:100%;padding:14px 16px;border:2px solid #E2E8F0;border-radius:10px;font-family:\'Sora\',sans-serif;font-size:13.5px;color:var(--navy);transition:all 0.2s;outline:none;background:#FAFBFF;appearance:none;cursor:pointer;margin-bottom:18px">
      <option value="" disabled selected>Select Semester</option>
      <option value="H1">H1</option>
      <option value="H2">H2</option>
      <option value="H3">H3</option>
      <option value="H4">H4</option>
    </select>
    <button type="submit" class="btn" id="loginBtn">Login to Dashboard</button>
  </form>
  <div id="msg"></div>
  <div class="links">Don't have an account? <a href="/register">Register here</a></div>
  <div class="admin-link"><a href="/admin/login">Admin Panel →</a></div>
</div>
<script>
document.getElementById('loginForm').addEventListener('submit', async(e)=>{
  e.preventDefault();
  const btn=document.getElementById('loginBtn');
  const msg=document.getElementById('msg');
  btn.disabled=true; btn.innerHTML='<span class="spinner"></span>Logging in...';
  msg.style.display='none';
  const res = await fetch('/api/login',{
    method:'POST',headers:{'Content-Type':'application/json'},
    body:JSON.stringify({user_id:document.getElementById('uid').value.trim(),password:document.getElementById('pwd').value,semester:document.getElementById('semester').value})
  });
  const data = await res.json();
  if(res.ok){window.location.href='/dashboard';}
  else{msg.className='error';msg.textContent='❌ '+data.error;btn.disabled=false;btn.innerHTML='Login to Dashboard';}
});
</script>
</body></html>'''

DASHBOARD_HTML = '''<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Dashboard — ITI Assessment System</title>
<link href="https://fonts.googleapis.com/css2?family=Sora:wght@300;400;600;700;800&display=swap" rel="stylesheet">
<style>
*{margin:0;padding:0;box-sizing:border-box}
:root{--navy:#0B1D3A;--blue:#1A56DB;--accent:#F59E0B;--green:#16A34A;--light:#EFF6FF;--radius:14px}
body{font-family:'Sora',sans-serif;background:#F1F5F9;min-height:100vh}
.topbar{background:var(--navy);padding:0 32px;height:64px;display:flex;align-items:center;justify-content:space-between;position:sticky;top:0;z-index:100;box-shadow:0 2px 16px rgba(0,0,0,0.2)}
.logo{display:flex;align-items:center;gap:10px;color:#fff;font-weight:700;font-size:15px}
.logo span{font-size:22px}
.user-info{display:flex;align-items:center;gap:16px}
.user-badge{background:rgba(255,255,255,0.1);border:1px solid rgba(255,255,255,0.15);color:#fff;padding:6px 14px;border-radius:20px;font-size:12.5px;font-weight:600}
.logout-btn{background:rgba(239,68,68,0.15);border:1px solid rgba(239,68,68,0.3);color:#FCA5A5;padding:7px 14px;border-radius:8px;font-family:'Sora',sans-serif;font-size:12.5px;font-weight:600;cursor:pointer;transition:all 0.2s}
.logout-btn:hover{background:rgba(239,68,68,0.25)}
.main{max-width:860px;margin:0 auto;padding:40px 24px}

/* ── HOME VIEW ── */
.welcome{background:linear-gradient(135deg,var(--navy),var(--blue));border-radius:20px;padding:36px 40px;color:#fff;margin-bottom:36px;position:relative;overflow:hidden}
.welcome::after{content:'';position:absolute;right:32px;top:50%;transform:translateY(-50%);width:100px;height:100px;opacity:0.15;background-image:url('data:image/jpeg;base64,/9j/4AAQSkZJRgABAQEASABIAAD/2wBDAAYEBQYFBAYGBQYHBwYIChAKCgkJChQODwwQFxQYGBcUFhYaHSUfGhsjHBYWICwgIyYnKSopGR8tMC0oMCUoKSj/2wBDAQcHBwoIChMKChMoGhYaKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCj/wgARCAKbBLADASIAAhEBAxEB/8QAHAABAAMBAAMBAAAAAAAAAAAAAAUGBwEDBAgC/8QAGgEBAAMBAQEAAAAAAAAAAAAAAAEDBAIFBv/aAAwDAQACEAMQAAAB1QAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAH45TYzzPQ0VnTjvRWdDRWdDRWdDRWdDRWdDRuZ3w0XucjRWdDRWdDRWdDRWdDRe5xa7qZ7vO+hhBIAAAAAAAAAAAAAAAAA4jvPShors3KHC8Uat2Jlu9XRPQAAAAAAAAAAAAAAA4d568RFc+pEPxn07tGvPWjo6s5+f1SZXF88dt5+hnzyl9DPnkfQz554j6HfPJP0M+eR9DPnnkPod88j6GfPKX0M+eUPoZ88pfQz55H0M+eR9D+XFdqqkOZApMXKRfzH0IUXAAAAAAAAAAActlTte/HYB73iBIAAAAAAAAAAAAcO8567n2lfh+arv48ziuM2nw1Ec57HE+k5zDnNffz3hqsxDzGn6EOrAAAAAAAAAADg6/PrI9rtfh+Kbv+Mziuc2oQ9A7zns0THuM/6/POxUc6i2X+gX+/3As1KPeKP1GT953ZWAEAAAAkEAAAkEALHtmJ7Zm7CvoCkxcpF/MfQhRcAAAAAAAAAABy11S178dgHveIEgAAAAAAABwPFFxxMqbDc59K9DLvW4z6DEVVzmko7jmjnXI464AAAHO8TqsxDTOn6EOrAAAAAABw7z1YmK7Ao8PxRpsdl/g5zX+IrDnPIehxzRzvEcdcAAADvOlsv9Av9/uBZqUe8UfqMn7zuysIAAAAAAAAAABKx7Zie2Zewr6ApMXKRfzH0IUXAAAAAAAAAAActdUte/HYB73iBIAAA50OflH67Fw3Nds5nsPzn1CJzfnGa4Q8PznP5fF3nNDvCAAAAAAAAHO8TqsxDzGn6EOrDnQ4O8erHPtK7ERTefxmcVxn1CGoPec9niY5xn7+eoq4EAdcAAAAAADvOlsv9Av9/uBZqUe8UfqMn7zuysIAAAAAAAAAABKx7Zie2Zewr6ApMXKRfzH0IUXAAAAAAAAAAActdUte/HYB73iDku856cc+6rMPzTfvDmMZzm0uHpLnPOxPgcZ+d52K3BHeAAAAAAAAAAAA50nU5ivtHvWHlGiYq0yNzDw8579D1jvOf3/R/PeaDiOO8AAAAAAAAAAAAB3nS2X+gX+/3As1KPeKP1GT953ZWEAAAAAAAAAAAlY9sxPbMvYV9AUmLlIv5j6EKLgAAAAAAAAAAOWuqWvfjsA97xAlA5lpOa0eP3hXh64HeAAAAAAAAAAAAAAAB3g7wT1wgAAAAAAAAAAAAAAAAB3nS2X+gX+/3As1KPeKP1GT953ZWEAAAAAAAAAAAlY9sxPbMvYV9AUmLk4z5j6IKLQAAAAAAAAAAOWyp2vfjsA97xAlXs00vNKPHCvCAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA7zpbL/QL/f7gWalHvFH6jJ+87srCAAAAAAAAAAASse2YntmbsKugKv6F2YddKXVx3Se3UUpdRSl1FKXUUpdRSl1FKXUUpdRSl1FKXUUpdRSZ+X5bV3vO68wSr2Z6ZmdHj9FeEAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB3nS2X+gX+/wBwLNSrWnhjHdnWc4xzaEsXbQMXbQMXbQMX7s4xdtCGMc2hLF20DF20DF20DF20IYxzaOGZaccSESBxWvRxarn2luO7mphFzUwXNTBc1MFzUwXNTBc1MFzUwXNTBc1MFzUwXNS562uW6as4Sr2aaXmdHj9FeEAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB3nS2X+gX+/wBwLNQrBZ2N9s52Njg2Njg2Njg2Njg2Njg2Njg2Njg2Njg2Njg2Njg2Njg2Njg2NmOm8T0RIFJi5SL+Z+hDPcAAAAAAAAAABy11S178dgHveIEq9mml5pR44V4QAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAHedLZf6Bf7/cCzUo94o/UZP3ndlYQAAAAAAAAAAAn9ww/cM/YU9BKkxcpF/MfQhRcAAAAAAAAAABy11S178dgHveIEq9memZpR44V4QAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAHedLZf6Bf7/cCzUo94o/UZP3ndlYQAAAAAAAAAAAntxw7cc/YVdAUmLlIv5j6EKLgAAAAAAAAAAOWuqWvfjsA97xAlXs00vM6PH6K8IAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAADvOlsv9Av9/uBZqUe8UfqMn7zuysIAAAAAAAAAAAT244duOfsKugKTFykX8x9CFFwAAAAAAAAAAHLXVLXvx2Ae94gSrua6VmtHjhXhAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAd50tl/oF/v9wLNSj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vEHJV7NdXqdPmVRbHGWprYKmtiVTWwVNbEKmtiVTWxCprYKmtiVT7a0KolYmKQjgAAAAAAAAAAAAAAB3nQWWbK0tabamtiVTWxCprYlU1sQqa2CprYlVFrQqa2Cp9tY7fq1ZbvV6O9Cj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECXOgAAAAAAAAABRqZdKZn8Pg4ygAAAAAAAAAAAAAAO86NayXW7fRkBd6oAAAAAAAAAACj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECQAAAAAAAAAAFHpd0pefww4ygAAAAAAAAAAAAAAO86NbyTW7fRkBd6oAAAAAAAAAACj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECQAAAAAAAAAAFHpd0pefww4ygAAAAAAAAAAAAAAO86NbyTW7fRkBd6oAAAAAAAAAACj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECQAAAAAAAAAAFHpd0pefww4ygAAAAAAAAAAAAAAO86NbyTW7fRkBd6oAAAAAAAAAACj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECQAAAAAAAAAAFHpd0pefww4ygAAAAAAAAAAAAAAO86NbyTW7fRkBd6oAAAAAAAAAACj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECQAAAAAAAAAAFHpd0pefww4ygAAAAAAAAAAAAAAO86NbyTW7fRkBd6oAAAAAAAAAACj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECQAAAAAAAAAAFHpd0pmfw+DjKAAAAAAAAAAAAAAA7zo1rJdat9GRF3qgAAAAAAAAAAKPeKP1GT953ZWEAAAAAAAAAAAJ7ccO3HP2FXQFJi5yP+c9703uqbfSe7w9N7o9J7o9J7o9J7nT0nudPSe6PSe6PSe6PSe6PSe509G1wNj25JrvO+344SAAAAAAAAAAHCkUy91ejxozko4zRaUEWkxGJMRiTEYkxGJMRiTEYkxGJMRiTEYkxGJMRiTEYkxGJMRnZMRmtZxpVm/3hd6YAAAAAAAAAACj3eozGPdmO6+IZMiGTIhkyIZMiGTIhkyIZMiGTIhkyIZMiGTIhkyPLuORa5n66K+gOdIBIAAAAAAAAAABzqHOkgAAAAAAAAAAAOdAAAAAAAAAAAAAAAABzo46QCQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAHp+5FFC1D53+hz9vU9ck0bClsV6fP0A8EAWZVZ8zmxUu0l6eGMJkAhyYeHzBGQ5a1fmzy1Waxs1+Vza0lger6hKvW8R71f8AxNGYaxCTYeOIJtVPaLC8fkIbMtIyA30Dx+QZX4dWGUeLQMaLW1YU+6ej4SUeGALK8Ppkkq36LO9P3D847skAezLeLxnsorzHvqn7RYn4/ZTYa++uSxCk08EIWJVJw98ByIJhU/eJ54okm3OgELT53LzfSOJF6HrkuqciTbnQ/ECWFVZYlPF+64Rl2+cNuLH4YmaMb1qMmz802QxY+j/1XvOTTwxZNAEQS/58UWQNz+cdqLUgZw/TkMTSpyJNudAAAEVKxRhP0R87/RBR8v1Gpk9N20ZHEbf87n0TU7Dg57t7nrAVCtaoPnbUKhbyZx7YsePoADBd6wY1iLlP0ZdoudSRYXhuZH4/9AYqe7YPYupF5RteKHPb9/UijSllrhlG9YLvRWIGVqpZor1ulP3DCtUJTINfyA30AAEFi+0YufQoKHnWjZ4S8/ov7IqkXSgFg8UQKxolK/Zs/wA6fRfzsbvAz0EZBotZ3AqVW1b8mD7r897sVWsWesmxfO30T87lotc3NlGzb6Cz89615HdCgS0LuRUvSvQx6o77gx9Efv8AH7AKll+oZeb7lOrZSVi9+lp5T6LtUaZlr3zn9FGJ2j81o0Gv+tZSVrt69Y+dtXzH6IKdcPN+D503/APoEpmU/Q+Amn+5b/2QWL7Vix9CAYRu+Empen7s+fNeqZp9AlUtfvZ2VSx0q7Fhrn6Fa2jBNNLsAABFSseYJ9EYxs5SKtdoA0wH4+dvonFzXsd2UUbQsr9E2P08w9opWmQdqJLDvorOy/8AmxORNQ+fL/EmiS8fDFx8WUecteR2ufL1jO01Y9O7YjKGtYrqtDLDcK1ZRXLHCGQbTmetnzj9DUCINg/OZdKvolOvh7OQbNmxsAAAILF9vy42cFEql5rpp4MC0/y0M2Vlv7PcrUl5jT/nb6JxY1eCsMSZxtmWamPz+vyfPO65VrJUazdII0v53+iMXNSl42SFBv1PKZqtB0w+fNGulTLF6tRkSv0raMjPoD9/n9AFSy/W6Ea/lOrZ6etpdFvQj5D0j5/+isW2Q8ft5d6xpmZ+X1D3tH9GYPnP6IzD0DZPxmukHzt9A5BsR4sA+jMpNU8mOTZcsH+jKCXf2cSkTVPn68RhoM/ETZ84b/R68bTUq1pphe2VunGtMs/Joctm+jnlAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA/P6GVeDXOGW+HVemLaZYAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB//8QAMxAAAAUCBAQFBQEAAgMBAAAAAAIDBAUBBhUzNEASFCA1EBETFjEwMkFQYCEisCMkcJD/2gAIAQEAAQUC/wCg24qUHEUcRRxFHFQcVBxUHEUcRRxFHEUcRRxFHEUcRRxFHEUcRRxFHEUcRRxFHEUcRRxUHFQcZR5+f81WtKD1SD1SD1SD1SD1CD1CD1SD1CD1CD1SD1SD1SD1CD1CD1CD1CD1CD1CD1CD1CD1SD1SD1CD1SD1SD1SChy16pHz9fzqPOo86jzqPOo86jzqPOo86jzqPOo86jzqPOo86jzqPOo86jzqPOo86jzqPOo86jzqK1qInK/UHdIphSYblCs0pULyDlURla1afozrpECsq2JVWbCsm5UpAqHUV6LprWjTiOOI44zjjOOI44zjjOOI44jjiOOM44zjiOOI44jjiOOI44jjiOOI44jjjOOM44jjiOOI4gDG53pks/dRGT+jMcpQrINkwpNp0Csu4OFHKyg+fGoitHvjGoUKv26QUmkqBWZXMFHjhQVrWvRbub0XVpKbq39b0yWfuojK351kyBaVbJhWb/1STcnB1Dn+hURWj3Si6SYVlWyYVmwpJuTg6qin0bdzei6tJurf1vTJZ+6iMndVr5BV4gnRWZQKFZpWoVfOFBWta/UqIrR7WpqFCz9ulRWaSoFJlaoUeOFBWta/Vt3N6Lq0m6t/W9Mln7qIydsY5ShSSbECs2UKyzk4O4WUrsKiK0exOukQKyrYgVm6hSTcnB1VD7K3c3ourSbq39b0yWfuojJ2SjhJMKyzYgUmjVCsi5UBznPtaiK0f1KmoUKP26YVmkqBWZXMFHjhQV867a3c3ourSbq39b0yWfuojJ+rU1KBZ83SorNJUCkysYKO11BX/dzURWj+gdwkQKyzYgVm61Csm5ODqnPXdW7m9F1aTdW/remSz91EZP0TqkIFZNsSis3QKyrk9TrKn31RFaPxqalAo/bJhSaSoFZhcwVeLqiv+131u5vRdWk3Vv63pks/dRGT0qOkUwrMNyhWaUqFJByoDGqav6KPXTTaKyrYgUm6hWTcnB1FD/pLdzei6tJurf1vTJZ+6iMnomTmIz+f4+3c3ourSbq39b0yWfuojJ6JzRfx9u5vRdWk3Vv63pks/c1ERldE5ov4+3c3ourSbq39b0yWfuaiIyuic0X8fbub0XVpN1Aa3petVVFeRXHILjkFxyC45FcciuORXHIrjkVxyK45FcciuORXHIrjkVxyK45FcciuORXHIrjkVxyK45FcciuORXHILiPROin0Tmi/j7dzeieaKu2+BPhgT4YE+GBPhgT4YE+GBPhgT4YE+GBPhgT4YE+GBPhgT4YE+GBPhgT4YE+GBPhgT4YE+GBPhgT4YE+GBPhgT4RES6buf0s5ov4+3c3+FnNF/H27m/pXjtVJXEFxiC4xBcYguMQXGILjEFxiC4xBcYguMQXGILjEFxiC4xBcYguMQXGILjEFxiC4xBcYguMQXGILjEFxiC4j1jrJ9E5ov4+3c3onHirND3A+HuB8PcD4e4Hw9wPh7gfD3A+HuB8PcD4e4Hw9wPh7gfD3A+HuB8PcD4e4Hw9wPh7gfD3A+HuB8PcD4e4Hw9wPh7gfD3A+HuB8IyZduHXTJZ+5qInK6JzRfx9u5vRdOk3UFrumSz91E5PROaL+Pt3N6Lq0m6gtd0yWfuojJ6JzRfx9u5vRdWk3UFrumSz91EZPROaL+Pt3N6Lq0m6gtd0yWfuojJ6JzRfx9u5vRdWk3UFrumSz91EZPROaL+Pt3N6Lq0m6gtd0yWfuojJ6JzRfx9u5vRdWk3UFrumSz91EZPROaL+Pt3N6Lq0m6gtd0yWfuojJ6JzRfx9u5vRdWk3UFrumSz91EZPRIN6uUMEVGCKjBFRgiowRUYIqMEVGCKjBFRgiowRUYIqMEVGCKjBFRgiowRUYIqMEVGCKjBFQ9ZGafpUohRVPBFRgiowRUYIqMEVGCKjBFRgiowRUYIqMEVGCKjBFRgiowRUYKqMEVGCKjBFRgioi2B2h+i6tJuoLXdMln7qIyd1cH6T8sNLubq0m6gtd0yWfuojJ3VwfpPyw0u5urSbqC13TJZ+6iMndXB+k/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyt1cHz+j/LDS7m6tJuoLXdMln7qJyd1cH6T8sNLubp0m6gtd0v0FDrcqsOVWHKrDlVhyqw5VYcqsOVWHKrDlVhyqw5VYcqsOVWHKrDlVhyqw5VYcqsOVWHKrDlVhyqw5VYcqsOVWEamZNPdTSCqw5ByOQdDkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkXXmzLUrfc3Egq4bYW9GFvRhb0YW9GFvRhb0YW9GFvRhb0YW9GFvRhb0YW9GFvRhb0YW9GFvRhb0YW9GFvRhb0YW9GFvRhb0YW9GFvRDx7pJ3/+LztWqLdpcS6zun+0+pc75w2Wtpwq4Z/VkppFgpHvSPUeqacLNmkTLSC7zrllDJMoKRdLP/FSvCRa4XZFvcb0VuR5QUulxWvuN6Pcb0Qcou/U+jX4dzMkR1GKqLNOu4XjlolAST1062Uq7MzbRM+s8d/VUPRMhbibmcFr5lC1alSNMyfNNDmUbqqUSTSuNuo4p/tPoVr5UVuNumukeiif15PQxvcy/a4cJNytnzdzVy9bthj7AITDJatK0NTwVVIkVScYEqWeYGqg4RXpeGfaWgVUKkRGTaLn8V5JqgdJQipA4fNm4x9gEZhiqCKEPRy8QbC5nCThzb0g1QZ4wxDddNwRy/bNjILpuCOHSDeiMyzWWWUTTTbv2aygMYpaKyjNIY+wCUyxVBDlPQTfb7b7l0cBRwFEyUtGEJ3LgKOAoXWRalbP2zmqihEqHm2JDpqFUTWfNkQadYFqScYGqi5RXp4KSLEiiZimKuumgRs/bOTOHaDalZ9h5ozLFapTUNQOl0UCNHjVxUKSjNM9FSVTXmWKNaT7CtW7tBx0OJJo3HuBgEJZkuKnLQhJRmdTouTt9s9yDp4g1DZ2i6K5kWrYY+wDeUZuPExqFovMMURSeYVqi/arDzp5PJNn6SKhKSJJdjQhJVkofzp5VkWFFCmpUr6SZ+i0OUsgWXY8KUmzVOooVMiMm0WU8V5FqgdFUixHMi0SDo5TSDWVYlbllmRjUrStK/4HEmzQHuBgG8ozXFP9+jJ6GN7mT7bs0UWu5IdtblFBW32FSSdt+inByqzRyU1DFnJcrAjdB9NLNrZakClvMDlcwK7Oso+Wd1tLQTnb7e7j43J3KC7fcz9Ro2hIvFDUgGFKOLZaHETCKsnElFoyAnmKTFeFhGzxr7YZhgyTZJXdrGsuZvHNrfdvAyt1BstPdvtnuQmYysgVG1m9Ke3GPk6tdGpEl3cO7YOaO2032+2+5dU12+E7l4XbooWQLH1IxkJlRK1kaBRpSrJK1y1MS22RQrbLQ1JCPcxCtuylXqVfiS7pG6O59BGOF0VW1vVXHt9hwyVs0InESS7F0meihLu0NnamvxL/APGRZpvphNC22SYcW2zUo/ZOIlxb0nV8gsoVFOUmXD5ZhbVVCkt9gUri2mhw6SkIgkdXjkSfZ43J2+2e5C8/mIVdno3tpEKW8wOWXgDsy23LqFV/E/ILOHkZbrf0T2+wNRxa5RGMzNWr232taJIlM9JbLOpEbcaJKmp5JOO6NKebR/b7XhbIlUe0tlnWjW32rdWY7fBdz8bm7jAdvfwTZyo4RKm8b22zUQTtpmQ6qhGjd/KupJwwtknDgLDyf2yThjpVzHOEFSrJdcnoY3uZPtuzRWgWlXXgp9jjurX/ABrNqerJQiRUmHjdqBE3NpaCc7fbvcfG5O5QXb5aOJIIYdJxiqdwvUQ3uhucNXiDoou7V2tofC7tZaTEvD4T/b7Z7kJOdbsq45JOq81PAr+bTrMulXS9saCb7fbfcuqa7fCdy8Ls0VutCunpS0KUSEggxIpcq6xivJ0wq7naCYkHi7a0NXX4ku6x2jufQWmQpn/gb7ZKnlKRuiu7Q2dqa/Ev3KHLQjDwu+n/AKdnai5zGpGWoUhn/jIFKZoxpSkoT7PG5O32z3IXmLNJT0/CRp5s4/8AyTL9k1AGcKIrS8aE7oUKGlwM1wQ1DleaZDuiWWFPscd0Z6Z9pGPcyfYJnQQXc/G5u4wHb1Mt53NlpRd6pitoR6Vkp7nqPc9R7nqJdzzy1rKHMy65PQxvcyfbdmis/UeB/scd1aaa52R0HdtyqaiPg5cJNk5t/wA85tLQTBKnYRS1GsgkoVUgcLpt05RzzT2C7equkiC1oah0kzh/Cs102ZjtJROvmS7tXataVY+F3au1+2eE/wBvtnuT+tStGlCKSDciaaQrWhaXM6TcPLX0M32+2+5dU12+E7l4XborP1XhcZzmkbdRblZeF1u06NbP1dfiS7rHaO5tBaOu8DfbJ90jdFd2hs7U1+JfuUVofC79HZ2oetyum7ls5iXkdcbdYtH7WoXlmSJZKWXkSxtKlkCfZ43J2+2e5C8xZmR4P9Ix7oX7CuUTnr/oVZt1aXHENmyNnrqmDzTI18pNHKB/scd0Z6Z9pWdeGTT/ANIJROqjKPV5SRQVIsmFlSIpzDorp9AdvU+x9/kmxrSrQXEyq7ZwT0rF0lRBUnopj0UgdwxIqShaU635DKNWMO9Tfl+24mqrprbUc5aLeBv9KtDvayDYtSoOEE3Cb62TUOQs40oVxPKBCCdOlJuGUqLdaqtGdaedJa3SuDJNZljUrmeOEIV48UlIRfmIlI6DOXhCvjUjZdlXmp4gVTnHtIiAK1OJqJLIJpx0swP6s+qI0ixG1xxrp05gG6jZh4S6J12cFFPG76tPOkrblFjokm2Io/mz0LHyr6spAKENAtlWzSVSOszhIp43fdUokdZnFRD1F94XE1WdNbajXLVfwl4dKQoWMlo9QkhNFHDOuw9t41GttRrpq4r8Pod6o/ZEMm2nm6jlnbkY6auvA3w/h3qj9iQybW42izprbUc5aL1+JGHeqvo9MyTXwuRos7bW1HOWiwVImrRzbjJWvtVLzQtpmnWRM3j4+HTM5ky/4XxnG6jhnBRLxu+FzsHDsWwyXZpeDwlTt2kO+I/On6iLq21aKFJPNhzc7UYRJyCsZHpMEa086S1uqGVQxxrRoWaWc1p/41oZ8Z+2LUiBqedJa3VarIY62KwLMqufxL2+R0ZJlMMKldTxglDvnqkpBrevDonQZibgKuFG6U2yCOOLqp0rROUgkHopGy8ebnpwoNSddhjbX/NMlEyfz9ftfRcmRyWVl2o9yOKBSffqikZJyasVGJR6f/xGtPMcBBw0/wChr//EAC8RAAAEBAYABQMFAQAAAAAAAAABAgQDMDNREBESExQxFTJAQVAgYWIFIUJSkCL/2gAIAQMBAT8B/wADobDWnVmPDvyHh35Dw78h4d9x4d9x4d9x4d+Q8O/IeHfceHfceHfceHfkI7PaTqz9KqKhPZhT2GXQSeos/RKiJT2YU9hkIcTcTqx1DUMxqGoahqGoahqGoahqBYt6ZTX1L0BnkFOIafcKfpLogp8s+gqMtXZ4EIflKbnkFOIafcKfp9gp8s+gqOtXZjPBpTLA5ycW9Mpr6lMNaU9hTuGQU/8A6kFPIiga1K7+ohC8pSVREp7MKeQ0hT/+pBTuIr3BrUrv62lMsDnJxb0ymvqUhUZCezCnyC6Cn6j6IKcxFe4MzOUQh+QvoMyIKcw0+4U/SXRBT5Z9BUdauzGctpTLA5ycW9Mpr6liaiLsKdQ0+4U/L+JBT2IfQVFWrs55BERKUlmYU9hl0FPz/iQU7iK9wa1H36BpTLA5ycW9Mpr6li+P/v4hpTLA5ycW9Mpr6li+8/xDSmWBzk4oexElpIeIRBz4g58Qc+IOfEHPiDnxBz4g58Qc+IOfEHPiCK7XFLSrF95/iGlMschpGkaRpGQyGkaRpGkZDT9CWcRRZkOBFHAijgRRwIo4EUcCKOBFHAijgRRwIo4EUcCKIrVcMtSsX3n+IaUyxzGYzGYzGYzGYzGYzGYz+iBTKa+pYvvP8Q0plgc5OLemU19Sxfef4hpTLA5ycW9Mpr6li+8/xDSmWBzk4t6ZTX1LF95/iGlMsDnJxb0ymvqWLuCtas0kOLFsOLFsOLFsOLFsOLFsOLFsOLFsOLFsOLFsOLFsFpNB5H6JENUTyjixbDixbDixbDixbDixbDixbDixbDixbDixbDixbBuk0oyPA5ycW9Mpr6lOd1D9Ew7Oac5OLemU19SnO6h+iYdnNOcnFvTKa+pTndQ/RMOzmnOTi3plNfUpzuofomHZzTnJxb0ymvqU53UP0TDs5pzk4t6ZTX1Kc7qH6Jh2c05ycW9Mpr6lOd1D9F+n9nNOcnGBFQUMv3G8i43kXG8i43kXG8i43kXG8i43kXG6i43UXG8i43kXDyIlUPIjnOoajifsQ2l2G0uw2l2G0uw2l2G0uw2l2G0uw2l2G0uw2l2G0uw2l2G0uw2l2G0uw2l2DFBpM85pzi/3R//EACgRAAADBwQDAQEBAQAAAAAAAAABAwIQERITMDEEBRUzQFFhUCEgkP/aAAgBAgEBPwH/AIHKbhI1LAcn8HJ/Byfwcn8HJ/Byfwcn8HJ/Byfwcn8HJ/ByfwIa2s1LDxYkDULw4kJyBHFxFExRFEURRFEURRFEURRFEURRDTMr9T2Hd2/t8GYhUFQxE3li9MQNQGoYib2MOYzeVy/U9h3dv7bkYA2yFQG2Yj/ssWYkQNQhUBtnYTw5jN5XL9T2Hd2/tsTEKhCoYmO4WP8AM5A1BUMTGdxPDmM3lcv1PYd3b+3/AAbZCoDUMRPwI/wG2QqCc/BTw5jN5XL9T2Hd2/tepn8hPDmM3lcv1PYd3b+16mfyE8OYzeVy9vQJtnExxyQ45IcckOOSHHJDjkhxyQ45IcckOOSHHJDjkglo2EmpmXqZ/ITw4jgKpiqYqmKpiqYqmKpiqYqmKpiqYqmKphpqZ7WuTZOBjkEhyKQ5FIcikORSHIpDkUhyKQ5FIcikORSHIpBLVsKnKy9TP5CeHEURSMUjFJoUjFIxSMUmhSaFJoUmhSMUmhSMGzK/U9h3dv7XqZ/ITw5jN5XL9T2Hd2/tepn8hPDmM3lcv1PYd3b+16mfyE8OYzeVy/U9h3dv7XqZ/IYw5jN5XL9T2Hd2/te2RxEpiUxKYlMSmJTEhiUxKYlMQ8IiiJTEpiUxKYlMSmJTEpiUxKYZL+OYzeVy/U9h3dv7bzefCTusZvK5fqew7u39t5vPhJ3WM3lcv1PYd3b+283nwk7rGbyuX6nsO7t/bebz4Sd1jN5XL9T2Hd2/tvN58JO6xm8rl+p7Du7f23m8+EndYzeVy/U9h3dv7bzefCTusZvK5fqEmzUOBCip6FFT0KKnoUVPQoqehRU9Cip6FFT0KKnoUVPQoqehRU9DQptMqf0rzZHEQMQMQMQMQMQMQMQMQMQMQMQMQMQMQMQMQMJldYyJiExCYhMQmITEJiExCYhMQmITEJiCp/3/ALo//8QAPBAAAgEBBAULAwMDAwUAAAAAAQIAAxESMTIEITNxkRATICIjNEBBUFFhQmCBMFJyBRQkYoKwcJCSocH/2gAIAQEABj8C/wCA21mYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiZhMRMRNX21rIEzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeM1MOPSxMzGYmYmYmZjMxmYzEzEzEzEzEzEzEzMZmMzGYmYmYmYmYmYmZjMxmJ9K61RZ1bWM7NLN811LN0W3H0TruohFtp+J2dOWXru6Pzjlt/RFhImduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxgtZj+fsXrMBNdQQ3EJnVsTdOvUM19BfH6yBOtUE6ikzqqFnWqGayT0H6K+LH2D13Ama9unZ09XzM90fE67sf0V8X13Ama9unZ0+MzXd067sf0X6K+LHrmudaoJ2YLzs0A3zXUNntNZP6q+G1kCdaoPxOzUtOooE61VrJrP6r9FfFj1jrMBM9u6dkhJ+Z1bEnXqN4JfBdZ1E1G9unZU+MzXZ1nY/nwT9FfFj1Tr1FEsBLbp2dOyZ7N067E+GX9XWZrqA7p2SFp1AFmuqfxNZJ8M/RXxY9P1mdaoJ2alp1FCzr1TNevxS/o9Z1E1G9unZU+MzXN0tZyfz4t+ivix6X1mAme9unZU+M1EKJ16jHx69DWRNdQWzqKSZ1AFnWqH8S0m3x9Tor4sej9aosNy1jOzQLvm0I3S1iT6Gt9wJZbaZ2dOzfM93dLXcn0R+ivix6MShsMtOv7Pfor4sejH7QqdFfFj0Y/aFTor4sejH7QqdFfFjpWqNUwEwmEwmEwmEwmEwmEwmEwmEwmEwmEwmEwmEwmEwlj9E/aFTohaOMyjjMo4zKvGZV4zKvGZV4zKvGZV4zKvGZV4zKvGZV4zKsyrMomVeMyrxmVeMyrxmVeMyrxmUcZlWZVmUcZlHGX6oFnox+0H+xj9oVPRrFssnlPpn0z6Z9M+mfTPpn0z6Z9M+mfTPpn0z6Z9M+mfTPpn0z6Z9M+mfTPKWv0T9oVOiHo2W/MwpzCnMKcwpzCnMKcwpzCnMKcwpzCnMKcwpzCnMKcwpzCnMKcwpzCnMKcwpzCnMKcwpzCnAlQJd+PRj9oP0V8Wvox+0KnRXxa+jH7QqdFfFr6MftB+ivi19GP2g/RXxa+jH7QqdFfFr6MftB+ivi19GP2g/RXxa+jH7QqdFfFr6NcU2GbQTaCbQTaCbQTaCbQTaCbQTaCbQTaCbQTOJnEzibQTaCbQTaCZxOs1voocONc2gmcTaCbQTOJtBNoJtBNoJtBM4m0E2gm0E2gmcTOJtBNoJtBGLMDb0V8Wvrg9EETxS+LX1weiCJ4pfFr64PRBE8Uvi19cHogieKXxa+uD0QRPFL4tfXB6IInil8Wvrg9EETxS+LX1weiCJ4pfFr64PRBE8Uvi19cHogieKXxa+uD0QRPFL4tfXB6IInil8Wvrg9EETxS+LX1weiCJ4pfFr64PRBE8Uvi16VqoTMhmzMyGZDMhmQzIZkMyGZDMhmQzZmZDMhmQzIZszMhmQzZmZDMhmQzIZkM64s8WOaQtNk02LTZNNkZsjNk02TTZNNk02TTZNNk02TTZNNk02TTZNNk02TTZNNk02TTZNNk02TTZNNk02TTZNNk02TTZNNk02TTZGDsWiBhYfFAUULGd3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3abBoGqUWUf9l96gFtkFI0lstg/VUUahUQtWa8f1rlRWJ+JzlMED56d/R88VK2Q/6f0HambGirUqkr0CR5RlGjahO6zXoolgoLO6zusZa1G4B+kY6plB/bFetm/QB0bHdbLmk5P4+DNRVtMFJ6agH2/WLHATmrj222QHkYrjZCv02/tiM+YwucBOaCPbbZ+lbDSKPbbZA4wPgKu6Lvgl6s10SyhUDTtqqrNrLFrLb8y0G0cttRwo+ZZzwO6Wc9LaVRW3RYd8vVDYsuU6yluhcq1grS9TNo5O2rKs2s6tdfzLUYHdO3qBN8BouHEu1aoVp3hZeosGEu1qoUy/Sa8vvO2qKs5unUtaXqpAX5l2jUQt8cnWIG+daus2s6tdfzLUYEfHJU3RejlHCZRwlSxRF3zKOEyjhL1UqgllCqGMtqMFHzLvPAn4gdT1TO0rIPzNsJZzwnZVFbdy3HqJegKYS9Va6Jdo1QxnbVVWbWWLWFvzLVII5L1dgF+ZZo7qT8clx6yhpfvC77yxq4t+JtZ2NVW6HaVlBm1nUrrb8y9b1feXFrLe9ui0Xk7eoFltBwwna1lBm1nZ1lt5bWNgljV1t+Jtp2ddD+ZbbqlRDXW9ZhL5PVvYxf8hcIFWupJlvlLhqJelowlRDWW97QOx6t7GD/ACFlynWUtLzmxZcp1lLdC7VqhWl6mbVjJUrKG9ozKerexiA6QuEsFdbZaMOTtK62zazs662/pVd0XfByc3og67TnNOqs9Q+Us5v8w1NEcmz6TBRqkmmdVh8oCMDLqa6pwlrObvucJ2pZzLBTu7pz39PrHV9MC6QtjrDvlSJ0Ul2lqZvOGppFU2f+5ZzU7Mshl99ILJ7CDni2r2gSiTZ8znKpa34mapObpW2fMWLo2iC2sZzumVbtvvjFqc4xIjReQXapQztajsfiWWND/b1GD/Muljq8veLUHnKm6L06kXfyiO9285wnO1WKUzhbBfrMd0/t1YgWYwnSK7Numu8d86hZTL9Nzd8mEuVdoIY/8pT3QyzRR2jTnf6hWZnOuwSzm/zC+iOdX0mCm5JS2wgwMuB5Ghjn5ioW5rRlnXvOfmdnbTPxAQxs8mEu1dosLubAJzOjWqltmrzl/TKh1+QlnN275bTLUz8RlD87o5iH3MXd0Gi8lOHRtD6t7Fpe0qo1R5YEK/InO6OxZBxn9vXa8vlLZzAcqmFkWppB5xjNlZului1ip+ZzVSoahlSqS97Gc0cttkBvPxgdWe0Qj4h/nEHxKla172MFI5bbINb8YKiF7R7ypuib+ikaq5a9DSGW2yKxL2kQMC+r5lrHqqJzOjWqnsIH0xyzewmxhfQ3Ib2M5nSCSlthB8orpgf0Ku6Lvg5CbNfK26f7om6Nb7xLg6AdRraHfKkXopLpNjeRl6haR/plmk0LfxZO2RkltCoG5B0BDpDrafLleLyFB2lT2E/xdH1bpsl4S2pQvD2sltalzbQSpui9OpF38ogv4LLF1DktrNr9pd0TR/8A7LRSH/jNdIcJzelaNc+Y0Mb+Up7oZaeUw2e8p7uRoY++UwvKseNZOv0KgfCyCzC9F3dBovIkdvPlqbotn7oIa2inrftMu3GKD3Fssr6PLCxQ/MtU2iVN0/3Rd3I26H+Up7pU3Rf5QbuSpuib+ikbdG/lKe7kCDAwvzBqH4ndHndHndHnOJQZJY9ur3/Qq7ou+DkPK26H+UTdOdA6h84KFUhagw5S9ZgohZcgwhlQL7RTU1C3XAyG0HkL1WsAjOuW3VEg52oq2+5lqm0Trop/EY3BTPuJdovbY1mqKTjZBNR5Vi8rxZUK42Qc+dV7ziikAF+OS06hOyNtnnyVN0Xp1Iu/lEflYPbYMIrUQC3meXmbwLnyjQxv5Snuh6Bh3ynu5Ghjb5T3cojxqT4GXgDYDqMC6R2bzvFPjLWrqd0aloNNua82iBsbYu7oNF5Ej8tTdF/lBLi1VLe1s1ztKSH8TnqJuH9samxJQSpumv8AdF3cjbof5SnulTdFvfug3clQLjZFap5HXA9M2g8heo1iiF0yxI26Nb+6U7PbkNzMIU0hRdOrWMIGpqjA/E2a8Js14Tm2NIP7TqAWfH6FRUzGB2pWLbjBLtBbTC1endHKZfFLq3sYobGyFKqhll/Qn/BlihyvGWBCPxA/9Rrmz9sRdBodUS5XW601w1NGa658jOxD2SwIeED/ANSrG7+2D+0odnFSoLGl8VWV5/j1bw+DLGpE/iXKgKIfxOd0hr9Tk1G7UHnD/bhv9su3WHzZFGktbU84GoU7wgSsLG5WSkLWgerSurLDDU0Rrrexl1FZ0H5li6LYfeD+9rc3T8wIo0KmXHmZdrLY0dKYtYxXq0rqjpulMWsYHqUiFBx5btBbxjNXp3Ry25anvP8AFN4fEunRbx3SxuxUy8hatpBhavTujkZ1pdW3GIr4iFKItaX69O6vKYaiUrVtxiK+YS7QW8YzV6d0cjOlK1bcYiviOULQW8YzV6d0HkuVArfBlqg0901V2slrlqm+OFCpq1CAqPO2AdApRFrRXq0rq8icwl6yONIS6Tyuq4mB2pdW9bbLh1aoamiaQbfmWLa44ywUTbul7TXuj5l2nj5mWGGrohx+mXFRyIp0glKcs87JfFLq3rbYitjZCDDV0Q22+UuKr2fMVtJJWn7TXDU0c3KntOwDWfEsFM8IG/qVY3P2wDQ6PZiKlUWNyGtox63tLtJXu8YvOWokAbGX07Op7if4rlh8Sw0TbuljXkU/ic5ptW83sIFXAfcBsjVqFcvuMsr0C28TXohllHRrv4l7SbVH+qWLrfzb/olrmUcJgP8Aga//xAAtEAACAQIFAwQCAgMBAQAAAAAAAREhMUBBUWGhECDxMFBxkWCBsfBwweGwgP/aAAgBAQABPyH/AMDZ2Av2eaPNHmjzx5488eaPNHmjzR5o80eaPNHmjzR5o80eaPNHmjzR5o84eePJCS4n7anNq+0Su1bKjdweEHhB4QeEHjB4weEHjHTHhB4UeFHjR4weMd2MYxjwg8KPGjwg8IPChvDb27jpQj9nlDy55M8meVPKnkTyZ5M8meTPJnkzy55U8qeVPJnkzyZ5M8meTPKnlSD/ALhm3lt/PtDaV2TUZrKRQ6ZsJaf3DBNm2kcjttmxexScyjEIdWgY5X72ymhdg+/knbJOTkxRYnWMbxjet4xvWta1resY3rUjMISTn3LmKdi4P2NTKHdjmJmiIg+ybFkJvVD+r+HA6pZt6vrk9gqaVm7M6ew/a0diFpXk7i2H/ihzOOSPQfK9hK5inYu4/wCRfP7YsDGyXxmFbUGDlxZHfkxdbhVYllP8lsY2S7eTC9qKMfzGkLz6p5QrewK5inYuYtCVJfIzun0ciT/TDKh/sS8dshdB8v1MmFs9F0rN2SRj7DP9vYduDJxTGSLgP5eCfKFb2BXMU7FzD8+TJmVZZBFNCF5I2HBtfsu5dXgMnzg6tbjNxbVbYPaerD90o8kixa9QlFsM+UK3sCuYp2LmCb1ockDGjZNhPJO5sv8AH+g1lvuyMJk+fUr6KpUluxU3swU1Legwp/2Pm/isM1thuyMU+UK3sCuYp2LnrKZUluSj9eo3a3cmkm1F8MrSwzuP5OcTk+fRrcEzTNxI52WQyJLekWgZGpjPHnlCt7ArmKdi56SJtctyjpNBmhhynWSM9DKSMbk7Ta98seNPJkhS0eWweynjhnMm1b9kXKFb2BXMU7FzukojWUiH4NB/tcDiuxoGRs3fsTyP3UzGzR1aIdTU7Vuhl7I92X9ofKFb2BXMU7FztYp1kOa2Tzfs7repb2w8oVvYFcxTsXO4qy/FFyhW9gVzFWMu9vLFb8UXKFb3pFjLvbyxW/FFyBWxXO7oVmPNGw+zYfZsfs2ptfs2v2bX7Nr9m1+za/Ztfs2v2bX7Nr9m1+za/Ztfs2v2bX7Nr9m1+za/Ztfs2v2OC37HBaT7eWK34olZpq8xReopQhCEIQhCHmDzB5n0IQhCEJeYPOdZSR0e1+zcsVv8Kvlit+ah4/eYogAAAAAAAAAAAOAtUc7dvLFb8US+XdhehMzMzM+JPE4CZmZmZmZmZnxJOsr+wCwv9vLFb8UfMFbFcgWPHYvdvLFb8UHKFbFcgXbcxTsXO3lit+KLlCtiuQLtuYp2LnbyxW/FDyhWxXIF23MU7Fzt5Yrfij5QrYrkC7bmKdi528sVvxRcoVsVyBdtzFOxc7eQKy/FHyhWxXIF23MU7FzuCsvxQ8oVsVyBdtzFOxc7irfig5QrYrkC7bmKdi52v6uoxQYVERERERVVczERE4KKensmaFGd0YUTETERETVVWzER2U67eUK2K5Au25inYuYu45exrmOHiuUK2K5Au25inYuYu72SuY4eK5QrYrkC7bmKdi5i7jl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkCx47F338rmOHiuUK2K5Au24ZYl2Zexdxy9jXMcPFcwWK5Au2CIxTI3MzMzMyMzMjMjMzOFk6pr4ttQGgv+MeA9hZ3ZmZmZmZmZmZmZmZmZmZvDnzARdFVWKWObshf83Eqqqqqqqqqqqqqqqq+PEei7x3Ri4IxLI7II95j/7lmRFmBJK0JTqPK1XqsjjQY3lu/WUfQOSAMCJ+EiLm2Vs9CKbKjH506p9jUaWkwKjtkWZtfpi2VF+ybC9E2bH6ZtfpkICSnWvpUuV4FdNRVeLUh70j0JhZXYPsjhejBqfKyZnJYb1rRxIuJnAQos1PSyMkhZFOzGktRlS0aSKyJqg8D1r6M22Q2jSRlNVMrBRvHRCzrsfmXYTV7Rsc8fxIs2igWEtRdd2aNBMzgpUXyiEp3dXj8qq7ZtFJdmxxGLSr810vJaSNLt+iwN6BSR3qIerDX3BG13mNY7spQmaLsUUtASj5ViJNrUN1QjY2adJUt3QS1fomPR/RSEGF9xZt1Lk9sr6LzDkxoQc6T3Ch6L27S4gc1Nkh51sYVBhcZooxlDURyT4I0V8k4+fdGMqNXDTQxpN7NEC+qzZK6J5tLYnCv9EYMFdhZp9NyLBn2GDoVmRECksmsiQzoqIJT+hPOxp9W0lLcIcNLmScn9SHiWwNA2OsKwZKanATlSu+uGqOo6Kq7RLJUymT+5DxKUydBNNSqrowoIzYwlNFQuKC/RGpy8hlDcIIhYBAg1gonZFsMSMyskvs8RAivWeRanVr5Efmt9hHTRaZWQzKqrJswpdm3vGICrs0UklEieK8gtxaRNC8/awtslrMZJLaS3GrSKZJyf1IbpSWTcDJJTTXpxvDQsoOeS50M0ctGRDF6HUjc6swYhSoDk5Cky2aD29nVlAnTzXoTANqxJndcriRdG30zhCp/N2JUIlDoIDeIJnquqmoVau7ZKvL0OhFQLOY2ZB6g1BYfuRw/wDuCGxiVE4OzEwMExcioJLk9i09hPv6JNS+mcByEqtPWRoZSVpDNTVnQXZ0qtOhcnv4/TlboqZMlkEcvhsxshyzpBvpkWZCaJNBZnzh+0dWgqZlBRzr6nGEUX916AkzlQWkmZutMsbn3VGTrqF6WyKxaOJQgRTnGHUNaZbrqmYoo3UetlyC8aU0hMDUepHeUjKxiBXhbWurOJAn1Yk8lZULXD1rBOKirBxnfWikz6HbU7lq60JgTUGO/DQIrDbndE/yDskXQdBlgRmJoMKUurCwdlTF2yfaaqErGKppakOWWypYlRZOhEBWQMswqq4fi5E+r6DeK1IVQBF8IVuqhYJG61XDGGuBVBtXJFNpFC7PFXAcLNJArVuLaBFxqlE9wJiUrlelG8NdNtUhZ9eUESZKiLyNC2t0hUVKFKl9iyyuR0zgHJ7oPEbBvEqZ1T+i/ms5B0k1+z4lKa9/68jajrwDldJyU/B9isFqrE1BgJzkK0+rJoeegXJ7+P05W666N0wISoLJdKD2RbsdZzyzCODT1FXbWxlkXacY43ZVNBTLVuvHFf0tepHBOMVodimCFHVTY1VDs6y4dyVRlKk9inE9QUNqg4j0K2l19+qFGKOOY0FgU98wW7EQB6m92qDdPMpi77NPoFs8d05Q/u79cj+7ucV2Fit3yieP6MEi4NY3+nYZnWz50uOS0yw9KN4a7d8x1WqglWBm5wZirbotqdXchw0484SpLFIchLIQcWlNdFPFsxJocEfIxsoCSgyzQlf7UQ0FTagTnrhmQy+CdPTI2n2bn9eAcgnhWgIatq/YhB2nQ8tSLtkWQl6FnoXJ7+P05W7kGeyrhvCZFZ6xZ9KdFxO+q5wey6OCcTo/Bdr9mRFpRITZjCdV5MRSoPifAlPIlmzauWykGcR6EbyeyP7u5xR1St0qESQiaIkIYfnO44UlScugIneEWTaq0dOc6ziNxKoQ9R/2Mm1VQ6VP5CwBRhb1aaPotEpLbLSpwnr0XKGii5WTUej4rrk0dSzBZyU0h4WeBii/NCKJDaPQS1KKIdPRIJCHeBO2mNsO16rDV2mNzrrBC0JBjAwyY4R85oj4ItEEdXZxHxl8HI9QZ0ZjpUFhDFUmLPVjtMakwujlH3bxMyMxZH3RzRGbLohfFk6oq6VZSfyKsL8OkgBcQyFkhKFQVmqy1Duha1JkBIrsI1eUH7F6Ee/dus3gshq5F2xTkSnRiSrqu0L9EFh+hUE2Pyoz0NHVijtohcCiHhmqt9+Q4CKhlDCt0U9ojOjtetao7KPTduZX0IvlLDZbl+sP1yzbGPH6l4bXbTQWVCqolk2RJzvdUlSvAxA2QStCKo+BxHQnrMqYcRbNBBcKhrrBTch9i0no3pj3BianrH/EYWdlUCM44Iux9qATYFHZJ5shiZN2+lTedA0g5RdUcyqENOpJP2FMe5QbQtAdzXHIhp0Q4jVAiGMyX+hCVOdzFORKdxR9yltQhUjZNSL2B1ToMldQHF1/QCwISoTbDoxA5WmdHIyT5bJExJJXahFtQsqy7WY20PVKLyvxFZF7V3F2t5hB4MuiMaa75in4o6DPYVlJDBZjUfPgRm2BzlFzuorogDErk0yf7ERQqF+QK2hobVGPiLcqgyAQVmFJX/DEE9nnISZ7m1v0Q1muf4SRYTG65kmyv1/4Nf8A/9oADAMBAAIAAwAAABDzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzxTz3X3zxSzz3zxv3zzzzzzzzzzzzzzzzjDUVZLzzzzzzzzzzzzzzDDRUfbzzIDDXDDbLCLCCCLzx8MMMMMMMMMMMOxbzzzzzzzzzzzzzR0jlAGZ4PzzzzzzzzzzjwyqdmYpoPx033/8A/wD/AH//AP8Afy6/PPwwwwwwwwwwww/FvPPPPPPPOIOfWsQJ39//APt/zzzzjzSRHbXRC8/f/wDrD899/wD/AP8A/wD/AP8A/wD/AOmvzz8MMMMMMMMMMMPxbzzzjjmZDIae/wD/AP8A/wD/AP8A+sPzjxO0N4or/wDP/wD/AP8A/wD6w/Pff/8A/wD/AP8A/wD/AP8A+mvzz8MMMMMMMMMMMPxbhv6BCDp9/wD/AP8A/wD/AP8A/wD/AP8ArJtugh9//wD/AP8A/wD/AP8A/wD/AP6w/Pff/wD/AP8A/wD/AP8A/wD+mvzz8MMMMMMMMMMMPxagt8+//wD/AP8A/wD/AP8A/wD/AP8A/wD/AL3j/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wCsPxT3/wD/AP8A/wD/AP8A/wD/AKa/PCiwwwwwwwwwww7VuAv/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AOsPzXX/AP8A/wD/AP8A/wD/AP8A4YfLGTrjDDDDDDDDDP8A6hb/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/APrD8GHc889t/c889/HH84xuzzzzzzzzzzz/ANqBv/8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP6w/Fu8/wDPPP8A/wD/APPPN3ywXPPPPPPPPPPPP1agL/8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8ArD8V9/8A/wD/AP8A/wD/AP8A/wCoH7z8MMMMMMMMMMMPxahL/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wDrD899/wD/AP8A/wD/AP8A/wD/AOqnzz8MMMMMMMMMMMPxagb/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/APrD8U9//wD/AP8A/wD/AP8A/wD6qfPPwwwwwwwwwwww/FqKv/8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP6g/FPf/wD/AP8A/wD/AP8A/wD+qnzz8MMMMMMMMMMMPxaz/LLDDLDLLCIz/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP6zwsMsMssAsol/Pff/AP8A/wD/AP8A/wD/AP8Aqp88/DDDDDDDDDDDD8Wss888888888sy/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A7Yc88888888sM899/wD/AP8A/wD/AP8A/wD/AOqnzz8MMMMMMMMMMMPxbzzzzzzzzzzzxf8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AOqjzzzzzzzzzzzz33//AP8A/wD/AP8A/wD/APqp88/DDDDDDDDDDDD8W888888888888X//AP8A/wD/AP8A/wD/AP8A/wD/AP8A+qnzzzzzzzzzzzxT3/8A/wD/AP8A/wD/AP8A/qp88/DDDDDDDDDDDD8W888888888888X/8A/wD/AP8A/wD/AP8A/wD/AP8A/wD+qnzzzzzzzzzzzxT3/wD/AP8A/wD/AP8A/wD/AKqfPPwwwwwwwwwwww/FvPPPPPPPPPPPF/8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AKqfPPPPPPPPPPPFPf8A/wD/AP8A/wD/AP8A/wDqp88/DDDDDDDDDDDD8W888888888888X/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wDqo88888888888899//wD/AP8A/wD/AP8A/wD6qfPNwwwwwwwwwwww/FvPPPPPPPPPPPE//wD/AP8A/wD/AP8A/wD/AP8A/wD/APqU88888888888819//AP8A/wD/AP8A/wD/AP6qfODgxAAACCQwAACfdvPPPPPPPPPPPCE8/wD/AP8A/wD/AP8A/wD/AP8A/wD+PPzzzzzzzzzzzx/7rPPPPPPPPPJKfzy4xzzzzyxzzzzz7zzzzzzzzzzzyywwwwwwwwwwwwwwwwxz3zzzzzzzzzzzwzzzzzzzzzzzzwxzzzzzzzzzzzjzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzTzzzzzzzzzzzyzTDzjDDizDTzTjgBTTTRjjDQDzhAiTjTDjSRjjjDjDjDDzDDzxxTjDjzCASAQjTzTDxzjDDTzzzzxAjhxxgDBSgABzRzwBhhQiwDzzyhSxxiRBBiQhAgAyixyxxzzyxyggzTwBgiTyBTiCTjwRDzzzzwjwixijgTiBxjRxCDyjRiwQDzzyBSSxDCgChSiChBygBwBjzzDjywCxQxASCDBzyhByyTgTzzyzzxzwzwxzyzyzxzyyzyzzzyzzzzzzzzzxzzzzzzwzzzywiTTzyyzzyzzzzzyyyxwwzwyzxxzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzz//EACYRAAEDAgUFAQEBAAAAAAAAAAABETEQMGFxkaGxIUBBUPFRIJD/2gAIAQMBAT8Q/wADlEPLgP8Agf8AA/4H/I/5H/I/4H/A/wCR/wAj/kf8Hk4Ts1ViFhQbqGz9dlDSEGriTBKKrdamW2ACnrsxbnN2CMlYnR5qPylLFVVkmba6qZKxOhLFyMMLYKpZX+NC9CuzFuclyYIh596Z0pFbInD/ANTQ2VmGEIxXFQ/OZCw6ns6F6FdmLc5P7VUSaOftjxUSQlFtTQ2n8SikwPPRFEShiqVeq3dC9CuzFuclUl1MTQTh+aPKwqqs3podGjofsqY6V0EyfstC9CuzFuclVEY/aOq9toXoV2d3mrx+4yFUGwyGToZZlmWZZlmWZZlmWZZlnQI1eP1WVHGjBgwYNGjBgwYNGCI1U/EpkamRqZGpkamRqZGpkamRqZGpkamRqZGp4eV4/VZRo0aNGjRo0aNGjRoivXb3eavH6rQudKQrsxbnJXj9VoXoV2Ytzkrx+q0L0K7MW5yet9C9CuzFucleqxbttttsZiy9kuMh7dlstlibT60hehXZi3OT1u0L0K7MW5yet2hehXZi3OT1u0L0K7MW5yet2hehXZi3OT1u0L0K7MW5yet2hehXZ3eT0G2V2F6FUAVMfpg9TB6mD1MHqYPUwepg9TB6mH1MPqYPUweo1Qt5bKopjDEmJMSYkxJiTEmJMSYkxJiTEmJMSY3QYo12AyjKMoyjKMoyjKMoyjKMoyiVb2Df7Q//xAAoEQACAAQEBgMBAQAAAAAAAAAAAREwMWEQQEFxIIGRofDxIVBRkLH/2gAIAQIBAT8Q/gcl2jchMPuQ+5D7kPuQ+5CYfch9yH3Ifch9zSTKtVWICcVHJNFWNRCxw+II/wBN5uNxuNxuNxuN5uNxuIyGPdzf8BT4wGurGuiGyg3a40praVRr1EqiGKDdVlyBTwo4R4Y8EcNeD7udKW0qYmP8oZGzrxMpSXUMRH+UM6jbdeJlHCjkfu50uNoqxqoPSQ3alakJLKXA2lUadROiGyg6hy2UcKM3Xg+7nSwbSqI6jXRDFBuq57EqSLEx/lDbqNt1NpzKOFGbrwfdzpcJDIwybKOFGbrwfdzf8hfUjKOFHI7C5/JefnIvPzkXn5yLz85F5+ci8/ORefnIvPzkXn5yLz85F5+ci8/ORGQ4iwofTso4MaKnf/8A/wDsaLxd40UbxuG4bhuG4bhuG4bhuG4a9BYUPp2UcIiCnff1Xbmg8e7m0BYUPp2UcKORu7nSwofTso4Ucj93OlhQ+nZRwo5H7udLCl9OynhRyP3c6WD0SLZbLZbLZbLJbLZbGycHkbjKC2Wy2Wy2Wy2Wy2WxygeFHI/dzpTVyOxW5tHI/dzpZ7ZW5tHI/dzpZ7ZW5tHI/dzpZ7ZW5tHI/dzpZ7ZW5tHI/dzpZ7ZW5tHI/dzpZ/rc2jNfBtVtfwvejL3oy96MvejL3oy96Mvehe9C96MvejL3oy96MXW5LacxHAtFotFotFotFotFotFotFotFotDlGM1oLEuF0ul0ul0ul0ul0ul0uCH8OCBDggQIcMP7g//xAAuEAACAQIEBAYDAQEAAwAAAAAAAREhMRBBUfFAYXGhIIGRscHwMFDh0WCAoLD/2gAIAQEAAT8Q/wDeenCSeLn9XOE/pp8M8XJOE4sdQ25ojahtc2sbcNuG3DahtY2sbWNrYU2gbQNqG0DaptQ2obWNrG1jaxto28UPiC6UnJki/TMyG2J00G2nJcMk5MTm1kL9G3zG2TNTIlToTWMZOfmCav3ufheee2+/fJ9e+T798m2/A8UV96+cdvtvm2/9Nt/6bbNv/wC4DfSvnCr6d84HfWvk5QWQ34XZj4kmigUklCHnesN+G/cIMhvk34b8N+G/Dfhvw3r4IQhvw34b8N+G/DfmMISkW9Q7Xs3YVBfoW4JJJJEkoXVkYucyz6CxWjJKC9SMP9WCCdMWjL61GTOCOMbWZI2Tg0VW0lzYxhjyUc/Dgn5iR6nl8YbKovE/UY4UqSOOhZ4wkxo5qRF+9Yb98D323ebvNb1xv/C27zd+Ft//AOm/sP7/ADT9eL+/xt/d5u/Czjnvh21RG9phOfZisvA8ErLib53OBcczIyw5nliJW3WmVNkITGdUvcR5jY5a/wDgHbdQmX6szqkN2E7kdhhnxLZYmg6rkO+qSENknNLS+w7F6EJjF0suXehAGWr+A1lnet7iCskcoMx/R+Bi8A2FfLh3h3IsungdsBkuJvndjuLing62JhSyUasjY9WozhxySRHnFXzQnVhJKrzLXuj49CCvD6/6J5WRaxdWL3x9xHYfC4d2Jb6jbyEMaEubOsfA6ScPRPuQyTOPgWnQTIXQlvDoEhUr6hOOgqOplTF+Yjt/hDFhpxg3ciy6eB2wGS4m+d4O4uEbUwVJHUTgbmk5khayr0hvRChNcl8xhkNlUQyWmLbCgZzK1cjFCsOo+RPJUHW9/BGNXmXufR0FwEkjYtRIypDNCGevohp9hoSTKgJh3opUtE3KpNCRmJ1qP3EkqJIVLkuKHNl340dqMWGh3Hi7uRZdPA7YDJcTfO8HcXATh5jIzp5ls2MtEPySJlmXREoTH8Z1SSOC7ulFEKB5jZ1SSwgXkdDMrkLqZ+Oz6XPo6C/CycFhPMkcXbgU3Lk1n0Ojfgxh7CilSacKE9GSGnDBNhai5wJ2ukWs/DXX8CH9H4GLDQ7jxd3IsungdsBkuJvneDuL8slCXkKSllzNiSs8kQyNqifmL3TlO7C2FKsll6i3z/WNJsp8CmSdV+ez6XOw+ELxsuPozmFLQ+o6raKDCQrslixUWUsE51VAgomjYSjQYDFXUEidIROs+B6wTr+ZHajFhodx4u7kWXTwO2AyXE3zvB3F48ybid5JJkb1NdtAjaBDyNJpIoQmcWzNtof3u5g/gXXud7hJLJE0KZ3K+COBs+lzsMGOzE1kTQT1RM0EJVpdStsSWmhohcXE0dQ1OxNl2IU/cucKlo9CJ1akwS6SirFC8EYR+dE/L+Biw0O48XdyLLp4HbAZLib53g7iwkpqSSkUumMqFQcoOTWtw1tivmrePYytiXbzZIpE0tqGQmsvq5FRQ0oG3RYSeYrcPb1Ir6JMGVD7Y7eIjnK4+9u5agRN0QrZtbu+5UrWhvgNjVuwyFS0YN5458QjtBiw04wbuRZdPA7YDJcTfO8HcWDvXMbhVaXUgF0IZElE2SCfmJuqN/gPjSPkoGdbdsIU6Yt4THGa8uor3aVUlP0MlH1J+ZX8lJy9BLVFdoQzntxg0mrLfNlJuKi1wnCeMR2oxYaHcOLu5Fl08DtgMlxN87wdxYOxLGEO4PUFU2NuSaQoET0FQeiJ/QdBncac3MCihe51qTyJepXw5mfGIn5fwMWD4yLuRZdPA7YDJcTfO8HcQx2O+8C6/wDFo7QYsHxg3ciy6eLZcUu6HdYMdsIt9P8AjkXgxYO/GDdyLLpihl+BUQ+JXd4sdj7HIt9P+OR2gxDFxaz+MKWXQzx/wf3hzgSf8BRr6bEyt94WUpSlKUpSlKUpSlKUpJKzAYRLmjNbCsMdj7HIt9P+OR2gxYyi+NNAnUl9Bto24bcNuG3Dbhtw24bcNuG3DZJsk1/Rm3Dbhtw24bcNuGzjZJqehNtE/wDiFui6tLFbw0giEJRqVKlSpUqVKlSpUqVKlSpUqVKlSpUdoFztghjsfY5Fvp/xyO1GIY62kVcqi4h8i9hF/A3jXCCCCCCCCCCCCCCCCCCCCCCCCajnkIY7H2ORb6f8cjTZfAxYLBcT5YZeBlHODue0+xM+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5MTaMuhCxQ5wIY7H2ORb6f8cjtBiHhE4g0tBtT2mbMzYmbMzZmbMzZmP8AjM3YbsNmZszNmZszNmZszNmZszNmZszNmZsTNiZszF/GZd94KTGRBDFafDdFNI4mE+hOvqPBjsfY5Fr/AI5HajFxu2p2T4LF4GXhWrfib53g7rBjsfY5Fvp/xyNNl8DFxo2uGWLwO2AyXE3zvB3EMdj7HIt9P+OR2gxYacWNrgli6CxdsBkuJvneDuIY7H2ORa/45E/L+Biw0O4cVa4JYugsXbAZLib53g7iGOx9jkWv+OQ/o/AxYPi4tcEsXTwO2AyXE3zvB3EMdsIt9P8AjkdoMWD4sbXBLF0Fi7YDJcTfO8HcQx2/5BsjtRiwfFxa4JYungdsBkuJvneDuIY7Hf8A/HpXwbFg+Li1wSxdBYu2AyXE3zvB3EMdjuy2v3EDXI8/z0KeBSdt8FrCw04sbXBLF0Fi7YDJcTfO8HcWDl0kd+TmAXg5ejNlZsrNlZsrNlZshshsrNlZsLNhZsLNjZsxszNkZsrNkNkNmFMS0hWlfolUScCUttITIUSMnRGyM2c2Q2Q2dmys2VmyGyD/AIjNmZsLNhZsLNhZtxszNkZsjNkYu0UpJEeDQ7jxVrgli6eB2wGS4m+d4O4sGqjVcxIggjwQR+Sso9qIeX6GaQoguUOlxwQR4Y8MeFoqLV46HceKtcEsXQWLtgMlxN87wdxcT7QVqYZ/oMheKV6HceKtcEsXQWLtgMlxN87wdxcSvaF+iZcXV6HceKtcEsXTwO2AyXE3zvB3FxPsBfomR9LnxR6cWNrgli6CxdsBkuJvneDuLifYC/RMhhdn4l8XFrgli6CxdsBkuJvneDuLifYC/RMsK7PxL4uLXBLF08DtgMlxN87wdxcT7AX6JkILs/EvixtcEsXQWLtgMlxN87wdxcT7AX6JkfS5nZ+JfFxa4JYungdsBkuJvneDuLifYC/RMsK7PxL4uLXBLF08DtgMlxN87wdxcT7AX6JkILs/EvixtcEsXQWLtgMlxN87wdxcT7AX6JkfS58Ueh3DirXBLF0Fi7YDJcTfO8HcXE+wF+iZcXR6HceKtcEsXTwO2AyXE3zvB3XFewF+iZGT0cU+nFja4JYugsXYumXE3zux3XFewFb9FlxjP3DirJnZCxeBjcg+uZpTiFLS5j4r2QhywkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkVqlzyFVXi+W8RnmZObDKk8pYvA7DNqmhdPwzazY8P282828282828282z+lf4zbjbjbjbDazazbTbzbzbl/ptw+pT9CeuZSZNxcQ3DIzSE3QQc4MBsxt5t5txtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxCs8RObM1LDG0GzIy4dxmK/InuIaZgVSQNsNsNs/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/pT+AoOZuoRXJLWLUVEsYIwJPMggj8cEfiZzD4kkigJakEYkEEEEEEEEEEEEEEEEEEEEEEEEEEEEDWjIpcjqLiFCxGnE+eEc8I/8M5/LXQf6CXp+NfqHhP58/H0/FPApYQpqSPXumIKYJxUND8c+JvBqqH3IXhcUi1xYz+H2G5zMQocvHJD8LMh1lDaUW+hCaZl11geM4yN8M2hWPMmeT9vBIYlNTQ3+lG5GpvRH0H4GXONtIl2ERt2YMdP6PkXPs+Qq1BFQ5KozI/A7UKGa6wOHZXcS5wNsJ0P2PDOMNawUvIV7VUyV8GkFraZAmWkzG1+ZZbcolcnb8EVEk6UJPPBXEvUJqQIqOvdrAghqSsRDFCttMkRDr3SibVEr6JO7wz4F2IjYjiNVKJnLsIEaBDvXgO8mT1+591oKGXjaCuPq0WJpSzD9EJtUWsxGfDhXA4lFRsr1GeYzINS0Q5LSjg2hImW4liXqJT4tC36XM4qydi+RFwZskKjacKaXkJZmZkLjVcVpLLh7TH1J1INZzfoNimjNMQKYHEkMSyajS/YcKwtlcRTEzZVCuLE3ppCKX67EBl58oaxFro6v3CbQhVtK/oJisxChjkiKtEoXGI4SKc+h1HpWVbT7h0i6UchiLjvB0mU4SaBDLpSDT9DM757eHmpVajmG9u7gP8AihbDmRlUFQLHWknMaT2xsom0FP4C6gaUXTaloJDouHE1KepI5I1xQaE9dOT7DS7qzSRPY9VaRFs8mfS+DJJt2IAXAynaLXEcqlGjQ6Ic26MZrjxAub9BighmmckxucJUFkOCTT9MOSlmUJp+6iYySXZDyu5HyZF3UHQo6sbEvRtoFiGWbsLRQnd+lz3wekkZtwXZYTHag0OK+chdazhTA8rVJBqNTJlQmZFJZKaTXn4O0Z2xlQqn7dZ0h/kZXFlm9EVv9h3e6E+X+C01Ms05QhGTpbYS8xrRjiSRVgM2GVgig+5VzEUtHKH6oXZkUmTyESK82kNNukIeDqEO5I0DuORGXyjM2iwzIdQS0aixtIpuVAXmW7NZKoSjpoobcqhZMwrqxZ4ZPZIRHi2lNLOmFRTdlK0lkujrjVe2ZlUIyKV1KRr54NymMnvCS6iFyMozQ0JDNoSKKYHMDl/2HpDoU5PsJLws05Q/wd5O7e591oKafMX2ngqqjzFaywbTeTZPDFAFYAWyXJoXJC7zlyFNkE09UxtWjqOnOysUp5cskKfxUNpekcOSWaG64mtNHSj9BgJqKG41R2L5FTl1jAaqkivMVsVLnP5QokpAYhRiNGlVQqMKRNk+3oMQrpReMkPNiGnkSKDbD6zQt2EpSXyCSkSkL0iFDaojmVQpm0skCVYabKQkCzanNcy5z57dNeiKmfFAk2JWWTpeghsaqqPXB1S61cHRCk85SSE1axrFlsm2Jb/BorL7NRFOkXrDvnt+DvsdT6vMtdMJBVm8xgdWIaJjPySjNI2hcbsdNwSTYkHodxzB4BSSc+rF4KzWVWSQbdzkWo5NWJRsmGuVZs7x7FMKNJCpLVknsKhpoRtWPlqcxYqkiXZSOn2oAhg0mTK5NCsQH5quJQ/qVgU0bumUUrEdjvHsIKrqFrDIGETNWqcuYhzq7oENIlyZc0ILFGbR9URyNaXvRIhkzfLIkzcaKkVZdku5Tq2KzarzGyNJLQPIZ2QbOaI9c0RSVBMpaIUPol4O0Z2wigHUL7StQVUnzFkxyGhm+tRpjPJMqduRuPNRS3XWs06DVOlkhYlLJm3CbFzJQ0Z9yaMWrENFOMw8pdGqj+lcNtpKMpEJaHBXhMU3ydYkaCFs7Koj+AKJ7D2IFkpUh7jRWi7DPNE1WtiV1lmnWE4H0nQ/QO+dKSoSj0TrsPTLQ5VLHRYqVK186iIeiE9edDXoOpF0bvEwRiwaIqWMqlEToHVSIUrxkTin+Q1mJb0IKUoRLk27mucXmooiO0dKfJO6K/Zgbuum8hMdYC8feTu3ufdaFjzHzi2mZUxRNLSj4BBiyK6itFH2JoQjrQ2JqqoM3zFhHIqqGtHYvk7/AOx9fmK2Nzr8o7dDkKutD5j0jSnwVahWnqVs550HTK4bVKIKGKqkOqPU9wW+q9se1fsM4ijrKWkGZB3bCm0lLdBUaCts+YiSnvMX1/ogNQxNV/0R0fWDpyL89nFI5Btu0HfPb8HfY6n1eZa6Yd4ylRAz5ipqcZBFrIjtYs43kNVMo7T8xKxeKhJSINHZJ/0V1I5cqTMs3Yd49iXk/c7admJCr3Kw7rCqTX7CFSFH1kO3K6xxI7x7Cq6EHzqJluVtLVoWh6icDZRxYgKbiESJRlScDah5XP5eBpCvKSgh5SFDRM+o08HaM7IVi/rQ++jmXIsbYrb2GbhSSXqVpzLPoVycORLk9SpJKv8AlPQQqOzDb0ZSttCwvURuuZZM76dw9z6rTD6TR4C7T7He/Y7oPpNMO6HcvcsdMfv6HZo7l7HZ/c+g0weCamDFr0hNOwcsLvFf/QUynuC0Pms3uJKuhWKn4O8ndvc+60LHR+C/faM+jzK1uIiNzbLgUTmalMyrOkl5DJJaU8FiZnO5Esx78YkeaKOj+RmVbtJLkX5POYKr2YSmsGk+NtonoP6baDcpMpo9iiV5VSE5MlmUxUSd42Og/EjXqyFCqtHE3K4q2Gj6tEqiGRIBKTmDMzO2+B25sXdvbBHWIvDusPdtubUehyLLJaKhrBQ0JYhJDw87SFH55nszvnt+DvsdT6PMtdMO4FTmSssHalx5PARKJTWBM+04zXqvjPsVlWgnpHfPY7H7nbTtz2Me++x2f3OzYkd89jtvufYaLHuDtl7IQjLCT0epHRvlNOzaHe44Yq3UXWrWJyimUd+SFX5NAadcj3BOTqfUaeDt/hnbCsXdaO5e+PfvY7mGSY3Cj7FMNqRp5HJgTUoc1V3cZ9RdVTq1HSh6zjVpXQI6WsVULxvKoxxI6y6YfaaMp6X3O0ihY5aeQ2CKdOaRIaKmQNdMFBKSS1Fqt8MlNxBiiOTBkzDGCXxu6JN0dmhG1JS37BJhWiVuaUkvg0JzkMQ7Dgs2lUeSSBDqRNRJCpbFHoffPgzvudBiconPPoIzR0qBJ+n4Kfe9RjUV2ZqGm5GWsVPyQzovIIxFSROD7CJeg/tI5KImSPYyNHFhFFWmkXprhMcjDtT4pSRpdamRtm6InmJfJp3WfJqyGbVNozzMfKYrTMQEkIaeaHVkzS5aFvLVb6IlF7dGyIHPSOTrnk1ZEqCim92QyvAMVRn1KOTDhegY5GSys0w1CSHrFRa51Td65kAkSSougpusxU6h2IrqzqoYm4Kj9wyKGDkFsSabZUGf65jjC/qDlR1QiklA80MOXYmfJ5EjhUZKDlNYIWoIo/6JZ4dXDLRJf6Lj5XRkHOeXIYHerUK9GCJ8UYbdS5MdIwoicihJ0waLS5aXKWulJPB3kkdlaS/UvpbeFzbMXmrERdky6FiUoLqqjcsS7ieUie9QtYkrV2mhL76jpCcyKUKVoOCr3kBxgS+Jl0GS6wJszRGrJlI9egyZ9bVDqU4oJVsK0ldpoh7aDpEyR7dxDTjktHUVmJKScF7MhqTehVac6inoyyuVoyQgjMg/Qbg4rSb7sePMzPRSSldPoXgiUiELnWQsJUosNUdQA2OuMm6UeQ1D5hqICyMpzDVIKNoljJ8k+R+fn9EjzFlRBqlLnJdEoNjOWoYX6ShHmivuSJJ8mPHSkjiuTdSVSpdI10VZH33U9TgYcqpKlRBQZJoxHUrfQKksGh5GUwOQKlpnA5udCkfkhK0ioh8yuvFCuIVtmU/yYz0lLUTZdcPJCKakourdKshK+ocMXwqaeXJjmzVwlJnSSdx6OKR5CYilT6seZOrHqIpkboonRkx8uzYRaQm4SLq1UVQlVOXzuHvoAbli8a5YPkZlTPw10FhDM/BD1Jeh0KiM8ameEY1w6PwIz8HTwsX4o8GpUzwyph1xqVMvHHjnwYtB6jUxEsQtIdBBRRbLb81Il0W9WW97TPfKJQ6JhLpkeiC8knSVb5chTm/Ejl4M8eRkVIeuNcPMfLDr4ankVyxrPirOCGlyHJ6/9GihJzUl0PVCvvuSrD1/+DT/AP/Z');background-size:contain;background-repeat:no-repeat;background-position:center}
.welcome h1{font-size:26px;font-weight:800;margin-bottom:8px}
.welcome p{color:rgba(255,255,255,0.65);font-size:14px}
.info-grid{display:grid;grid-template-columns:repeat(3,1fr);gap:14px;margin-bottom:36px}
.info-card{background:#fff;border-radius:var(--radius);padding:20px 22px;box-shadow:0 1px 6px rgba(0,0,0,0.06)}
.info-card .label{font-size:11px;font-weight:600;color:#94A3B8;letter-spacing:0.8px;text-transform:uppercase;margin-bottom:6px}
.info-card .val{font-size:15px;font-weight:700;color:var(--navy)}
.section-label{font-size:13px;font-weight:700;color:#64748B;letter-spacing:1.2px;text-transform:uppercase;margin-bottom:16px}
.options-grid{display:grid;grid-template-columns:1fr 1fr;gap:20px}
.option-card{background:#fff;border-radius:20px;padding:36px 32px;box-shadow:0 2px 12px rgba(0,0,0,0.07);cursor:pointer;transition:all 0.25s;border:2px solid transparent;position:relative;overflow:hidden}
.option-card:hover{transform:translateY(-4px);box-shadow:0 12px 32px rgba(0,0,0,0.12);border-color:var(--blue)}
.option-card.green:hover{border-color:var(--green)}
.option-card .opt-icon{font-size:44px;margin-bottom:18px;display:block}
.option-card h3{font-size:18px;font-weight:800;color:var(--navy);margin-bottom:8px}
.option-card p{font-size:13px;color:#64748B;line-height:1.6}
.option-card .arrow{position:absolute;bottom:24px;right:24px;width:36px;height:36px;background:var(--light);border-radius:50%;display:flex;align-items:center;justify-content:center;font-size:16px;transition:all 0.2s}
.option-card:hover .arrow{background:var(--blue);color:#fff}
.option-card.green .arrow{background:#DCFCE7}
.option-card.green:hover .arrow{background:var(--green);color:#fff}
.coming-soon-badge{background:#FEF3C7;color:#92400E;font-size:10.5px;font-weight:700;padding:3px 10px;border-radius:20px;letter-spacing:0.5px;display:inline-block;margin-bottom:12px}

/* ── INNER VIEWS ── */
.view{display:none}
.view.active{display:block}
.back-btn{display:inline-flex;align-items:center;gap:8px;color:var(--blue);font-size:13.5px;font-weight:700;cursor:pointer;margin-bottom:24px;padding:8px 16px;background:#fff;border-radius:10px;box-shadow:0 1px 6px rgba(0,0,0,0.07);transition:all 0.2s;border:none;font-family:'Sora',sans-serif}
.back-btn:hover{background:var(--light);transform:translateX(-2px)}

/* ── FORMATIVE ASSESSMENT ── */
.upload-card{background:#fff;border-radius:20px;padding:40px;box-shadow:0 2px 12px rgba(0,0,0,0.07)}
.upload-card h2{font-size:19px;font-weight:700;color:var(--navy);margin-bottom:6px}
.upload-card p{color:#64748B;font-size:13.5px;margin-bottom:30px}
.upload-grid{display:grid;grid-template-columns:1fr 1fr;gap:18px;margin-bottom:24px}
.file-box{border:2px dashed #CBD5E1;border-radius:var(--radius);padding:24px 20px;text-align:center;cursor:pointer;transition:all 0.2s;position:relative;background:#FAFBFF}
.file-box:hover,.file-box.active{border-color:var(--blue);background:#EFF6FF}
.file-box input{position:absolute;inset:0;opacity:0;cursor:pointer;width:100%}
.file-box .ico{font-size:32px;margin-bottom:10px}
.file-box .title{font-size:13px;font-weight:600;color:var(--navy);margin-bottom:4px}
.file-box .hint{font-size:11.5px;color:#94A3B8}
.file-box .fname{font-size:12px;color:#16A34A;font-weight:600;margin-top:8px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap}
.gen-btn{width:100%;padding:17px;background:linear-gradient(135deg,var(--blue),var(--navy));color:#fff;border:none;border-radius:var(--radius);font-family:'Sora',sans-serif;font-size:15px;font-weight:700;cursor:pointer;transition:all 0.25s;box-shadow:0 6px 20px rgba(26,86,219,0.3);letter-spacing:0.5px}
.gen-btn:hover{transform:translateY(-2px);box-shadow:0 10px 28px rgba(26,86,219,0.4)}
.gen-btn:disabled{opacity:0.6;cursor:not-allowed;transform:none}
#status{margin-top:18px;padding:14px 18px;border-radius:10px;font-size:13.5px;font-weight:500;display:none;text-align:center}
#status.loading{background:#EFF6FF;color:var(--blue);display:block}
#status.error{background:#FEF2F2;color:#991B1B;display:block}
#status.success{background:#ECFDF5;color:#065F46;display:block}

/* ── PROGRESS CARD ── */
.progress-placeholder{background:#fff;border-radius:20px;padding:60px 40px;box-shadow:0 2px 12px rgba(0,0,0,0.07);text-align:center}
.progress-placeholder .big-icon{font-size:72px;margin-bottom:20px}
.progress-placeholder h2{font-size:22px;font-weight:800;color:var(--navy);margin-bottom:10px}
.progress-placeholder p{color:#64748B;font-size:14px;max-width:400px;margin:0 auto}
.coming-pill{display:inline-block;margin-top:20px;background:#FEF3C7;color:#92400E;font-size:13px;font-weight:700;padding:8px 20px;border-radius:30px}

.spinner{display:inline-block;width:14px;height:14px;border:2px solid rgba(26,86,219,0.3);border-top-color:var(--blue);border-radius:50%;animation:spin 0.8s linear infinite;margin-right:8px;vertical-align:middle}
@keyframes spin{to{transform:rotate(360deg)}}
@media(max-width:600px){.info-grid{grid-template-columns:1fr}.upload-grid{grid-template-columns:1fr}.options-grid{grid-template-columns:1fr}}

/* ── CHANGE PASSWORD BOX ── */
.chpwd-box{background:#fff;border-radius:14px;padding:18px 20px;box-shadow:0 1px 6px rgba(0,0,0,0.06);border:1.5px solid #E2E8F0;margin-top:28px;max-width:340px;margin-left:auto}
.chpwd-box .chpwd-title{font-size:12px;font-weight:700;color:#64748B;letter-spacing:0.8px;text-transform:uppercase;margin-bottom:12px;display:flex;align-items:center;gap:6px}
.chpwd-box input{width:100%;padding:9px 12px;border:1.5px solid #E2E8F0;border-radius:8px;font-family:\'Sora\',sans-serif;font-size:12.5px;color:var(--navy);outline:none;background:#FAFBFF;margin-bottom:8px;transition:border 0.2s}
.chpwd-box input:focus{border-color:var(--blue);background:#fff}
.chpwd-box .chpwd-btn{width:100%;padding:9px;background:linear-gradient(135deg,var(--blue),var(--navy));color:#fff;border:none;border-radius:8px;font-family:\'Sora\',sans-serif;font-size:12.5px;font-weight:700;cursor:pointer;transition:all 0.2s;margin-top:2px}
.chpwd-box .chpwd-btn:hover{opacity:0.88;transform:translateY(-1px)}
.chpwd-box .chpwd-btn:disabled{opacity:0.6;cursor:not-allowed;transform:none}
.chpwd-msg{margin-top:8px;padding:8px 10px;border-radius:7px;font-size:11.5px;font-weight:600;display:none;text-align:center}
.chpwd-msg.success{background:#ECFDF5;color:#065F46;border:1px solid #A7F3D0;display:block}
.chpwd-msg.error{background:#FEF2F2;color:#991B1B;border:1px solid #FECACA;display:block}

/* ── IRREGULARITY REPORT ── */
.irr-upload-card{background:#fff;border-radius:20px;padding:40px;box-shadow:0 2px 12px rgba(0,0,0,0.07);margin-bottom:24px}
.irr-upload-card h2{font-size:19px;font-weight:700;color:var(--navy);margin-bottom:6px}
.irr-upload-card p{color:#64748B;font-size:13.5px;margin-bottom:24px}
.irr-upload-grid{display:grid;grid-template-columns:1fr 1fr;gap:18px;margin-bottom:20px}
.irr-trainee-table-wrap{background:#fff;border-radius:20px;padding:32px 28px;box-shadow:0 2px 12px rgba(0,0,0,0.07)}
.irr-trainee-table-wrap h3{font-size:16px;font-weight:700;color:var(--navy);margin-bottom:4px}
.irr-trainee-table-wrap .sub{font-size:12.5px;color:#64748B;margin-bottom:20px}
.irr-table{width:100%;border-collapse:collapse;font-size:12.5px}
.irr-table th{background:#0B1D3A;color:#fff;padding:10px 10px;text-align:center;font-weight:700;white-space:nowrap;font-size:11.5px;letter-spacing:0.3px}
.irr-table th:first-child{text-align:left;border-radius:8px 0 0 0}
.irr-table th:last-child{border-radius:0 8px 0 0}
.irr-table td{padding:9px 8px;border-bottom:1px solid #F1F5F9;vertical-align:middle;text-align:center}
.irr-table td:first-child{text-align:left;font-weight:600;color:#0B1D3A}
.irr-table tr:hover td{background:#EFF6FF}
.irr-table input[type=number],.irr-table input[type=date],.irr-table input[type=text]{width:100%;padding:6px 8px;border:1.5px solid #CBD5E1;border-radius:7px;font-family:'Sora',sans-serif;font-size:11.5px;text-align:center;outline:none;transition:border 0.2s}
.irr-table input:focus{border-color:var(--blue)}
.irr-table .pct-cell{font-weight:700;font-size:13px}
.irr-table .pct-ok{color:#16A34A}
.irr-table .pct-bad{color:#DC2626}
.irr-no-data{text-align:center;padding:48px 20px;color:#94A3B8;font-size:14px}
#irrStatus{margin-top:16px;padding:12px 16px;border-radius:10px;font-size:13px;font-weight:500;display:none;text-align:center}
#irrStatus.loading{background:#EFF6FF;color:var(--blue);display:block}
#irrStatus.error{background:#FEF2F2;color:#991B1B;display:block}
#irrStatus.success{background:#ECFDF5;color:#065F46;display:block}
</style>
</head>
<body>
<div class="topbar">
  <div class="logo"><span style="display:inline-flex;align-items:center;width:28px;height:28px;overflow:hidden;border-radius:4px;"><img src="data:image/jpeg;base64,/9j/4AAQSkZJRgABAQEASABIAAD/2wBDAAYEBQYFBAYGBQYHBwYIChAKCgkJChQODwwQFxQYGBcUFhYaHSUfGhsjHBYWICwgIyYnKSopGR8tMC0oMCUoKSj/2wBDAQcHBwoIChMKChMoGhYaKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCgoKCj/wgARCAKbBLADASIAAhEBAxEB/8QAHAABAAMBAAMBAAAAAAAAAAAAAAUGBwEDBAgC/8QAGgEBAAMBAQEAAAAAAAAAAAAAAAEDBAIFBv/aAAwDAQACEAMQAAAB1QAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAH45TYzzPQ0VnTjvRWdDRWdDRWdDRWdDRWdDRuZ3w0XucjRWdDRWdDRWdDRWdDRe5xa7qZ7vO+hhBIAAAAAAAAAAAAAAAAA4jvPShors3KHC8Uat2Jlu9XRPQAAAAAAAAAAAAAAA4d568RFc+pEPxn07tGvPWjo6s5+f1SZXF88dt5+hnzyl9DPnkfQz554j6HfPJP0M+eR9DPnnkPod88j6GfPKX0M+eUPoZ88pfQz55H0M+eR9D+XFdqqkOZApMXKRfzH0IUXAAAAAAAAAAActlTte/HYB73iBIAAAAAAAAAAAAcO8567n2lfh+arv48ziuM2nw1Ec57HE+k5zDnNffz3hqsxDzGn6EOrAAAAAAAAAADg6/PrI9rtfh+Kbv+Mziuc2oQ9A7zns0THuM/6/POxUc6i2X+gX+/3As1KPeKP1GT953ZWAEAAAAkEAAAkEALHtmJ7Zm7CvoCkxcpF/MfQhRcAAAAAAAAAABy11S178dgHveIEgAAAAAAABwPFFxxMqbDc59K9DLvW4z6DEVVzmko7jmjnXI464AAAHO8TqsxDTOn6EOrAAAAAABw7z1YmK7Ao8PxRpsdl/g5zX+IrDnPIehxzRzvEcdcAAADvOlsv9Av9/uBZqUe8UfqMn7zuysIAAAAAAAAAABKx7Zie2Zewr6ApMXKRfzH0IUXAAAAAAAAAAActdUte/HYB73iBIAAA50OflH67Fw3Nds5nsPzn1CJzfnGa4Q8PznP5fF3nNDvCAAAAAAAAHO8TqsxDzGn6EOrDnQ4O8erHPtK7ERTefxmcVxn1CGoPec9niY5xn7+eoq4EAdcAAAAAADvOlsv9Av9/uBZqUe8UfqMn7zuysIAAAAAAAAAABKx7Zie2Zewr6ApMXKRfzH0IUXAAAAAAAAAAActdUte/HYB73iDku856cc+6rMPzTfvDmMZzm0uHpLnPOxPgcZ+d52K3BHeAAAAAAAAAAAA50nU5ivtHvWHlGiYq0yNzDw8579D1jvOf3/R/PeaDiOO8AAAAAAAAAAAAB3nS2X+gX+/3As1KPeKP1GT953ZWEAAAAAAAAAAAlY9sxPbMvYV9AUmLlIv5j6EKLgAAAAAAAAAAOWuqWvfjsA97xAlA5lpOa0eP3hXh64HeAAAAAAAAAAAAAAAB3g7wT1wgAAAAAAAAAAAAAAAAB3nS2X+gX+/3As1KPeKP1GT953ZWEAAAAAAAAAAAlY9sxPbMvYV9AUmLk4z5j6IKLQAAAAAAAAAAOWyp2vfjsA97xAlXs00vNKPHCvCAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA7zpbL/QL/f7gWalHvFH6jJ+87srCAAAAAAAAAAASse2YntmbsKugKv6F2YddKXVx3Se3UUpdRSl1FKXUUpdRSl1FKXUUpdRSl1FKXUUpdRSZ+X5bV3vO68wSr2Z6ZmdHj9FeEAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB3nS2X+gX+/wBwLNSrWnhjHdnWc4xzaEsXbQMXbQMXbQMX7s4xdtCGMc2hLF20DF20DF20DF20IYxzaOGZaccSESBxWvRxarn2luO7mphFzUwXNTBc1MFzUwXNTBc1MFzUwXNTBc1MFzUwXNS562uW6as4Sr2aaXmdHj9FeEAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB3nS2X+gX+/wBwLNQrBZ2N9s52Njg2Njg2Njg2Njg2Njg2Njg2Njg2Njg2Njg2Njg2Njg2Njg2NmOm8T0RIFJi5SL+Z+hDPcAAAAAAAAAABy11S178dgHveIEq9mml5pR44V4QAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAHedLZf6Bf7/cCzUo94o/UZP3ndlYQAAAAAAAAAAAn9ww/cM/YU9BKkxcpF/MfQhRcAAAAAAAAAABy11S178dgHveIEq9memZpR44V4QAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAHedLZf6Bf7/cCzUo94o/UZP3ndlYQAAAAAAAAAAAntxw7cc/YVdAUmLlIv5j6EKLgAAAAAAAAAAOWuqWvfjsA97xAlXs00vM6PH6K8IAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAADvOlsv9Av9/uBZqUe8UfqMn7zuysIAAAAAAAAAAAT244duOfsKugKTFykX8x9CFFwAAAAAAAAAAHLXVLXvx2Ae94gSrua6VmtHjhXhAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAd50tl/oF/v9wLNSj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vEHJV7NdXqdPmVRbHGWprYKmtiVTWwVNbEKmtiVTWxCprYKmtiVT7a0KolYmKQjgAAAAAAAAAAAAAAB3nQWWbK0tabamtiVTWxCprYlU1sQqa2CprYlVFrQqa2Cp9tY7fq1ZbvV6O9Cj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECXOgAAAAAAAAABRqZdKZn8Pg4ygAAAAAAAAAAAAAAO86NayXW7fRkBd6oAAAAAAAAAACj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECQAAAAAAAAAAFHpd0pefww4ygAAAAAAAAAAAAAAO86NbyTW7fRkBd6oAAAAAAAAAACj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECQAAAAAAAAAAFHpd0pefww4ygAAAAAAAAAAAAAAO86NbyTW7fRkBd6oAAAAAAAAAACj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECQAAAAAAAAAAFHpd0pefww4ygAAAAAAAAAAAAAAO86NbyTW7fRkBd6oAAAAAAAAAACj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECQAAAAAAAAAAFHpd0pefww4ygAAAAAAAAAAAAAAO86NbyTW7fRkBd6oAAAAAAAAAACj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECQAAAAAAAAAAFHpd0pefww4ygAAAAAAAAAAAAAAO86NbyTW7fRkBd6oAAAAAAAAAACj3ij9Rk/ed2VhAAAAAAAAAAACe3HDtxz9hV0BSYuUi/mPoQouAAAAAAAAAAA5a6pa9+OwD3vECQAAAAAAAAAAFHpd0pmfw+DjKAAAAAAAAAAAAAAA7zo1rJdat9GRF3qgAAAAAAAAAAKPeKP1GT953ZWEAAAAAAAAAAAJ7ccO3HP2FXQFJi5yP+c9703uqbfSe7w9N7o9J7o9J7o9J7nT0nudPSe6PSe6PSe6PSe6PSe509G1wNj25JrvO+344SAAAAAAAAAAHCkUy91ejxozko4zRaUEWkxGJMRiTEYkxGJMRiTEYkxGJMRiTEYkxGJMRiTEYkxGJMRnZMRmtZxpVm/3hd6YAAAAAAAAAACj3eozGPdmO6+IZMiGTIhkyIZMiGTIhkyIZMiGTIhkyIZMiGTIhkyPLuORa5n66K+gOdIBIAAAAAAAAAABzqHOkgAAAAAAAAAAAOdAAAAAAAAAAAAAAAABzo46QCQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAHp+5FFC1D53+hz9vU9ck0bClsV6fP0A8EAWZVZ8zmxUu0l6eGMJkAhyYeHzBGQ5a1fmzy1Waxs1+Vza0lger6hKvW8R71f8AxNGYaxCTYeOIJtVPaLC8fkIbMtIyA30Dx+QZX4dWGUeLQMaLW1YU+6ej4SUeGALK8Ppkkq36LO9P3D847skAezLeLxnsorzHvqn7RYn4/ZTYa++uSxCk08EIWJVJw98ByIJhU/eJ54okm3OgELT53LzfSOJF6HrkuqciTbnQ/ECWFVZYlPF+64Rl2+cNuLH4YmaMb1qMmz802QxY+j/1XvOTTwxZNAEQS/58UWQNz+cdqLUgZw/TkMTSpyJNudAAAEVKxRhP0R87/RBR8v1Gpk9N20ZHEbf87n0TU7Dg57t7nrAVCtaoPnbUKhbyZx7YsePoADBd6wY1iLlP0ZdoudSRYXhuZH4/9AYqe7YPYupF5RteKHPb9/UijSllrhlG9YLvRWIGVqpZor1ulP3DCtUJTINfyA30AAEFi+0YufQoKHnWjZ4S8/ov7IqkXSgFg8UQKxolK/Zs/wA6fRfzsbvAz0EZBotZ3AqVW1b8mD7r897sVWsWesmxfO30T87lotc3NlGzb6Cz89615HdCgS0LuRUvSvQx6o77gx9Efv8AH7AKll+oZeb7lOrZSVi9+lp5T6LtUaZlr3zn9FGJ2j81o0Gv+tZSVrt69Y+dtXzH6IKdcPN+D503/APoEpmU/Q+Amn+5b/2QWL7Vix9CAYRu+Empen7s+fNeqZp9AlUtfvZ2VSx0q7Fhrn6Fa2jBNNLsAABFSseYJ9EYxs5SKtdoA0wH4+dvonFzXsd2UUbQsr9E2P08w9opWmQdqJLDvorOy/8AmxORNQ+fL/EmiS8fDFx8WUecteR2ufL1jO01Y9O7YjKGtYrqtDLDcK1ZRXLHCGQbTmetnzj9DUCINg/OZdKvolOvh7OQbNmxsAAAILF9vy42cFEql5rpp4MC0/y0M2Vlv7PcrUl5jT/nb6JxY1eCsMSZxtmWamPz+vyfPO65VrJUazdII0v53+iMXNSl42SFBv1PKZqtB0w+fNGulTLF6tRkSv0raMjPoD9/n9AFSy/W6Ea/lOrZ6etpdFvQj5D0j5/+isW2Q8ft5d6xpmZ+X1D3tH9GYPnP6IzD0DZPxmukHzt9A5BsR4sA+jMpNU8mOTZcsH+jKCXf2cSkTVPn68RhoM/ETZ84b/R68bTUq1pphe2VunGtMs/Joctm+jnlAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA/P6GVeDXOGW+HVemLaZYAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB//8QAMxAAAAUCBAQFBQEAAgMBAAAAAAIDBAUBBhUzNEASFCA1EBETFjEwMkFQYCEisCMkcJD/2gAIAQEAAQUC/wCg24qUHEUcRRxFHFQcVBxUHEUcRRxFHEUcRRxFHEUcRRxFHEUcRRxFHEUcRRxFHEUcRRxUHFQcZR5+f81WtKD1SD1SD1SD1SD1CD1CD1SD1CD1CD1SD1SD1SD1CD1CD1CD1CD1CD1CD1CD1CD1SD1SD1CD1SD1SD1SChy16pHz9fzqPOo86jzqPOo86jzqPOo86jzqPOo86jzqPOo86jzqPOo86jzqPOo86jzqPOo86jzqK1qInK/UHdIphSYblCs0pULyDlURla1afozrpECsq2JVWbCsm5UpAqHUV6LprWjTiOOI44zjjOOI44zjjOOI44jjiOOM44zjiOOI44jjiOOI44jjiOOI44jjjOOM44jjiOOI4gDG53pks/dRGT+jMcpQrINkwpNp0Csu4OFHKyg+fGoitHvjGoUKv26QUmkqBWZXMFHjhQVrWvRbub0XVpKbq39b0yWfuojK351kyBaVbJhWb/1STcnB1Dn+hURWj3Si6SYVlWyYVmwpJuTg6qin0bdzei6tJurf1vTJZ+6iMndVr5BV4gnRWZQKFZpWoVfOFBWta/UqIrR7WpqFCz9ulRWaSoFJlaoUeOFBWta/Vt3N6Lq0m6t/W9Mln7qIydsY5ShSSbECs2UKyzk4O4WUrsKiK0exOukQKyrYgVm6hSTcnB1VD7K3c3ourSbq39b0yWfuojJ2SjhJMKyzYgUmjVCsi5UBznPtaiK0f1KmoUKP26YVmkqBWZXMFHjhQV867a3c3ourSbq39b0yWfuojJ+rU1KBZ83SorNJUCkysYKO11BX/dzURWj+gdwkQKyzYgVm61Csm5ODqnPXdW7m9F1aTdW/remSz91EZP0TqkIFZNsSis3QKyrk9TrKn31RFaPxqalAo/bJhSaSoFZhcwVeLqiv+131u5vRdWk3Vv63pks/dRGT0qOkUwrMNyhWaUqFJByoDGqav6KPXTTaKyrYgUm6hWTcnB1FD/pLdzei6tJurf1vTJZ+6iMnomTmIz+f4+3c3ourSbq39b0yWfuojJ6JzRfx9u5vRdWk3Vv63pks/c1ERldE5ov4+3c3ourSbq39b0yWfuaiIyuic0X8fbub0XVpN1Aa3petVVFeRXHILjkFxyC45FcciuORXHIrjkVxyK45FcciuORXHIrjkVxyK45FcciuORXHIrjkVxyK45FcciuORXHILiPROin0Tmi/j7dzeieaKu2+BPhgT4YE+GBPhgT4YE+GBPhgT4YE+GBPhgT4YE+GBPhgT4YE+GBPhgT4YE+GBPhgT4YE+GBPhgT4YE+GBPhgT4RES6buf0s5ov4+3c3+FnNF/H27m/pXjtVJXEFxiC4xBcYguMQXGILjEFxiC4xBcYguMQXGILjEFxiC4xBcYguMQXGILjEFxiC4xBcYguMQXGILjEFxiC4j1jrJ9E5ov4+3c3onHirND3A+HuB8PcD4e4Hw9wPh7gfD3A+HuB8PcD4e4Hw9wPh7gfD3A+HuB8PcD4e4Hw9wPh7gfD3A+HuB8PcD4e4Hw9wPh7gfD3A+HuB8IyZduHXTJZ+5qInK6JzRfx9u5vRdOk3UFrumSz91E5PROaL+Pt3N6Lq0m6gtd0yWfuojJ6JzRfx9u5vRdWk3UFrumSz91EZPROaL+Pt3N6Lq0m6gtd0yWfuojJ6JzRfx9u5vRdWk3UFrumSz91EZPROaL+Pt3N6Lq0m6gtd0yWfuojJ6JzRfx9u5vRdWk3UFrumSz91EZPROaL+Pt3N6Lq0m6gtd0yWfuojJ6JzRfx9u5vRdWk3UFrumSz91EZPRIN6uUMEVGCKjBFRgiowRUYIqMEVGCKjBFRgiowRUYIqMEVGCKjBFRgiowRUYIqMEVGCKjBFQ9ZGafpUohRVPBFRgiowRUYIqMEVGCKjBFRgiowRUYIqMEVGCKjBFRgiowRUYKqMEVGCKjBFRgioi2B2h+i6tJuoLXdMln7qIyd1cH6T8sNLubq0m6gtd0yWfuojJ3VwfpPyw0u5urSbqC13TJZ+6iMndXB+k/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyd1cHz+j/LDS7m6tJuoLXdMln7qIyt1cHz+j/LDS7m6tJuoLXdMln7qJyd1cH6T8sNLubp0m6gtd0v0FDrcqsOVWHKrDlVhyqw5VYcqsOVWHKrDlVhyqw5VYcqsOVWHKrDlVhyqw5VYcqsOVWHKrDlVhyqw5VYcqsOVWEamZNPdTSCqw5ByOQdDkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkHI5ByOQcjkXXmzLUrfc3Egq4bYW9GFvRhb0YW9GFvRhb0YW9GFvRhb0YW9GFvRhb0YW9GFvRhb0YW9GFvRhb0YW9GFvRhb0YW9GFvRhb0YW9GFvRDx7pJ3/+LztWqLdpcS6zun+0+pc75w2Wtpwq4Z/VkppFgpHvSPUeqacLNmkTLSC7zrllDJMoKRdLP/FSvCRa4XZFvcb0VuR5QUulxWvuN6Pcb0Qcou/U+jX4dzMkR1GKqLNOu4XjlolAST1062Uq7MzbRM+s8d/VUPRMhbibmcFr5lC1alSNMyfNNDmUbqqUSTSuNuo4p/tPoVr5UVuNumukeiif15PQxvcy/a4cJNytnzdzVy9bthj7AITDJatK0NTwVVIkVScYEqWeYGqg4RXpeGfaWgVUKkRGTaLn8V5JqgdJQipA4fNm4x9gEZhiqCKEPRy8QbC5nCThzb0g1QZ4wxDddNwRy/bNjILpuCOHSDeiMyzWWWUTTTbv2aygMYpaKyjNIY+wCUyxVBDlPQTfb7b7l0cBRwFEyUtGEJ3LgKOAoXWRalbP2zmqihEqHm2JDpqFUTWfNkQadYFqScYGqi5RXp4KSLEiiZimKuumgRs/bOTOHaDalZ9h5ozLFapTUNQOl0UCNHjVxUKSjNM9FSVTXmWKNaT7CtW7tBx0OJJo3HuBgEJZkuKnLQhJRmdTouTt9s9yDp4g1DZ2i6K5kWrYY+wDeUZuPExqFovMMURSeYVqi/arDzp5PJNn6SKhKSJJdjQhJVkofzp5VkWFFCmpUr6SZ+i0OUsgWXY8KUmzVOooVMiMm0WU8V5FqgdFUixHMi0SDo5TSDWVYlbllmRjUrStK/4HEmzQHuBgG8ozXFP9+jJ6GN7mT7bs0UWu5IdtblFBW32FSSdt+inByqzRyU1DFnJcrAjdB9NLNrZakClvMDlcwK7Oso+Wd1tLQTnb7e7j43J3KC7fcz9Ro2hIvFDUgGFKOLZaHETCKsnElFoyAnmKTFeFhGzxr7YZhgyTZJXdrGsuZvHNrfdvAyt1BstPdvtnuQmYysgVG1m9Ke3GPk6tdGpEl3cO7YOaO2032+2+5dU12+E7l4XbooWQLH1IxkJlRK1kaBRpSrJK1y1MS22RQrbLQ1JCPcxCtuylXqVfiS7pG6O59BGOF0VW1vVXHt9hwyVs0InESS7F0meihLu0NnamvxL/APGRZpvphNC22SYcW2zUo/ZOIlxb0nV8gsoVFOUmXD5ZhbVVCkt9gUri2mhw6SkIgkdXjkSfZ43J2+2e5C8/mIVdno3tpEKW8wOWXgDsy23LqFV/E/ILOHkZbrf0T2+wNRxa5RGMzNWr232taJIlM9JbLOpEbcaJKmp5JOO6NKebR/b7XhbIlUe0tlnWjW32rdWY7fBdz8bm7jAdvfwTZyo4RKm8b22zUQTtpmQ6qhGjd/KupJwwtknDgLDyf2yThjpVzHOEFSrJdcnoY3uZPtuzRWgWlXXgp9jjurX/ABrNqerJQiRUmHjdqBE3NpaCc7fbvcfG5O5QXb5aOJIIYdJxiqdwvUQ3uhucNXiDoou7V2tofC7tZaTEvD4T/b7Z7kJOdbsq45JOq81PAr+bTrMulXS9saCb7fbfcuqa7fCdy8Ls0VutCunpS0KUSEggxIpcq6xivJ0wq7naCYkHi7a0NXX4ku6x2jufQWmQpn/gb7ZKnlKRuiu7Q2dqa/Ev3KHLQjDwu+n/AKdnai5zGpGWoUhn/jIFKZoxpSkoT7PG5O32z3IXmLNJT0/CRp5s4/8AyTL9k1AGcKIrS8aE7oUKGlwM1wQ1DleaZDuiWWFPscd0Z6Z9pGPcyfYJnQQXc/G5u4wHb1Mt53NlpRd6pitoR6Vkp7nqPc9R7nqJdzzy1rKHMy65PQxvcyfbdmis/UeB/scd1aaa52R0HdtyqaiPg5cJNk5t/wA85tLQTBKnYRS1GsgkoVUgcLpt05RzzT2C7equkiC1oah0kzh/Cs102ZjtJROvmS7tXataVY+F3au1+2eE/wBvtnuT+tStGlCKSDciaaQrWhaXM6TcPLX0M32+2+5dU12+E7l4XborP1XhcZzmkbdRblZeF1u06NbP1dfiS7rHaO5tBaOu8DfbJ90jdFd2hs7U1+JfuUVofC79HZ2oetyum7ls5iXkdcbdYtH7WoXlmSJZKWXkSxtKlkCfZ43J2+2e5C8xZmR4P9Ix7oX7CuUTnr/oVZt1aXHENmyNnrqmDzTI18pNHKB/scd0Z6Z9pWdeGTT/ANIJROqjKPV5SRQVIsmFlSIpzDorp9AdvU+x9/kmxrSrQXEyq7ZwT0rF0lRBUnopj0UgdwxIqShaU635DKNWMO9Tfl+24mqrprbUc5aLeBv9KtDvayDYtSoOEE3Cb62TUOQs40oVxPKBCCdOlJuGUqLdaqtGdaedJa3SuDJNZljUrmeOEIV48UlIRfmIlI6DOXhCvjUjZdlXmp4gVTnHtIiAK1OJqJLIJpx0swP6s+qI0ixG1xxrp05gG6jZh4S6J12cFFPG76tPOkrblFjokm2Io/mz0LHyr6spAKENAtlWzSVSOszhIp43fdUokdZnFRD1F94XE1WdNbajXLVfwl4dKQoWMlo9QkhNFHDOuw9t41GttRrpq4r8Pod6o/ZEMm2nm6jlnbkY6auvA3w/h3qj9iQybW42izprbUc5aL1+JGHeqvo9MyTXwuRos7bW1HOWiwVImrRzbjJWvtVLzQtpmnWRM3j4+HTM5ky/4XxnG6jhnBRLxu+FzsHDsWwyXZpeDwlTt2kO+I/On6iLq21aKFJPNhzc7UYRJyCsZHpMEa086S1uqGVQxxrRoWaWc1p/41oZ8Z+2LUiBqedJa3VarIY62KwLMqufxL2+R0ZJlMMKldTxglDvnqkpBrevDonQZibgKuFG6U2yCOOLqp0rROUgkHopGy8ebnpwoNSddhjbX/NMlEyfz9ftfRcmRyWVl2o9yOKBSffqikZJyasVGJR6f/xGtPMcBBw0/wChr//EAC8RAAAEBAYABQMFAQAAAAAAAAABAgQDMDNREBESExQxFTJAQVAgYWIFIUJSkCL/2gAIAQMBAT8B/wADobDWnVmPDvyHh35Dw78h4d9x4d9x4d9x4d+Q8O/IeHfceHfceHfceHfkI7PaTqz9KqKhPZhT2GXQSeos/RKiJT2YU9hkIcTcTqx1DUMxqGoahqGoahqGoahqBYt6ZTX1L0BnkFOIafcKfpLogp8s+gqMtXZ4EIflKbnkFOIafcKfp9gp8s+gqOtXZjPBpTLA5ycW9Mpr6lMNaU9hTuGQU/8A6kFPIiga1K7+ohC8pSVREp7MKeQ0hT/+pBTuIr3BrUrv62lMsDnJxb0ymvqUhUZCezCnyC6Cn6j6IKcxFe4MzOUQh+QvoMyIKcw0+4U/SXRBT5Z9BUdauzGctpTLA5ycW9Mpr6liaiLsKdQ0+4U/L+JBT2IfQVFWrs55BERKUlmYU9hl0FPz/iQU7iK9wa1H36BpTLA5ycW9Mpr6li+P/v4hpTLA5ycW9Mpr6li+8/xDSmWBzk4oexElpIeIRBz4g58Qc+IOfEHPiDnxBz4g58Qc+IOfEHPiCK7XFLSrF95/iGlMschpGkaRpGQyGkaRpGkZDT9CWcRRZkOBFHAijgRRwIo4EUcCKOBFHAijgRRwIo4EUcCKIrVcMtSsX3n+IaUyxzGYzGYzGYzGYzGYzGYz+iBTKa+pYvvP8Q0plgc5OLemU19Sxfef4hpTLA5ycW9Mpr6li+8/xDSmWBzk4t6ZTX1LF95/iGlMsDnJxb0ymvqWLuCtas0kOLFsOLFsOLFsOLFsOLFsOLFsOLFsOLFsOLFsOLFsFpNB5H6JENUTyjixbDixbDixbDixbDixbDixbDixbDixbDixbDixbBuk0oyPA5ycW9Mpr6lOd1D9Ew7Oac5OLemU19SnO6h+iYdnNOcnFvTKa+pTndQ/RMOzmnOTi3plNfUpzuofomHZzTnJxb0ymvqU53UP0TDs5pzk4t6ZTX1Kc7qH6Jh2c05ycW9Mpr6lOd1D9F+n9nNOcnGBFQUMv3G8i43kXG8i43kXG8i43kXG8i43kXG6i43UXG8i43kXDyIlUPIjnOoajifsQ2l2G0uw2l2G0uw2l2G0uw2l2G0uw2l2G0uw2l2G0uw2l2G0uw2l2G0uw2l2DFBpM85pzi/3R//EACgRAAADBwQDAQEBAQAAAAAAAAABAwIQERITMDEEBRUzQFFhUCEgkP/aAAgBAgEBPwH/AIHKbhI1LAcn8HJ/Byfwcn8HJ/Byfwcn8HJ/Byfwcn8HJ/ByfwIa2s1LDxYkDULw4kJyBHFxFExRFEURRFEURRFEURRFEURRDTMr9T2Hd2/t8GYhUFQxE3li9MQNQGoYib2MOYzeVy/U9h3dv7bkYA2yFQG2Yj/ssWYkQNQhUBtnYTw5jN5XL9T2Hd2/tsTEKhCoYmO4WP8AM5A1BUMTGdxPDmM3lcv1PYd3b+3/AAbZCoDUMRPwI/wG2QqCc/BTw5jN5XL9T2Hd2/tepn8hPDmM3lcv1PYd3b+16mfyE8OYzeVy9vQJtnExxyQ45IcckOOSHHJDjkhxyQ45IcckOOSHHJDjkglo2EmpmXqZ/ITw4jgKpiqYqmKpiqYqmKpiqYqmKpiqYqmKphpqZ7WuTZOBjkEhyKQ5FIcikORSHIpDkUhyKQ5FIcikORSHIpBLVsKnKy9TP5CeHEURSMUjFJoUjFIxSMUmhSaFJoUmhSMUmhSMGzK/U9h3dv7XqZ/ITw5jN5XL9T2Hd2/tepn8hPDmM3lcv1PYd3b+16mfyE8OYzeVy/U9h3dv7XqZ/IYw5jN5XL9T2Hd2/te2RxEpiUxKYlMSmJTEhiUxKYlMQ8IiiJTEpiUxKYlMSmJTEpiUxKYZL+OYzeVy/U9h3dv7bzefCTusZvK5fqew7u39t5vPhJ3WM3lcv1PYd3b+283nwk7rGbyuX6nsO7t/bebz4Sd1jN5XL9T2Hd2/tvN58JO6xm8rl+p7Du7f23m8+EndYzeVy/U9h3dv7bzefCTusZvK5fqEmzUOBCip6FFT0KKnoUVPQoqehRU9Cip6FFT0KKnoUVPQoqehRU9DQptMqf0rzZHEQMQMQMQMQMQMQMQMQMQMQMQMQMQMQMQMQMJldYyJiExCYhMQmITEJiExCYhMQmITEJiCp/3/ALo//8QAPBAAAgEBBAULAwMDAwUAAAAAAQIAAxESMTIEITNxkRATICIjNEBBUFFhQmCBMFJyBRQkYoKwcJCSocH/2gAIAQEABj8C/wCA21mYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiYiZhMRMRNX21rIEzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeMzrxmdeM1MOPSxMzGYmYmYmZjMxmYzEzEzEzEzEzEzEzMZmMzGYmYmYmYmYmYmZjMxmJ9K61RZ1bWM7NLN811LN0W3H0TruohFtp+J2dOWXru6Pzjlt/RFhImduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxmduMztxgtZj+fsXrMBNdQQ3EJnVsTdOvUM19BfH6yBOtUE6ikzqqFnWqGayT0H6K+LH2D13Ama9unZ09XzM90fE67sf0V8X13Ama9unZ0+MzXd067sf0X6K+LHrmudaoJ2YLzs0A3zXUNntNZP6q+G1kCdaoPxOzUtOooE61VrJrP6r9FfFj1jrMBM9u6dkhJ+Z1bEnXqN4JfBdZ1E1G9unZU+MzXZ1nY/nwT9FfFj1Tr1FEsBLbp2dOyZ7N067E+GX9XWZrqA7p2SFp1AFmuqfxNZJ8M/RXxY9P1mdaoJ2alp1FCzr1TNevxS/o9Z1E1G9unZU+MzXN0tZyfz4t+ivix6X1mAme9unZU+M1EKJ16jHx69DWRNdQWzqKSZ1AFnWqH8S0m3x9Tor4sej9aosNy1jOzQLvm0I3S1iT6Gt9wJZbaZ2dOzfM93dLXcn0R+ivix6MShsMtOv7Pfor4sejH7QqdFfFj0Y/aFTor4sejH7QqdFfFjpWqNUwEwmEwmEwmEwmEwmEwmEwmEwmEwmEwmEwmEwmEwlj9E/aFTohaOMyjjMo4zKvGZV4zKvGZV4zKvGZV4zKvGZV4zKvGZV4zKsyrMomVeMyrxmVeMyrxmVeMyrxmUcZlWZVmUcZlHGX6oFnox+0H+xj9oVPRrFssnlPpn0z6Z9M+mfTPpn0z6Z9M+mfTPpn0z6Z9M+mfTPpn0z6Z9M+mfTPKWv0T9oVOiHo2W/MwpzCnMKcwpzCnMKcwpzCnMKcwpzCnMKcwpzCnMKcwpzCnMKcwpzCnMKcwpzCnMKcwpzCnAlQJd+PRj9oP0V8Wvox+0KnRXxa+jH7QqdFfFr6MftB+ivi19GP2g/RXxa+jH7QqdFfFr6MftB+ivi19GP2g/RXxa+jH7QqdFfFr6NcU2GbQTaCbQTaCbQTaCbQTaCbQTaCbQTaCbQTOJnEzibQTaCbQTaCZxOs1voocONc2gmcTaCbQTOJtBNoJtBNoJtBM4m0E2gm0E2gmcTOJtBNoJtBGLMDb0V8Wvrg9EETxS+LX1weiCJ4pfFr64PRBE8Uvi19cHogieKXxa+uD0QRPFL4tfXB6IInil8Wvrg9EETxS+LX1weiCJ4pfFr64PRBE8Uvi19cHogieKXxa+uD0QRPFL4tfXB6IInil8Wvrg9EETxS+LX1weiCJ4pfFr64PRBE8Uvi16VqoTMhmzMyGZDMhmQzIZkMyGZDMhmQzZmZDMhmQzIZszMhmQzZmZDMhmQzIZkM64s8WOaQtNk02LTZNNkZsjNk02TTZNNk02TTZNNk02TTZNNk02TTZNNk02TTZNNk02TTZNNk02TTZNNk02TTZNNk02TTZNNk02TTZGDsWiBhYfFAUULGd3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3ad3abBoGqUWUf9l96gFtkFI0lstg/VUUahUQtWa8f1rlRWJ+JzlMED56d/R88VK2Q/6f0HambGirUqkr0CR5RlGjahO6zXoolgoLO6zusZa1G4B+kY6plB/bFetm/QB0bHdbLmk5P4+DNRVtMFJ6agH2/WLHATmrj222QHkYrjZCv02/tiM+YwucBOaCPbbZ+lbDSKPbbZA4wPgKu6Lvgl6s10SyhUDTtqqrNrLFrLb8y0G0cttRwo+ZZzwO6Wc9LaVRW3RYd8vVDYsuU6yluhcq1grS9TNo5O2rKs2s6tdfzLUYHdO3qBN8BouHEu1aoVp3hZeosGEu1qoUy/Sa8vvO2qKs5unUtaXqpAX5l2jUQt8cnWIG+daus2s6tdfzLUYEfHJU3RejlHCZRwlSxRF3zKOEyjhL1UqgllCqGMtqMFHzLvPAn4gdT1TO0rIPzNsJZzwnZVFbdy3HqJegKYS9Va6Jdo1QxnbVVWbWWLWFvzLVII5L1dgF+ZZo7qT8clx6yhpfvC77yxq4t+JtZ2NVW6HaVlBm1nUrrb8y9b1feXFrLe9ui0Xk7eoFltBwwna1lBm1nZ1lt5bWNgljV1t+Jtp2ddD+ZbbqlRDXW9ZhL5PVvYxf8hcIFWupJlvlLhqJelowlRDWW97QOx6t7GD/ACFlynWUtLzmxZcp1lLdC7VqhWl6mbVjJUrKG9ozKerexiA6QuEsFdbZaMOTtK62zazs662/pVd0XfByc3og67TnNOqs9Q+Us5v8w1NEcmz6TBRqkmmdVh8oCMDLqa6pwlrObvucJ2pZzLBTu7pz39PrHV9MC6QtjrDvlSJ0Ul2lqZvOGppFU2f+5ZzU7Mshl99ILJ7CDni2r2gSiTZ8znKpa34mapObpW2fMWLo2iC2sZzumVbtvvjFqc4xIjReQXapQztajsfiWWND/b1GD/Muljq8veLUHnKm6L06kXfyiO9285wnO1WKUzhbBfrMd0/t1YgWYwnSK7Numu8d86hZTL9Nzd8mEuVdoIY/8pT3QyzRR2jTnf6hWZnOuwSzm/zC+iOdX0mCm5JS2wgwMuB5Ghjn5ioW5rRlnXvOfmdnbTPxAQxs8mEu1dosLubAJzOjWqltmrzl/TKh1+QlnN275bTLUz8RlD87o5iH3MXd0Gi8lOHRtD6t7Fpe0qo1R5YEK/InO6OxZBxn9vXa8vlLZzAcqmFkWppB5xjNlZului1ip+ZzVSoahlSqS97Gc0cttkBvPxgdWe0Qj4h/nEHxKla172MFI5bbINb8YKiF7R7ypuib+ikaq5a9DSGW2yKxL2kQMC+r5lrHqqJzOjWqnsIH0xyzewmxhfQ3Ib2M5nSCSlthB8orpgf0Ku6Lvg5CbNfK26f7om6Nb7xLg6AdRraHfKkXopLpNjeRl6haR/plmk0LfxZO2RkltCoG5B0BDpDrafLleLyFB2lT2E/xdH1bpsl4S2pQvD2sltalzbQSpui9OpF38ogv4LLF1DktrNr9pd0TR/8A7LRSH/jNdIcJzelaNc+Y0Mb+Up7oZaeUw2e8p7uRoY++UwvKseNZOv0KgfCyCzC9F3dBovIkdvPlqbotn7oIa2inrftMu3GKD3Fssr6PLCxQ/MtU2iVN0/3Rd3I26H+Up7pU3Rf5QbuSpuib+ikbdG/lKe7kCDAwvzBqH4ndHndHndHnOJQZJY9ur3/Qq7ou+DkPK26H+UTdOdA6h84KFUhagw5S9ZgohZcgwhlQL7RTU1C3XAyG0HkL1WsAjOuW3VEg52oq2+5lqm0Trop/EY3BTPuJdovbY1mqKTjZBNR5Vi8rxZUK42Qc+dV7ziikAF+OS06hOyNtnnyVN0Xp1Iu/lEflYPbYMIrUQC3meXmbwLnyjQxv5Snuh6Bh3ynu5Ghjb5T3cojxqT4GXgDYDqMC6R2bzvFPjLWrqd0aloNNua82iBsbYu7oNF5Ej8tTdF/lBLi1VLe1s1ztKSH8TnqJuH9samxJQSpumv8AdF3cjbof5SnulTdFvfug3clQLjZFap5HXA9M2g8heo1iiF0yxI26Nb+6U7PbkNzMIU0hRdOrWMIGpqjA/E2a8Js14Tm2NIP7TqAWfH6FRUzGB2pWLbjBLtBbTC1endHKZfFLq3sYobGyFKqhll/Qn/BlihyvGWBCPxA/9Rrmz9sRdBodUS5XW601w1NGa658jOxD2SwIeED/ANSrG7+2D+0odnFSoLGl8VWV5/j1bw+DLGpE/iXKgKIfxOd0hr9Tk1G7UHnD/bhv9su3WHzZFGktbU84GoU7wgSsLG5WSkLWgerSurLDDU0Rrrexl1FZ0H5li6LYfeD+9rc3T8wIo0KmXHmZdrLY0dKYtYxXq0rqjpulMWsYHqUiFBx5btBbxjNXp3Ry25anvP8AFN4fEunRbx3SxuxUy8hatpBhavTujkZ1pdW3GIr4iFKItaX69O6vKYaiUrVtxiK+YS7QW8YzV6d0cjOlK1bcYiviOULQW8YzV6d0HkuVArfBlqg0901V2slrlqm+OFCpq1CAqPO2AdApRFrRXq0rq8icwl6yONIS6Tyuq4mB2pdW9bbLh1aoamiaQbfmWLa44ywUTbul7TXuj5l2nj5mWGGrohx+mXFRyIp0glKcs87JfFLq3rbYitjZCDDV0Q22+UuKr2fMVtJJWn7TXDU0c3KntOwDWfEsFM8IG/qVY3P2wDQ6PZiKlUWNyGtox63tLtJXu8YvOWokAbGX07Op7if4rlh8Sw0TbuljXkU/ic5ptW83sIFXAfcBsjVqFcvuMsr0C28TXohllHRrv4l7SbVH+qWLrfzb/olrmUcJgP8Aga//xAAtEAACAQIFAwQCAgMBAQAAAAAAAREhMUBBUWGhECDxMFBxkWCBsfBwweGwgP/aAAgBAQABPyH/AMDZ2Av2eaPNHmjzx5488eaPNHmjzR5o80eaPNHmjzR5o80eaPNHmjzR5o84eePJCS4n7anNq+0Su1bKjdweEHhB4QeEHjB4weEHjHTHhB4UeFHjR4weMd2MYxjwg8KPGjwg8IPChvDb27jpQj9nlDy55M8meVPKnkTyZ5M8meTPJnkzy55U8qeVPJnkzyZ5M8meTPKnlSD/ALhm3lt/PtDaV2TUZrKRQ6ZsJaf3DBNm2kcjttmxexScyjEIdWgY5X72ymhdg+/knbJOTkxRYnWMbxjet4xvWta1resY3rUjMISTn3LmKdi4P2NTKHdjmJmiIg+ybFkJvVD+r+HA6pZt6vrk9gqaVm7M6ew/a0diFpXk7i2H/ihzOOSPQfK9hK5inYu4/wCRfP7YsDGyXxmFbUGDlxZHfkxdbhVYllP8lsY2S7eTC9qKMfzGkLz6p5QrewK5inYuYtCVJfIzun0ciT/TDKh/sS8dshdB8v1MmFs9F0rN2SRj7DP9vYduDJxTGSLgP5eCfKFb2BXMU7FzD8+TJmVZZBFNCF5I2HBtfsu5dXgMnzg6tbjNxbVbYPaerD90o8kixa9QlFsM+UK3sCuYp2LmCb1ockDGjZNhPJO5sv8AH+g1lvuyMJk+fUr6KpUluxU3swU1Legwp/2Pm/isM1thuyMU+UK3sCuYp2LnrKZUluSj9eo3a3cmkm1F8MrSwzuP5OcTk+fRrcEzTNxI52WQyJLekWgZGpjPHnlCt7ArmKdi56SJtctyjpNBmhhynWSM9DKSMbk7Ta98seNPJkhS0eWweynjhnMm1b9kXKFb2BXMU7FzukojWUiH4NB/tcDiuxoGRs3fsTyP3UzGzR1aIdTU7Vuhl7I92X9ofKFb2BXMU7FztYp1kOa2Tzfs7repb2w8oVvYFcxTsXO4qy/FFyhW9gVzFWMu9vLFb8UXKFb3pFjLvbyxW/FFyBWxXO7oVmPNGw+zYfZsfs2ptfs2v2bX7Nr9m1+za/Ztfs2v2bX7Nr9m1+za/Ztfs2v2bX7Nr9m1+za/Ztfs2v2OC37HBaT7eWK34olZpq8xReopQhCEIQhCHmDzB5n0IQhCEJeYPOdZSR0e1+zcsVv8Kvlit+ah4/eYogAAAAAAAAAAAOAtUc7dvLFb8US+XdhehMzMzM+JPE4CZmZmZmZmZnxJOsr+wCwv9vLFb8UfMFbFcgWPHYvdvLFb8UHKFbFcgXbcxTsXO3lit+KLlCtiuQLtuYp2LnbyxW/FDyhWxXIF23MU7Fzt5Yrfij5QrYrkC7bmKdi528sVvxRcoVsVyBdtzFOxc7eQKy/FHyhWxXIF23MU7FzuCsvxQ8oVsVyBdtzFOxc7irfig5QrYrkC7bmKdi52v6uoxQYVERERERVVczERE4KKensmaFGd0YUTETERETVVWzER2U67eUK2K5Au25inYuYu45exrmOHiuUK2K5Au25inYuYu72SuY4eK5QrYrkC7bmKdi5i7jl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkC7bmKdi5jDl7GuY4eK5QrYrkCx47F338rmOHiuUK2K5Au24ZYl2Zexdxy9jXMcPFcwWK5Au2CIxTI3MzMzMyMzMjMjMzOFk6pr4ttQGgv+MeA9hZ3ZmZmZmZmZmZmZmZmZmZvDnzARdFVWKWObshf83Eqqqqqqqqqqqqqqqq+PEei7x3Ri4IxLI7II95j/7lmRFmBJK0JTqPK1XqsjjQY3lu/WUfQOSAMCJ+EiLm2Vs9CKbKjH506p9jUaWkwKjtkWZtfpi2VF+ybC9E2bH6ZtfpkICSnWvpUuV4FdNRVeLUh70j0JhZXYPsjhejBqfKyZnJYb1rRxIuJnAQos1PSyMkhZFOzGktRlS0aSKyJqg8D1r6M22Q2jSRlNVMrBRvHRCzrsfmXYTV7Rsc8fxIs2igWEtRdd2aNBMzgpUXyiEp3dXj8qq7ZtFJdmxxGLSr810vJaSNLt+iwN6BSR3qIerDX3BG13mNY7spQmaLsUUtASj5ViJNrUN1QjY2adJUt3QS1fomPR/RSEGF9xZt1Lk9sr6LzDkxoQc6T3Ch6L27S4gc1Nkh51sYVBhcZooxlDURyT4I0V8k4+fdGMqNXDTQxpN7NEC+qzZK6J5tLYnCv9EYMFdhZp9NyLBn2GDoVmRECksmsiQzoqIJT+hPOxp9W0lLcIcNLmScn9SHiWwNA2OsKwZKanATlSu+uGqOo6Kq7RLJUymT+5DxKUydBNNSqrowoIzYwlNFQuKC/RGpy8hlDcIIhYBAg1gonZFsMSMyskvs8RAivWeRanVr5Efmt9hHTRaZWQzKqrJswpdm3vGICrs0UklEieK8gtxaRNC8/awtslrMZJLaS3GrSKZJyf1IbpSWTcDJJTTXpxvDQsoOeS50M0ctGRDF6HUjc6swYhSoDk5Cky2aD29nVlAnTzXoTANqxJndcriRdG30zhCp/N2JUIlDoIDeIJnquqmoVau7ZKvL0OhFQLOY2ZB6g1BYfuRw/wDuCGxiVE4OzEwMExcioJLk9i09hPv6JNS+mcByEqtPWRoZSVpDNTVnQXZ0qtOhcnv4/TlboqZMlkEcvhsxshyzpBvpkWZCaJNBZnzh+0dWgqZlBRzr6nGEUX916AkzlQWkmZutMsbn3VGTrqF6WyKxaOJQgRTnGHUNaZbrqmYoo3UetlyC8aU0hMDUepHeUjKxiBXhbWurOJAn1Yk8lZULXD1rBOKirBxnfWikz6HbU7lq60JgTUGO/DQIrDbndE/yDskXQdBlgRmJoMKUurCwdlTF2yfaaqErGKppakOWWypYlRZOhEBWQMswqq4fi5E+r6DeK1IVQBF8IVuqhYJG61XDGGuBVBtXJFNpFC7PFXAcLNJArVuLaBFxqlE9wJiUrlelG8NdNtUhZ9eUESZKiLyNC2t0hUVKFKl9iyyuR0zgHJ7oPEbBvEqZ1T+i/ms5B0k1+z4lKa9/68jajrwDldJyU/B9isFqrE1BgJzkK0+rJoeegXJ7+P05W666N0wISoLJdKD2RbsdZzyzCODT1FXbWxlkXacY43ZVNBTLVuvHFf0tepHBOMVodimCFHVTY1VDs6y4dyVRlKk9inE9QUNqg4j0K2l19+qFGKOOY0FgU98wW7EQB6m92qDdPMpi77NPoFs8d05Q/u79cj+7ucV2Fit3yieP6MEi4NY3+nYZnWz50uOS0yw9KN4a7d8x1WqglWBm5wZirbotqdXchw0484SpLFIchLIQcWlNdFPFsxJocEfIxsoCSgyzQlf7UQ0FTagTnrhmQy+CdPTI2n2bn9eAcgnhWgIatq/YhB2nQ8tSLtkWQl6FnoXJ7+P05W7kGeyrhvCZFZ6xZ9KdFxO+q5wey6OCcTo/Bdr9mRFpRITZjCdV5MRSoPifAlPIlmzauWykGcR6EbyeyP7u5xR1St0qESQiaIkIYfnO44UlScugIneEWTaq0dOc6ziNxKoQ9R/2Mm1VQ6VP5CwBRhb1aaPotEpLbLSpwnr0XKGii5WTUej4rrk0dSzBZyU0h4WeBii/NCKJDaPQS1KKIdPRIJCHeBO2mNsO16rDV2mNzrrBC0JBjAwyY4R85oj4ItEEdXZxHxl8HI9QZ0ZjpUFhDFUmLPVjtMakwujlH3bxMyMxZH3RzRGbLohfFk6oq6VZSfyKsL8OkgBcQyFkhKFQVmqy1Duha1JkBIrsI1eUH7F6Ee/dus3gshq5F2xTkSnRiSrqu0L9EFh+hUE2Pyoz0NHVijtohcCiHhmqt9+Q4CKhlDCt0U9ojOjtetao7KPTduZX0IvlLDZbl+sP1yzbGPH6l4bXbTQWVCqolk2RJzvdUlSvAxA2QStCKo+BxHQnrMqYcRbNBBcKhrrBTch9i0no3pj3BianrH/EYWdlUCM44Iux9qATYFHZJ5shiZN2+lTedA0g5RdUcyqENOpJP2FMe5QbQtAdzXHIhp0Q4jVAiGMyX+hCVOdzFORKdxR9yltQhUjZNSL2B1ToMldQHF1/QCwISoTbDoxA5WmdHIyT5bJExJJXahFtQsqy7WY20PVKLyvxFZF7V3F2t5hB4MuiMaa75in4o6DPYVlJDBZjUfPgRm2BzlFzuorogDErk0yf7ERQqF+QK2hobVGPiLcqgyAQVmFJX/DEE9nnISZ7m1v0Q1muf4SRYTG65kmyv1/4Nf8A/9oADAMBAAIAAwAAABDzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzxTz3X3zxSzz3zxv3zzzzzzzzzzzzzzzzjDUVZLzzzzzzzzzzzzzzDDRUfbzzIDDXDDbLCLCCCLzx8MMMMMMMMMMMOxbzzzzzzzzzzzzzR0jlAGZ4PzzzzzzzzzzjwyqdmYpoPx033/8A/wD/AH//AP8Afy6/PPwwwwwwwwwwww/FvPPPPPPPOIOfWsQJ39//APt/zzzzjzSRHbXRC8/f/wDrD899/wD/AP8A/wD/AP8A/wD/AOmvzz8MMMMMMMMMMMPxbzzzjjmZDIae/wD/AP8A/wD/AP8A+sPzjxO0N4or/wDP/wD/AP8A/wD6w/Pff/8A/wD/AP8A/wD/AP8A+mvzz8MMMMMMMMMMMPxbhv6BCDp9/wD/AP8A/wD/AP8A/wD/AP8ArJtugh9//wD/AP8A/wD/AP8A/wD/AP6w/Pff/wD/AP8A/wD/AP8A/wD+mvzz8MMMMMMMMMMMPxagt8+//wD/AP8A/wD/AP8A/wD/AP8A/wD/AL3j/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wCsPxT3/wD/AP8A/wD/AP8A/wD/AKa/PCiwwwwwwwwwww7VuAv/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AOsPzXX/AP8A/wD/AP8A/wD/AP8A4YfLGTrjDDDDDDDDDP8A6hb/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/APrD8GHc889t/c889/HH84xuzzzzzzzzzzz/ANqBv/8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP6w/Fu8/wDPPP8A/wD/APPPN3ywXPPPPPPPPPPPP1agL/8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8ArD8V9/8A/wD/AP8A/wD/AP8A/wCoH7z8MMMMMMMMMMMPxahL/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wDrD899/wD/AP8A/wD/AP8A/wD/AOqnzz8MMMMMMMMMMMPxagb/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/APrD8U9//wD/AP8A/wD/AP8A/wD6qfPPwwwwwwwwwwww/FqKv/8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP6g/FPf/wD/AP8A/wD/AP8A/wD+qnzz8MMMMMMMMMMMPxaz/LLDDLDLLCIz/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP6zwsMsMssAsol/Pff/AP8A/wD/AP8A/wD/AP8Aqp88/DDDDDDDDDDDD8Wss888888888sy/wD/AP8A/wD/AP8A/wD/AP8A/wD/AP8A7Yc88888888sM899/wD/AP8A/wD/AP8A/wD/AOqnzz8MMMMMMMMMMMPxbzzzzzzzzzzzxf8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AOqjzzzzzzzzzzzz33//AP8A/wD/AP8A/wD/APqp88/DDDDDDDDDDDD8W888888888888X//AP8A/wD/AP8A/wD/AP8A/wD/AP8A+qnzzzzzzzzzzzxT3/8A/wD/AP8A/wD/AP8A/qp88/DDDDDDDDDDDD8W888888888888X/8A/wD/AP8A/wD/AP8A/wD/AP8A/wD+qnzzzzzzzzzzzxT3/wD/AP8A/wD/AP8A/wD/AKqfPPwwwwwwwwwwww/FvPPPPPPPPPPPF/8A/wD/AP8A/wD/AP8A/wD/AP8A/wD/AKqfPPPPPPPPPPPFPf8A/wD/AP8A/wD/AP8A/wDqp88/DDDDDDDDDDDD8W888888888888X/AP8A/wD/AP8A/wD/AP8A/wD/AP8A/wDqo88888888888899//wD/AP8A/wD/AP8A/wD6qfPNwwwwwwwwwwww/FvPPPPPPPPPPPE//wD/AP8A/wD/AP8A/wD/AP8A/wD/APqU88888888888819//AP8A/wD/AP8A/wD/AP6qfODgxAAACCQwAACfdvPPPPPPPPPPPCE8/wD/AP8A/wD/AP8A/wD/AP8A/wD+PPzzzzzzzzzzzx/7rPPPPPPPPPJKfzy4xzzzzyxzzzzz7zzzzzzzzzzzyywwwwwwwwwwwwwwwwxz3zzzzzzzzzzzwzzzzzzzzzzzzwxzzzzzzzzzzzjzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzTzzzzzzzzzzzyzTDzjDDizDTzTjgBTTTRjjDQDzhAiTjTDjSRjjjDjDjDDzDDzxxTjDjzCASAQjTzTDxzjDDTzzzzxAjhxxgDBSgABzRzwBhhQiwDzzyhSxxiRBBiQhAgAyixyxxzzyxyggzTwBgiTyBTiCTjwRDzzzzwjwixijgTiBxjRxCDyjRiwQDzzyBSSxDCgChSiChBygBwBjzzDjywCxQxASCDBzyhByyTgTzzyzzxzwzwxzyzyzxzyyzyzzzyzzzzzzzzzxzzzzzzwzzzywiTTzyyzzyzzzzzyyyxwwzwyzxxzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzz//EACYRAAEDAgUFAQEBAAAAAAAAAAABETEQMGFxkaGxIUBBUPFRIJD/2gAIAQMBAT8Q/wADlEPLgP8Agf8AA/4H/I/5H/I/4H/A/wCR/wAj/kf8Hk4Ts1ViFhQbqGz9dlDSEGriTBKKrdamW2ACnrsxbnN2CMlYnR5qPylLFVVkmba6qZKxOhLFyMMLYKpZX+NC9CuzFuclyYIh596Z0pFbInD/ANTQ2VmGEIxXFQ/OZCw6ns6F6FdmLc5P7VUSaOftjxUSQlFtTQ2n8SikwPPRFEShiqVeq3dC9CuzFuclUl1MTQTh+aPKwqqs3podGjofsqY6V0EyfstC9CuzFuclVEY/aOq9toXoV2d3mrx+4yFUGwyGToZZlmWZZlmWZZlmWZZlnQI1eP1WVHGjBgwYNGjBgwYNGCI1U/EpkamRqZGpkamRqZGpkamRqZGpkamRqZGp4eV4/VZRo0aNGjRo0aNGjRoivXb3eavH6rQudKQrsxbnJXj9VoXoV2Ytzkrx+q0L0K7MW5yet9C9CuzFucleqxbttttsZiy9kuMh7dlstlibT60hehXZi3OT1u0L0K7MW5yet2hehXZi3OT1u0L0K7MW5yet2hehXZi3OT1u0L0K7MW5yet2hehXZ3eT0G2V2F6FUAVMfpg9TB6mD1MHqYPUwepg9TB6mH1MPqYPUweo1Qt5bKopjDEmJMSYkxJiTEmJMSYkxJiTEmJMSY3QYo12AyjKMoyjKMoyjKMoyjKMoyiVb2Df7Q//xAAoEQACAAQEBgMBAQAAAAAAAAAAAREwMWEQQEFxIIGRofDxIVBRkLH/2gAIAQIBAT8Q/gcl2jchMPuQ+5D7kPuQ+5CYfch9yH3Ifch9zSTKtVWICcVHJNFWNRCxw+II/wBN5uNxuNxuNxuN5uNxuIyGPdzf8BT4wGurGuiGyg3a40praVRr1EqiGKDdVlyBTwo4R4Y8EcNeD7udKW0qYmP8oZGzrxMpSXUMRH+UM6jbdeJlHCjkfu50uNoqxqoPSQ3alakJLKXA2lUadROiGyg6hy2UcKM3Xg+7nSwbSqI6jXRDFBuq57EqSLEx/lDbqNt1NpzKOFGbrwfdzpcJDIwybKOFGbrwfdzf8hfUjKOFHI7C5/JefnIvPzkXn5yLz85F5+ci8/ORefnIvPzkXn5yLz85F5+ci8/ORGQ4iwofTso4MaKnf/8A/wDsaLxd40UbxuG4bhuG4bhuG4bhuG4a9BYUPp2UcIiCnff1Xbmg8e7m0BYUPp2UcKORu7nSwofTso4Ucj93OlhQ+nZRwo5H7udLCl9OynhRyP3c6WD0SLZbLZbLZbLJbLZbGycHkbjKC2Wy2Wy2Wy2Wy2WxygeFHI/dzpTVyOxW5tHI/dzpZ7ZW5tHI/dzpZ7ZW5tHI/dzpZ7ZW5tHI/dzpZ7ZW5tHI/dzpZ7ZW5tHI/dzpZ/rc2jNfBtVtfwvejL3oy96MvejL3oy96Mvehe9C96MvejL3oy96MXW5LacxHAtFotFotFotFotFotFotFotFotDlGM1oLEuF0ul0ul0ul0ul0ul0uCH8OCBDggQIcMP7g//xAAuEAACAQIEBAYDAQEAAwAAAAAAAREhMRBBUfFAYXGhIIGRscHwMFDh0WCAoLD/2gAIAQEAAT8Q/wDeenCSeLn9XOE/pp8M8XJOE4sdQ25ojahtc2sbcNuG3DahtY2sbWNrYU2gbQNqG0DaptQ2obWNrG1jaxto28UPiC6UnJki/TMyG2J00G2nJcMk5MTm1kL9G3zG2TNTIlToTWMZOfmCav3ufheee2+/fJ9e+T798m2/A8UV96+cdvtvm2/9Nt/6bbNv/wC4DfSvnCr6d84HfWvk5QWQ34XZj4kmigUklCHnesN+G/cIMhvk34b8N+G/Dfhvw3r4IQhvw34b8N+G/DfmMISkW9Q7Xs3YVBfoW4JJJJEkoXVkYucyz6CxWjJKC9SMP9WCCdMWjL61GTOCOMbWZI2Tg0VW0lzYxhjyUc/Dgn5iR6nl8YbKovE/UY4UqSOOhZ4wkxo5qRF+9Yb98D323ebvNb1xv/C27zd+Ft//AOm/sP7/ADT9eL+/xt/d5u/Czjnvh21RG9phOfZisvA8ErLib53OBcczIyw5nliJW3WmVNkITGdUvcR5jY5a/wDgHbdQmX6szqkN2E7kdhhnxLZYmg6rkO+qSENknNLS+w7F6EJjF0suXehAGWr+A1lnet7iCskcoMx/R+Bi8A2FfLh3h3IsungdsBkuJvndjuLing62JhSyUasjY9WozhxySRHnFXzQnVhJKrzLXuj49CCvD6/6J5WRaxdWL3x9xHYfC4d2Jb6jbyEMaEubOsfA6ScPRPuQyTOPgWnQTIXQlvDoEhUr6hOOgqOplTF+Yjt/hDFhpxg3ciy6eB2wGS4m+d4O4uEbUwVJHUTgbmk5khayr0hvRChNcl8xhkNlUQyWmLbCgZzK1cjFCsOo+RPJUHW9/BGNXmXufR0FwEkjYtRIypDNCGevohp9hoSTKgJh3opUtE3KpNCRmJ1qP3EkqJIVLkuKHNl340dqMWGh3Hi7uRZdPA7YDJcTfO8HcXATh5jIzp5ls2MtEPySJlmXREoTH8Z1SSOC7ulFEKB5jZ1SSwgXkdDMrkLqZ+Oz6XPo6C/CycFhPMkcXbgU3Lk1n0Ojfgxh7CilSacKE9GSGnDBNhai5wJ2ukWs/DXX8CH9H4GLDQ7jxd3IsungdsBkuJvneDuL8slCXkKSllzNiSs8kQyNqifmL3TlO7C2FKsll6i3z/WNJsp8CmSdV+ez6XOw+ELxsuPozmFLQ+o6raKDCQrslixUWUsE51VAgomjYSjQYDFXUEidIROs+B6wTr+ZHajFhodx4u7kWXTwO2AyXE3zvB3F48ybid5JJkb1NdtAjaBDyNJpIoQmcWzNtof3u5g/gXXud7hJLJE0KZ3K+COBs+lzsMGOzE1kTQT1RM0EJVpdStsSWmhohcXE0dQ1OxNl2IU/cucKlo9CJ1akwS6SirFC8EYR+dE/L+Biw0O48XdyLLp4HbAZLib53g7iwkpqSSkUumMqFQcoOTWtw1tivmrePYytiXbzZIpE0tqGQmsvq5FRQ0oG3RYSeYrcPb1Ir6JMGVD7Y7eIjnK4+9u5agRN0QrZtbu+5UrWhvgNjVuwyFS0YN5458QjtBiw04wbuRZdPA7YDJcTfO8HcWDvXMbhVaXUgF0IZElE2SCfmJuqN/gPjSPkoGdbdsIU6Yt4THGa8uor3aVUlP0MlH1J+ZX8lJy9BLVFdoQzntxg0mrLfNlJuKi1wnCeMR2oxYaHcOLu5Fl08DtgMlxN87wdxYOxLGEO4PUFU2NuSaQoET0FQeiJ/QdBncac3MCihe51qTyJepXw5mfGIn5fwMWD4yLuRZdPA7YDJcTfO8HcQx2O+8C6/wDFo7QYsHxg3ciy6eLZcUu6HdYMdsIt9P8AjkXgxYO/GDdyLLpihl+BUQ+JXd4sdj7HIt9P+OR2gxDFxaz+MKWXQzx/wf3hzgSf8BRr6bEyt94WUpSlKUpSlKUpSlKUpJKzAYRLmjNbCsMdj7HIt9P+OR2gxYyi+NNAnUl9Bto24bcNuG3Dbhtw24bcNuG3DZJsk1/Rm3Dbhtw24bcNuGzjZJqehNtE/wDiFui6tLFbw0giEJRqVKlSpUqVKlSpUqVKlSpUqVKlSpUdoFztghjsfY5Fvp/xyO1GIY62kVcqi4h8i9hF/A3jXCCCCCCCCCCCCCCCCCCCCCCCCajnkIY7H2ORb6f8cjTZfAxYLBcT5YZeBlHODue0+xM+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5M+5MTaMuhCxQ5wIY7H2ORb6f8cjtBiHhE4g0tBtT2mbMzYmbMzZmbMzZmP8AjM3YbsNmZszNmZszNmZszNmZszNmZszNmZsTNiZszF/GZd94KTGRBDFafDdFNI4mE+hOvqPBjsfY5Fr/AI5HajFxu2p2T4LF4GXhWrfib53g7rBjsfY5Fvp/xyNNl8DFxo2uGWLwO2AyXE3zvB3EMdj7HIt9P+OR2gxYacWNrgli6CxdsBkuJvneDuIY7H2ORa/45E/L+Biw0O4cVa4JYugsXbAZLib53g7iGOx9jkWv+OQ/o/AxYPi4tcEsXTwO2AyXE3zvB3EMdsIt9P8AjkdoMWD4sbXBLF0Fi7YDJcTfO8HcQx2/5BsjtRiwfFxa4JYungdsBkuJvneDuIY7Hf8A/HpXwbFg+Li1wSxdBYu2AyXE3zvB3EMdjuy2v3EDXI8/z0KeBSdt8FrCw04sbXBLF0Fi7YDJcTfO8HcWDl0kd+TmAXg5ejNlZsrNlZsrNlZshshsrNlZsLNhZsLNjZsxszNkZsrNkNkNmFMS0hWlfolUScCUttITIUSMnRGyM2c2Q2Q2dmys2VmyGyD/AIjNmZsLNhZsLNhZtxszNkZsjNkYu0UpJEeDQ7jxVrgli6eB2wGS4m+d4O4sGqjVcxIggjwQR+Sso9qIeX6GaQoguUOlxwQR4Y8MeFoqLV46HceKtcEsXQWLtgMlxN87wdxcT7QVqYZ/oMheKV6HceKtcEsXQWLtgMlxN87wdxcSvaF+iZcXV6HceKtcEsXTwO2AyXE3zvB3FxPsBfomR9LnxR6cWNrgli6CxdsBkuJvneDuLifYC/RMhhdn4l8XFrgli6CxdsBkuJvneDuLifYC/RMsK7PxL4uLXBLF08DtgMlxN87wdxcT7AX6JkILs/EvixtcEsXQWLtgMlxN87wdxcT7AX6JkfS5nZ+JfFxa4JYungdsBkuJvneDuLifYC/RMsK7PxL4uLXBLF08DtgMlxN87wdxcT7AX6JkILs/EvixtcEsXQWLtgMlxN87wdxcT7AX6JkfS58Ueh3DirXBLF0Fi7YDJcTfO8HcXE+wF+iZcXR6HceKtcEsXTwO2AyXE3zvB3XFewF+iZGT0cU+nFja4JYugsXYumXE3zux3XFewFb9FlxjP3DirJnZCxeBjcg+uZpTiFLS5j4r2QhywkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkkVqlzyFVXi+W8RnmZObDKk8pYvA7DNqmhdPwzazY8P282828282828282z+lf4zbjbjbjbDazazbTbzbzbl/ptw+pT9CeuZSZNxcQ3DIzSE3QQc4MBsxt5t5txtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxtxCs8RObM1LDG0GzIy4dxmK/InuIaZgVSQNsNsNs/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/ptn9Ns/pT+AoOZuoRXJLWLUVEsYIwJPMggj8cEfiZzD4kkigJakEYkEEEEEEEEEEEEEEEEEEEEEEEEEEEEDWjIpcjqLiFCxGnE+eEc8I/8M5/LXQf6CXp+NfqHhP58/H0/FPApYQpqSPXumIKYJxUND8c+JvBqqH3IXhcUi1xYz+H2G5zMQocvHJD8LMh1lDaUW+hCaZl11geM4yN8M2hWPMmeT9vBIYlNTQ3+lG5GpvRH0H4GXONtIl2ERt2YMdP6PkXPs+Qq1BFQ5KozI/A7UKGa6wOHZXcS5wNsJ0P2PDOMNawUvIV7VUyV8GkFraZAmWkzG1+ZZbcolcnb8EVEk6UJPPBXEvUJqQIqOvdrAghqSsRDFCttMkRDr3SibVEr6JO7wz4F2IjYjiNVKJnLsIEaBDvXgO8mT1+591oKGXjaCuPq0WJpSzD9EJtUWsxGfDhXA4lFRsr1GeYzINS0Q5LSjg2hImW4liXqJT4tC36XM4qydi+RFwZskKjacKaXkJZmZkLjVcVpLLh7TH1J1INZzfoNimjNMQKYHEkMSyajS/YcKwtlcRTEzZVCuLE3ppCKX67EBl58oaxFro6v3CbQhVtK/oJisxChjkiKtEoXGI4SKc+h1HpWVbT7h0i6UchiLjvB0mU4SaBDLpSDT9DM757eHmpVajmG9u7gP8AihbDmRlUFQLHWknMaT2xsom0FP4C6gaUXTaloJDouHE1KepI5I1xQaE9dOT7DS7qzSRPY9VaRFs8mfS+DJJt2IAXAynaLXEcqlGjQ6Ic26MZrjxAub9BighmmckxucJUFkOCTT9MOSlmUJp+6iYySXZDyu5HyZF3UHQo6sbEvRtoFiGWbsLRQnd+lz3wekkZtwXZYTHag0OK+chdazhTA8rVJBqNTJlQmZFJZKaTXn4O0Z2xlQqn7dZ0h/kZXFlm9EVv9h3e6E+X+C01Ms05QhGTpbYS8xrRjiSRVgM2GVgig+5VzEUtHKH6oXZkUmTyESK82kNNukIeDqEO5I0DuORGXyjM2iwzIdQS0aixtIpuVAXmW7NZKoSjpoobcqhZMwrqxZ4ZPZIRHi2lNLOmFRTdlK0lkujrjVe2ZlUIyKV1KRr54NymMnvCS6iFyMozQ0JDNoSKKYHMDl/2HpDoU5PsJLws05Q/wd5O7e591oKafMX2ngqqjzFaywbTeTZPDFAFYAWyXJoXJC7zlyFNkE09UxtWjqOnOysUp5cskKfxUNpekcOSWaG64mtNHSj9BgJqKG41R2L5FTl1jAaqkivMVsVLnP5QokpAYhRiNGlVQqMKRNk+3oMQrpReMkPNiGnkSKDbD6zQt2EpSXyCSkSkL0iFDaojmVQpm0skCVYabKQkCzanNcy5z57dNeiKmfFAk2JWWTpeghsaqqPXB1S61cHRCk85SSE1axrFlsm2Jb/BorL7NRFOkXrDvnt+DvsdT6vMtdMJBVm8xgdWIaJjPySjNI2hcbsdNwSTYkHodxzB4BSSc+rF4KzWVWSQbdzkWo5NWJRsmGuVZs7x7FMKNJCpLVknsKhpoRtWPlqcxYqkiXZSOn2oAhg0mTK5NCsQH5quJQ/qVgU0bumUUrEdjvHsIKrqFrDIGETNWqcuYhzq7oENIlyZc0ILFGbR9URyNaXvRIhkzfLIkzcaKkVZdku5Tq2KzarzGyNJLQPIZ2QbOaI9c0RSVBMpaIUPol4O0Z2wigHUL7StQVUnzFkxyGhm+tRpjPJMqduRuPNRS3XWs06DVOlkhYlLJm3CbFzJQ0Z9yaMWrENFOMw8pdGqj+lcNtpKMpEJaHBXhMU3ydYkaCFs7Koj+AKJ7D2IFkpUh7jRWi7DPNE1WtiV1lmnWE4H0nQ/QO+dKSoSj0TrsPTLQ5VLHRYqVK186iIeiE9edDXoOpF0bvEwRiwaIqWMqlEToHVSIUrxkTin+Q1mJb0IKUoRLk27mucXmooiO0dKfJO6K/Zgbuum8hMdYC8feTu3ufdaFjzHzi2mZUxRNLSj4BBiyK6itFH2JoQjrQ2JqqoM3zFhHIqqGtHYvk7/AOx9fmK2Nzr8o7dDkKutD5j0jSnwVahWnqVs550HTK4bVKIKGKqkOqPU9wW+q9se1fsM4ijrKWkGZB3bCm0lLdBUaCts+YiSnvMX1/ogNQxNV/0R0fWDpyL89nFI5Btu0HfPb8HfY6n1eZa6Yd4ylRAz5ipqcZBFrIjtYs43kNVMo7T8xKxeKhJSINHZJ/0V1I5cqTMs3Yd49iXk/c7admJCr3Kw7rCqTX7CFSFH1kO3K6xxI7x7Cq6EHzqJluVtLVoWh6icDZRxYgKbiESJRlScDah5XP5eBpCvKSgh5SFDRM+o08HaM7IVi/rQ++jmXIsbYrb2GbhSSXqVpzLPoVycORLk9SpJKv8AlPQQqOzDb0ZSttCwvURuuZZM76dw9z6rTD6TR4C7T7He/Y7oPpNMO6HcvcsdMfv6HZo7l7HZ/c+g0weCamDFr0hNOwcsLvFf/QUynuC0Pms3uJKuhWKn4O8ndvc+60LHR+C/faM+jzK1uIiNzbLgUTmalMyrOkl5DJJaU8FiZnO5Esx78YkeaKOj+RmVbtJLkX5POYKr2YSmsGk+NtonoP6baDcpMpo9iiV5VSE5MlmUxUSd42Og/EjXqyFCqtHE3K4q2Gj6tEqiGRIBKTmDMzO2+B25sXdvbBHWIvDusPdtubUehyLLJaKhrBQ0JYhJDw87SFH55nszvnt+DvsdT6PMtdMO4FTmSssHalx5PARKJTWBM+04zXqvjPsVlWgnpHfPY7H7nbTtz2Me++x2f3OzYkd89jtvufYaLHuDtl7IQjLCT0epHRvlNOzaHe44Yq3UXWrWJyimUd+SFX5NAadcj3BOTqfUaeDt/hnbCsXdaO5e+PfvY7mGSY3Cj7FMNqRp5HJgTUoc1V3cZ9RdVTq1HSh6zjVpXQI6WsVULxvKoxxI6y6YfaaMp6X3O0ihY5aeQ2CKdOaRIaKmQNdMFBKSS1Fqt8MlNxBiiOTBkzDGCXxu6JN0dmhG1JS37BJhWiVuaUkvg0JzkMQ7Dgs2lUeSSBDqRNRJCpbFHoffPgzvudBiconPPoIzR0qBJ+n4Kfe9RjUV2ZqGm5GWsVPyQzovIIxFSROD7CJeg/tI5KImSPYyNHFhFFWmkXprhMcjDtT4pSRpdamRtm6InmJfJp3WfJqyGbVNozzMfKYrTMQEkIaeaHVkzS5aFvLVb6IlF7dGyIHPSOTrnk1ZEqCim92QyvAMVRn1KOTDhegY5GSys0w1CSHrFRa51Td65kAkSSougpusxU6h2IrqzqoYm4Kj9wyKGDkFsSabZUGf65jjC/qDlR1QiklA80MOXYmfJ5EjhUZKDlNYIWoIo/6JZ4dXDLRJf6Lj5XRkHOeXIYHerUK9GCJ8UYbdS5MdIwoicihJ0waLS5aXKWulJPB3kkdlaS/UvpbeFzbMXmrERdky6FiUoLqqjcsS7ieUie9QtYkrV2mhL76jpCcyKUKVoOCr3kBxgS+Jl0GS6wJszRGrJlI9egyZ9bVDqU4oJVsK0ldpoh7aDpEyR7dxDTjktHUVmJKScF7MhqTehVac6inoyyuVoyQgjMg/Qbg4rSb7sePMzPRSSldPoXgiUiELnWQsJUosNUdQA2OuMm6UeQ1D5hqICyMpzDVIKNoljJ8k+R+fn9EjzFlRBqlLnJdEoNjOWoYX6ShHmivuSJJ8mPHSkjiuTdSVSpdI10VZH33U9TgYcqpKlRBQZJoxHUrfQKksGh5GUwOQKlpnA5udCkfkhK0ioh8yuvFCuIVtmU/yYz0lLUTZdcPJCKakourdKshK+ocMXwqaeXJjmzVwlJnSSdx6OKR5CYilT6seZOrHqIpkboonRkx8uzYRaQm4SLq1UVQlVOXzuHvoAbli8a5YPkZlTPw10FhDM/BD1Jeh0KiM8ameEY1w6PwIz8HTwsX4o8GpUzwyph1xqVMvHHjnwYtB6jUxEsQtIdBBRRbLb81Il0W9WW97TPfKJQ6JhLpkeiC8knSVb5chTm/Ejl4M8eRkVIeuNcPMfLDr4ankVyxrPirOCGlyHJ6/9GihJzUl0PVCvvuSrD1/+DT/AP/Z" style="width:100%;height:100%;object-fit:contain;border-radius:8px;"></span> ITI Assessment System</div>
  <div class="user-info">
    <div class="user-badge" id="uidBadge">Loading...</div>
    <button class="logout-btn" onclick="logout()">Logout</button>
  </div>
</div>
<div class="main">

  <!-- ── HOME VIEW ── -->
  <div class="view active" id="viewHome">
    <div class="welcome">
      <h1 id="welcomeName">Welcome 👋</h1>
      <p>Choose a module below to get started.</p>
    </div>
    <div class="info-grid">
      <div class="info-card"><div class="label">SI Name</div><div class="val" id="dSI">—</div></div>
      <div class="info-card"><div class="label">Trade Name</div><div class="val" id="dTrade">—</div></div>
      <div class="info-card"><div class="label">ITI Name</div><div class="val" id="dITI">—</div></div>
    </div>
    <div class="section-label">Select Module</div>
    <div class="options-grid">
      <div class="option-card" onclick="showView('viewFormative')">
        <span class="opt-icon">📊</span>
        <h3>Formative Assessment</h3>
        <p>Upload trainee details and LO syllabus to auto-generate a complete Excel assessment report with marks for all practicals.</p>
        <div class="arrow">→</div>
      </div>
      <div class="option-card" style="border-color:transparent" onclick="showView('viewIrregularity')">
        <span class="opt-icon">⚠️</span>
        <h3>Irregularity Report &amp; Warning Letter</h3>
        <p>Upload ITI letterhead and trainee details to generate attendance irregularity reports and warning letters for absent trainees.</p>
        <div class="arrow" style="background:#FEF3C7">→</div>
      </div>
    </div>

    <!-- ── CHANGE PASSWORD ── -->
    <div style="margin-top:28px;max-width:340px;margin-left:auto">
      <button id="toggleChpwdBtn" onclick="toggleChpwd()" style="width:100%;padding:9px 16px;background:#fff;border:1.5px solid #CBD5E1;border-radius:10px;font-family:\\'Sora\\',sans-serif;font-size:12.5px;font-weight:700;color:#475569;cursor:pointer;display:flex;align-items:center;justify-content:center;gap:8px;transition:all 0.2s" onmouseover="this.style.background=\\'#F8FAFC\\'" onmouseout="this.style.background=\\'#fff\\'">🔒 Change Password</button>
      <div class="chpwd-box" id="chpwdBox" style="display:none;margin-top:8px">
        <div class="chpwd-title">🔒 Change Password</div>
        <input type="password" id="cpCurrent" placeholder="Current password">
        <input type="password" id="cpNew" placeholder="New password (min 4 chars)">
        <input type="password" id="cpConfirm" placeholder="Confirm new password">
        <button class="chpwd-btn" id="cpBtn" onclick="changeUserPassword()">Update Password</button>
        <div class="chpwd-msg" id="cpMsg"></div>
      </div>
    </div>
  </div>

  <!-- ── FORMATIVE ASSESSMENT VIEW ── -->
  <div class="view" id="viewFormative">
    <button class="back-btn" onclick="showView('viewHome')">← Back to Dashboard</button>
    <div class="upload-card">
      <h2>📊 Formative Assessment Generator</h2>
      <p>Download the blank templates, fill them in, then upload both files and click Generate.</p>
      <div style="display:grid;grid-template-columns:1fr 1fr;gap:12px;margin-bottom:24px">
        <a href="/download/trainee-template" style="display:flex;align-items:center;gap:10px;background:#EFF6FF;border:2px solid #BFDBFE;border-radius:12px;padding:14px 18px;text-decoration:none;transition:all 0.2s" onmouseover="this.style.background='#DBEAFE'" onmouseout="this.style.background='#EFF6FF'">
          <span style="font-size:24px">📥</span>
          <div>
            <div style="font-size:13px;font-weight:700;color:#1B4F8A">Download Trainee Template</div>
            <div style="font-size:11px;color:#64748B">Fill trainee details + marks (out of 70)</div>
          </div>
        </a>
        <a href="/download/lo-template" style="display:flex;align-items:center;gap:10px;background:#EFF6FF;border:2px solid #BFDBFE;border-radius:12px;padding:14px 18px;text-decoration:none;transition:all 0.2s" onmouseover="this.style.background='#DBEAFE'" onmouseout="this.style.background='#EFF6FF'">
          <span style="font-size:24px">📥</span>
          <div>
            <div style="font-size:13px;font-weight:700;color:#1B4F8A">Download LO Template</div>
            <div style="font-size:11px;color:#64748B">Fill LO numbers and practical range</div>
          </div>
        </a>
      </div>
      <div style="background:#FFF7ED;border:1px solid #FDE68A;border-radius:10px;padding:12px 16px;margin-bottom:22px;font-size:12.5px;color:#92400E;font-weight:500">
        ⚠️ <strong>Important:</strong> Column <code>All_LO_average_base_on_70_SEM_I</code> in Trainee Template must be filled with marks <strong>out of 70</strong>. The system will automatically scale to 100.
      </div>
      <form id="genForm">
        <div class="upload-grid">
          <div class="file-box" id="wb1">
            <input type="file" accept=".xlsx,.xls" id="tFile">
            <div class="ico">📋</div>
            <div class="title">Trainee Details</div>
            <div class="hint">Template_Trainee_Details.xlsx</div>
            <div class="fname" id="fn1"></div>
          </div>
          <div class="file-box" id="wb2">
            <input type="file" accept=".xlsx,.xls" id="lFile">
            <div class="ico">📚</div>
            <div class="title">LO Details</div>
            <div class="hint">Template_LO_Details.xlsx</div>
            <div class="fname" id="fn2"></div>
          </div>
        </div>
        <button type="submit" class="gen-btn" id="gBtn">⚡ Generate Assessment Report</button>
      </form>
      <div id="status"></div>

      <!-- ── PDF DOWNLOAD BUTTON ── -->
      <div style="margin-top:14px;padding-top:14px;border-top:2px solid #E2E8F0">
        <div style="font-size:13px;font-weight:700;color:#7C1D1D;margin-bottom:4px">🖨️ Download Print-Ready PDF</div>
        <div style="font-size:11.5px;color:#64748B;margin-bottom:10px">All trainees in one PDF — A4 landscape, print-ready. Uses the same files uploaded above.</div>
        <button type="button" onclick="dlPDF()" id="pdfBtn" style="display:flex;align-items:center;gap:12px;width:100%;background:linear-gradient(135deg,#7C1D1D,#B91C1C);border:none;border-radius:12px;padding:14px 20px;cursor:pointer;font-family:inherit;transition:opacity 0.2s" onmouseover="this.style.opacity=\'0.88\'" onmouseout="this.style.opacity=\'1\'">
          <span style="font-size:26px">🖨️</span>
          <div style="text-align:left">
            <div style="font-size:13px;font-weight:800;color:#FFFFFF">Download PDF Report</div>
            <div style="font-size:10.5px;color:rgba(255,255,255,0.70)">All trainees — one file — A4 landscape — ready to print</div>
          </div>
        </button>
        <div id="pdfStatus" style="margin-top:8px;font-size:12px;text-align:center;min-height:18px;font-weight:500"></div>
      </div>

      <!-- ── ANNEXURE DOWNLOAD BUTTONS ── -->
      <div style="margin-top:22px;padding-top:20px;border-top:2px solid #E2E8F0">
        <div style="font-size:13px;font-weight:700;color:#1B4F8A;margin-bottom:4px">📎 Download Annexures</div>
        <div style="font-size:11.5px;color:#64748B;margin-bottom:12px">Uses the same two files above. Click any button after selecting both files.</div>
        <div style="display:grid;grid-template-columns:1fr 1fr 1fr;gap:10px">
          <div style="display:flex;flex-direction:column;gap:6px">
            <button type="button" onclick="dlAnnexure('ES')"  style="display:flex;flex-direction:column;align-items:center;gap:4px;background:#EFF6FF;border:2px solid #BFDBFE;border-radius:12px;padding:10px 6px;cursor:pointer;font-family:inherit;transition:background 0.2s" onmouseover="this.style.background='#DBEAFE'" onmouseout="this.style.background='#EFF6FF'">
              <span style="font-size:18px">📥</span>
              <span style="font-size:10px;font-weight:700;color:#1B4F8A">Annexure III ES</span>
              <span style="font-size:9px;color:#64748B">Excel</span>
            </button>
            <button type="button" onclick="dlAnnexurePDF('ES')" style="display:flex;flex-direction:column;align-items:center;gap:4px;background:#FEF2F2;border:2px solid #FECACA;border-radius:12px;padding:10px 6px;cursor:pointer;font-family:inherit;transition:background 0.2s" onmouseover="this.style.background='#FEE2E2'" onmouseout="this.style.background='#FEF2F2'">
              <span style="font-size:18px">🖨️</span>
              <span style="font-size:10px;font-weight:700;color:#991B1B">Annexure III ES</span>
              <span style="font-size:9px;color:#64748B">PDF (Print-Ready)</span>
            </button>
          </div>
          <div style="display:flex;flex-direction:column;gap:6px">
            <button type="button" onclick="dlAnnexure('WCS')" style="display:flex;flex-direction:column;align-items:center;gap:4px;background:#EFF6FF;border:2px solid #BFDBFE;border-radius:12px;padding:10px 6px;cursor:pointer;font-family:inherit;transition:background 0.2s" onmouseover="this.style.background='#DBEAFE'" onmouseout="this.style.background='#EFF6FF'">
              <span style="font-size:18px">📥</span>
              <span style="font-size:10px;font-weight:700;color:#1B4F8A">Annexure III WCS</span>
              <span style="font-size:9px;color:#64748B">Excel</span>
            </button>
            <button type="button" onclick="dlAnnexurePDF('WCS')" style="display:flex;flex-direction:column;align-items:center;gap:4px;background:#FEF2F2;border:2px solid #FECACA;border-radius:12px;padding:10px 6px;cursor:pointer;font-family:inherit;transition:background 0.2s" onmouseover="this.style.background='#FEE2E2'" onmouseout="this.style.background='#FEF2F2'">
              <span style="font-size:18px">🖨️</span>
              <span style="font-size:10px;font-weight:700;color:#991B1B">Annexure III WCS</span>
              <span style="font-size:9px;color:#64748B">PDF (Print-Ready)</span>
            </button>
          </div>
          <div style="display:flex;flex-direction:column;gap:6px">
            <button type="button" onclick="dlAnnexure('ED')"  style="display:flex;flex-direction:column;align-items:center;gap:4px;background:#EFF6FF;border:2px solid #BFDBFE;border-radius:12px;padding:10px 6px;cursor:pointer;font-family:inherit;transition:background 0.2s" onmouseover="this.style.background='#DBEAFE'" onmouseout="this.style.background='#EFF6FF'">
              <span style="font-size:18px">📥</span>
              <span style="font-size:10px;font-weight:700;color:#1B4F8A">Annexure III ED</span>
              <span style="font-size:9px;color:#64748B">Excel</span>
            </button>
            <button type="button" onclick="dlAnnexurePDF('ED')" style="display:flex;flex-direction:column;align-items:center;gap:4px;background:#FEF2F2;border:2px solid #FECACA;border-radius:12px;padding:10px 6px;cursor:pointer;font-family:inherit;transition:background 0.2s" onmouseover="this.style.background='#FEE2E2'" onmouseout="this.style.background='#FEF2F2'">
              <span style="font-size:18px">🖨️</span>
              <span style="font-size:10px;font-weight:700;color:#991B1B">Annexure III ED</span>
              <span style="font-size:9px;color:#64748B">PDF (Print-Ready)</span>
            </button>
          </div>
        </div>
        <div id="annStatus" style="margin-top:8px;font-size:12px;text-align:center;min-height:18px;font-weight:500"></div>
      </div>

      <!-- ── LO-WISE SUMMARY DOWNLOAD ── -->
      <div style="margin-top:18px;padding-top:18px;border-top:2px solid #E2E8F0">
        <div style="font-size:13px;font-weight:700;color:#1B4F8A;margin-bottom:4px">📊 Download LO-Wise Summary (Annexure II)</div>
        <div style="font-size:11.5px;color:#64748B;margin-bottom:12px">Requires both Trainee &amp; LO files. Generates one sheet per LO with all trainees and their category-wise marks.</div>
        <div style="display:grid;grid-template-columns:1fr 1fr;gap:10px">
          <button type="button" onclick="dlLOSummary()" style="display:flex;align-items:center;gap:10px;background:linear-gradient(135deg,#0B1D3A,#1B4F8A);border:none;border-radius:12px;padding:12px 16px;cursor:pointer;font-family:inherit;transition:opacity 0.2s" onmouseover="this.style.opacity=\'0.88\'" onmouseout="this.style.opacity=\'1\'">
            <span style="font-size:22px">📥</span>
            <div style="text-align:left">
              <div style="font-size:12px;font-weight:800;color:#FFFFFF">LO Wise Summary</div>
              <div style="font-size:9.5px;color:rgba(255,255,255,0.65)">Excel — one sheet per LO</div>
            </div>
          </button>
          <button type="button" onclick="dlLOSummaryPDF()" style="display:flex;align-items:center;gap:10px;background:linear-gradient(135deg,#7C1D1D,#B91C1C);border:none;border-radius:12px;padding:12px 16px;cursor:pointer;font-family:inherit;transition:opacity 0.2s" onmouseover="this.style.opacity=\'0.88\'" onmouseout="this.style.opacity=\'1\'">
            <span style="font-size:22px">🖨️</span>
            <div style="text-align:left">
              <div style="font-size:12px;font-weight:800;color:#FFFFFF">LO Wise Summary</div>
              <div style="font-size:9.5px;color:rgba(255,255,255,0.70)">PDF — print-ready</div>
            </div>
          </button>
        </div>
        <div id="loSumStatus" style="margin-top:8px;font-size:12px;text-align:center;min-height:18px;font-weight:500"></div>
      </div>

      <!-- ── PER TRAINEE PER LO REPORT ── -->
      <div style="margin-top:18px;padding-top:18px;border-top:2px solid #E2E8F0">
        <div style="font-size:13px;font-weight:700;color:#1B4F8A;margin-bottom:4px">📋 Download Per Trainee Per LO Assessment</div>
        <div style="font-size:11.5px;color:#64748B;margin-bottom:12px">Requires both Trainee &amp; LO files. Generates one sheet per trainee per LO with complete practical-wise marks matching the official proforma.</div>
        <div style="display:grid;grid-template-columns:1fr 1fr;gap:10px">
          <button type="button" onclick="dlTraineeLO()" style="display:flex;align-items:center;gap:10px;background:linear-gradient(135deg,#064E3B,#065F46);border:none;border-radius:12px;padding:12px 16px;cursor:pointer;font-family:inherit;transition:opacity 0.2s" onmouseover="this.style.opacity=\'0.88\'" onmouseout="this.style.opacity=\'1\'">
            <span style="font-size:22px">📋</span>
            <div style="text-align:left">
              <div style="font-size:12px;font-weight:800;color:#FFFFFF">Per Trainee Per LO</div>
              <div style="font-size:9.5px;color:rgba(255,255,255,0.65)">Excel — official proforma</div>
            </div>
          </button>
          <button type="button" onclick="dlTraineeLOPDF()" style="display:flex;align-items:center;gap:10px;background:linear-gradient(135deg,#7C1D1D,#B91C1C);border:none;border-radius:12px;padding:12px 16px;cursor:pointer;font-family:inherit;transition:opacity 0.2s" onmouseover="this.style.opacity=\'0.88\'" onmouseout="this.style.opacity=\'1\'">
            <span style="font-size:22px">🖨️</span>
            <div style="text-align:left">
              <div style="font-size:12px;font-weight:800;color:#FFFFFF">Per Trainee Per LO</div>
              <div style="font-size:9.5px;color:rgba(255,255,255,0.70)">PDF — print-ready</div>
            </div>
          </button>
        </div>
        <div id="traineeLOStatus" style="margin-top:8px;font-size:12px;text-align:center;min-height:18px;font-weight:500"></div>
      </div>

      <!-- ── ALL TRAINEE ALL LO SUMMARY ── -->
      <div style="margin-top:18px;padding-top:18px;border-top:2px solid #E2E8F0">
        <div style="font-size:13px;font-weight:700;color:#1B4F8A;margin-bottom:4px">📊 Download All Trainee All LO Summary</div>
        <div style="font-size:11.5px;color:#64748B;margin-bottom:12px">Generates a summary showing all trainees and all LO scores — one row per trainee with columns for each LO (as per uploaded LO file) and average. Requires both Trainee &amp; LO files.</div>
        <div style="display:grid;grid-template-columns:1fr 1fr;gap:10px">
          <button type="button" onclick="dlTraineeSummary()" style="display:flex;align-items:center;gap:10px;background:linear-gradient(135deg,#1B4F8A,#2E75B6);border:none;border-radius:12px;padding:12px 16px;cursor:pointer;font-family:inherit;transition:opacity 0.2s" onmouseover="this.style.opacity=\'0.88\'" onmouseout="this.style.opacity=\'1\'">
            <span style="font-size:22px">📊</span>
            <div style="text-align:left">
              <div style="font-size:12px;font-weight:800;color:#FFFFFF">Trainee Wise Summary</div>
              <div style="font-size:9.5px;color:rgba(255,255,255,0.65)">Excel — LO scores per trainee</div>
            </div>
          </button>
          <button type="button" onclick="dlTraineeSummaryPDF()" style="display:flex;align-items:center;gap:10px;background:linear-gradient(135deg,#7C1D1D,#B91C1C);border:none;border-radius:12px;padding:12px 16px;cursor:pointer;font-family:inherit;transition:opacity 0.2s" onmouseover="this.style.opacity=\'0.88\'" onmouseout="this.style.opacity=\'1\'">
            <span style="font-size:22px">🖨️</span>
            <div style="text-align:left">
              <div style="font-size:12px;font-weight:800;color:#FFFFFF">Trainee Wise Summary</div>
              <div style="font-size:9.5px;color:rgba(255,255,255,0.70)">PDF — print-ready</div>
            </div>
          </button>
        </div>
        <div id="traineeSumStatus" style="margin-top:8px;font-size:12px;text-align:center;min-height:18px;font-weight:500"></div>
      </div>

      <!-- ── PROGRESS CARD DOWNLOAD ── -->
      <div style="margin-top:18px;padding-top:18px;border-top:2px solid #E2E8F0">
        <div style="font-size:13px;font-weight:700;color:#065F46;margin-bottom:4px">📈 Download Progress Cards</div>
        <div style="font-size:11.5px;color:#64748B;margin-bottom:12px">Generates official ITI Progress Cards for all trainees — one sheet per trainee. Uses the Trainee Details file already uploaded above.</div>
        <div style="display:grid;grid-template-columns:1fr 1fr;gap:10px">
          <button type="button" onclick="dlProgressCard()" style="display:flex;align-items:center;gap:10px;background:linear-gradient(135deg,#065F46,#047857);border:none;border-radius:12px;padding:12px 16px;cursor:pointer;font-family:inherit;transition:opacity 0.2s" onmouseover="this.style.opacity=\'0.88\'" onmouseout="this.style.opacity=\'1\'">
            <span style="font-size:22px">📥</span>
            <div style="text-align:left">
              <div style="font-size:12px;font-weight:800;color:#FFFFFF">Progress Cards</div>
              <div style="font-size:9.5px;color:rgba(255,255,255,0.65)">Excel — one sheet per trainee</div>
            </div>
          </button>
          <button type="button" onclick="dlProgressCardPDF()" style="display:flex;align-items:center;gap:10px;background:linear-gradient(135deg,#7C1D1D,#B91C1C);border:none;border-radius:12px;padding:12px 16px;cursor:pointer;font-family:inherit;transition:opacity 0.2s" onmouseover="this.style.opacity=\'0.88\'" onmouseout="this.style.opacity=\'1\'">
            <span style="font-size:22px">🖨️</span>
            <div style="text-align:left">
              <div style="font-size:12px;font-weight:800;color:#FFFFFF">Progress Cards</div>
              <div style="font-size:9.5px;color:rgba(255,255,255,0.70)">PDF — print-ready</div>
            </div>
          </button>
        </div>
        <div id="pcStatus" style="margin-top:8px;font-size:12px;text-align:center;min-height:18px;font-weight:500"></div>
      </div>
    </div>
  </div>

  <!-- ── IRREGULARITY REPORT & WARNING LETTER VIEW ── -->
  <div class="view" id="viewIrregularity">
    <button class="back-btn" onclick="showView(\'viewHome\')">← Back to Dashboard</button>

    <div class="irr-upload-card">
      <h2>⚠️ Irregularity Report &amp; Warning Letter</h2>
      <p>Upload the ITI letterhead (<strong>.docx only</strong>, with letterhead in the Header section) and trainee details (Excel) to begin. Trainee names will appear below where you can fill attendance data.</p>
      <a href="/download/irr-trainee-template" style="display:inline-flex;align-items:center;gap:8px;background:#F0FDF4;border:2px solid #86EFAC;border-radius:10px;padding:10px 18px;text-decoration:none;color:#16A34A;font-weight:700;font-size:13px;margin-bottom:18px;transition:all 0.2s" onmouseover="this.style.background='#DCFCE7'" onmouseout="this.style.background='#F0FDF4'">
        ⬇️ &nbsp;Download Blank Trainee Proforma
      </a>
      <div class="irr-upload-grid">
        <div class="file-box" id="irrLhBox">
          <input type="file" accept=".docx" id="irrLhFile" onchange="irrFileChange(\'irrLhFile\',\'irrLhBox\',\'irrLhName\')">
          <div class="ico">📄</div>
          <div class="title">ITI Letter Head</div>
          <div class="hint">Upload .docx letterhead only</div>
          <div style="margin-top:6px;padding:6px 10px;background:#FFF7ED;border:1px solid #F97316;border-radius:6px;font-size:11px;color:#92400E;line-height:1.5;text-align:left;">
            ⚠️ <strong>Important:</strong> Letterhead must be in <strong>.docx format only</strong>. The letterhead content should be placed in the <strong>Header section</strong> of the Word document (Insert → Header).
          </div>
          <div class="fname" id="irrLhName"></div>
        </div>
        <div class="file-box" id="irrTrBox">
          <input type="file" accept=".xlsx,.xls" id="irrTrFile" onchange="irrFileChange(\'irrTrFile\',\'irrTrBox\',\'irrTrName\');loadIrrTrainees()">
          <div class="ico">📋</div>
          <div class="title">Trainee Details</div>
          <div class="hint">Same trainee Excel used in assessment</div>
          <div class="fname" id="irrTrName"></div>
        </div>
      </div>
      <div id="irrStatus"></div>
    </div>

    <!-- Trainee attendance table -->
    <div class="irr-trainee-table-wrap" id="irrTableSection" style="display:none">
      <h3>📋 Trainee Attendance Details</h3>
      <div class="sub">Fill attendance data for each trainee. Percentage is auto-calculated from Working Days and Present Days.</div>
      <div style="overflow-x:auto">
        <table class="irr-table" id="irrTable">
          <thead>
            <tr>
              <th style="min-width:180px">Trainee Name</th>
              <th style="min-width:110px">Working Days</th>
              <th style="min-width:110px">Present Days</th>
              <th style="min-width:80px">% Attendance</th>
              <th style="min-width:150px">Absent From Date</th>
              <th style="min-width:150px">Attendance Upto Date</th>
              <th style="min-width:160px">Previous Informed Date</th>
            </tr>
          </thead>
          <tbody id="irrTableBody">
          </tbody>
        </table>
      </div>
      <div style="margin-top:24px;display:flex;gap:12px;flex-wrap:wrap">
        <button onclick="generateIrrDocReport()" style="flex:1;min-width:200px;padding:14px 20px;background:linear-gradient(135deg,#1B4F8A,#2563EB);color:#fff;border:none;border-radius:var(--radius);font-family:\'Sora\',sans-serif;font-size:14px;font-weight:700;cursor:pointer;transition:all 0.2s;box-shadow:0 4px 16px rgba(37,99,235,0.3)">
          📄 Generate Irregularity Report
        </button>
        <button onclick="generateIrrReport()" style="flex:1;min-width:200px;padding:14px 20px;background:linear-gradient(135deg,#92400E,#F59E0B);color:#fff;border:none;border-radius:var(--radius);font-family:\'Sora\',sans-serif;font-size:14px;font-weight:700;cursor:pointer;transition:all 0.2s;box-shadow:0 4px 16px rgba(245,158,11,0.3)">
          ⚡ Generate Warning Letters
        </button>
        <button onclick="clearIrrTable()" style="padding:14px 20px;background:#F1F5F9;color:#64748B;border:2px solid #E2E8F0;border-radius:var(--radius);font-family:\'Sora\',sans-serif;font-size:13px;font-weight:600;cursor:pointer;transition:all 0.2s">
          🗑️ Clear All
        </button>
      </div>
      <div id="irrGenStatus" style="margin-top:12px;font-size:12.5px;text-align:center;min-height:18px;font-weight:500"></div>
    </div>

    <div class="irr-trainee-table-wrap" id="irrPlaceholder">
      <div style="text-align:center;padding:40px 20px">
        <div style="font-size:64px;margin-bottom:16px">📋</div>
        <div style="font-size:17px;font-weight:700;color:var(--navy);margin-bottom:8px">Upload Trainee Details to Begin</div>
        <div style="font-size:13px;color:#94A3B8">Once you upload the trainee Excel file above, all trainee names will appear here with fields to enter attendance data.</div>
      </div>
    </div>
  </div>

</div>
<script>
function showView(id){
  document.querySelectorAll('.view').forEach(v=>v.classList.remove('active'));
  document.getElementById(id).classList.add('active');
  window.scrollTo(0,0);
}

async function loadUser(){
  const r=await fetch('/api/me');
  if(!r.ok){window.location.href='/login';return;}
  const d=await r.json();
  document.getElementById('uidBadge').textContent='ID: '+d.user_id;
  document.getElementById('welcomeName').textContent='Welcome, '+d.si_name+' 👋';
  document.getElementById('dSI').textContent=d.si_name;
  document.getElementById('dTrade').textContent=d.trade_name;
  document.getElementById('dITI').textContent=d.iti_name;
}
loadUser();

function bindFile(inputId,boxId,nameId){
  const inp=document.getElementById(inputId);
  const box=document.getElementById(boxId);
  const nm=document.getElementById(nameId);
  inp.addEventListener('change',()=>{if(inp.files[0]){nm.textContent='📁 '+inp.files[0].name;box.classList.add('active');}});
}
bindFile('tFile','wb1','fn1'); bindFile('lFile','wb2','fn2');

// Load saved filenames from previous session
async function loadSavedFiles(){
  try{
    const r=await fetch('/api/saved-files');
    if(!r.ok) return;
    const d=await r.json();
    if(d.trainee){
      const nm=document.getElementById('fn1');
      const box=document.getElementById('wb1');
      nm.textContent='✅ Saved: '+d.trainee;
      nm.style.color='#16A34A';
      box.classList.add('active');
    }
    if(d.lo){
      const nm=document.getElementById('fn2');
      const box=document.getElementById('wb2');
      nm.textContent='✅ Saved: '+d.lo;
      nm.style.color='#16A34A';
      box.classList.add('active');
    }
    // IRR letterhead — show saved status in upload box
    if(d.irr_letterhead){
      const nm=document.getElementById('irrLhName');
      const box=document.getElementById('irrLhBox');
      if(nm) nm.textContent='✅ Saved: '+d.irr_letterhead;
      if(nm) nm.style.color='#16A34A';
      if(box) box.classList.add('active');
    }
    // IRR trainee — show saved status and auto-load trainee table
    if(d.irr_trainee){
      const nm=document.getElementById('irrTrName');
      const box=document.getElementById('irrTrBox');
      if(nm) nm.textContent='✅ Saved: '+d.irr_trainee;
      if(nm) nm.style.color='#16A34A';
      if(box) box.classList.add('active');
      // Auto-load the saved trainee table
      try{
        const tr=await fetch('/api/irr-load-saved-trainees');
        if(tr.ok){
          const td=await tr.json();
          if(td.trainees&&td.trainees.length>0) buildIrrTable(td.trainees);
        }
      }catch(e){}
    }
  }catch(e){}
}
loadSavedFiles();

document.getElementById('genForm').addEventListener('submit',async(e)=>{
  e.preventDefault();
  const status=document.getElementById('status');
  const btn=document.getElementById('gBtn');
  btn.disabled=true; status.className='loading';
  status.innerHTML='<span class="spinner"></span>Generating report for all students...';
  const fd=new FormData();
  const tF=document.getElementById('tFile').files[0];
  const lF=document.getElementById('lFile').files[0];
  if(tF) fd.append('trainee',tF);
  if(lF) fd.append('lo',lF);
  try{
    const r=await fetch('/generate',{method:'POST',body:fd});
    if(r.ok){
      const blob=await r.blob();
      const url=URL.createObjectURL(blob);
      const a=document.createElement('a');a.href=url;a.download='ITI_Assessment_Report.xlsx';
      document.body.appendChild(a);a.click();a.remove();URL.revokeObjectURL(url);
      status.className='success';status.textContent='✅ Report generated! Download started automatically.';
    }else{
      const err=await r.json();
      status.className='error';status.textContent='❌ '+(err.error||'Error occurred.');
    }
  }catch(e){status.className='error';status.textContent='❌ Connection error.';}
  finally{btn.disabled=false;}
});

async function dlAnnexure(type){
  const st=document.getElementById('annStatus');
  const tFileAnn=document.getElementById('tFile').files[0];
  st.style.color='#1B4F8A';st.textContent='⏳ Generating Annexure III '+type+'...';
  const fd=new FormData();
  if(tFileAnn) fd.append('trainee',tFileAnn);
  try{
    const r=await fetch('/generate-annexure-'+type.toLowerCase(),{method:'POST',body:fd});
    if(r.ok){
      const blob=await r.blob();
      const url=URL.createObjectURL(blob);
      const a=document.createElement('a');a.href=url;a.download='ANNEXURE_III_'+type+'.xlsx';
      document.body.appendChild(a);a.click();a.remove();URL.revokeObjectURL(url);
      st.style.color='#16A34A';st.textContent='✅ ANNEXURE_III_'+type+'.xlsx downloaded!';
    }else{
      let msg='Error';try{const e=await r.json();msg=e.error||msg;}catch(e){}
      st.style.color='#DC2626';st.textContent='❌ '+msg;
    }
  }catch(ex){st.style.color='#DC2626';st.textContent='❌ '+ex.message;}
}

async function dlLOSummary(){
  const st=document.getElementById('loSumStatus');
  const tFile=document.getElementById('tFile').files[0];
  const lFile=document.getElementById('lFile').files[0];
  st.style.color='#1B4F8A';st.textContent='⏳ Generating LO Wise Summary...';
  const fd=new FormData();
  if(tFile) fd.append('trainee',tFile);
  if(lFile) fd.append('lo',lFile);
  try{
    const r=await fetch('/generate-lo-summary',{method:'POST',body:fd});
    if(r.ok){
      const blob=await r.blob();
      const url=URL.createObjectURL(blob);
      const a=document.createElement('a');a.href=url;a.download='LO_Wise_Summary.xlsx';
      document.body.appendChild(a);a.click();a.remove();URL.revokeObjectURL(url);
      st.style.color='#16A34A';st.textContent='✅ LO_Wise_Summary.xlsx downloaded!';
    }else{
      let msg='Error';try{const e=await r.json();msg=e.error||msg;}catch(e){}
      st.style.color='#DC2626';st.textContent='❌ '+msg;
    }
  }catch(ex){st.style.color='#DC2626';st.textContent='❌ '+ex.message;}
}

async function dlTraineeSummary(){
  const st=document.getElementById('traineeSumStatus');
  const tFile=document.getElementById('tFile').files[0];
  const lFile=document.getElementById('lFile').files[0];
  st.style.color='#1B4F8A';st.textContent='⏳ Generating Trainee Wise Summary...';
  const fd=new FormData();
  if(tFile) fd.append('trainee',tFile);
  if(lFile) fd.append('lo',lFile);
  try{
    const r=await fetch('/generate-trainee-summary',{method:'POST',body:fd});
    if(r.ok){
      const blob=await r.blob();
      const url=URL.createObjectURL(blob);
      const a=document.createElement('a');a.href=url;a.download='Trainee_Wise_Summary.xlsx';
      document.body.appendChild(a);a.click();a.remove();URL.revokeObjectURL(url);
      st.style.color='#16A34A';st.textContent='✅ Trainee_Wise_Summary.xlsx downloaded!';
    }else{
      let msg='Error';try{const e=await r.json();msg=e.error||msg;}catch(e){}
      st.style.color='#DC2626';st.textContent='❌ '+msg;
    }
  }catch(ex){st.style.color='#DC2626';st.textContent='❌ '+ex.message;}
}

async function dlTraineeSummaryPDF(){
  const st=document.getElementById('traineeSumStatus');
  const tFile=document.getElementById('tFile').files[0];
  const lFile=document.getElementById('lFile').files[0];
  st.style.color='#7C1D1D';st.textContent='⏳ Generating Trainee Wise Summary PDF...';
  const fd=new FormData();
  if(tFile) fd.append('trainee',tFile);
  if(lFile) fd.append('lo',lFile);
  try{
    const r=await fetch('/generate-trainee-summary-pdf',{method:'POST',body:fd});
    if(r.ok){
      const blob=await r.blob();
      const url=URL.createObjectURL(blob);
      const a=document.createElement('a');a.href=url;a.download='Trainee_Wise_Summary.pdf';
      document.body.appendChild(a);a.click();a.remove();URL.revokeObjectURL(url);
      st.style.color='#16A34A';st.textContent='✅ Trainee_Wise_Summary.pdf downloaded!';
    }else{
      let msg='Error';try{const e=await r.json();msg=e.error||msg;}catch(e){}
      st.style.color='#DC2626';st.textContent='❌ '+msg;
    }
  }catch(ex){st.style.color='#DC2626';st.textContent='❌ '+ex.message;}
}

async function dlProgressCard(){
  const st=document.getElementById('pcStatus');
  const tFile=document.getElementById('tFile').files[0];
  const fd=new FormData();
  if(tFile) fd.append('trainee',tFile);
  st.style.color='#1B4F8A';st.textContent='⏳ Generating Progress Cards...';
  try{
    const r=await fetch('/generate-progress-card',{method:'POST',body:fd});
    if(r.ok){
      const blob=await r.blob();
      const url=URL.createObjectURL(blob);
      const a=document.createElement('a');a.href=url;a.download='Progress_Cards.xlsx';
      document.body.appendChild(a);a.click();a.remove();URL.revokeObjectURL(url);
      st.style.color='#16A34A';st.textContent='✅ Progress_Cards.xlsx downloaded!';
    }else{
      let msg='Error';try{const e=await r.json();msg=e.error||msg;}catch(e){}
      st.style.color='#DC2626';st.textContent='❌ '+msg;
    }
  }catch(ex){st.style.color='#DC2626';st.textContent='❌ '+ex.message;}
}

async function dlTraineeLO(){
  const st=document.getElementById('traineeLOStatus');
  const tFile=document.getElementById('tFile').files[0];
  const lFile=document.getElementById('lFile').files[0];
  st.style.color='#1B4F8A';st.textContent='⏳ Generating Per Trainee Per LO Report...';
  const fd=new FormData();
  if(tFile) fd.append('trainee',tFile);
  if(lFile) fd.append('lo',lFile);
  try{
    const r=await fetch('/generate-trainee-lo-report',{method:'POST',body:fd});
    if(r.ok){
      const blob=await r.blob();
      const url=URL.createObjectURL(blob);
      const a=document.createElement('a');a.href=url;a.download='Per_Trainee_Per_LO_Report.xlsx';
      document.body.appendChild(a);a.click();a.remove();URL.revokeObjectURL(url);
      st.style.color='#16A34A';st.textContent='✅ Per_Trainee_Per_LO_Report.xlsx downloaded!';
    }else{
      let msg='Error';try{const e=await r.json();msg=e.error||msg;}catch(e){}
      st.style.color='#DC2626';st.textContent='❌ '+msg;
    }
  }catch(ex){st.style.color='#DC2626';st.textContent='❌ '+ex.message;}
}

async function dlPDF(){
  const st=document.getElementById('pdfStatus');
  const btn=document.getElementById('pdfBtn');
  const tFile=document.getElementById('tFile').files[0];
  const lFile=document.getElementById('lFile').files[0];
  st.style.color='#7C1D1D';st.textContent='⏳ Generating print-ready PDF for all trainees...';
  btn.disabled=true;
  const fd=new FormData();
  if(tFile) fd.append('trainee',tFile);
  if(lFile) fd.append('lo',lFile);
  try{
    const r=await fetch('/generate-pdf',{method:'POST',body:fd});
    if(r.ok){
      const blob=await r.blob();
      const url=URL.createObjectURL(blob);
      const a=document.createElement('a');a.href=url;a.download='ITI_Assessment_Report.pdf';
      document.body.appendChild(a);a.click();a.remove();URL.revokeObjectURL(url);
      st.style.color='#16A34A';st.textContent='✅ ITI_Assessment_Report.pdf downloaded! Ready to print.';
    }else{
      let msg='Error';try{const e=await r.json();msg=e.error||msg;}catch(e){}
      st.style.color='#DC2626';st.textContent='❌ '+msg;
    }
  }catch(ex){st.style.color='#DC2626';st.textContent='❌ '+ex.message;}
  finally{btn.disabled=false;}
}

async function dlAnnexurePDF(type){
  const st=document.getElementById('annStatus');
  const tFile=document.getElementById('tFile').files[0];
  st.style.color='#7C1D1D';st.textContent='⏳ Generating Annexure III '+type+' PDF...';
  const fd=new FormData();
  if(tFile) fd.append('trainee',tFile);
  try{
    const r=await fetch('/generate-annexure-'+type.toLowerCase()+'-pdf',{method:'POST',body:fd});
    if(r.ok){
      const blob=await r.blob();
      const url=URL.createObjectURL(blob);
      const a=document.createElement('a');a.href=url;a.download='ANNEXURE_III_'+type+'.pdf';
      document.body.appendChild(a);a.click();a.remove();URL.revokeObjectURL(url);
      st.style.color='#16A34A';st.textContent='✅ ANNEXURE_III_'+type+'.pdf downloaded!';
    }else{
      let msg='Error';try{const e=await r.json();msg=e.error||msg;}catch(e){}
      st.style.color='#DC2626';st.textContent='❌ '+msg;
    }
  }catch(ex){st.style.color='#DC2626';st.textContent='❌ '+ex.message;}
}

async function dlLOSummaryPDF(){
  const st=document.getElementById('loSumStatus');
  const tFile=document.getElementById('tFile').files[0];
  const lFile=document.getElementById('lFile').files[0];
  st.style.color='#7C1D1D';st.textContent='⏳ Generating LO Wise Summary PDF...';
  const fd=new FormData();
  if(tFile) fd.append('trainee',tFile);
  if(lFile) fd.append('lo',lFile);
  try{
    const r=await fetch('/generate-lo-summary-pdf',{method:'POST',body:fd});
    if(r.ok){
      const blob=await r.blob();
      const url=URL.createObjectURL(blob);
      const a=document.createElement('a');a.href=url;a.download='LO_Wise_Summary.pdf';
      document.body.appendChild(a);a.click();a.remove();URL.revokeObjectURL(url);
      st.style.color='#16A34A';st.textContent='✅ LO_Wise_Summary.pdf downloaded!';
    }else{
      let msg='Error';try{const e=await r.json();msg=e.error||msg;}catch(e){}
      st.style.color='#DC2626';st.textContent='❌ '+msg;
    }
  }catch(ex){st.style.color='#DC2626';st.textContent='❌ '+ex.message;}
}

async function dlTraineeLOPDF(){
  const st=document.getElementById('traineeLOStatus');
  const tFile=document.getElementById('tFile').files[0];
  const lFile=document.getElementById('lFile').files[0];
  st.style.color='#7C1D1D';st.textContent='⏳ Generating Per Trainee Per LO PDF...';
  const fd=new FormData();
  if(tFile) fd.append('trainee',tFile);
  if(lFile) fd.append('lo',lFile);
  try{
    const r=await fetch('/generate-trainee-lo-pdf',{method:'POST',body:fd});
    if(r.ok){
      const blob=await r.blob();
      const url=URL.createObjectURL(blob);
      const a=document.createElement('a');a.href=url;a.download='Per_Trainee_Per_LO_Report.pdf';
      document.body.appendChild(a);a.click();a.remove();URL.revokeObjectURL(url);
      st.style.color='#16A34A';st.textContent='✅ Per_Trainee_Per_LO_Report.pdf downloaded!';
    }else{
      let msg='Error';try{const e=await r.json();msg=e.error||msg;}catch(e){}
      st.style.color='#DC2626';st.textContent='❌ '+msg;
    }
  }catch(ex){st.style.color='#DC2626';st.textContent='❌ '+ex.message;}
}

async function dlProgressCardPDF(){
  const st=document.getElementById('pcStatus');
  const tFile=document.getElementById('tFile').files[0];
  st.style.color='#7C1D1D';st.textContent='⏳ Generating Progress Cards PDF...';
  const fd=new FormData();
  if(tFile) fd.append('trainee',tFile);
  try{
    const r=await fetch('/generate-progress-card-pdf',{method:'POST',body:fd});
    if(r.ok){
      const blob=await r.blob();
      const url=URL.createObjectURL(blob);
      const a=document.createElement('a');a.href=url;a.download='Progress_Cards.pdf';
      document.body.appendChild(a);a.click();a.remove();URL.revokeObjectURL(url);
      st.style.color='#16A34A';st.textContent='✅ Progress_Cards.pdf downloaded!';
    }else{
      let msg='Error';try{const e=await r.json();msg=e.error||msg;}catch(e){}
      st.style.color='#DC2626';st.textContent='❌ '+msg;
    }
  }catch(ex){st.style.color='#DC2626';st.textContent='❌ '+ex.message;}
}

async function logout(){
  await fetch('/api/logout',{method:'POST'});
  window.location.href='/login';
}

function toggleChpwd(){
  const box=document.getElementById('chpwdBox');
  const btn=document.getElementById('toggleChpwdBtn');
  const isOpen=box.style.display!=='none';
  box.style.display=isOpen?'none':'block';
  btn.innerHTML=isOpen?'🔒 Change Password':'✖ Hide Password Form';
  btn.style.borderColor=isOpen?'#CBD5E1':'#1A56DB';
  btn.style.color=isOpen?'#475569':'#1A56DB';
  if(isOpen){document.getElementById('cpCurrent').value='';document.getElementById('cpNew').value='';document.getElementById('cpConfirm').value='';const msg=document.getElementById('cpMsg');msg.className='chpwd-msg';msg.style.display='none';}
}

async function changeUserPassword(){
  const cur=document.getElementById('cpCurrent').value.trim();
  const nw=document.getElementById('cpNew').value.trim();
  const cf=document.getElementById('cpConfirm').value.trim();
  const msg=document.getElementById('cpMsg');
  const btn=document.getElementById('cpBtn');
  msg.className='chpwd-msg';msg.style.display='none';
  if(!cur||!nw||!cf){msg.className='chpwd-msg error';msg.textContent='❌ Please fill in all fields.';return;}
  if(nw.length<4){msg.className='chpwd-msg error';msg.textContent='❌ Min 4 characters required.';return;}
  if(nw!==cf){msg.className='chpwd-msg error';msg.textContent='❌ Passwords do not match.';return;}
  btn.disabled=true;btn.textContent='Updating...';
  const r=await fetch('/api/user/change-password',{method:'POST',headers:{'Content-Type':'application/json'},
    body:JSON.stringify({current_password:cur,new_password:nw})});
  const d=await r.json();
  btn.disabled=false;btn.textContent='Update Password';
  if(r.ok){
    msg.className='chpwd-msg success';msg.textContent='✅ Password updated!';
    document.getElementById('cpCurrent').value='';
    document.getElementById('cpNew').value='';
    document.getElementById('cpConfirm').value='';
  }else{msg.className='chpwd-msg error';msg.textContent='❌ '+(d.error||'Failed');}
}

/* ── IRREGULARITY REPORT JS ── */
function irrFileChange(inputId, boxId, nameId){
  const inp=document.getElementById(inputId);
  const box=document.getElementById(boxId);
  const nm=document.getElementById(nameId);
  if(inp.files[0]){
    nm.textContent='📁 '+inp.files[0].name;
    box.classList.add('active');
    // Auto-save letterhead immediately so it persists across sessions
    if(inputId==='irrLhFile'){
      const fd=new FormData();
      fd.append('letterhead',inp.files[0]);
      fetch('/api/irr-save-letterhead',{method:'POST',body:fd})
        .then(r=>r.json())
        .then(d=>{ if(d.saved){ nm.textContent='✅ Saved: '+inp.files[0].name; nm.style.color='#16A34A'; }})
        .catch(()=>{});
    }
  }
}

async function loadIrrTrainees(){
  const tFile=document.getElementById('irrTrFile').files[0];
  if(!tFile) return;
  const st=document.getElementById('irrStatus');
  st.className='loading';st.textContent='⏳ Reading trainee details...';
  const fd=new FormData();
  fd.append('trainee',tFile);
  try{
    const r=await fetch('/api/irr-get-trainees',{method:'POST',body:fd});
    const d=await r.json();
    if(!r.ok){st.className='error';st.textContent='❌ '+(d.error||'Failed to read file');return;}
    st.style.display='none';st.className='';
    buildIrrTable(d.trainees);
  }catch(ex){st.className='error';st.textContent='❌ '+ex.message;}
}

function buildIrrTable(trainees){
  const tbody=document.getElementById('irrTableBody');
  tbody.innerHTML='';
  if(!trainees||trainees.length===0){
    tbody.innerHTML='<tr><td colspan="6" class="irr-no-data">No trainees found in the uploaded file.</td></tr>';
    document.getElementById('irrTableSection').style.display='block';
    document.getElementById('irrPlaceholder').style.display='none';
    return;
  }
  trainees.forEach((t,i)=>{
    const name = typeof t === 'object' ? t.name : t;
    const tr=document.createElement('tr');
    tr.dataset.name=name;
    tr.dataset.idx=i;
    // Store full trainee data as JSON for later use in warning letter
    if(typeof t === 'object') tr.dataset.info=JSON.stringify(t);
    tr.innerHTML=`
      <td><span style="font-size:12px">${name}</span></td>
      <td><input type="number" min="0" id="wd_${i}" placeholder="e.g. 120" oninput="calcPct(${i})"></td>
      <td><input type="number" min="0" id="pd_${i}" placeholder="e.g. 95" oninput="calcPct(${i})"></td>
      <td><span class="pct-cell" id="pct_${i}">—</span></td>
      <td><input type="date" id="af_${i}"></td>
      <td><input type="date" id="au_${i}"></td>
      <td><input type="date" id="pi_${i}" title="Optional — leave blank if no previous letter was sent"></td>
    `;
    tbody.appendChild(tr);
  });
  document.getElementById('irrTableSection').style.display='block';
  document.getElementById('irrPlaceholder').style.display='none';
}

function calcPct(i){
  const wd=parseFloat(document.getElementById('wd_'+i).value)||0;
  const pd=parseFloat(document.getElementById('pd_'+i).value)||0;
  const pctEl=document.getElementById('pct_'+i);
  if(wd>0){
    const pct=Math.round((pd/wd)*100*10)/10;
    pctEl.textContent=pct+'%';
    pctEl.className='pct-cell '+(pct>=80?'pct-ok':'pct-bad');
  }else{
    pctEl.textContent='—';
    pctEl.className='pct-cell';
  }
}

function clearIrrTable(){
  const rows=document.getElementById('irrTableBody').querySelectorAll('tr');
  rows.forEach((tr,i)=>{
    const inp=tr.querySelectorAll('input');
    inp.forEach(el=>{el.value='';});
    const pct=tr.querySelector('.pct-cell');
    if(pct){pct.textContent='—';pct.className='pct-cell';}
  });
}

async function generateIrrDocReport(){
  const st=document.getElementById('irrGenStatus');
  const rows=document.getElementById('irrTableBody').querySelectorAll('tr[data-idx]');
  const tFile=document.getElementById('irrTrFile').files[0];
  if(!tFile && rows.length===0){
    st.style.color='#DC2626';st.textContent='❌ Please upload trainee details first.';return;
  }
  const data=[];
  rows.forEach((tr,i)=>{
    const name=tr.dataset.name;
    const info=tr.dataset.info?JSON.parse(tr.dataset.info):{};
    const wd=document.getElementById('wd_'+i)?document.getElementById('wd_'+i).value:'';
    const pd=document.getElementById('pd_'+i)?document.getElementById('pd_'+i).value:'';
    const au=document.getElementById('au_'+i)?document.getElementById('au_'+i).value:'';
    if(!wd||!pd){return;}
    const pct=(parseFloat(wd)>0)?Math.round((parseFloat(pd||0)/parseFloat(wd))*100*10)/10:null;
    if(pct===null||pct>=80){return;}
    data.push({name,info,working_days:wd,present_days:pd,percentage:pct,attendance_upto:au});
  });
  if(data.length===0){
    st.style.color='#F59E0B';st.textContent='ℹ️ No trainees with less than 80% attendance found.';return;
  }
  st.style.color='#1B4F8A';st.textContent='⏳ Generating Irregularity Report...';
  const fd=new FormData();
  fd.append('attendance_data',JSON.stringify(data));
  try{
    const r=await fetch('/api/irr-generate-report',{method:'POST',body:fd});
    if(r.ok){
      const blob=await r.blob();
      const url=URL.createObjectURL(blob);
      const a=document.createElement('a');a.href=url;a.download='Irregularity_Report.docx';
      document.body.appendChild(a);a.click();a.remove();URL.revokeObjectURL(url);
      st.style.color='#16A34A';st.textContent='✅ Irregularity Report generated! Download started.';
    }else{
      let msg='Error';try{const e=await r.json();msg=e.error||msg;}catch(e){}
      st.style.color='#DC2626';st.textContent='❌ '+msg;
    }
  }catch(ex){st.style.color='#DC2626';st.textContent='❌ '+ex.message;}
}

async function generateIrrReport(){
  const st=document.getElementById('irrGenStatus');
  const lhFile=document.getElementById('irrLhFile').files[0];
  const tFile=document.getElementById('irrTrFile').files[0];

  // Collect table data — only include trainees with <80% attendance
  const rows=document.getElementById('irrTableBody').querySelectorAll('tr[data-idx]');

  // Allow generation if table is already loaded from saved file (rows exist)
  if(!tFile && rows.length===0){
    st.style.color='#DC2626';
    st.textContent='❌ Please upload trainee details first.';
    return;
  }
  const data=[];
  let skipped=0;
  rows.forEach((tr,i)=>{
    const name=tr.dataset.name;
    const info=tr.dataset.info?JSON.parse(tr.dataset.info):{};
    const wd=document.getElementById('wd_'+i)?document.getElementById('wd_'+i).value:'';
    const pd=document.getElementById('pd_'+i)?document.getElementById('pd_'+i).value:'';
    const af=document.getElementById('af_'+i)?document.getElementById('af_'+i).value:'';
    const au=document.getElementById('au_'+i)?document.getElementById('au_'+i).value:'';
    const pi=document.getElementById('pi_'+i)?document.getElementById('pi_'+i).value:'';
    // absent_from is optional — trainee may still be attending but % is low
    if(!wd||!pd){return;} // skip if core fields missing
    const pct=(parseFloat(wd)>0)?Math.round((parseFloat(pd||0)/parseFloat(wd))*100*10)/10:null;
    if(pct===null||pct>=80){skipped++;return;} // only <80%
    data.push({name,info,working_days:wd,present_days:pd,percentage:pct,absent_from:af,attendance_upto:au,prev_informed:pi});
  });

  if(data.length===0){
    st.style.color='#F59E0B';
    st.textContent=skipped>0
      ?'ℹ️ No trainees with less than 80% attendance found. Warning letters not needed.'
      :'❌ Please fill Working Days and Present Days for trainees.';
    return;
  }

  st.style.color='#1B4F8A';
  st.textContent='⏳ Generating warning letters for '+data.length+' trainee(s)...';

  const fd=new FormData();
  if(lhFile) fd.append('letterhead',lhFile);
  fd.append('attendance_data',JSON.stringify(data));

  try{
    const r=await fetch('/api/irr-generate-letters',{method:'POST',body:fd});
    if(r.ok){
      const blob=await r.blob();
      const url=URL.createObjectURL(blob);
      const a=document.createElement('a');a.href=url;a.download='Warning_Letters.docx';
      document.body.appendChild(a);a.click();a.remove();URL.revokeObjectURL(url);
      st.style.color='#16A34A';
      st.textContent='✅ Warning letters generated for '+data.length+' trainee(s)! Download started.'+(skipped?' ('+skipped+' skipped — attendance ≥80%.)':'');
    }else{
      let msg='Error';try{const e=await r.json();msg=e.error||msg;}catch(e){}
      st.style.color='#DC2626';st.textContent='❌ '+msg;
    }
  }catch(ex){st.style.color='#DC2626';st.textContent='❌ '+ex.message;}
}
</script>
</body></html>'''

ADMIN_LOGIN_HTML = '''<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Admin Login — ITI Assessment</title>
<link href="https://fonts.googleapis.com/css2?family=Sora:wght@400;600;700;800&display=swap" rel="stylesheet">
<style>
*{margin:0;padding:0;box-sizing:border-box}
body{font-family:'Sora',sans-serif;min-height:100vh;background:#0B1D3A;display:flex;align-items:center;justify-content:center;padding:20px}
.card{background:#fff;border-radius:20px;padding:48px 44px;width:100%;max-width:420px;box-shadow:0 40px 100px rgba(0,0,0,0.4)}
.badge{display:inline-block;background:#FEF3C7;color:#92400E;font-size:11px;font-weight:700;padding:4px 12px;border-radius:20px;letter-spacing:1px;text-transform:uppercase;margin-bottom:20px}
h2{font-size:24px;font-weight:800;color:#0B1D3A;margin-bottom:6px}
p{color:#64748B;font-size:13.5px;margin-bottom:32px}
label{display:block;font-size:12px;font-weight:600;color:#0B1D3A;margin-bottom:7px;text-transform:uppercase;letter-spacing:0.5px}
input{width:100%;padding:13px 16px;border:2px solid #E2E8F0;border-radius:10px;font-family:'Sora',sans-serif;font-size:13.5px;outline:none;transition:all 0.2s;margin-bottom:18px}
input:focus{border-color:#1A56DB;box-shadow:0 0 0 4px rgba(26,86,219,0.08)}
.btn{width:100%;padding:15px;background:linear-gradient(135deg,#92400E,#F59E0B);color:#fff;border:none;border-radius:10px;font-family:'Sora',sans-serif;font-size:14px;font-weight:700;cursor:pointer;transition:all 0.2s}
.btn:hover{transform:translateY(-1px);box-shadow:0 8px 20px rgba(245,158,11,0.35)}
#msg{margin-top:14px;padding:12px;border-radius:8px;font-size:13px;text-align:center;display:none}
#msg.error{background:#FEF2F2;color:#991B1B;display:block}
.back{text-align:center;margin-top:20px;font-size:13px;color:#64748B}
.back a{color:#1A56DB;font-weight:600;text-decoration:none}
</style>
</head>
<body>
<div class="card">
  <div class="badge">🔐 Admin Access</div>
  <h2>Admin Panel</h2>
  <p>Login to manage registration requests and approve users.</p>
  <form id="aForm">
    <label>Admin Password</label>
    <input type="password" id="apwd" placeholder="Enter admin password" required>
    <button type="submit" class="btn">Login as Admin</button>
  </form>
  <div id="msg"></div>
  <div class="back"><a href="/login">← Back to user login</a></div>
</div>
<script>
document.getElementById('aForm').addEventListener('submit',async(e)=>{
  e.preventDefault();
  const msg=document.getElementById('msg');
  const r=await fetch('/api/admin/login',{method:'POST',headers:{'Content-Type':'application/json'},
    body:JSON.stringify({password:document.getElementById('apwd').value})});
  const d=await r.json();
  if(r.ok){window.location.href='/admin/dashboard';}
  else{msg.className='error';msg.textContent='❌ '+d.error;}
});
</script>
</body></html>'''

ADMIN_DASHBOARD_HTML = '''<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Admin Dashboard — ITI Assessment</title>
<link href="https://fonts.googleapis.com/css2?family=Sora:wght@300;400;600;700;800&display=swap" rel="stylesheet">
<style>
*{margin:0;padding:0;box-sizing:border-box}
:root{--navy:#0B1D3A;--blue:#1A56DB;--accent:#F59E0B;--green:#16A34A;--red:#DC2626}
body{font-family:'Sora',sans-serif;background:#F1F5F9;min-height:100vh}
.topbar{background:linear-gradient(135deg,#92400E,#F59E0B);padding:0 32px;height:64px;display:flex;align-items:center;justify-content:space-between}
.topbar h1{color:#fff;font-size:17px;font-weight:700}
.topbar .right{display:flex;align-items:center;gap:12px}
.badge{background:rgba(255,255,255,0.2);color:#fff;padding:5px 14px;border-radius:20px;font-size:12px;font-weight:600}
.logout{background:rgba(255,255,255,0.15);border:1px solid rgba(255,255,255,0.3);color:#fff;padding:7px 14px;border-radius:8px;font-family:'Sora',sans-serif;font-size:12.5px;cursor:pointer;font-weight:600}
.logout:hover{background:rgba(255,255,255,0.25)}
.main{max-width:1100px;margin:0 auto;padding:36px 24px}
.stats{display:grid;grid-template-columns:repeat(3,1fr);gap:16px;margin-bottom:32px}
.stat{background:#fff;border-radius:14px;padding:22px 24px;box-shadow:0 1px 6px rgba(0,0,0,0.06)}
.stat .num{font-size:32px;font-weight:800;color:var(--navy)}
.stat .lbl{font-size:12.5px;color:#64748B;margin-top:4px;font-weight:600}
.stat.pending .num{color:#D97706}
.stat.approved .num{color:var(--green)}
.section-title{font-size:17px;font-weight:700;color:var(--navy);margin-bottom:16px}
.table-wrap{background:#fff;border-radius:16px;overflow:hidden;box-shadow:0 1px 8px rgba(0,0,0,0.07);margin-bottom:32px}
table{width:100%;border-collapse:collapse}
th{background:#F8FAFC;padding:12px 16px;text-align:left;font-size:11.5px;font-weight:700;color:#475569;text-transform:uppercase;letter-spacing:0.5px;border-bottom:1px solid #E2E8F0}
td{padding:14px 16px;font-size:13.5px;color:var(--navy);border-bottom:1px solid #F1F5F9;vertical-align:middle}
tr:last-child td{border-bottom:none}
tr:hover td{background:#F8FAFC}
.status-badge{padding:4px 12px;border-radius:20px;font-size:11.5px;font-weight:700;letter-spacing:0.3px}
.status-badge.pending{background:#FEF3C7;color:#92400E}
.status-badge.approved{background:#DCFCE7;color:#15803D}
.status-badge.disapproved{background:#FEE2E2;color:#991B1B}
.action-btn{padding:7px 14px;border-radius:8px;font-family:'Sora',sans-serif;font-size:12.5px;font-weight:700;cursor:pointer;border:none;transition:all 0.2s}
.approve-btn{background:#16A34A;color:#fff}
.approve-btn:hover{background:#15803D}
.disapprove-btn{background:#DC2626;color:#fff}
.disapprove-btn:hover{background:#B91C1C}
.creds-box{background:#F0FDF4;border:1px solid #BBF7D0;border-radius:10px;padding:12px 16px;font-size:12.5px;color:#15803D;font-weight:600;margin-top:6px;display:none}
.empty{text-align:center;padding:40px;color:#94A3B8;font-size:14px}
.spinner{display:inline-block;width:18px;height:18px;border:2px solid #E2E8F0;border-top-color:var(--blue);border-radius:50%;animation:spin 0.8s linear infinite}
@keyframes spin{to{transform:rotate(360deg)}}
@media(max-width:700px){.stats{grid-template-columns:1fr};table{font-size:12px}}
.pwd-card{background:#fff;border-radius:16px;box-shadow:0 1px 8px rgba(0,0,0,0.07);padding:28px 32px;max-width:480px;margin-bottom:40px}
.pwd-card h3{font-size:15px;font-weight:700;color:var(--navy);margin-bottom:18px}
.pwd-card label{display:block;font-size:12.5px;font-weight:600;color:#475569;margin-bottom:5px;margin-top:14px}
.pwd-card input{width:100%;padding:10px 14px;border:1.5px solid #CBD5E1;border-radius:8px;font-family:'Sora',sans-serif;font-size:13px;color:var(--navy);outline:none;transition:border 0.2s}
.pwd-card input:focus{border-color:var(--blue);box-shadow:0 0 0 3px rgba(26,86,219,0.08)}
.pwd-save-btn{margin-top:18px;padding:10px 28px;background:linear-gradient(135deg,#92400E,#F59E0B);color:#fff;border:none;border-radius:8px;font-family:'Sora',sans-serif;font-size:13px;font-weight:700;cursor:pointer;transition:all 0.2s}
.pwd-save-btn:hover{transform:translateY(-1px);box-shadow:0 6px 16px rgba(245,158,11,0.35)}
#pwdMsg{margin-top:12px;padding:10px 14px;border-radius:8px;font-size:12.5px;display:none}
#pwdMsg.success{background:#F0FDF4;color:#15803D;display:block}
#pwdMsg.error{background:#FEF2F2;color:#991B1B;display:block}
</style>
</head>
<body>
<div class="topbar">
  <h1>🔐 Admin Dashboard — ITI Assessment System</h1>
  <div class="right">
    <div class="badge">Administrator</div>
    <button class="logout" onclick="adminLogout()">Logout</button>
  </div>
</div>
<div class="main">
  <div class="stats">
    <div class="stat pending"><div class="num" id="cPending">—</div><div class="lbl">⏳ Pending Requests</div></div>
    <div class="stat approved"><div class="num" id="cApproved">—</div><div class="lbl">✅ Approved Users</div></div>
    <div class="stat"><div class="num" id="cTotal">—</div><div class="lbl">📋 Total Registrations</div></div>
  </div>

  <div class="section-title">⏳ Pending Approvals</div>
  <div class="table-wrap">
    <table>
      <thead><tr><th>ID</th><th>SI Name</th><th>Trade</th><th>ITI Name</th><th>Mobile</th><th>Registered</th><th>Action</th></tr></thead>
      <tbody id="pendingBody"><tr><td colspan="7" class="empty"><span class="spinner"></span></td></tr></tbody>
    </table>
  </div>

  <div class="section-title">✅ Approved Users</div>
  <div class="table-wrap">
    <table>
      <thead><tr><th>ID</th><th>SI Name</th><th>Trade</th><th>ITI Name</th><th>Mobile</th><th>User ID</th><th>Approved On</th><th>Status</th><th>Action</th></tr></thead>
      <tbody id="approvedBody"><tr><td colspan="9" class="empty"><span class="spinner"></span></td></tr></tbody>
    </table>
  </div>

  <div class="section-title" style="display:flex;align-items:center;justify-content:space-between">
    <span>🚫 Deregistered Users</span>
    <button class="action-btn" style="background:#64748B;color:#fff;font-size:12px" onclick="clearDeregistered()">🗑️ Clear All Deregistered</button>
  </div>
  <div class="table-wrap">
    <table>
      <thead><tr><th>ID</th><th>SI Name</th><th>Trade</th><th>ITI Name</th><th>Mobile</th><th>Registered On</th><th>Approved On</th><th>Deregistered On</th><th>Status</th></tr></thead>
      <tbody id="disapprovedBody"><tr><td colspan="9" class="empty"><span class="spinner"></span></td></tr></tbody>
    </table>
  </div>

  <div class="section-title" style="display:flex;align-items:center;justify-content:space-between">
    <span>🔑 Change Admin Password</span>
    <button id="toggleAdminPwdBtn" onclick="toggleAdminPwd()" style="padding:7px 18px;background:#fff;border:1.5px solid #CBD5E1;border-radius:8px;font-family:\\'Sora\\',sans-serif;font-size:12.5px;font-weight:700;color:#475569;cursor:pointer;transition:all 0.2s">🔒 Change Password</button>
  </div>
  <div class="pwd-card" id="adminPwdCard" style="display:none">
    <h3>Update your admin password</h3>
    <label>Current Password</label>
    <input type="password" id="pwdCurrent" placeholder="Enter current password">
    <label>New Password</label>
    <input type="password" id="pwdNew" placeholder="Enter new password (min 6 chars)">
    <label>Confirm New Password</label>
    <input type="password" id="pwdConfirm" placeholder="Re-enter new password">
    <br>
    <button class="pwd-save-btn" onclick="changeAdminPassword()">🔒 Update Password</button>
    <div id="pwdMsg"></div>
  </div>
</div>
<script>
async function loadData(){
  const r=await fetch('/api/admin/users');
  if(!r.ok){window.location.href='/admin/login';return;}
  const users=await r.json();
  const pending=users.filter(u=>u.status==='pending');
  const approved=users.filter(u=>u.status==='approved');
  const disapproved=users.filter(u=>u.status==='disapproved');
  document.getElementById('cPending').textContent=pending.length;
  document.getElementById('cApproved').textContent=approved.length;
  document.getElementById('cTotal').textContent=users.length;

  const pb=document.getElementById('pendingBody');
  pb.innerHTML=pending.length===0?'<tr><td colspan="7" class="empty">No pending requests 🎉</td></tr>':
    pending.map(u=>`<tr>
      <td>${u.id}</td><td><strong>${u.si_name}</strong></td><td>${u.trade_name}</td>
      <td>${u.iti_name}</td><td>${u.mobile}</td>
      <td style="font-size:12px;color:#64748B">${u.created_at?.split('T')[0]||u.created_at||'—'}</td>
      <td>
        <button class="action-btn approve-btn" onclick="approveUser(${u.id},this)">✅ Approve</button>
        <div class="creds-box" id="cred_${u.id}"></div>
      </td>
    </tr>`).join('');

  const ab=document.getElementById('approvedBody');
  ab.innerHTML=approved.length===0?'<tr><td colspan="9" class="empty">No approved users yet</td></tr>':
    approved.map(u=>`<tr>
      <td>${u.id}</td><td><strong>${u.si_name}</strong></td><td>${u.trade_name}</td>
      <td>${u.iti_name}</td><td>${u.mobile}</td>
      <td><strong style="color:#1A56DB">${u.user_id}</strong></td>
      <td style="font-size:12px;color:#16A34A">${u.approved_at?u.approved_at.split('T')[0]||u.approved_at:'—'}</td>
      <td><span class="status-badge approved">Approved</span></td>
      <td>
        <button class="action-btn disapprove-btn" onclick="disapproveUser(${u.id},this)">🚫 Deregister</button>
      </td>
    </tr>`).join('');

  const db2=document.getElementById('disapprovedBody');
  db2.innerHTML=disapproved.length===0?'<tr><td colspan="9" class="empty">No deregistered users</td></tr>':
    disapproved.map(u=>`<tr>
      <td>${u.id}</td><td><strong>${u.si_name}</strong></td><td>${u.trade_name}</td>
      <td>${u.iti_name}</td><td>${u.mobile}</td>
      <td style="font-size:12px;color:#64748B">${u.created_at?u.created_at.split('T')[0]||u.created_at:'—'}</td>
      <td style="font-size:12px;color:#16A34A">${u.approved_at?u.approved_at.split('T')[0]||u.approved_at:'—'}</td>
      <td style="font-size:12px;color:#DC2626">${u.deregistered_at?u.deregistered_at.split('T')[0]||u.deregistered_at:'—'}</td>
      <td><span class="status-badge disapproved">Deregistered</span></td>
    </tr>`).join('');
}

async function approveUser(id,btn){
  btn.disabled=true;btn.textContent='Approving...';
  const r=await fetch('/api/admin/approve/'+id,{method:'POST'});
  const d=await r.json();
  if(r.ok){
    const box=document.getElementById('cred_'+id);
    box.style.display='block';
    box.innerHTML=`✅ Approved!<br>🆔 User ID: <strong>${d.user_id}</strong><br>🔑 Password: <strong>${d.password}</strong><br><small>Share these credentials with the user.</small>`;
    btn.style.display='none';
    setTimeout(loadData,2000);
  }else{btn.disabled=false;btn.textContent='✅ Approve';alert('Error: '+d.error);}
}

async function disapproveUser(id,btn){
  if(!confirm('Are you sure you want to deregister this user? Their login access will be revoked and they cannot re-register.')) return;
  btn.disabled=true;btn.textContent='Revoking...';
  const r=await fetch('/api/admin/disapprove/'+id,{method:'POST'});
  const d=await r.json();
  if(r.ok){ setTimeout(loadData,500); }
  else{ btn.disabled=false;btn.textContent='🚫 Deregister';alert('Error: '+d.error); }
}

async function clearDeregistered(){
  if(!confirm('Clear ALL deregistered users from the list? This cannot be undone.')) return;
  const r=await fetch('/api/admin/clear-deregistered',{method:'POST'});
  const d=await r.json();
  if(r.ok){ loadData(); }
  else{ alert('Error: '+d.error); }
}

async function adminLogout(){
  await fetch('/api/admin/logout',{method:'POST'});
  window.location.href='/admin/login';
}

function toggleAdminPwd(){
  const card=document.getElementById('adminPwdCard');
  const btn=document.getElementById('toggleAdminPwdBtn');
  const isOpen=card.style.display!=='none';
  card.style.display=isOpen?'none':'block';
  btn.textContent=isOpen?'🔒 Change Password':'✖ Hide';
  btn.style.borderColor=isOpen?'#CBD5E1':'#1A56DB';
  btn.style.color=isOpen?'#475569':'#1A56DB';
  if(isOpen){document.getElementById('pwdCurrent').value='';document.getElementById('pwdNew').value='';document.getElementById('pwdConfirm').value='';const msg=document.getElementById('pwdMsg');msg.className='';msg.style.display='none';}
}

async function changeAdminPassword(){
  const cur=document.getElementById('pwdCurrent').value.trim();
  const nw=document.getElementById('pwdNew').value.trim();
  const cf=document.getElementById('pwdConfirm').value.trim();
  const msg=document.getElementById('pwdMsg');
  msg.className='';msg.style.display='none';
  if(!cur||!nw||!cf){msg.className='error';msg.textContent='❌ Please fill in all fields.';return;}
  if(nw.length<6){msg.className='error';msg.textContent='❌ New password must be at least 6 characters.';return;}
  if(nw!==cf){msg.className='error';msg.textContent='❌ New passwords do not match.';return;}
  const r=await fetch('/api/admin/change-password',{method:'POST',headers:{'Content-Type':'application/json'},
    body:JSON.stringify({current_password:cur,new_password:nw})});
  const d=await r.json();
  if(r.ok){
    msg.className='success';msg.textContent='✅ Password updated successfully!';
    document.getElementById('pwdCurrent').value='';
    document.getElementById('pwdNew').value='';
    document.getElementById('pwdConfirm').value='';
  }else{msg.className='error';msg.textContent='❌ '+d.error;}
}

loadData();
setInterval(loadData,15000);
</script>
</body></html>'''

# ─── AUTH ROUTES ──────────────────────────────────────────────────
@app.route('/')
def home(): return redirect('/login')

@app.route('/register')
def register_page(): return REGISTER_HTML

@app.route('/login')
def login_page(): return LOGIN_HTML

@app.route('/dashboard')
def dashboard_page():
    if 'user_id' not in session: return redirect('/login')
    return DASHBOARD_HTML

@app.route('/admin/login')
def admin_login_page(): return ADMIN_LOGIN_HTML

@app.route('/admin/dashboard')
def admin_dashboard_page():
    if not session.get('is_admin'): return redirect('/admin/login')
    return ADMIN_DASHBOARD_HTML

@app.route('/api/register', methods=['POST'])
def api_register():
    d = request.json
    for f in ['si_name','trade_name','iti_name','mobile']:
        if not d.get(f): return jsonify({'error': f'{f} is required'}), 400
    if not d['mobile'].isdigit() or len(d['mobile']) != 10:
        return jsonify({'error': 'Mobile must be 10 digits'}), 400
    try:
        with get_db() as db:
            existing = db.execute('SELECT status FROM users WHERE mobile=? ORDER BY id DESC LIMIT 1', (d['mobile'],)).fetchone()
            if existing:
                if existing['status'] == 'approved':
                    return jsonify({'error': 'This mobile number is already registered and approved. Contact admin if you need help.'}), 400
                if existing['status'] == 'disapproved':
                    return jsonify({'error': 'This mobile number has been deregistered by admin. Contact admin for access.'}), 400
            db.execute(
                'INSERT INTO users (si_name,trade_name,iti_name,mobile,year_of_assessment,assessment_location,near_trade,trade_duration,semester,batch) VALUES (?,?,?,?,?,?,?,?,?,?)',
                (d['si_name'], d['trade_name'], d['iti_name'], d['mobile'],
                 d.get('year_of_assessment',''), d.get('assessment_location',''),
                 d.get('near_trade',''), d.get('trade_duration',''),
                 d.get('semester',''), d.get('batch',''))
            )
            db.commit()
        return jsonify({'message': 'Registration submitted. Awaiting admin approval.'}), 201
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/login', methods=['POST'])
def api_login():
    d = request.json
    uid = d.get('user_id','').strip()
    pwd = d.get('password','')
    with get_db() as db:
        user = db.execute('SELECT * FROM users WHERE user_id=? AND status=?', (uid,'approved')).fetchone()
    if not user or not check_password_hash(user['password_hash'], pwd):
        return jsonify({'error': 'Invalid User ID or Password'}), 401
    session['user_id'] = uid
    session['si_name'] = user['si_name']
    session['trade_name'] = user['trade_name']
    session['iti_name'] = user['iti_name']
    session['year_of_assessment'] = user['year_of_assessment'] or ''
    session['assessment_location'] = user['assessment_location'] or ''
    session['near_trade'] = user['near_trade'] or ''
    session['trade_duration'] = user['trade_duration'] or ''
    session['semester'] = d.get('semester') or user['semester'] or ''
    session['batch'] = user['batch'] or ''
    return jsonify({'message': 'Login successful'})

@app.route('/api/me')
def api_me():
    if 'user_id' not in session: return jsonify({'error': 'Not logged in'}), 401
    return jsonify({
        'user_id': session['user_id'],
        'si_name': session['si_name'],
        'trade_name': session['trade_name'],
        'iti_name': session['iti_name'],
        'year_of_assessment': session.get('year_of_assessment',''),
        'assessment_location': session.get('assessment_location',''),
        'near_trade': session.get('near_trade',''),
        'trade_duration': session.get('trade_duration',''),
        'semester': session.get('semester',''),
        'batch': session.get('batch',''),
    })

@app.route('/api/saved-files')
def api_saved_files():
    if 'user_id' not in session: return jsonify({'trainee': None, 'lo': None, 'irr_trainee': None, 'irr_letterhead': None})
    uid = session['user_id']
    def fname(key):
        path = _load_user_file(uid, key)
        return os.path.basename(path) if path else None
    def fname_ext(key, ext):
        path = _load_user_file_ext(uid, key, ext)
        return os.path.basename(path) if path else None
    return jsonify({
        'trainee':         fname('trainee'),
        'lo':              fname('lo'),
        'irr_trainee':     fname_ext('irr_trainee', 'xlsx'),
        'irr_letterhead':  fname_ext('irr_letterhead', 'docx'),
    })

@app.route('/api/logout', methods=['POST'])
def api_logout():
    session.clear(); return jsonify({'message': 'Logged out'})

@app.route('/api/user/change-password', methods=['POST'])
def api_user_change_password():
    if 'user_id' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
    d = request.json
    current_pwd = d.get('current_password', '')
    new_pwd = d.get('new_password', '')
    if not current_pwd or not new_pwd:
        return jsonify({'error': 'Both fields are required'}), 400
    if len(new_pwd) < 4:
        return jsonify({'error': 'New password must be at least 4 characters'}), 400
    uid = session['user_id']
    with get_db() as db:
        user = db.execute('SELECT * FROM users WHERE user_id=?', (uid,)).fetchone()
        if not user or not check_password_hash(user['password_hash'], current_pwd):
            return jsonify({'error': 'Current password is incorrect'}), 401
        db.execute('UPDATE users SET password_hash=? WHERE user_id=?',
                   (generate_password_hash(new_pwd), uid))
        db.commit()
    return jsonify({'message': 'Password updated successfully'})

# ─── IRREGULARITY REPORT ROUTES (NEW FEATURE) ────────────────────
@app.route('/api/irr-save-letterhead', methods=['POST'])
@login_required
def api_irr_save_letterhead():
    """Save letterhead file immediately when user selects it."""
    uid = session['user_id']
    lh_file = request.files.get('letterhead')
    if not lh_file or not lh_file.filename:
        return jsonify({'error': 'No file'}), 400
    try:
        import io as _io
        lh_file.seek(0)
        raw = lh_file.read()
        fname = lh_file.filename.lower()
        if fname.endswith('.doc') and not fname.endswith('.docx'):
            raw = _doc_to_docx_bytes(raw)
            if not raw:
                return jsonify({'error': 'Could not convert .doc — please save as .docx in Word and re-upload.'}), 500
        _save_user_file_ext(uid, 'irr_letterhead', _io.BytesIO(raw), 'docx')
        return jsonify({'saved': True, 'filename': 'irr_letterhead.docx'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/irr-load-saved-trainees', methods=['GET'])
@login_required
def api_irr_load_saved_trainees():
    """Return trainee list from the saved IRR trainee file (used on page load after login)."""
    uid = session['user_id']
    path = _load_user_file_ext(uid, 'irr_trainee', 'xlsx')
    if not path:
        return jsonify({'trainees': [], 'count': 0, 'saved': False})
    try:
        df = pd.read_excel(path)
        df.columns = [str(c).strip() for c in df.columns]
        col_map = {str(c).strip(): c for c in df.columns}
        def safe(key):
            col = col_map.get(key.strip())
            if col:
                return df[col].fillna('').astype(str).str.strip()
            return pd.Series([''] * len(df))
        def safe_contains(substring):
            """Match first column whose name contains the given substring."""
            for k in col_map:
                if substring in k:
                    return df[col_map[k]].fillna('').astype(str).str.strip()
            return pd.Series([''] * len(df))
        surname   = safe('અટક')
        firstname = safe('તાલીમાર્થીનું નામ')
        father    = safe('પિતાનું નામ')
        addr1     = safe('એડ્રેસ-૧')
        addr2     = safe('એડ્રેસ-૨')
        pincode   = safe('પીનકોડ')
        trade     = safe('ટ્રેડ નું નામ અને બેચ')
        si_name   = safe_contains('સુ.')
        iti_name  = safe('આઈ.ટી.આઈ નું નામ')
        trainees = []
        for i in range(len(df)):
            parts = [p for p in [surname[i], firstname[i], father[i]] if p]
            full_name = ' '.join(parts) if parts else f'Trainee {i+1}'
            trainees.append({'name': full_name, 'surname': surname[i], 'firstname': firstname[i],
                             'father': father[i], 'addr1': addr1[i], 'addr2': addr2[i],
                             'pincode': pincode[i], 'trade': trade[i], 'si_name': si_name[i],
                             'iti_name': iti_name[i]})
        trainees = [t for t in trainees if t['name'].strip()]
        return jsonify({'trainees': trainees, 'count': len(trainees), 'saved': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/irr-get-trainees', methods=['POST'])
@login_required
def api_irr_get_trainees():
    """
    Parse the Irregularity Report trainee Excel (Gujarati format) and return
    trainee details. This file is SEPARATE from the assessment trainee file —
    it is read directly from the upload and never saved to disk.

    Expected columns (Gujarati):
      અટક           = Surname
      તાલીમાર્થીનું નામ = First Name
      પિતાનું નામ     = Father's Name
      એડ્રેસ-૧        = Address line 1
      એડ્રેસ-૨        = Address line 2
      પીનકોડ          = Pincode
      ટ્રેડ નું નામ અને બેચ = Trade & Batch
      સુ.ઈ નું નામ)   = SI Name
      આઈ.ટી.આઈ નું નામ = ITI Name
    """
    try:
        tfile = request.files.get('trainee')
        if not tfile or not tfile.filename:
            return jsonify({'error': 'No file uploaded.'}), 400

        uid = session['user_id']
        # Save to disk so it persists across logout/login
        _save_user_file_ext(uid, 'irr_trainee', tfile, 'xlsx')
        tfile.seek(0)

        df = pd.read_excel(tfile)
        df.columns = [str(c).strip() for c in df.columns]

        # Column name maps (strip spaces for safe matching)
        col_map = {str(c).strip(): c for c in df.columns}

        def safe(key):
            """Return series for a column, matched by stripped name."""
            col = col_map.get(key.strip())
            if col:
                return df[col].fillna('').astype(str).str.strip()
            return pd.Series([''] * len(df))

        def safe_contains(substring):
            """Match first column whose name contains the given substring."""
            for k in col_map:
                if substring in k:
                    return df[col_map[k]].fillna('').astype(str).str.strip()
            return pd.Series([''] * len(df))

        # Gujarati column keys (stripped)
        SURNAME_COL   = 'અટક'
        FIRSTNAME_COL = 'તાલીમાર્થીનું નામ'
        FATHER_COL    = 'પિતાનું નામ'
        ADDR1_COL     = 'એડ્રેસ-૧'
        ADDR2_COL     = 'એડ્રેસ-૨'
        PINCODE_COL   = 'પીનકોડ'
        TRADE_COL     = 'ટ્રેડ નું નામ અને બેચ'
        ITI_COL       = 'આઈ.ટી.આઈ નું નામ'

        surname   = safe(SURNAME_COL)
        firstname = safe(FIRSTNAME_COL)
        father    = safe(FATHER_COL)
        addr1     = safe(ADDR1_COL)
        addr2     = safe(ADDR2_COL)
        pincode   = safe(PINCODE_COL)
        trade     = safe(TRADE_COL)
        si_name   = safe_contains('સુ.')
        iti_name  = safe(ITI_COL)

        trainees = []
        for i in range(len(df)):
            # Full name: Surname + First Name + Father's Name
            parts = [p for p in [surname[i], firstname[i], father[i]] if p]
            full_name = ' '.join(parts) if parts else f'Trainee {i+1}'

            trainees.append({
                'name':     full_name,
                'surname':  surname[i],
                'firstname':firstname[i],
                'father':   father[i],
                'addr1':    addr1[i],
                'addr2':    addr2[i],
                'pincode':  pincode[i],
                'trade':    trade[i],
                'si_name':  si_name[i],
                'iti_name': iti_name[i],
            })

        # Remove rows where name is completely empty
        trainees = [t for t in trainees if t['name'].strip()]
        return jsonify({'trainees': trainees, 'count': len(trainees)})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/irr-generate-report', methods=['POST'])
@login_required
def api_irr_generate_report():
    """
    Generate the Irregularity Report Word document (અનિયમિતતા રિપોર્ટ).
    One document with header info + table of trainees with <80% attendance.
    Data comes from the attendance table filled by the user.
    """
    import json
    from docx import Document as DocxDocument
    from docx.shared import Pt, Emu, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.oxml import OxmlElement

    try:
        raw = request.form.get('attendance_data', '[]')
        trainees = json.loads(raw)
        if not trainees:
            return jsonify({'error': 'No trainees with less than 80% attendance.'}), 400

        # Pick common info from first trainee
        info0     = trainees[0].get('info', {})
        si_name   = info0.get('si_name', '')
        iti_name  = info0.get('iti_name', '')
        trade     = info0.get('trade', '')
        att_upto  = trainees[0].get('attendance_upto', '')
        # Format att_upto for display (already DD/MM/YYYY from frontend)
        def fmt_date(s):
            if not s: return ''
            try:
                from datetime import datetime as dt
                return dt.strptime(s, '%Y-%m-%d').strftime('%d/%m/%Y')
            except:
                return s
        att_upto_disp = fmt_date(att_upto) if '-' in (att_upto or '') else att_upto

        today_str = datetime.now().strftime('%d/%m/%Y')

        FONT = 'Nirmala UI'
        SIZE = Pt(11)

        doc = DocxDocument()
        sec = doc.sections[0]
        sec.page_width    = Emu(7560945)
        sec.page_height   = Emu(10692765)
        sec.top_margin    = Emu(914400)
        sec.bottom_margin = Emu(914400)
        sec.left_margin   = Emu(457200)
        sec.right_margin  = Emu(629920)

        def add_p(text='', bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=0):
            p = doc.add_paragraph()
            p.alignment = align
            pf = p.paragraph_format
            pf.space_before = Pt(sb)
            pf.space_after  = Pt(sa)
            if text:
                r = p.add_run(text)
                r.bold = bold
                r.font.name = FONT
                r.font.size = SIZE
            return p

        def add_mixed_p(segments, align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=0):
            p = doc.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_before = Pt(sb)
            p.paragraph_format.space_after  = Pt(sa)
            for text, bold in segments:
                r = p.add_run(text)
                r.bold = bold
                r.font.name = FONT
                r.font.size = SIZE
            return p

        # Gujarati month names
        GUJ_MONTHS = {
            1: 'જાન્યુઆરી', 2: 'ફેબ્રુઆરી', 3: 'માર્ચ', 4: 'એપ્રિલ',
            5: 'મે', 6: 'જૂન', 7: 'જુલાઈ', 8: 'ઓગસ્ટ',
            9: 'સપ્ટેમ્બર', 10: 'ઓક્ટોબર', 11: 'નવેમ્બર', 12: 'ડિસેમ્બર'
        }
        GUJ_YEAR_DIGITS = {'0':'૦','1':'૧','2':'૨','3':'૩','4':'૪',
                           '5':'૫','6':'૬','7':'૭','8':'૮','9':'૯'}
        def to_guj_year(yr):
            return ''.join(GUJ_YEAR_DIGITS[c] for c in str(yr))

        def att_month_label(date_str):
            """Return e.g. 'જૂન - ૨૦૨૬' from a date string."""
            try:
                from datetime import datetime as dt
                d = dt.strptime(date_str, '%d/%m/%Y') if '/' in (date_str or '') else \
                    dt.strptime(date_str, '%Y-%m-%d')
                return f"{GUJ_MONTHS[d.month]} - {to_guj_year(d.year)}"
            except Exception:
                return date_str

        # ── Sender block ─────────────────────────────────────────────
        add_p(si_name, bold=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
        add_p(f'{iti_name}', align=WD_ALIGN_PARAGRAPH.RIGHT)
        add_p(f'તા.  {today_str}', align=WD_ALIGN_PARAGRAPH.RIGHT)
        add_p()

        # ── Recipient ────────────────────────────────────────────────
        add_p('પ્રતિ,')
        add_p('આચાર્યશ્રી')
        add_p(f'{iti_name}')

        # ── Subject ──────────────────────────────────────────────────
        month_label = att_month_label(att_upto_disp or att_upto)
        add_mixed_p([
            ('\tવિષય – ', False),
            (f'{month_label}', False),
            (' મહિના અંતિત હાજરી બાબત.', False),
        ], sb=4)

        # ── Body para ────────────────────────────────────────────────
        add_mixed_p([
            ('	સવિનય ઉપરોક્ત વિષય અનુસંધાને જણાવવાનુ કે, સદર આઈ.ટી.આઈ.', False),
            (' ના ટ્રેડ – ', False),
            (f'{trade}', False),
            (' ના નીચે મુજબ ના તાલીમાર્થીઓની માહે- ', False),
            (f'{month_label}', False),
            (' અંતિત હાજરી નીચે મુજબ છે. જે બાબતે તેઓના વાલીને જાણ કરતો પત્ર મોકલી આપવા યોગ્ય કરવા આપ સાહેબશ્રીને નમ્ર વિનંતી.', False),
        ], align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=4, sa=4)

        # ── Table ────────────────────────────────────────────────────
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = 'Table Grid'
        # Header row
        hdr = tbl.rows[0]
        for ci, (txt, w) in enumerate([('ક્રમ', 1000000), ('તાલીમાર્થીનુ નામ', 4500000), ('હાજરી (%)', 2000000)]):
            cell = hdr.cells[ci]
            cell.width = Emu(w)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(txt)
            r.bold = True
            r.font.name = FONT
            r.font.size = SIZE
            # No fill (white for print)
            # White font → change to black for print
            r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        # Gujarati number map
        GUJ_NUMS = ['૦','૧','૨','૩','૪','૫','૬','૭','૮','૯']
        def to_guj(n):
            return ''.join(GUJ_NUMS[int(d)] for d in str(n).zfill(2))

        for idx, t in enumerate(trainees, start=1):
            name = t.get('name', '')
            pct  = t.get('percentage', '')
            pct_str = f"{pct}%" if pct is not None else ''
            row = tbl.add_row()
            for ci, txt in enumerate([to_guj(idx), name, pct_str]):
                cell = row.cells[ci]
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after  = Pt(1)
                r = p.add_run(txt)
                r.font.name = FONT
                r.font.size = SIZE

        # ── Closing ──────────────────────────────────────────────────
        add_p()
        add_p('આભાર સહ.', align=WD_ALIGN_PARAGRAPH.CENTER)
        add_p('આપનો વિશ્વાસુ,', align=WD_ALIGN_PARAGRAPH.RIGHT)
        add_p()
        add_p(si_name, align=WD_ALIGN_PARAGRAPH.RIGHT)

        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        return send_file(out,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True,
                         download_name='Irregularity_Report.docx')

    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'detail': traceback.format_exc()}), 500


@app.route('/api/irr-generate-letters', methods=['POST'])
@login_required
def api_irr_generate_letters():  # FORMATTING UPDATED TO MATCH DEMO WARNING_LETTERS.DOCX
    """
    Generate warning letters (ચેતવણી પત્ર) as a single .docx file,
    one letter per page, only for trainees with <80% attendance.
    Formatting exactly replicates the demo Warning_Letters.docx:
      - Page: A4, margins top=112.5pt bottom=28.35pt left=81.35pt right=73.45pt
      - Table 0: ક્રમાંક (bold left) | તા. DATE (normal right), equal cols, borderless
      - ચેતવણી પત્ર: bold, 14pt Nirmala UI, centered, space_before=38100
      - પ્રતિ, block: bold Nirmala UI, space_before=25400
      - વિષય: all bold, centered, space_before=38100
      - Opening para: first_line_indent=457200 (1 inch), space_before=38100, space_after=25400
      - Points 1-6: left_indent=228600, first_line_indent=-228600, space_before=50800, space_after=25400
      - Signature table: left cell empty (space_before=355600), right cell center-aligned multi-line
      - CC block: space_before=50800; CC line: first_line_indent=457200
    """
    import json
    from docx import Document as DocxDocument
    from docx.shared import Pt, Emu, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    try:
        raw = request.form.get('attendance_data', '[]')
        trainees = json.loads(raw)
        if not trainees:
            return jsonify({'error': 'No trainees with less than 80% attendance to generate letters for.'}), 400

        lh_file = request.files.get('letterhead')
        uid = session['user_id']
        lh_doc_template = None

        if lh_file and lh_file.filename:
            import io as _io
            lh_file.seek(0)
            raw_bytes = lh_file.read()
            fname = lh_file.filename.lower()
            if fname.endswith('.doc') and not fname.endswith('.docx'):
                raw_bytes = _doc_to_docx_bytes(raw_bytes)
            if raw_bytes:
                _save_user_file_ext(uid, 'irr_letterhead', _io.BytesIO(raw_bytes), 'docx')
                try:
                    lh_doc_template = DocxDocument(_io.BytesIO(raw_bytes))
                except Exception:
                    lh_doc_template = None
        else:
            saved_lh_path = _load_user_file_ext(uid, 'irr_letterhead', 'docx')
            if saved_lh_path:
                try:
                    import io as _io
                    lh_doc_template = DocxDocument(_io.BytesIO(open(saved_lh_path, 'rb').read()))
                except Exception:
                    lh_doc_template = None

        # ── Exact measurements from demo ──────────────────────────────────
        FONT_GUJ    = 'Nirmala UI'
        FONT_SHRUTI = 'Shruti'
        # Demo body font size: inherited (None = 11pt default from Normal style)
        # Title font: 14pt (177800 EMU = 14pt)
        SIZE_TITLE  = Pt(14)

        def fmt_date(iso_str):
            if not iso_str:
                return ''
            try:
                from datetime import datetime as dt
                return dt.strptime(iso_str, '%Y-%m-%d').strftime('%d/%m/%Y')
            except Exception:
                return iso_str

        def _set_para_fmt(p, sb=None, sa=None, left=None, first=None, align=None):
            pf = p.paragraph_format
            if sb is not None:   pf.space_before      = Emu(sb)
            if sa is not None:   pf.space_after        = Emu(sa)
            if left is not None: pf.left_indent        = Emu(left)
            if first is not None:pf.first_line_indent  = Emu(first)
            if align is not None: p.alignment = align

        def add_run(p, text, bold=None, font=FONT_GUJ, size=None):
            r = p.add_run(text)
            r.bold = bold
            r.font.name = font
            if size:
                r.font.size = size
            return r

        def _remove_table_borders(tbl):
            """Remove all borders from a table (borderless style like demo)."""
            tbl_pr = tbl._tbl.tblPr
            if tbl_pr is None:
                tbl_pr = OxmlElement('w:tblPr')
                tbl._tbl.insert(0, tbl_pr)
            # Remove existing tblBorders if any
            existing = tbl_pr.find(qn('w:tblBorders'))
            if existing is not None:
                tbl_pr.remove(existing)
            tbl_borders = OxmlElement('w:tblBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'none')
                b.set(qn('w:sz'), '0')
                b.set(qn('w:space'), '0')
                b.set(qn('w:color'), 'auto')
                tbl_borders.append(b)
            tbl_pr.append(tbl_borders)

        def _remove_cell_borders(cell):
            """Remove all borders from a single cell."""
            tc_pr = cell._tc.get_or_add_tcPr()
            existing = tc_pr.find(qn('w:tcBorders'))
            if existing is not None:
                tc_pr.remove(existing)
            tc_borders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'none')
                b.set(qn('w:sz'), '0')
                b.set(qn('w:space'), '0')
                b.set(qn('w:color'), 'auto')
                tc_borders.append(b)
            tc_pr.append(tc_borders)

        def page_break(doc):
            p = doc.add_paragraph()
            run = p.add_run()
            br = OxmlElement('w:br')
            br.set(qn('w:type'), 'page')
            run._r.append(br)

        def _ensure_table_grid_style(doc):
            from docx.oxml.ns import qn as _qn
            from lxml import etree as _etree
            styles_el = doc.part.styles._element
            for s in styles_el.findall(_qn('w:style')):
                if s.get(_qn('w:styleId')) == 'TableGrid':
                    return
            _NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            xml = (
                '<w:style xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                ' w:type="table" w:styleId="TableGrid">'
                '<w:name w:val="Table Grid"/>'
                '<w:basedOn w:val="TableNormal"/>'
                '<w:uiPriority w:val="59"/>'
                '<w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
                '<w:tblPr>'
                '<w:tblBorders>'
                '<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                '<w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                '<w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                '</w:tblBorders>'
                '</w:tblPr>'
                '</w:style>'
            )
            styles_el.append(_etree.fromstring(xml))

        # ── Base document setup ───────────────────────────────────────────
        if lh_doc_template is not None:
            doc = lh_doc_template
            body = doc.element.body
            for child in list(body):
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if tag in ('p', 'tbl', 'sdt'):
                    body.remove(child)
            body.append(OxmlElement('w:p'))
            _ensure_table_grid_style(doc)
        else:
            doc = DocxDocument()
            sec = doc.sections[0]
            # Demo page: A4 portrait, margins from measurement
            sec.page_width    = Emu(7560945)   # 595.35pt
            sec.page_height   = Emu(10692765)  # 841.95pt
            sec.top_margin    = Emu(1600200)   # 112.5pt (112.5*12700=1428750 emu) — use measured 112.5pt
            sec.bottom_margin = Emu(403225)    # 28.35pt
            sec.left_margin   = Emu(1157415)   # 81.35pt
            sec.right_margin  = Emu(1044450)   # 73.45pt

        today_str = datetime.now().strftime('%d/%m/%Y')

        for idx, t in enumerate(trainees):
            if idx > 0:
                page_break(doc)

            info          = t.get('info', {})
            surname       = info.get('surname', '')
            firstname     = info.get('firstname', '')
            father        = info.get('father', '')
            addr1         = info.get('addr1', '')
            addr2         = info.get('addr2', '')
            pincode       = info.get('pincode', '')
            trade         = info.get('trade', '')
            si_name       = info.get('si_name', '')
            iti_name      = info.get('iti_name', '')

            pct           = t.get('percentage', '')
            absent_from   = fmt_date(t.get('absent_from', ''))
            att_upto      = fmt_date(t.get('attendance_upto', ''))
            prev_informed = fmt_date(t.get('prev_informed', ''))
            pct_str       = f"{pct}%" if pct is not None else ''
            _BLANK = '__________'
            absent_from_d   = absent_from   if absent_from   else _BLANK
            att_upto_d      = att_upto      if att_upto      else _BLANK
            prev_informed_d = prev_informed if prev_informed else _BLANK

            # ── TABLE 0: ક્રમાંક (bold, left) | તા. DATE (normal, right) ──────
            # Demo: 2-col equal width, borderless, space_before=38100 on both cells
            tbl0 = doc.add_table(rows=1, cols=2)
            tbl0.style = 'Table Grid'
            _remove_table_borders(tbl0)
            # Set equal column widths ~half of content width each (2797175 EMU each from demo)
            tbl0.columns[0].width = Emu(2797175)
            tbl0.columns[1].width = Emu(2797175)
            _remove_cell_borders(tbl0.rows[0].cells[0])
            _remove_cell_borders(tbl0.rows[0].cells[1])
            # Left cell: ક્રમાંક: — bold, left-aligned
            lc = tbl0.rows[0].cells[0]
            lp = lc.paragraphs[0]
            lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_para_fmt(lp, sb=38100)
            add_run(lp, 'ક્રમાંક', bold=True)
            add_run(lp, ':', bold=True)
            # Right cell: તા. DATE — normal, right-aligned
            rc = tbl0.rows[0].cells[1]
            rp = rc.paragraphs[0]
            rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_para_fmt(rp, sb=38100)
            add_run(rp, 'તા.  ', bold=None)
            add_run(rp, today_str, bold=None, font=FONT_SHRUTI)

            # ── ચેતવણી પત્ર — bold 14pt, centered, space_before=38100 ───────
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_para_fmt(p_title, sb=38100)
            add_run(p_title, 'ચેતવણી', bold=True, size=SIZE_TITLE)
            add_run(p_title, ' ', bold=True, size=SIZE_TITLE)
            add_run(p_title, 'પત્ર', bold=True, size=SIZE_TITLE)

            # ── પ્રતિ, block — bold, left, space_before=25400 ────────────────
            father_full = f'{surname} {father}'.strip()
            p_prati = doc.add_paragraph()
            p_prati.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_para_fmt(p_prati, sb=25400)
            add_run(p_prati, 'પ્રતિ', bold=True)
            add_run(p_prati, ',', bold=True)

            p_name = doc.add_paragraph()
            p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run(p_name, 'શ્રી ' + father_full, bold=True)

            if addr1:
                p_addr1 = doc.add_paragraph()
                p_addr1.alignment = WD_ALIGN_PARAGRAPH.LEFT
                add_run(p_addr1, 'એડ્રેસ:- ' + addr1, bold=True)
            if addr2:
                p_addr2 = doc.add_paragraph()
                p_addr2.alignment = WD_ALIGN_PARAGRAPH.LEFT
                add_run(p_addr2, addr2, bold=True)
            if pincode:
                p_pin = doc.add_paragraph()
                p_pin.alignment = WD_ALIGN_PARAGRAPH.LEFT
                add_run(p_pin, 'પીન- ' + pincode, bold=True)

            # ── વિષય: — all bold, centered, space_before=38100 ──────────────
            p_visay = doc.add_paragraph()
            p_visay.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_para_fmt(p_visay, sb=38100)
            add_run(p_visay, 'વિષય', bold=True)
            add_run(p_visay, ':  ', bold=True)
            add_run(p_visay, 'આપના', bold=True)
            add_run(p_visay, ' ', bold=True)
            add_run(p_visay, 'આશ્રિત', bold=True)
            add_run(p_visay, ' ', bold=True)
            add_run(p_visay, 'ચિરંજીવીની', bold=True)
            add_run(p_visay, ' ', bold=True)
            add_run(p_visay, 'સંસ્થામાં', bold=True)
            add_run(p_visay, ' ', bold=True)
            add_run(p_visay, 'તાલીમ', bold=True)
            add_run(p_visay, ' ', bold=True)
            add_run(p_visay, 'બાબત', bold=True)

            # ── શ્રીમાન, — normal, left ──────────────────────────────────────
            p_shri = doc.add_paragraph()
            p_shri.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run(p_shri, 'શ્રીમાન,', bold=None)

            # ── Opening paragraph ─────────────────────────────────────────────
            # Demo: first_line_indent=457200 (1 inch), space_before=38100, space_after=25400
            # left_indent=0
            trainee_display = f'{surname} {firstname}'.strip()
            p_open = doc.add_paragraph()
            p_open.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_para_fmt(p_open, sb=38100, sa=25400, left=0, first=457200)
            add_run(p_open, 'આપના સુપુત્ર/સુપુત્રી/આશ્રિત   ', bold=None)
            add_run(p_open, 'શ્રી', bold=True)
            add_run(p_open, '  ', bold=True)
            add_run(p_open, trainee_display, bold=True)
            add_run(p_open, '  ', bold=True)
            add_run(p_open, 'આ સંસ્થાના ', bold=None)
            add_run(p_open, 'ટ્રેડ', bold=True)
            add_run(p_open, ' ', bold=True)
            add_run(p_open, trade, bold=True)
            add_run(p_open, ' ', bold=None)
            add_run(p_open, 'મા   તાલીમ લે છે, તે માટે તેઓના ટ્રેડ ઇન્સ્ટ્રક્ટરશ્રીએ ', bold=None)
            add_run(p_open, 'તા', bold=True)
            add_run(p_open, '. ', bold=True)
            add_run(p_open, today_str, bold=True, font=FONT_SHRUTI)
            add_run(p_open, '  ના રોજ નીચે સહી કરનારને રિપોર્ટ કરેલ છે.', bold=None)

            # ── Points 1-6 ────────────────────────────────────────────────────
            # Demo: left_indent=228600, first_line_indent=-228600, sb=50800, sa=25400
            def add_point(num_text, segments):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _set_para_fmt(p, sb=50800, sa=25400, left=228600, first=-228600)
                add_run(p, num_text, bold=True)
                for text, bold, font in segments:
                    add_run(p, text, bold=bold, font=font or FONT_GUJ)
                return p

            # Point 1
            add_point('1.  ', [
                ('આપના સુપુત્ર/સુપુત્રી/આશ્રિત ', None, None),
                ('તારીખ', True, None),
                (' ', True, None),
                (absent_from_d, True, FONT_SHRUTI),
                (' ના રોજ થી સતત ગેરહાજર છે. જે આજ દિન સુધી સંસ્થામાં હાજર થયા નથી.'
                 ' જાણ કર્યા સિવાય તાલીમમાં ગેરહાજર રહેવું તે સંસ્થાના નિયમોનો શિસ્ત ભંગ થાય છે.'
                 ' આ પત્ર મળ્યે દિન-૪ માં હાજર થવાની સૂચના આપવામાં આવે છે.'
                 ' જો તેમ કરવામાં કસુર થશે તો ટ્રેનિંગ મેન્યુઅલના નિયમ-22 મુજબ'
                 ' આપના સુપુત્ર/સુપુત્રી/આશ્રિતનું નામ ભાગેડુ જાહેર કરી કમી કરવામાં આવશે. જેની ગંભીર નોંધ લેશો.', None, None),
            ])

            # Point 2
            add_point('2.  ', [
                ('આપના આશ્રિતનું ', None, None),
                ('તાલીમ', True, None),
                (' ', True, None),
                ('અંગેનું', True, None),
                (' ', True, None),
                ('સેશનલ', True, None),
                (' ', True, None),
                ('વર્ક', True, None),
                (' ', True, None),
                ('તથા', True, None),
                (' ', True, None),
                ('થીયરી', True, None),
                (', ', True, None),
                ('પ્રેક્ટીકલ', True, None),
                (' ', True, None),
                ('કાર્ય', True, None),
                (' ', True, None),
                ('બાકી', True, None),
                (' ', True, None),
                ('છે', True, None),
                (', જે પૂર્ણ કરવા જણાવવામાં આવે છે.', None, None),
            ])

            # Point 3
            add_point('3.  ', [
                ('આપના સુપુત્ર/સુપુત્રી/આશ્રિત માટે સંસ્થામાં અત્યાર સુધી તારીખ ', None, None),
                (att_upto_d, True, FONT_SHRUTI),
                ('  સુધી હાજરીના ', None, None),
                ('ટકા', True, None),
                ('  ', True, None),
                (pct_str + '  ', True, FONT_SHRUTI),
                ('છે, સંસ્થાના નિયમ મુજબ ', None, None),
                ('૮૦ ટકાથી  ઓછી હાજરી થાય', True, None),
                (' તો અંતિમ ', None, None),
                ('અ.ભા.વ્ય. કસોટીમાં', True, None),
                (' આપના સુપુત્ર/સુપુત્રી/આશ્રિતને ', None, None),
                ('બેસવા દેવામાં આવશે નહીં', True, None),
                ('.', None, None),
            ])

            # Point 4
            add_point('4.  ', [
                ('આપના સુપુત્ર/સુપુત્રી/આશ્રિત તાલીમ માટે સરકારી સાધનો વગેરે આપેલ છે,'
                 ' જે હજી સુધી પરત કરેલ નથી, જે તાત્કાલિક પરત કરવા,', None, None),
            ])

            # Point 5
            add_point('5.  ', [
                ('આ અગાઉ ', None, None),
                ('તારીખ', True, None),
                (' ', True, None),
                (prev_informed_d, True, FONT_SHRUTI),
                (' ના રોજ આપને જણાવવામાં આવેલ છે, આપના સુપુત્ર/સુપુત્રી/આશ્રિતને'
                 ' હાજર કરવા અંગેની આ છેલ્લી ચેતવણી જેવો પત્ર જેની નોંધ લેવી.', None, None),
            ])

            # Point 6
            add_point('6.  ', [
                ('તેઓ તાલીમમાં ખૂબ જ અનિયમિત છે.', None, None),
            ])

            # ── TABLE 1: Signature block ──────────────────────────────────────
            # Demo: 2-col borderless table, left col empty (space_before=355600 ~25pt),
            # right col center-aligned with 3 lines: આચાર્ય / ઔદ્યોગિક તાલીમ સંસ્થા / iti_name
            # Left col width=3900170, right col width=2700020
            tbl1 = doc.add_table(rows=1, cols=2)
            tbl1.style = 'Table Grid'
            _remove_table_borders(tbl1)
            tbl1.columns[0].width = Emu(3900170)
            tbl1.columns[1].width = Emu(2700020)
            _remove_cell_borders(tbl1.rows[0].cells[0])
            _remove_cell_borders(tbl1.rows[0].cells[1])
            # Left cell: empty spacer with large space_before for signature room
            lc1 = tbl1.rows[0].cells[0]
            lp1 = lc1.paragraphs[0]
            _set_para_fmt(lp1, sb=355600)
            # Right cell: centered text with 3 lines
            rc1 = tbl1.rows[0].cells[1]
            rp1 = rc1.paragraphs[0]
            rp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_para_fmt(rp1, sb=355600)
            add_run(rp1, 'આચાર્ય', bold=None)
            # Second line in right cell
            rp2 = rc1.add_paragraph()
            rp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(rp2, 'ઔદ્યોગિક તાલીમ સંસ્થા', bold=None)
            # Third line: ITI name (city/location from iti_name or addr)
            if iti_name:
                rp3 = rc1.add_paragraph()
                rp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run(rp3, iti_name, bold=None)

            # ── નકલ રવાના: ────────────────────────────────────────────────────
            # Demo: space_before=50800, normal, left-aligned
            p_cc_hdr = doc.add_paragraph()
            p_cc_hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_para_fmt(p_cc_hdr, sb=50800)
            add_run(p_cc_hdr, 'નકલ રવાના:', bold=None)

            # ── CC line: first_line_indent=457200, SI name bold, rest normal ──
            # Demo: "શ્રી [bold SI name]   સુ..ઈ. [normal] જાણ તથા...સારુ"
            p_cc = doc.add_paragraph()
            p_cc.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_para_fmt(p_cc, left=0, first=457200)
            add_run(p_cc, 'શ્રી ', bold=None)
            add_run(p_cc, si_name, bold=True)
            add_run(p_cc, '   સુ..ઈ. ', bold=None)
            add_run(p_cc, ' ', bold=None)
            add_run(p_cc, 'જાણ તથા તાલીમાર્થી આપેલ મુદતમાં હાજર ન થાય તો'
                          ' તાલીમાર્થીની વિગતો સાથે જરૂરી રિપોર્ટ કરવા સારુ', bold=None)

        # ── Save ─────────────────────────────────────────────────────────────
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        return send_file(out,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True,
                         download_name='Warning_Letters.docx')

    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'detail': traceback.format_exc()}), 500



@app.route('/api/admin/login', methods=['POST'])
def api_admin_login():
    pwd = request.json.get('password','')
    with get_db() as db:
        admin = db.execute('SELECT * FROM admin WHERE id=1').fetchone()
    if admin and check_password_hash(admin['password_hash'], pwd):
        session['is_admin'] = True
        return jsonify({'message': 'Admin login successful'})
    return jsonify({'error': 'Invalid admin password'}), 401

@app.route('/api/admin/logout', methods=['POST'])
def api_admin_logout():
    session.pop('is_admin', None); return jsonify({'message': 'Logged out'})

@app.route('/api/admin/change-password', methods=['POST'])
def api_admin_change_password():
    if not session.get('is_admin'):
        return jsonify({'error': 'Unauthorized'}), 401
    d = request.json
    current_pwd = d.get('current_password', '')
    new_pwd = d.get('new_password', '')
    if not current_pwd or not new_pwd:
        return jsonify({'error': 'Both current and new password are required'}), 400
    if len(new_pwd) < 6:
        return jsonify({'error': 'New password must be at least 6 characters'}), 400
    with get_db() as db:
        admin = db.execute('SELECT * FROM admin WHERE id=1').fetchone()
        if not admin or not check_password_hash(admin['password_hash'], current_pwd):
            return jsonify({'error': 'Current password is incorrect'}), 401
        db.execute('UPDATE admin SET password_hash=? WHERE id=1', (generate_password_hash(new_pwd),))
        db.commit()
    return jsonify({'message': 'Password updated successfully'})

@app.route('/api/admin/users')
def api_admin_users():
    if not session.get('is_admin'): return jsonify({'error': 'Unauthorized'}), 401
    with get_db() as db:
        users = db.execute('SELECT id,si_name,trade_name,iti_name,mobile,status,user_id,created_at,approved_at,deregistered_at FROM users ORDER BY id DESC').fetchall()
    return jsonify([dict(u) for u in users])

@app.route('/api/admin/approve/<int:uid>', methods=['POST'])
def api_admin_approve(uid):
    if not session.get('is_admin'): return jsonify({'error': 'Unauthorized'}), 401
    with get_db() as db:
        user = db.execute('SELECT * FROM users WHERE id=?', (uid,)).fetchone()
        if not user: return jsonify({'error': 'User not found'}), 404
        if user['status'] == 'approved': return jsonify({'error': 'Already approved'}), 400
        new_uid = gen_user_id(user['mobile'])
        pwd = gen_password(user['mobile'])
        pwd_hash = generate_password_hash(pwd)
        db.execute('UPDATE users SET status=?, user_id=?, password_hash=?, approved_at=? WHERE id=?',
                   ('approved', new_uid, pwd_hash, datetime.now().strftime('%Y-%m-%d %H:%M:%S'), uid))
        db.commit()
    return jsonify({'user_id': new_uid, 'password': pwd, 'message': 'User approved'})

@app.route('/api/admin/disapprove/<int:uid>', methods=['POST'])
def api_admin_disapprove(uid):
    if not session.get('is_admin'): return jsonify({'error': 'Unauthorized'}), 401
    with get_db() as db:
        user = db.execute('SELECT * FROM users WHERE id=?', (uid,)).fetchone()
        if not user: return jsonify({'error': 'User not found'}), 404
        if user['status'] != 'approved': return jsonify({'error': 'User is not approved'}), 400
        db.execute("UPDATE users SET status='disapproved', user_id=NULL, password_hash=NULL, deregistered_at=? WHERE id=?",
                   (datetime.now().strftime('%Y-%m-%d %H:%M:%S'), uid))
        db.commit()
    return jsonify({'message': 'User disapproved successfully'})

@app.route('/api/admin/clear-deregistered', methods=['POST'])
def api_admin_clear_deregistered():
    if not session.get('is_admin'): return jsonify({'error': 'Unauthorized'}), 401
    with get_db() as db:
        db.execute("DELETE FROM users WHERE status='disapproved'")
        db.commit()
    return jsonify({'message': 'Deregistered users cleared successfully'})

# ─── GENERATE ROUTE (protected) ───────────────────────────────────

# ─── ANNEXURE GENERATION ──────────────────────────────────────────────────────

def _build_annexure_wb(ann_type, trainee_df, user_info):
    """
    Build ANNEXURE_III_{ann_type} Excel workbook from scratch.
    No template file needed — structure is fully self-contained in code.

    Layout (openpyxl column numbers):
        Col 1 (A) = Roll No
        Col 2-3 (B:C) = Name (merged)
        Col 4 (D) = Attendance (max 5)
        Col 5-6 (E:F) = Speed/Accuracy/Communication (max 5, merged)
        Col 7-8 (G:H) = Creative Work (max 10, merged)
        Col 9 (I) = Quarterly-1 (max 20)
        Col 10 (J) = Quarterly-2 (max 20)
        Col 11 (K) = Total (max 60)
        Col 12-13 (L:M) = Total/6 out of 10 (merged)
        Col 14 (N) = Sign of Trainee

    Data columns in trainee file:
        ES   -> Employability Skills mark (out of 10)
        WC_SC -> Workshop Calc & Science (out of 10)
        ED   -> Engineering Drawing (out of 10)

    Logic: t10 = mark from file (0-10), t60 = t10*6, distribute t60 across sub-cols.
    """
    import math as _m
    from openpyxl import Workbook as _WB
    from openpyxl.styles import (Font as _Font, PatternFill as _Fill,
                                  Alignment as _Align, Border as _Border, Side as _Side)
    from openpyxl.utils import get_column_letter as _gcl

    mark_col = {'ES': 'ES', 'WCS': 'WC_SC', 'ED': 'ED'}[ann_type]
    title_map = {
        'ES':  'FORMAT FOR INTERNAL ASSESSMENT FOR EMPLOYABILITY SKILLS',
        'WCS': 'FORMAT FOR INTERNAL ASSESSMENT FOR WORK SHOP CALCULATION AND SCIENCE',
        'ED':  'FORMAT FOR INTERNAL ASSESSMENT FOR ENGINEERING DRAWING',
    }
    speed_label = {
        'ES':  'Communication Skill for ES',
        'WCS': 'Speed for WC & Sc',
        'ED':  'Accuracy of ED',
    }

    ui = user_info or {}

    # ── Styles ────────────────────────────────────────────────────────
    NAVY   = '0B1D3A'
    BLUE   = '1B4F8A'
    HDR2   = '2E75B6'
    YELLOW = 'FFF2CC'
    WHITE  = 'FFFFFF'
    LGRAY  = 'D9E1F2'
    thin   = _Side(style='thin', color='000000')
    bdr    = _Border(left=thin, right=thin, top=thin, bottom=thin)

    def _sf(size=9, bold=False, color='000000'):
        return _Font(name='Calibri', size=size, bold=bold, color=color)

    def _hf(size=9, color='000000'):
        return _Font(name='Calibri', size=size, bold=True, color=color)

    def _ca(h='center', v='center', wrap=True):
        return _Align(horizontal=h, vertical=v, wrap_text=wrap)

    def _fill(hex_color):
        return _Fill("none")

    def _set(ws, row, col, val, font=None, fill=None, align=None, border=None):
        c = ws.cell(row=row, column=col, value=val)
        if font:   c.font   = font
        if fill:   c.fill   = fill
        if align:  c.alignment = align
        if border: c.border = border
        return c

    def _merge_set(ws, row, c1, c2, val, font=None, fill=None, align=None, border=None):
        ws.merge_cells(f'{_gcl(c1)}{row}:{_gcl(c2)}{row}')
        _set(ws, row, c1, val, font, fill, align, border)
        # Apply border to all cells in merged range for clean look
        for c in range(c1, c2+1):
            cell = ws.cell(row=row, column=c)
            if border: cell.border = border

    # ── Build workbook ────────────────────────────────────────────────
    wb = _WB()
    ws = wb.active
    ws.title = f'ANNEXURE_III_{ann_type}'

    # ── Row 1: Main title ─────────────────────────────────────────────
    _merge_set(ws, 1, 1, 14, 'ANNEXURE-III (FAR-2)',
               font=_hf(12, '000000'), fill=_fill('EBF3FB'), align=_ca(), border=bdr)

    # ── Row 2: Sub-title ─────────────────────────────────────────────
    _merge_set(ws, 2, 1, 14, 'Internal Assessment',
               font=_hf(11, '000000'), fill=_fill('EBF3FB'), align=_ca(), border=bdr)

    # ── Row 3: Format description ─────────────────────────────────────
    _merge_set(ws, 3, 1, 14, title_map[ann_type],
               font=_hf(10, '000000'), fill=_fill(BLUE), align=_ca(), border=bdr)

    # ── Rows 4-8: Header info fields ──────────────────────────────────
    lbl_font  = _sf(9, bold=True, color='000000')
    val_font  = _sf(9, bold=False, color='1A1A1A')
    lbl_fill  = _fill(LGRAY)
    val_fill  = _fill('FAFBFF')

    def _hdr_row(row, lbl1, val1, lbl2, val2):
        _merge_set(ws, row, 1, 3,  lbl1, font=lbl_font, fill=lbl_fill, align=_ca('left'), border=bdr)
        _merge_set(ws, row, 4, 7,  val1, font=val_font, fill=val_fill, align=_ca('left'), border=bdr)
        _merge_set(ws, row, 8, 11, lbl2, font=lbl_font, fill=lbl_fill, align=_ca('left'), border=bdr)
        _merge_set(ws, row, 12, 14, val2, font=val_font, fill=val_fill, align=_ca('left'), border=bdr)

    _hdr_row(4, 'Name & Address of the Assessor', ui.get('si_name',''),
                'Year of Enrolment', ui.get('year_of_assessment',''))
    _hdr_row(5, 'Name & Address of ITI (Govt/Pvt)', ui.get('iti_name',''),
                'Date of Assessment', '')
    _hdr_row(6, 'Name & Address of the Industry', ui.get('iti_name',''),
                'Assessment Location', ui.get('assessment_location',''))

    # Row 7: Trade / Duration / Examination(Semester)
    _merge_set(ws, 7, 1, 2,  'Trade Name', font=lbl_font, fill=lbl_fill, align=_ca('left'), border=bdr)
    _merge_set(ws, 7, 3, 7,  ui.get('trade_name',''), font=val_font, fill=val_fill, align=_ca('left'), border=bdr)
    _merge_set(ws, 7, 8, 10, 'Duration Of Trade', font=lbl_font, fill=lbl_fill, align=_ca('left'), border=bdr)
    _merge_set(ws, 7, 11, 11, ui.get('trade_duration',''), font=val_font, fill=val_fill, align=_ca(), border=bdr)
    _merge_set(ws, 7, 12, 13, 'Examination', font=lbl_font, fill=lbl_fill, align=_ca('left'), border=bdr)
    _set(ws, 7, 14, ui.get('semester',''), font=val_font, align=_ca(), border=bdr)

    # Row 8: Learning Outcome / Batch
    _merge_set(ws, 8, 1, 7,  'Learning Outcome :', font=lbl_font, fill=lbl_fill, align=_ca('left'), border=bdr)
    _merge_set(ws, 8, 8, 11, 'Batch NO', font=lbl_font, fill=lbl_fill, align=_ca('left'), border=bdr)
    _merge_set(ws, 8, 12, 14, ui.get('batch',''), font=val_font, fill=val_fill, align=_ca('left'), border=bdr)

    # ── Row 9: Column headers ─────────────────────────────────────────
    hdr_fill = _fill(BLUE)
    hdr_font = _hf(8, '000000')
    hdrs = [
        (1,  1,  'Roll\nNo'),
        (2,  3,  'Name'),
        (4,  4,  'Attendance'),
        (5,  6,  speed_label[ann_type]),
        (7,  8,  'Creative Work\n(Chart, Model, Poster, Project work etc.)'),
        (9,  9,  'Quarterly -1'),
        (10, 10, 'Quarterly -2'),
        (11, 11, 'Total'),
        (12, 13, 'Convert Total Marks into 10 Marks = {(Col.K)/6}'),
        (14, 14, 'Sign of Trainee'),
    ]
    for (c1, c2, label) in hdrs:
        if c1 == c2:
            _set(ws, 9, c1, label, font=hdr_font, fill=hdr_fill, align=_ca(), border=bdr)
        else:
            _merge_set(ws, 9, c1, c2, label, font=hdr_font, fill=hdr_fill, align=_ca(), border=bdr)

    # ── Row 10: Maximum marks ─────────────────────────────────────────
    mx_font = _hf(8, '000000')
    mx_fill = _fill(YELLOW)
    _merge_set(ws, 10, 1, 3, 'Maximum Marks =>', font=mx_font, fill=mx_fill, align=_ca(), border=bdr)
    for col, val in [(4,'5'),(5,'5'),(7,'10'),(9,'20'),(10,'20'),(11,'60'),(12,''),(14,'')]:
        if col in (5,):
            _merge_set(ws, 10, 5, 6, '5', font=mx_font, fill=mx_fill, align=_ca(), border=bdr)
        elif col in (7,):
            _merge_set(ws, 10, 7, 8, '10', font=mx_font, fill=mx_fill, align=_ca(), border=bdr)
        elif col in (12,):
            _merge_set(ws, 10, 12, 13, '', font=mx_font, fill=mx_fill, align=_ca(), border=bdr)
        else:
            _set(ws, 10, col, val, font=mx_font, fill=mx_fill, align=_ca(), border=bdr)

    # ── Row 11: Column letter labels ─────────────────────────────────
    ltr_font = _hf(8, '000000')
    ltr_fill = _fill('D6E4F0')
    for col, lbl in [(1,'A'),(4,'B'),(5,'C'),(7,'D'),(9,'E'),(10,'F'),(11,'G'),(12,'H'),(14,'I')]:
        _set(ws, 11, col, lbl, font=ltr_font, fill=ltr_fill, align=_ca(), border=bdr)
    _merge_set(ws, 11, 2, 3, '', font=ltr_font, fill=ltr_fill, align=_ca(), border=bdr)
    _merge_set(ws, 11, 5, 6, 'C', font=ltr_font, fill=ltr_fill, align=_ca(), border=bdr)
    _merge_set(ws, 11, 7, 8, 'D', font=ltr_font, fill=ltr_fill, align=_ca(), border=bdr)
    _merge_set(ws, 11, 12, 13, 'H', font=ltr_font, fill=ltr_fill, align=_ca(), border=bdr)

    # ── Distribute helper ─────────────────────────────────────────────
    def _dist(target_60, maxes):
        """Return integers summing exactly to target_60, each bounded by maxes."""
        t = max(len(maxes), min(sum(maxes), target_60))
        vals = [max(1, int(t * m / sum(maxes))) for m in maxes]
        diff = t - sum(vals)
        idxs = list(range(len(maxes)))
        random.shuffle(idxs)
        for i in idxs:
            if diff == 0: break
            if diff > 0:
                add = min(diff, maxes[i] - vals[i]); vals[i] += add; diff -= add
            else:
                sub = min(-diff, vals[i] - 1);       vals[i] -= sub; diff += sub
        return vals

    # Sub-col maxes: Attendance(5), Speed(5), Creative(10), Q1(20), Q2(20)
    MAXES = [5, 5, 10, 20, 20]
    # openpyxl column indices for sub-marks: D=4, E(merged E:F)=5, G(merged G:H)=7, I=9, J=10
    WCOLS = [4, 5, 7, 9, 10]

    # ── Data rows ─────────────────────────────────────────────────────
    data_font = _sf(9)
    alt_fill  = _fill('EBF3FB')

    for idx, (_, trainee) in enumerate(trainee_df.iterrows()):
        r = 12 + idx
        row_fill = _fill(WHITE) if idx % 2 == 0 else alt_fill

        # Roll No
        roll = trainee.get('rollno', trainee.get('rollno ', idx + 2))
        _set(ws, r, 1, roll, font=data_font, fill=row_fill, align=_ca(), border=bdr)

        # Name (merged B:C)
        name = ' '.join(str(trainee.get(f, '') or '')
                        for f in ['Firstname', 'Fathername', 'Lastname']).strip()
        _merge_set(ws, r, 2, 3, name, font=data_font, fill=row_fill, align=_ca('left'), border=bdr)

        # Target mark out of 10 from trainee file
        raw = trainee.get(mark_col)
        if raw is None or (isinstance(raw, float) and _m.isnan(raw)):
            t10 = 5
        else:
            t10 = max(0, min(10, int(round(float(raw)))))

        t60  = t10 * 6
        subs = _dist(t60, MAXES)

        # Write sub-marks (E:F and G:H are merged, so write to E and G)
        col_map = list(zip(WCOLS, subs))
        for col, val in col_map:
            if col == 5:   # E:F merged
                _merge_set(ws, r, 5, 6, val, font=data_font, fill=row_fill, align=_ca(), border=bdr)
            elif col == 7: # G:H merged
                _merge_set(ws, r, 7, 8, val, font=data_font, fill=row_fill, align=_ca(), border=bdr)
            else:
                _set(ws, r, col, val, font=data_font, fill=row_fill, align=_ca(), border=bdr)

        # Total (K=col11)
        _set(ws, r, 11, t60, font=_sf(9, bold=True), fill=row_fill, align=_ca(), border=bdr)

        # H = Total/6 (L:M merged = col12:13)
        _merge_set(ws, r, 12, 13, t10, font=_sf(9, bold=True, color='000000'),
                   fill=row_fill, align=_ca(), border=bdr)

        # Sign column (N=col14) — blank
        _set(ws, r, 14, '', font=data_font, fill=row_fill, align=_ca(), border=bdr)

    # ── Column widths ─────────────────────────────────────────────────
    col_widths = {1:8, 2:18, 3:6, 4:10, 5:12, 6:6, 7:12, 8:6, 9:10, 10:10, 11:8, 12:12, 13:6, 14:12}
    for col, w in col_widths.items():
        ws.column_dimensions[_gcl(col)].width = w

    # Row heights
    for r, h in {1:20, 2:18, 3:18, 9:50, 10:18, 11:16}.items():
        ws.row_dimensions[r].height = h

    ws.freeze_panes = 'A12'

    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    return out


@app.route('/generate-annexure-es', methods=['POST'])
def generate_annexure_es():
    return _serve_annexure('ES')

@app.route('/generate-annexure-wcs', methods=['POST'])
def generate_annexure_wcs():
    return _serve_annexure('WCS')

@app.route('/generate-annexure-ed', methods=['POST'])
def generate_annexure_ed():
    return _serve_annexure('ED')

def _serve_annexure(ann_type):
    if 'user_id' not in session:
        return jsonify({'error': 'Please login first'}), 401
    try:
        uid = session.get('user_id')
        trainee_df = _get_df(uid, 'trainee', request.files.get('trainee'))
        if trainee_df is None:
            return jsonify({'error': 'No trainee file found. Please upload the Trainee Details file first.'}), 400
        user_info = {k: session.get(k,'') for k in
                     ['si_name','trade_name','iti_name','year_of_assessment',
                      'assessment_location','trade_duration','semester','batch']}
        out = _build_annexure_wb(ann_type, trainee_df, user_info)
        return send_file(out,
                         mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                         as_attachment=True,
                         download_name=f'ANNEXURE_III_{ann_type}.xlsx')
    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'detail': traceback.format_exc()}), 500

@app.route('/generate-lo-summary', methods=['POST'])
def generate_lo_summary():
    if 'user_id' not in session:
        return jsonify({'error': 'Please login first'}), 401
    try:
        uid = session.get('user_id')
        trainee_df = _get_df(uid, 'trainee', request.files.get('trainee'))
        lo_df      = _get_df(uid, 'lo',      request.files.get('lo'))
        if trainee_df is None or lo_df is None:
            return jsonify({'error': 'Both files are required. Please upload Trainee Details and LO Details files.'}), 400
        user_info = {k: session.get(k, '') for k in
                     ['si_name', 'trade_name', 'iti_name', 'year_of_assessment',
                      'assessment_location', 'near_trade', 'trade_duration', 'semester', 'batch']}
        # Reuse the marks cache from the main report generation so numbers match exactly.
        # Fall back to building fresh cache only if user skipped main report generation.
        t_hash = _file_hash(_load_user_file(uid, 'trainee'))
        l_hash = _file_hash(_load_user_file(uid, 'lo'))
        marks_cache = (_marks_cache_store.get(uid)
                       or _load_marks_cache_disk(uid, t_hash, l_hash))
        if marks_cache is None:
            marks_cache = build_all_marks(trainee_df, lo_df)
            _save_marks_cache_disk(uid, marks_cache, t_hash, l_hash)
            _marks_cache_store[uid] = marks_cache
        output = create_lo_summary_excel(trainee_df, lo_df, user_info, marks_cache=marks_cache)
        return send_file(output,
                         mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                         as_attachment=True,
                         download_name='LO_Wise_Summary.xlsx')
    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'detail': traceback.format_exc()}), 500


@app.route('/generate-trainee-summary', methods=['POST'])
def generate_trainee_summary():
    if 'user_id' not in session:
        return jsonify({'error': 'Please login first'}), 401
    try:
        uid        = session.get('user_id')
        trainee_df = _get_df(uid, 'trainee', request.files.get('trainee'))
        lo_df      = _get_df(uid, 'lo',      request.files.get('lo'))
        if trainee_df is None or lo_df is None:
            return jsonify({'error': 'Both Trainee and LO files are required.'}), 400
        user_info = {k: session.get(k, '') for k in
                     ['si_name', 'trade_name', 'iti_name', 'year_of_assessment',
                      'assessment_location', 'near_trade', 'trade_duration', 'semester', 'batch']}
        t_hash     = _file_hash(_load_user_file(uid, 'trainee'))
        l_hash     = _file_hash(_load_user_file(uid, 'lo'))
        marks_cache = (_marks_cache_store.get(uid)
                       or _load_marks_cache_disk(uid, t_hash, l_hash))
        if marks_cache is None:
            marks_cache = build_all_marks(trainee_df, lo_df)
            _marks_cache_store[uid] = marks_cache
            _save_marks_cache_disk(uid, marks_cache, t_hash, l_hash)
        output = create_trainee_summary_excel(trainee_df, lo_df, user_info, marks_cache=marks_cache)
        return send_file(output,
                         mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                         as_attachment=True,
                         download_name='Trainee_Wise_Summary.xlsx')
    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'detail': traceback.format_exc()}), 500


@app.route('/generate-trainee-summary-pdf', methods=['POST'])
def generate_trainee_summary_pdf():
    if 'user_id' not in session:
        return jsonify({'error': 'Please login first'}), 401
    if not _REPORTLAB_OK:
        return jsonify({'error': 'ReportLab not installed on server. Run: pip install reportlab'}), 500
    try:
        uid        = session.get('user_id')
        trainee_df = _get_df(uid, 'trainee', request.files.get('trainee'))
        lo_df      = _get_df(uid, 'lo',      request.files.get('lo'))
        if trainee_df is None or lo_df is None:
            return jsonify({'error': 'Both Trainee and LO files are required.'}), 400
        user_info = {k: session.get(k, '') for k in
                     ['si_name', 'trade_name', 'iti_name', 'year_of_assessment',
                      'assessment_location', 'near_trade', 'trade_duration', 'semester', 'batch']}
        t_hash     = _file_hash(_load_user_file(uid, 'trainee'))
        l_hash     = _file_hash(_load_user_file(uid, 'lo'))
        marks_cache = (_marks_cache_store.get(uid)
                       or _load_marks_cache_disk(uid, t_hash, l_hash))
        if marks_cache is None:
            marks_cache = build_all_marks(trainee_df, lo_df)
            _marks_cache_store[uid] = marks_cache
            _save_marks_cache_disk(uid, marks_cache, t_hash, l_hash)
        output = create_trainee_summary_pdf(trainee_df, lo_df, user_info, marks_cache=marks_cache)
        return send_file(output, mimetype='application/pdf',
                         as_attachment=True, download_name='Trainee_Wise_Summary.pdf')
    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'detail': traceback.format_exc()}), 500


@app.route('/generate-progress-card', methods=['POST'])
def generate_progress_card():
    if 'user_id' not in session:
        return jsonify({'error': 'Please login first'}), 401
    try:
        uid = session.get('user_id')
        trainee_df = _get_df(uid, 'trainee', request.files.get('trainee'))
        if trainee_df is None:
            return jsonify({'error': 'No trainee file found. Please upload the Trainee Details file first.'}), 400
        user_info = {k: session.get(k, '') for k in
                     ['si_name', 'trade_name', 'iti_name', 'year_of_assessment',
                      'assessment_location', 'near_trade', 'trade_duration', 'semester', 'batch']}
        output = create_progress_card(trainee_df, user_info)
        return send_file(output,
                         mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                         as_attachment=True,
                         download_name='Progress_Cards.xlsx')
    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'detail': traceback.format_exc()}), 500


@app.route('/generate-trainee-lo-report', methods=['POST'])
def generate_trainee_lo_report():
    if 'user_id' not in session:
        return jsonify({'error': 'Please login first'}), 401
    try:
        uid = session.get('user_id')
        trainee_df = _get_df(uid, 'trainee', request.files.get('trainee'))
        lo_df      = _get_df(uid, 'lo',      request.files.get('lo'))
        if trainee_df is None or lo_df is None:
            return jsonify({'error': 'Both files are required. Please upload Trainee Details and LO Details files.'}), 400
        user_info = {k: session.get(k, '') for k in
                     ['si_name', 'trade_name', 'iti_name', 'year_of_assessment',
                      'assessment_location', 'near_trade', 'trade_duration', 'semester', 'batch']}
        t_hash = _file_hash(_load_user_file(uid, 'trainee'))
        l_hash = _file_hash(_load_user_file(uid, 'lo'))
        marks_cache = (_marks_cache_store.get(uid)
                       or _load_marks_cache_disk(uid, t_hash, l_hash))
        if marks_cache is None:
            marks_cache = build_all_marks(trainee_df, lo_df)
            _save_marks_cache_disk(uid, marks_cache, t_hash, l_hash)
            _marks_cache_store[uid] = marks_cache
        output = create_trainee_lo_report(trainee_df, lo_df, user_info, marks_cache=marks_cache)
        return send_file(output,
                         mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                         as_attachment=True,
                         download_name='Per_Trainee_Per_LO_Report.xlsx')
    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'detail': traceback.format_exc()}), 500


# ─── PDF: ANNEXURE III ────────────────────────────────────────────────────────

def create_annexure_pdf(ann_type, trainee_df, user_info=None):
    """
    Generate print-ready A4 portrait PDF for Annexure III (ES / WCS / ED).
    Mirrors _build_annexure_wb exactly — same headers, same data, same layout.
    """
    if not _REPORTLAB_OK:
        raise RuntimeError("ReportLab is not installed. Run: pip install reportlab")

    import math as _m

    ui = user_info or {}
    mark_col   = {'ES': 'ES', 'WCS': 'WC_SC', 'ED': 'ED'}[ann_type]
    title_map  = {
        'ES':  'FORMAT FOR INTERNAL ASSESSMENT FOR EMPLOYABILITY SKILLS',
        'WCS': 'FORMAT FOR INTERNAL ASSESSMENT FOR WORK SHOP CALCULATION AND SCIENCE',
        'ED':  'FORMAT FOR INTERNAL ASSESSMENT FOR ENGINEERING DRAWING',
    }
    speed_label = {
        'ES':  'Communication\nSkill for ES',
        'WCS': 'Speed for\nWC & Sc',
        'ED':  'Accuracy\nof ED',
    }

    PAGE_W, PAGE_H = A4          # portrait
    L = R = 10 * mm
    T = B = 10 * mm
    USABLE_W = PAGE_W - L - R

    NAVY  = colors.HexColor('#0B1D3A')
    BLUE  = colors.HexColor('#1B4F8A')
    LBLUE = colors.HexColor('#D6E4F0')
    GOLD  = colors.HexColor('#FFF2CC')
    BLACK = colors.black
    LGRAY = colors.HexColor('#F8FAFC')

    def _ps(size=7, bold=False, align=TA_CENTER):
        return ParagraphStyle('_', fontSize=size, leading=size + 1.5,
                              fontName='Helvetica-Bold' if bold else 'Helvetica',
                              textColor=BLACK, alignment=align,
                              wordWrap='LTR', spaceBefore=0, spaceAfter=0)

    def P(txt, size=7, bold=False, align=TA_CENTER):
        return Paragraph(str(txt), _ps(size, bold, align))

    def PL(txt, size=7, bold=False):
        return Paragraph(str(txt), _ps(size, bold, TA_LEFT))

    # column widths: Roll | Name | Att | Speed | Creative | Q1 | Q2 | Total | /6 | Sign
    raw_w = [1.0, 3.5, 1.2, 1.5, 2.0, 1.5, 1.5, 1.2, 1.5, 1.8]
    tot = sum(raw_w)
    CW = [w / tot * USABLE_W for w in raw_w]

    def _dist(target_60, maxes):
        t = max(len(maxes), min(sum(maxes), target_60))
        vals = [max(1, int(t * m / sum(maxes))) for m in maxes]
        diff = t - sum(vals)
        idxs = list(range(len(maxes)))
        random.shuffle(idxs)
        for i in idxs:
            if diff == 0: break
            if diff > 0:
                add = min(diff, maxes[i] - vals[i]); vals[i] += add; diff -= add
            else:
                sub = min(-diff, vals[i] - 1);       vals[i] -= sub; diff += sub
        return vals

    MAXES = [5, 5, 10, 20, 20]

    buf = io.BytesIO()

    def _on_page(canvas, doc):
        canvas.saveState()
        canvas.setStrokeColor(BLUE)
        canvas.setLineWidth(0.6)
        canvas.rect(L - 2, B - 2, PAGE_W - L - R + 4, PAGE_H - T - B + 4)
        canvas.restoreState()

    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=L, rightMargin=R,
                            topMargin=T, bottomMargin=B)
    story = []

    # ── Title block ──────────────────────────────────────────────────────
    story.append(Table([[P('ANNEXURE-III (FAR-2)', size=10, bold=True)]],
                       colWidths=[USABLE_W],
                       style=TableStyle([('BOX',(0,0),(-1,-1),0.5,BLACK),
                                         ('BACKGROUND',(0,0),(-1,-1),LBLUE),
                                         ('TOPPADDING',(0,0),(-1,-1),3),
                                         ('BOTTOMPADDING',(0,0),(-1,-1),3)])))
    story.append(Table([[P('Internal Assessment', size=9, bold=True)]],
                       colWidths=[USABLE_W],
                       style=TableStyle([('BOX',(0,0),(-1,-1),0.5,BLACK),
                                         ('TOPPADDING',(0,0),(-1,-1),2),
                                         ('BOTTOMPADDING',(0,0),(-1,-1),2)])))
    story.append(Table([[P(title_map[ann_type], size=8, bold=True)]],
                       colWidths=[USABLE_W],
                       style=TableStyle([('BOX',(0,0),(-1,-1),0.5,BLACK),
                                         ('BACKGROUND',(0,0),(-1,-1),LBLUE),
                                         ('TOPPADDING',(0,0),(-1,-1),2),
                                         ('BOTTOMPADDING',(0,0),(-1,-1),2)])))

    # ── Header info (label | value | label | value) ──────────────────────
    hw = [USABLE_W*0.20, USABLE_W*0.30, USABLE_W*0.22, USABLE_W*0.28]
    info_rows = [
        ('Name & Address of the Assessor', ui.get('si_name',''),
         'Year of Enrolment', ui.get('year_of_assessment','')),
        ('Name & Address of ITI (Govt/Pvt)', ui.get('iti_name',''),
         'Date of Assessment', ''),
        ('Name & Address of the Industry', ui.get('iti_name',''),
         'Assessment Location', ui.get('assessment_location','')),
        ('Trade Name', ui.get('trade_name',''),
         'Duration Of Trade', ui.get('trade_duration','')),
        ('Learning Outcome :', '',
         'Batch NO', ui.get('batch','')),
    ]
    for (l1, v1, l2, v2) in info_rows:
        story.append(Table(
            [[PL(l1, bold=True), PL(v1), PL(l2, bold=True), PL(v2)]],
            colWidths=hw,
            style=TableStyle([('BOX',(0,0),(-1,-1),0.5,BLACK),
                               ('INNERGRID',(0,0),(-1,-1),0.3,BLACK),
                               ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
                               ('TOPPADDING',(0,0),(-1,-1),2),
                               ('BOTTOMPADDING',(0,0),(-1,-1),2),
                               ('LEFTPADDING',(0,0),(-1,-1),3),
                               ('RIGHTPADDING',(0,0),(-1,-1),3)])))

    # ── Column headers ────────────────────────────────────────────────────
    hdr_row = [P('Roll\nNo',7,True), P('Name',7,True),
               P('Attend-\nance',7,True), P(speed_label[ann_type],7,True),
               P('Creative Work\n(Chart/Model/\nPoster/Project)',7,True),
               P('Quarterly\n-1',7,True), P('Quarterly\n-2',7,True),
               P('Total',7,True),
               P('Total/6\n(out of 10)',7,True),
               P('Sign of\nTrainee',7,True)]
    max_row  = [P('',6,True), P('',6,True),
                P('5',6,True), P('5',6,True), P('10',6,True),
                P('20',6,True), P('20',6,True), P('60',6,True),
                P('10',6,True), P('',6,True)]
    lbl_row  = [P('A',6,True), P('',6,True),
                P('B',6,True), P('C',6,True), P('D',6,True),
                P('E',6,True), P('F',6,True), P('G',6,True),
                P('H',6,True), P('I',6,True)]

    hdr_tbl = Table([hdr_row, max_row, lbl_row], colWidths=CW)
    hdr_tbl.setStyle(TableStyle([
        ('GRID',(0,0),(-1,-1),0.4,BLACK),
        ('BACKGROUND',(0,0),(-1,0),LBLUE),
        ('BACKGROUND',(0,1),(-1,1),GOLD),
        ('BACKGROUND',(0,2),(-1,2),LGRAY),
        ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
        ('ALIGN',(0,0),(-1,-1),'CENTER'),
        ('TOPPADDING',(0,0),(-1,-1),2),
        ('BOTTOMPADDING',(0,0),(-1,-1),2),
    ]))
    story.append(hdr_tbl)

    # ── Data rows ─────────────────────────────────────────────────────────
    data_rows = []
    for idx, (_, trainee) in enumerate(trainee_df.iterrows()):
        roll = trainee.get('rollno', idx + 1)
        name = ' '.join(str(trainee.get(f,'') or '') for f in ['Firstname','Fathername','Lastname']).strip()
        raw  = trainee.get(mark_col)
        if raw is None or (isinstance(raw, float) and _m.isnan(raw)):
            t10 = 5
        else:
            t10 = max(0, min(10, int(round(float(raw)))))
        t60  = t10 * 6
        subs = _dist(t60, MAXES)
        bg   = LGRAY if idx % 2 == 0 else colors.white
        data_rows.append((roll, name, subs[0], subs[1], subs[2], subs[3], subs[4], t60, t10, bg))

    if data_rows:
        # Row height: 24 trainees fills one full portrait page exactly.
        # ≤24 → use fixed base height (page fills nicely, never oversized)
        # >24 → shrink proportionally so all rows still fit on one page
        HEADER_H_PT = 155
        BASELINE_N  = 24
        available_h = (PAGE_H - T - B) - HEADER_H_PT
        base_row_h  = available_h / BASELINE_N
        n = len(data_rows)
        row_h = base_row_h if n <= BASELINE_N else max(8, available_h / n)
        row_heights = [row_h] * n

        tbl_data = [[P(str(r[0]),6), PL(r[1],6), P(str(r[2]),6), P(str(r[3]),6),
                     P(str(r[4]),6), P(str(r[5]),6), P(str(r[6]),6),
                     P(str(r[7]),6,True), P(str(r[8]),6,True), P('',6)]
                    for r in data_rows]
        dt = Table(tbl_data, colWidths=CW, rowHeights=row_heights)
        style_cmds = [
            ('GRID',(0,0),(-1,-1),0.3,BLACK),
            ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
            ('ALIGN',(0,0),(-1,-1),'CENTER'),
            ('ALIGN',(1,0),(1,-1),'LEFT'),
            ('TOPPADDING',(0,0),(-1,-1),2),
            ('BOTTOMPADDING',(0,0),(-1,-1),2),
            ('LEFTPADDING',(0,0),(-1,-1),2),
            ('RIGHTPADDING',(0,0),(-1,-1),2),
        ]
        for i, r in enumerate(data_rows):
            style_cmds.append(('BACKGROUND',(0,i),(-1,i), r[9]))
        dt.setStyle(TableStyle(style_cmds))
        story.append(dt)

    doc.build(story, onFirstPage=_on_page, onLaterPages=_on_page)
    buf.seek(0)
    return buf


# ─── PDF: LO-WISE SUMMARY (ANNEXURE II) ──────────────────────────────────────

def create_lo_summary_pdf(trainee_df, lo_df, user_info=None, marks_cache=None):
    """
    Generate print-ready A4 PORTRAIT PDF for LO-Wise Summary (Annexure II).
    One page per LO; all trainees listed with category totals.
    Mirrors create_lo_summary_excel exactly.
    """
    if not _REPORTLAB_OK:
        raise RuntimeError("ReportLab is not installed. Run: pip install reportlab")

    ui = user_info or {}

    # ── Portrait A4 ──────────────────────────────────────────────────────────
    PAGE_W, PAGE_H = A4          # portrait: ~595 × 842 pt
    L = R = 8 * mm
    T = B = 8 * mm
    USABLE_W = PAGE_W - L - R

    BLUE  = colors.HexColor('#1B4F8A')
    LBLUE = colors.HexColor('#D6E4F0')
    GOLD  = colors.HexColor('#FFF2CC')
    BLACK = colors.black
    LGRAY = colors.HexColor('#F8FAFC')

    def _ps(size=7, bold=False, align=TA_CENTER):
        return ParagraphStyle('_', fontSize=size, leading=size + 1.5,
                              fontName='Helvetica-Bold' if bold else 'Helvetica',
                              textColor=BLACK, alignment=align,
                              wordWrap='LTR', spaceBefore=0, spaceAfter=0)

    def P(txt, size=7, bold=False, align=TA_CENTER):
        return Paragraph(str(txt), _ps(size, bold, align))

    def PL(txt, size=7, bold=False):
        return Paragraph(str(txt), _ps(size, bold, TA_LEFT))

    # ── Column widths (portrait) ─────────────────────────────────────────────
    # Order: Roll | Firstname | Fathername | Safety | Hygiene | Attend |
    #        Manuals | Know | Skills | Speed | Quality | VIVA | Total | Result
    raw_w = [0.7, 1.8, 1.8, 1.0, 1.0, 1.0, 0.9, 1.0, 1.0, 1.0, 1.0, 1.0, 1.2, 0.8]
    tot   = sum(raw_w)
    CW    = [w / tot * USABLE_W for w in raw_w]

    lo_rows_all = list(lo_df.iterrows())

    # Build trainee_lo_marks from cache (same as create_lo_summary_excel)
    trainee_lo_marks = {}
    for _, trainee in trainee_df.iterrows():
        roll = trainee['rollno']
        trainee_lo_marks[roll] = {}
        trainee_cache = (marks_cache or {}).get(roll, {})
        for lo_num, lo_data in trainee_cache.items():
            trainee_lo_marks[roll][lo_num] = lo_data['cat_avgs']

    all_lo_numbers = []
    seen = set()
    for _, lo in lo_rows_all:
        lo_num = lo['lo']
        if lo_num not in seen:
            seen.add(lo_num)
            all_lo_numbers.append(lo)

    buf = io.BytesIO()

    def _on_page(canvas, doc):
        canvas.saveState()
        canvas.setStrokeColor(BLUE)
        canvas.setLineWidth(0.6)
        canvas.rect(L - 2, B - 2, PAGE_W - L - R + 4, PAGE_H - T - B + 4)
        canvas.restoreState()

    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=L, rightMargin=R,
                            topMargin=T, bottomMargin=B)
    story = []

    for pg_idx, lo_row in enumerate(all_lo_numbers):
        lo_num  = lo_row['lo']
        lo_name = lo_row['lo_name']

        if pg_idx > 0:
            story.append(PageBreak())

        # ── Titles ───────────────────────────────────────────────────────────
        story.append(Table([[P('ANNEXURE-II', size=10, bold=True)]],
                           colWidths=[USABLE_W],
                           style=TableStyle([('BOX',(0,0),(-1,-1),0.5,BLACK),
                                             ('BACKGROUND',(0,0),(-1,-1),LBLUE),
                                             ('TOPPADDING',(0,0),(-1,-1),3),
                                             ('BOTTOMPADDING',(0,0),(-1,-1),3)])))
        story.append(Table([[P('Internal Assessment', size=9, bold=True)]],
                           colWidths=[USABLE_W],
                           style=TableStyle([('BOX',(0,0),(-1,-1),0.5,BLACK),
                                             ('TOPPADDING',(0,0),(-1,-1),2),
                                             ('BOTTOMPADDING',(0,0),(-1,-1),2)])))

        # ── Header info grid ─────────────────────────────────────────────────
        hw = [USABLE_W*0.20, USABLE_W*0.30, USABLE_W*0.20, USABLE_W*0.30]
        info_rows = [
            ('Name & Address of the Assessor', ui.get('si_name',''),
             'Year of Enrolment', ui.get('year_of_assessment','')),
            ('Name & Address of ITI (Govt/Pvt)', ui.get('iti_name',''),
             'Date of Assessment', ''),
            ('Name & Address of the Industry', ui.get('iti_name',''),
             'Assessment Location', ui.get('assessment_location','')),
            ('Trade Name', ui.get('trade_name',''),
             'Duration Of Trade / Examination',
             f"{ui.get('trade_duration','')}  /  {ui.get('semester','')}"),
        ]
        for (l1, v1, l2, v2) in info_rows:
            story.append(Table(
                [[PL(l1, bold=True), PL(v1), PL(l2, bold=True), PL(v2)]],
                colWidths=hw,
                style=TableStyle([('BOX',(0,0),(-1,-1),0.5,BLACK),
                                   ('INNERGRID',(0,0),(-1,-1),0.3,BLACK),
                                   ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
                                   ('TOPPADDING',(0,0),(-1,-1),2),
                                   ('BOTTOMPADDING',(0,0),(-1,-1),2),
                                   ('LEFTPADDING',(0,0),(-1,-1),3),
                                   ('RIGHTPADDING',(0,0),(-1,-1),3)])))

        # ── LO label ─────────────────────────────────────────────────────────
        story.append(Table(
            [[PL(f'Learning Outcome: {lo_num}  —  {lo_name}', size=8, bold=True)]],
            colWidths=[USABLE_W],
            style=TableStyle([('BOX',(0,0),(-1,-1),0.5,BLACK),
                               ('BACKGROUND',(0,0),(-1,-1),LBLUE),
                               ('TOPPADDING',(0,0),(-1,-1),3),
                               ('BOTTOMPADDING',(0,0),(-1,-1),3),
                               ('LEFTPADDING',(0,0),(-1,-1),4)])))

        # ── Max marks row + column headers ───────────────────────────────────
        # Column order: Roll | Firstname | Fathername | 9 scores | Total | Result
        max_row = [P('',6), P('',6), P('',6),
                   P('15',6,True), P('10',6,True), P('10',6,True),
                   P('5', 6,True), P('10',6,True), P('10',6,True),
                   P('10',6,True), P('15',6,True), P('15',6,True),
                   P('100',6,True), P('',6)]
        hdr_row = [
            P('Roll\nNo',     6, True),
            P('Firstname',    6, True),
            P('Fathername',   6, True),
            P('Safety\nConsc.',       6, True),
            P('Workplace\nHygiene',   6, True),
            P('Attend/\nPunct.',      6, True),
            P('Manuals/\nInstr.',     6, True),
            P('Appln\nKnow.',         6, True),
            P('Skills',               6, True),
            P('Speed',                6, True),
            P('Quality',              6, True),
            P('VIVA',                 6, True),
            P('Total\n(Max 100)',     6, True),
            P('Result\n(Y/N)',        6, True),
        ]
        hdr_tbl = Table([max_row, hdr_row], colWidths=CW)
        hdr_tbl.setStyle(TableStyle([
            ('GRID',(0,0),(-1,-1),0.4,BLACK),
            ('BACKGROUND',(0,0),(-1,0),GOLD),
            ('BACKGROUND',(0,1),(-1,1),LBLUE),
            ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
            ('ALIGN',(0,0),(-1,-1),'CENTER'),
            ('TOPPADDING',(0,0),(-1,-1),2),
            ('BOTTOMPADDING',(0,0),(-1,-1),2),
        ]))
        story.append(hdr_tbl)

        # ── Data rows ────────────────────────────────────────────────────────
        data_rows = []
        for idx, (_, trainee) in enumerate(trainee_df.iterrows()):
            roll      = trainee['rollno']
            fname     = str(trainee.get('Firstname',  '') or '')
            fathrname = str(trainee.get('Fathername', '') or '')
            lo_data   = trainee_lo_marks.get(roll, {}).get(lo_num)
            if lo_data:
                s  = lo_data['safety_total']
                h  = lo_data['hygiene_total']
                a  = lo_data['attendance_total']
                mn = lo_data['manuals_total']
                k  = lo_data['knowledge_total']
                sk = lo_data['skills_total']
                sp = lo_data['speed_total']
                q  = lo_data['quality_total']
                v  = lo_data['viva_total']
                gt = lo_data['grand_total']
                result = 'Y' if gt >= 40 else 'N'
            else:
                s=h=a=mn=k=sk=sp=q=v=gt=''; result=''
            bg = LGRAY if idx % 2 == 0 else colors.white
            # Order matches columns: roll|fname|fathrname|s|h|a|mn|k|sk|sp|q|v|gt|result|bg
            data_rows.append((roll, fname, fathrname, s, h, a, mn, k, sk, sp, q, v, gt, result, bg))

        if data_rows:
            # Row height: 24 trainees fills one full portrait page exactly.
            # ≤24 → fixed base height; >24 → shrink to fit one page
            HEADER_H_PT = 165
            BASELINE_N  = 24
            available_h = (PAGE_H - T - B) - HEADER_H_PT
            base_row_h  = available_h / BASELINE_N
            n = len(data_rows)
            row_h = base_row_h if n <= BASELINE_N else max(8, available_h / n)
            row_heights = [row_h] * n

            tbl_data = [
                [P(str(r[0]),6), PL(r[1],6), PL(r[2],6),
                 P(str(r[3]),6),  P(str(r[4]),6),  P(str(r[5]),6),
                 P(str(r[6]),6),  P(str(r[7]),6),  P(str(r[8]),6),
                 P(str(r[9]),6),  P(str(r[10]),6), P(str(r[11]),6),
                 P(str(r[12]),6,True), P(str(r[13]),6)]
                for r in data_rows
            ]
            dt = Table(tbl_data, colWidths=CW, rowHeights=row_heights)
            sc = [('GRID',(0,0),(-1,-1),0.3,BLACK),
                  ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
                  ('ALIGN',(0,0),(-1,-1),'CENTER'),
                  ('ALIGN',(1,0),(1,-1),'LEFT'),   # Firstname left-aligned
                  ('ALIGN',(2,0),(2,-1),'LEFT'),   # Fathername left-aligned
                  ('TOPPADDING',(0,0),(-1,-1),2),
                  ('BOTTOMPADDING',(0,0),(-1,-1),2),
                  ('LEFTPADDING',(0,0),(-1,-1),2),
                  ('RIGHTPADDING',(0,0),(-1,-1),2)]
            for i, r in enumerate(data_rows):
                sc.append(('BACKGROUND',(0,i),(-1,i), r[14]))
            dt.setStyle(TableStyle(sc))
            story.append(dt)

    doc.build(story, onFirstPage=_on_page, onLaterPages=_on_page)
    buf.seek(0)
    return buf


# ─── PDF: PER TRAINEE PER LO REPORT ──────────────────────────────────────────

def create_trainee_lo_pdf(trainee_df, lo_df, user_info=None, marks_cache=None):
    """
    Generate print-ready A4 landscape PDF for Per-Trainee Per-LO report.
    One page per trainee per LO with all practical marks.
    Mirrors create_trainee_lo_report (demo format) exactly:
    - No background colors (plain white)
    - Row 6: LO No and Pract No rotated 90 (col 0,1); category headers wrap text
    - Row 7: all sub-column headers rotated 90
    - Row 8: max marks plain
    - Info rows: 6 segments matching demo merges
    """
    if not _REPORTLAB_OK:
        raise RuntimeError("ReportLab is not installed. Run: pip install reportlab")

    ui = user_info or {}

    PAGE_W, PAGE_H = landscape(A4)
    L = R = 5 * mm
    T = B = 5 * mm
    USABLE_W = PAGE_W - L - R

    BLACK = colors.black
    WHITE = colors.white

    def _ps(size=5.5, bold=False, align=TA_CENTER):
        return ParagraphStyle('_', fontSize=size, leading=size+1.2,
                              fontName='Helvetica-Bold' if bold else 'Helvetica',
                              textColor=BLACK, alignment=align,
                              wordWrap='LTR', spaceBefore=0, spaceAfter=0)

    def P(txt, size=5.5, bold=False, align=TA_CENTER):
        return Paragraph(str(txt), _ps(size, bold, align))

    def PL(txt, size=5.5, bold=False):
        return Paragraph(str(txt), _ps(size, bold, TA_LEFT))

    # 40 columns proportional to demo widths
    # Demo col widths (in Excel units, used as proportions):
    demo_raw = [8.42,5.0,3.42,2.85,3.14,3.42,4.57,4.71,4.0,3.14,
                3.42,2.85,3.42,3.42,2.85,3.28,3.14,3.0,3.14,4.28,
                3.28,3.42,5.14,3.0,2.85,3.28,4.85,4.57,3.85,4.28,
                3.0,3.14,4.42,3.0,3.14,4.14,4.42,5.57,8.0,9.0]
    tot = sum(demo_raw)
    CW = [w / tot * USABLE_W for w in demo_raw]

    buf = io.BytesIO()

    def _on_page(canvas, doc):
        canvas.saveState()
        canvas.setStrokeColor(BLACK)
        canvas.setLineWidth(0.5)
        canvas.rect(L-2, B-2, PAGE_W-L-R+4, PAGE_H-T-B+4)
        canvas.restoreState()

    doc = SimpleDocTemplate(buf, pagesize=landscape(A4),
                            leftMargin=L, rightMargin=R,
                            topMargin=T, bottomMargin=B)
    story = []

    BASE_GRID = [
        ('GRID',(0,0),(-1,-1),0.3,BLACK),
        ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
        ('ALIGN',(0,0),(-1,-1),'CENTER'),
        ('TOPPADDING',(0,0),(-1,-1),1),
        ('BOTTOMPADDING',(0,0),(-1,-1),1),
        ('LEFTPADDING',(0,0),(-1,-1),1),
        ('RIGHTPADDING',(0,0),(-1,-1),1),
    ]

    page_num = 0
    for _, trainee in trainee_df.iterrows():
        roll      = trainee['rollno']
        fname     = str(trainee.get('Firstname','') or '')
        fathrname = str(trainee.get('Fathername','') or '')
        lname     = str(trainee.get('Lastname','') or '')
        full_name = f"{fname} {fathrname} {lname}".strip()
        trainee_cache = (marks_cache or {}).get(roll, {})

        for lo_num, lo_data in trainee_cache.items():
            pract_from  = lo_data['pract_from']
            pract_to    = lo_data['pract_to']
            lo_name     = lo_data['lo_name']
            pract_marks = lo_data['pract_marks']
            lo_avg      = lo_data['lo_avg']

            if page_num > 0:
                story.append(PageBreak())
            page_num += 1

            # ── Row 1: Title ──────────────────────────────────────
            story.append(Table(
                [[P('Internal Assessment', size=9, bold=True)]],
                colWidths=[USABLE_W],
                style=TableStyle([('BOX',(0,0),(-1,-1),0.5,BLACK),
                                   ('TOPPADDING',(0,0),(-1,-1),2),
                                   ('BOTTOMPADDING',(0,0),(-1,-1),2)])))

            # ── Info rows 2-5 (6 segments matching demo) ─────────
            # Segment proportions: cols 1-3 | 4-17 | 18-22 | 23-27 | 28-34 | 35-40
            # Proportional widths:
            s1 = sum(demo_raw[0:3]) / tot * USABLE_W   # cols 1-3
            s2 = sum(demo_raw[3:17]) / tot * USABLE_W  # cols 4-17
            s3 = sum(demo_raw[17:22]) / tot * USABLE_W # cols 18-22
            s4 = sum(demo_raw[22:27]) / tot * USABLE_W # cols 23-27
            s5 = sum(demo_raw[27:34]) / tot * USABLE_W # cols 28-34
            s6 = sum(demo_raw[34:40]) / tot * USABLE_W # cols 35-40
            iw = [s1, s2, s3, s4, s5, s6]

            info_rows_data = [
                ('Name of Trainee:', full_name,
                 'Roll NO:', str(roll),
                 'Year of Enrollment:', ui.get('year_of_assessment','')),
                ('Name of ITI:', ui.get('iti_name',''),
                 'Date of Assessment:', '',
                 'Batch:', ui.get('batch','')),
                ('Name of the Industry:', ui.get('iti_name',''),
                 'Assessment Location:', ui.get('assessment_location',''),
                 'Sem:', ui.get('semester','')),
                ('Trade Name:', ui.get('trade_name',''),
                 'Duration of the Trade:', ui.get('trade_duration',''),
                 'S.I. Name:', ui.get('si_name','')),
            ]
            info_style = TableStyle(BASE_GRID + [
                ('ALIGN',(0,0),(0,-1),'LEFT'),
                ('ALIGN',(2,0),(2,-1),'LEFT'),
                ('ALIGN',(4,0),(4,-1),'LEFT'),
                ('FONTNAME',(0,0),(0,-1),'Helvetica-Bold'),
                ('FONTNAME',(2,0),(2,-1),'Helvetica-Bold'),
                ('FONTNAME',(4,0),(4,-1),'Helvetica-Bold'),
            ])
            for (l1,v1,l2,v2,l3,v3) in info_rows_data:
                story.append(Table(
                    [[PL(l1,bold=True), PL(v1), PL(l2,bold=True), P(v2), PL(l3,bold=True), P(v3)]],
                    colWidths=iw,
                    style=info_style))

            # ── Header rows: category (row 6) + sub-cols (row 7) + max (row 8) ──
            # Row 6: LO No [rot90], Pract No [rot90], cat spans [wrap], Grand Total [rot90], Sig [rot90]
            # Row 7: all sub headers [rot90]
            # Row 8: max marks
            # Rows 6-7: col 0 spans 2 rows (LO No), col 1 spans 2 rows (Pract No)
            # Similarly col 38 (Grand Total) and col 39 (Sig) span rows 6-7

            from reportlab.platypus import Flowable as _Flowable

            class RotTxt(_Flowable):
                def __init__(self, text, cw, ch, size=5, bold=False):
                    _Flowable.__init__(self)
                    self.text = text; self.width = cw; self.height = ch
                    self._fn = 'Helvetica-Bold' if bold else 'Helvetica'
                    self._sz = size
                def draw(self):
                    c = self.canv; c.saveState()
                    c.setFont(self._fn, self._sz)
                    c.translate(self.width/2, self.height/2); c.rotate(90)
                    lines = self.text.split('\n')
                    lh = self._sz * 1.2; tot_h = len(lines)*lh
                    for i,ln in enumerate(lines):
                        tw = c.stringWidth(ln, self._fn, self._sz)
                        c.drawString(-tw/2, tot_h/2-(i+0.5)*lh, ln)
                    c.restoreState()
                def wrap(self, aW, aH): return self.width, self.height

            CAT_H = 18   # category row height (pt)
            SUB_H = 110  # sub-column row height (pt)
            MAX_H = 10   # max row height

            def rot(txt, col_i, h, sz=5, bold=False):
                return RotTxt(txt, CW[col_i], h, sz, bold)

            def wrap_p(txt, sz=5, bold=False):
                return P(txt, sz, bold)

            # Row 0 (cat row): 40 cells
            cat_r = [''] * 40
            cat_r[0]  = rot('Learning\nOutcome\nNumber', 0, CAT_H+SUB_H, sz=5)   # spans 2 rows
            cat_r[1]  = rot('Practical /\nProfessional\nSkill Number', 1, CAT_H+SUB_H, sz=5)
            cat_r[2]  = wrap_p('Safety\nconsciousness', 5)
            cat_r[6]  = wrap_p('Workplace hygiene &\nEconomical use of materials', 5)
            cat_r[10] = wrap_p('Attendance/\nPunctuality', 5)
            cat_r[14] = wrap_p('Ability to follow Manuals/\nWritten instructions', 5)
            cat_r[18] = wrap_p('Application of\nKnowledge', 5)
            cat_r[22] = wrap_p('Skills to handle\ntools & equipment', 5)
            cat_r[26] = wrap_p('Speed in\ndoing work', 5)
            cat_r[30] = wrap_p('Quality in\nworkmanship', 5)
            cat_r[34] = wrap_p('VIVA', 5)
            cat_r[38] = rot('Grand\nTotal', 38, CAT_H+SUB_H, sz=5)   # spans 2 rows
            cat_r[39] = rot('Signature\nof Trainee', 39, CAT_H+SUB_H, sz=5)

            # Row 1 (sub row): all rotated 90 — single line per label (no wrap)
            sub_texts = [
                '', '',
                'Dress code','Use PPE','Apply/ practice safety','Total',
                'Maintain personal & workplace cleanliness',
                'Dispose scrap as per standard practice',
                'Select appropriate material & minimize wastage','Total',
                'Initiative','Account- ability','Participative in work','Total',
                'Select right manual','Search for appropriate topic',
                'Read & interpret the manual','Total',
                'Plan the work','Select appropriate tools & equipment',
                'Review the work','Total',
                'Handle & use tools & equipment','Maintain safety in handling',
                'Care & maintain','Total',
                'Properly sequence the work','Use appropriate technique',
                'Review the work during execution','Total',
                'Achieve work with high accuracy','Conform to requirement',
                'Satisfy the purpose','Total',
                'Response with clarity','Technical understand.',
                'Conscious towards job role','Total',
                '', '',
            ]
            sub_r = [rot(t, i, SUB_H, sz=4.5) if t else P('', 4.5)
                     for i, t in enumerate(sub_texts)]

            # Row 2 (max row)
            max_vals = ['','', 2,5,8,15, 3,2,5,10, 3,3,4,10, 1,2,2,5, 4,3,3,10, 4,3,3,10,
                        3,5,2,10, 7,3,5,15, 7,5,3,15, 100, '']
            max_r = [P(str(v) if v else '', 4.5) for v in max_vals]

            hdr_data = [cat_r, sub_r, max_r]
            hdr_style = TableStyle(BASE_GRID + [
                ('SPAN',(2,0),(5,0)),   # Safety
                ('SPAN',(6,0),(9,0)),   # Hygiene
                ('SPAN',(10,0),(13,0)), # Attendance
                ('SPAN',(14,0),(17,0)), # Manuals
                ('SPAN',(18,0),(21,0)), # Knowledge
                ('SPAN',(22,0),(25,0)), # Skills
                ('SPAN',(26,0),(29,0)), # Speed
                ('SPAN',(30,0),(33,0)), # Quality
                ('SPAN',(34,0),(37,0)), # VIVA
                # Col 0,1,38,39 span rows 0-1
                ('SPAN',(0,0),(0,1)),
                ('SPAN',(1,0),(1,1)),
                ('SPAN',(38,0),(38,1)),
                ('SPAN',(39,0),(39,1)),
            ])
            hdr_tbl = Table(hdr_data, colWidths=CW, rowHeights=[CAT_H, SUB_H, MAX_H])
            hdr_tbl.setStyle(hdr_style)
            story.append(hdr_tbl)

            # ── Data rows ─────────────────────────────────────────
            data_rows_tbl = []
            row_styles = []
            for pract_offset, pract_num in enumerate(range(pract_from, pract_to + 1)):
                m = pract_marks[pract_offset]
                row_vals = [
                    f'LO-{lo_num}', pract_num,
                    m['dress'], m['ppe'], m['apply_safety'], m['safety_total'],
                    m['personal'], m['scrap'], m['material'], m['hygiene_total'],
                    m['initiative'], m['accountability'], m['participative'], m['attendance_total'],
                    m['select_manual'], m['search_topic'], m['read_manual'], m['manuals_total'],
                    m['plan_work'], m['select_tools'], m['review_work'], m['knowledge_total'],
                    m['handle_tools'], m['safety_handling'], m['care_maintain'], m['skills_total'],
                    m['sequence'], m['technique'], m['review_execution'], m['speed_total'],
                    m['accuracy'], m['conform'], m['satisfy'], m['quality_total'],
                    m['clarity'], m['technical'], m['conscious'], m['viva_total'],
                    m['grand_total'], '',
                ]
                data_rows_tbl.append([P(str(v),5) for v in row_vals])

            # LO average row (merged, no color)
            lo_summary = [P(f'{lo_name}     Average of LO{lo_num}: {lo_avg}', 5, False)] + [P('',5)] * 39
            data_rows_tbl.append(lo_summary)
            avg_idx = len(data_rows_tbl) - 1
            row_styles += [('SPAN',(0,avg_idx),(-1,avg_idx))]

            if data_rows_tbl:
                dt = Table(data_rows_tbl, colWidths=CW)
                dt.setStyle(TableStyle(BASE_GRID + row_styles))
                story.append(dt)

    doc.build(story, onFirstPage=_on_page, onLaterPages=_on_page)
    buf.seek(0)
    return buf



# ─── EXCEL: ALL TRAINEE ALL LO SUMMARY ───────────────────────────────────────

def create_trainee_summary_excel(trainee_df, lo_df, user_info=None, marks_cache=None):
    """
    Generate All Trainee All LO Summary Excel.
    Layout:
      Row 1  : Title — merged A1:last_col, Calibri bold 24pt, centered, height 31.5
      Row 2  : Info  — merged A2:last_col, Arial bold 11pt, centered,
                       bottom-medium border, height 21.75
      Row 3  : Header row — gray fill, Arial bold 10pt, medium border, height 51.75
      Row 4+ : One row per trainee, height 15.0, thin border
    Columns: A=rollno, B=name, C..=LO cols, avg_col, lo70_col
    Page: Landscape A4
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    ui      = user_info or {}
    u_iti   = ui.get('iti_name', '')
    u_trade = ui.get('trade_name', '')
    u_si    = ui.get('si_name', '')
    u_batch = ui.get('batch', '')
    u_sem   = ui.get('semester', '')

    # Collect unique LO numbers in order from LO file
    lo_nums = []
    seen_lo = set()
    for _, lo in lo_df.iterrows():
        n = lo['lo']
        if n not in seen_lo:
            seen_lo.add(n)
            lo_nums.append(n)

    # Build per-trainee LO averages from marks_cache
    trainee_lo_avgs = {}
    for _, tr in trainee_df.iterrows():
        roll = tr['rollno']
        trainee_lo_avgs[roll] = {}
        cache = (marks_cache or {}).get(roll, {})
        for lo_num in lo_nums:
            lo_data = cache.get(lo_num)
            trainee_lo_avgs[roll][lo_num] = lo_data['lo_avg'] if lo_data else ''

    # ── Styles ────────────────────────────────────────────────────────
    GRAY_FILL = PatternFill('solid', fgColor='FFD3D3D3')
    NO_FILL   = PatternFill('none')
    med       = Side(style='medium', color='000000')
    thin_s    = Side(style='thin',   color='000000')
    MED_BDR   = Border(left=med,    right=med,    top=med,    bottom=med)
    THIN_BDR  = Border(left=thin_s, right=thin_s, top=thin_s, bottom=thin_s)
    BOT_MED   = Border(bottom=med)

    def _font(size=10, bold=False, name='Arial'):
        return Font(name=name, size=size, bold=bold)

    def _aln(h='center', v='center', wrap=False):
        return Alignment(horizontal=h, vertical=v, wrap_text=wrap)

    wb = Workbook()
    ws = wb.active
    ws.title = 'Sheet1'

    # ── Column layout ─────────────────────────────────────────────────
    n_lo     = len(lo_nums)
    avg_col  = 3 + n_lo        # 1-based: "Average of Lo"
    lo70_col = avg_col + 1     # 1-based: "LO Average Out of 70"
    last_col = lo70_col

    ws.column_dimensions['A'].width = 6.28515625
    ws.column_dimensions['B'].width = 33.140625
    for i in range(n_lo):
        ws.column_dimensions[get_column_letter(3 + i)].width = 6.28515625
    ws.column_dimensions[get_column_letter(avg_col)].width  = 8.28515625
    ws.column_dimensions[get_column_letter(lo70_col)].width = 8.7109375
    ws.column_dimensions[get_column_letter(lo70_col + 1)].width = 4.28515625
    ws.column_dimensions[get_column_letter(lo70_col + 2)].width = 6.140625

    # ── Row heights ────────────────────────────────────────────────────
    ws.row_dimensions[1].height = 31.5
    ws.row_dimensions[2].height = 21.75
    ws.row_dimensions[3].height = 51.75

    last_col_letter = get_column_letter(last_col)

    # ── Row 1: Title — merged A1:last_col, centered ───────────────────
    ws.merge_cells(f'A1:{last_col_letter}1')
    r1 = ws['A1']
    r1.value     = f'INDUSTRIAL TRAINING INSTITUTE {u_iti.upper()}'
    r1.font      = Font(name='Calibri', size=24, bold=True)
    r1.alignment = _aln('center', 'center', False)

    # ── Row 2: Info — merged A2:last_col, centered, bottom-medium border
    ws.merge_cells(f'A2:{last_col_letter}2')
    r2 = ws['A2']
    r2.value     = (f'Trade Name :{u_trade}  ||  Name Of SI :{u_si}  ||  '
                    f'Batch :{u_batch}  ||  Sem :{u_sem}')
    r2.font      = _font(11, bold=True, name='Arial')
    r2.alignment = _aln('center', 'center', True)
    r2.border    = BOT_MED

    # ── Row 3: Header row ─────────────────────────────────────────────
    hdr_labels = (
        [(1, 'rollno'), (2, 'TRAINEE NAME')] +
        [(3 + i, f'LO - {lo_nums[i]}') for i in range(n_lo)] +
        [(avg_col, 'Average of Lo'), (lo70_col, 'LO Average Out of 70')]
    )
    for col, label in hdr_labels:
        cell = ws.cell(row=3, column=col, value=label)
        cell.font      = _font(10, bold=True, name='Arial')
        cell.fill      = GRAY_FILL
        cell.alignment = _aln('center', 'center', True)
        cell.border    = MED_BDR

    # ── Data rows (row 4 onwards, one per trainee) ────────────────────
    for t_idx, (_, tr) in enumerate(trainee_df.iterrows()):
        r    = 4 + t_idx
        roll = tr['rollno']
        nm   = ' '.join([str(tr.get(f, '') or '') for f in
                         ['Firstname', 'Fathername', 'Lastname']]).strip()

        lo_avgs    = trainee_lo_avgs.get(roll, {})
        valid_avgs = [lo_avgs[n] for n in lo_nums if isinstance(lo_avgs.get(n), (int, float))]
        avg_of_lo  = round(sum(valid_avgs) / len(valid_avgs)) if valid_avgs else ''
        raw_70 = tr.get('All_LO_average_base_on_70_SEM_I')
        try:
            lo_out70 = int(float(raw_70)) if raw_70 is not None and str(raw_70) != 'nan' else ''
        except Exception:
            lo_out70 = ''

        ws.row_dimensions[r].height = 15.0

        data = (
            [(1, roll), (2, nm)] +
            [(3 + i, lo_avgs.get(lo_nums[i], '')) for i in range(n_lo)] +
            [(avg_col, avg_of_lo), (lo70_col, lo_out70)]
        )
        for col, val in data:
            cell           = ws.cell(row=r, column=col, value=val)
            cell.font      = _font(10, bold=(col == 1), name='Arial')
            cell.fill      = NO_FILL
            cell.alignment = _aln('left' if col == 2 else 'center', 'center', False)
            cell.border    = THIN_BDR

    # ── Page setup: landscape A4 ───────────────────────────────────────
    ws.page_setup.orientation = 'landscape'
    ws.page_setup.paperSize   = 9
    ws.page_setup.fitToPage   = True
    ws.page_setup.fitToWidth  = 1
    ws.page_setup.fitToHeight = 0

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ─── PDF: TRAINEE WISE LO SUMMARY ────────────────────────────────────────────

def create_trainee_summary_pdf(trainee_df, lo_df, user_info=None, marks_cache=None):
    """
    Generate a landscape A4 PDF matching the Trainee Wise LO Summary proforma.
    Columns: Roll No | Trainee Name | LO-1 .. LO-N | Average of Lo | LO Avg out of 70
    """
    if not _REPORTLAB_OK:
        raise RuntimeError("ReportLab is not installed. Run: pip install reportlab")

    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.units import mm
    from reportlab.lib.utils import simpleSplit

    ui      = user_info or {}
    u_iti   = ui.get('iti_name', '')
    u_trade = ui.get('trade_name', '')
    u_si    = ui.get('si_name', '')
    u_batch = ui.get('batch', '')
    u_sem   = ui.get('semester', '')

    PAGE_W, PAGE_H = landscape(A4)   # ~841.9 x 595.3 pt
    ML = MR = 10 * mm
    MT = MB = 10 * mm
    TW = PAGE_W - ML - MR
    BLACK = colors.black
    GRAY  = colors.HexColor('#D3D3D3')
    ALT   = colors.HexColor('#EEF4FB')

    # Collect LO numbers
    lo_nums = []
    seen_lo = set()
    for _, lo in lo_df.iterrows():
        n = lo['lo']
        if n not in seen_lo:
            seen_lo.add(n)
            lo_nums.append(n)

    # Build per-trainee LO avgs
    trainee_lo_avgs = {}
    for _, tr in trainee_df.iterrows():
        roll = tr['rollno']
        trainee_lo_avgs[roll] = {}
        cache = (marks_cache or {}).get(roll, {})
        for lo_num in lo_nums:
            lo_data = cache.get(lo_num)
            trainee_lo_avgs[roll][lo_num] = lo_data['lo_avg'] if lo_data else ''

    # ── Column widths proportional ────────────────────────────────────
    # Roll(1) | Name(4) | LO×n(1 each) | Avg(1.2) | Out70(1.4)
    n_lo = len(lo_nums)
    raw  = [1.0, 4.0] + [1.0] * n_lo + [1.2, 1.4]
    tot  = sum(raw)
    CW   = [r / tot * TW for r in raw]

    def col_x(i):
        return ML + sum(CW[:i])

    ROW_H_TITLE = 18
    ROW_H_INFO  = 14
    ROW_H_HDR   = 30
    TARGET_ROWS_PP = 24   # Always show exactly 24 rows per page

    buf = io.BytesIO()
    c   = rl_canvas.Canvas(buf, pagesize=landscape(A4))

    def draw_cell_pdf(x, y, w, h, text='', sz=8, bold=False,
                      align='center', bg=None, border_color=BLACK, lw=0.4):
        c.saveState()
        if bg:
            c.setFillColor(bg)
            c.rect(x, y - h, w, h, fill=1, stroke=0)
        c.setStrokeColor(border_color)
        c.setLineWidth(lw)
        c.rect(x, y - h, w, h, fill=0, stroke=1)
        if text:
            c.setFillColor(BLACK)
            fn = 'Helvetica-Bold' if bold else 'Helvetica'
            c.setFont(fn, sz)
            PAD = 2.0
            avail_w = max(w - PAD * 2, 1)
            cur_sz = sz
            while cur_sz >= 3.5:
                lines = simpleSplit(str(text), fn, cur_sz, avail_w)
                if not lines: lines = ['']
                if len(lines) * cur_sz * 1.25 <= h - 1.0:
                    break
                cur_sz -= 0.5
            leading = cur_sz * 1.25
            lines   = simpleSplit(str(text), fn, cur_sz, avail_w)
            if not lines: lines = ['']
            total_th = len(lines) * leading
            start_y  = (y - h / 2) + total_th / 2 - cur_sz * 0.80
            for i, ln in enumerate(lines):
                lw2 = c.stringWidth(ln, fn, cur_sz)
                lx  = (x + (w - lw2) / 2 if align == 'center'
                        else x + w - lw2 - PAD if align == 'right'
                        else x + PAD)
                c.drawString(lx, start_y - i * leading, ln)
        c.restoreState()

    # ── Row height: fixed at usable_h/24; only shrink if > 24 trainees ─
    n_trainees   = len(trainee_df)
    fixed_h      = ROW_H_TITLE + ROW_H_INFO + ROW_H_HDR
    usable_h     = PAGE_H - MT - MB - fixed_h
    ROW_H_DATA   = usable_h / TARGET_ROWS_PP          # default: fills page for 24
    if n_trainees > TARGET_ROWS_PP:
        # Shrink so all fit on one page; minimum 4pt
        ROW_H_DATA = max(4, usable_h / n_trainees)
    rows_pp      = TARGET_ROWS_PP if n_trainees <= TARGET_ROWS_PP else n_trainees

    trainee_list = list(trainee_df.iterrows())
    page_num = 0

    for chunk_start in range(0, n_trainees, rows_pp):
        chunk = trainee_list[chunk_start:chunk_start + rows_pp]
        if page_num > 0:
            c.showPage()
        page_num += 1

        y = PAGE_H - MT

        # Title row
        c.setFillColor(BLACK)
        c.setFont('Helvetica-Bold', 12)
        title = f'INDUSTRIAL TRAINING INSTITUTE {u_iti.upper()}'
        c.drawCentredString(PAGE_W / 2, y - 13, title)
        y -= ROW_H_TITLE

        # Info row
        info = (f'Trade Name :{u_trade}  ||  Name Of SI :{u_si}  ||  '
                f'Batch :{u_batch}  ||  Sem :{u_sem}')
        c.setFont('Helvetica-Bold', 9)
        c.drawString(ML, y - 10, info)
        y -= ROW_H_INFO

        # Header row
        hdr_labels = ['Roll No', 'TRAINEE NAME'] + [f'LO-{n}' for n in lo_nums] + ['Avg of LO', 'LO Avg\n/70']
        for i, lbl in enumerate(hdr_labels):
            draw_cell_pdf(col_x(i), y, CW[i], ROW_H_HDR,
                          lbl, sz=8, bold=True, align='center', bg=GRAY, lw=0.8)
        y -= ROW_H_HDR

        # Data rows
        for t_idx, (_, tr) in enumerate(chunk):
            roll    = tr['rollno']
            nm      = ' '.join([str(tr.get(f, '') or '') for f in
                                ['Firstname', 'Fathername', 'Lastname']]).strip()
            lo_avgs = trainee_lo_avgs.get(roll, {})
            valid   = [lo_avgs[n] for n in lo_nums if isinstance(lo_avgs.get(n), (int, float))]
            avg_lo  = round(sum(valid) / len(valid)) if valid else ''
            raw_70  = tr.get('All_LO_average_base_on_70_SEM_I')
            try:
                lo70 = int(float(raw_70)) if raw_70 is not None and str(raw_70) != 'nan' else ''
            except Exception:
                lo70 = ''

            row_bg = ALT if t_idx % 2 == 1 else None
            vals   = ([str(roll), nm] +
                      [str(lo_avgs.get(n, '')) for n in lo_nums] +
                      [str(avg_lo), str(lo70)])
            aligns = ['center', 'left'] + ['center'] * (n_lo + 2)
            for i, (val, aln) in enumerate(zip(vals, aligns)):
                draw_cell_pdf(col_x(i), y, CW[i], ROW_H_DATA,
                              val, sz=8, bold=(i == 0), align=aln, bg=row_bg, lw=0.3)
            y -= ROW_H_DATA

    c.save()
    buf.seek(0)
    return buf


# ─── PDF: PROGRESS CARD ───────────────────────────────────────────────────────

def create_progress_card_pdf(trainee_df, user_info=None):
    """
    Generate portrait A4 PDF for ITI Progress Cards. One page per trainee.
    - 2-year trade: card fills full A4 page (margins ~10mm each side)
    - 1-year trade: card occupies ~75% of A4 height, centred vertically
    - No background colours; plain white with thin black borders
    - Year summary: Formative Total header AND data row both span cols 4-8
    - All text fits inside its cell (font sized to cell)
    """
    if not _REPORTLAB_OK:
        raise RuntimeError("ReportLab is not installed. Run: pip install reportlab")

    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.utils import simpleSplit

    ui        = user_info or {}
    u_iti     = ui.get('iti_name', '')
    u_trade   = ui.get('trade_name', '')
    u_batch   = ui.get('batch', '')
    u_semester= ui.get('semester', 'H1')
    is_2yr    = _is_2_year_trade(ui)
    sem_upper = u_semester.strip().upper()

    PAGE_W, PAGE_H = A4          # 595.3 x 841.9 pt  (portrait)
    ML = MR = 10 * mm
    TW = PAGE_W - ML - MR        # usable width

    # ── Total available card height ──────────────────────────────────
    # 2-year: use 96% of page height; 1-year: use 72% (≈ 75%)
    CARD_H_FRAC = 0.96 if is_2yr else 0.72
    CARD_H = PAGE_H * CARD_H_FRAC
    # Top margin: centre card vertically
    MT = (PAGE_H - CARD_H) / 2

    BLACK = colors.black

    # ── 13 column proportional widths ────────────────────────────────
    # Col 0 = No.Half/SEM  1=Actual  2=Possible  3=%  4=SLO(70)
    # 5=ES  6=WC_SC  7=ED  8=Total  9=SignTrainee  10=SignSI  11=SignFI  12=SignPrincipal
    raw = [2.0, 1.1, 1.3, 0.9, 2.8, 1.05, 1.15, 1.05, 1.6, 2.0, 1.8, 1.6, 2.0]
    tot = sum(raw)
    CW  = [r / tot * TW for r in raw]

    def col_x(c):
        return ML + sum(CW[:c])

    def span_w(c1, c2):
        return sum(CW[c1:c2+1])

    # ── Row-height budget ─────────────────────────────────────────────
    # Count how many "height units" the card needs, then divide CARD_H
    # 1-year rows:  title×2, info×2, date×2, sec×1, hdr×2, data×2, yr_hdr×1, yr_dat×1  = 13 units
    # 2-year rows:  same top section + (yr1_sec + hdr×2 + data×2 + yr_hdr + yr_dat) * 2 + yr2_sec
    #               = title×2 + info×2 + date×2 + yr1_sec×1 + (sec+hdr×2+data×2+yr_hdr+yr_dat)×2
    # Unit sizes (relative):
    U_TITLE  = 1.6   # institute / PROGRESS CARD rows
    U_INFO   = 1.8   # roll/name rows
    U_DATE   = 1.5   # date rows
    U_SEC    = 1.1   # section label (HALF YEARLY / YEAR-N ASSESSMENT)
    U_HDR1   = 1.2   # header row 1 (Attendance span / GLO span)
    U_HDR2   = 3.2   # header row 2 (Actual/Possible/% | ES/WC_SC/ED) — tallest
    U_DATA   = 1.5   # H1/H2/H3/H4 rows
    U_YRHDR  = 1.8   # Year summary header
    U_YRDAT  = 1.5   # Year summary data (YEAR-1 / YEAR-2)

    if not is_2yr:
        total_units = (U_TITLE*2 + U_INFO*2 + U_DATE*2
                       + U_SEC + U_HDR1 + U_HDR2 + U_DATA*2
                       + U_YRHDR + U_YRDAT)
    else:
        total_units = (U_TITLE*2 + U_INFO*2 + U_DATE*2
                       + U_SEC   # YEAR-1 ASSESSMENT label
                       + U_SEC + U_HDR1 + U_HDR2 + U_DATA*2 + U_YRHDR + U_YRDAT
                       + U_SEC   # YEAR-2 ASSESSMENT label
                       + U_SEC + U_HDR1 + U_HDR2 + U_DATA*2 + U_YRHDR + U_YRDAT)

    unit = CARD_H / total_units

    RH_TITLE = U_TITLE * unit
    RH_INFO  = U_INFO  * unit
    RH_DATE  = U_DATE  * unit
    RH_SEC   = U_SEC   * unit
    RH_HDR1  = U_HDR1  * unit
    RH_HDR2  = U_HDR2  * unit
    RH_DATA  = U_DATA  * unit
    RH_YRHDR = U_YRHDR * unit
    RH_YRDAT = U_YRDAT * unit

    # ── Font sizes derived from row heights so text always fits ───────
    SZ_TITLE = min(11, RH_TITLE * 0.52)
    SZ_INFO  = min(9,  RH_INFO  * 0.42)
    SZ_DATE  = min(8,  RH_DATE  * 0.42)
    SZ_SEC   = min(8,  RH_SEC   * 0.55)
    SZ_HDR   = min(6.5,RH_HDR2  * 0.13)   # HDR2 is tall; use small font
    SZ_DATA  = min(8,  RH_DATA  * 0.50)
    SZ_YR    = min(7,  RH_YRHDR * 0.28)

    def fmt_date(v):
        if v is None: return ''
        try:
            import math
            if isinstance(v, float) and math.isnan(v): return ''
        except Exception: pass
        if hasattr(v, 'strftime'): return v.strftime('%d-%m-%Y')
        return str(v)

    def to_int(v):
        try: return int(float(v))
        except Exception:
            return '' if (v is None or (isinstance(v, float) and __import__('math').isnan(v))) else v

    buf = io.BytesIO()
    c   = rl_canvas.Canvas(buf, pagesize=A4)

    def draw_cell(x, y, w, h, text='', sz=8, bold=False, align='center',
                  border=True, lw=0.4):
        """
        Draw one bordered cell. x,y = top-left. Text is word-wrapped and
        vertically centred; font is auto-reduced if lines still overflow.
        """
        c.saveState()
        if border:
            c.setStrokeColor(BLACK)
            c.setLineWidth(lw)
            c.rect(x, y - h, w, h, fill=0, stroke=1)
        if text:
            c.setFillColor(BLACK)
            fn = 'Helvetica-Bold' if bold else 'Helvetica'
            PAD = 2.0   # horizontal padding pts each side
            avail_w = max(w - PAD * 2, 1)
            # Try font size; shrink until lines fit vertically
            cur_sz = sz
            while cur_sz >= 3.5:
                c.setFont(fn, cur_sz)
                lines = simpleSplit(str(text), fn, cur_sz, avail_w)
                if not lines: lines = ['']
                leading = cur_sz * 1.25
                if len(lines) * leading <= h - 1.5:
                    break
                cur_sz -= 0.5
            leading = cur_sz * 1.25
            total_th = len(lines) * leading
            start_y = (y - h / 2) + total_th / 2 - cur_sz * 0.80
            for i, ln in enumerate(lines):
                lw2 = c.stringWidth(ln, fn, cur_sz)
                if align == 'center':
                    lx = x + (w - lw2) / 2
                elif align == 'right':
                    lx = x + w - lw2 - PAD
                else:
                    lx = x + PAD
                c.drawString(lx, start_y - i * leading, ln)
        c.restoreState()

    # ── Reusable block: assessment table header (3 sub-rows) ──────────
    def draw_table_header(y_top):
        y = y_top
        # Row A: section labels
        draw_cell(col_x(0), y, span_w(0,3), RH_SEC,
                  'HALF YEARLY ASSESSMENT', sz=SZ_SEC, bold=True)
        draw_cell(col_x(4), y, span_w(4,12), RH_SEC,
                  'FORMATIVE ASSESSMENT', sz=SZ_SEC, bold=True)
        y -= RH_SEC

        # Row B+C combined height for spanning cells (col 0, 4, 8-12 span both rows)
        span_h = RH_HDR1 + RH_HDR2
        draw_cell(col_x(0), y, CW[0], span_h,
                  'No.\nof\nHalf\n/\nSEM', sz=SZ_HDR)
        draw_cell(col_x(1), y, span_w(1,3), RH_HDR1,
                  'Attendance During\nSemester', sz=SZ_HDR)
        draw_cell(col_x(4), y, CW[4], span_h,
                  'Specific\nLearning\nOutcome\n(70 Marks)\nA', sz=SZ_HDR)
        draw_cell(col_x(5), y, span_w(5,7), RH_HDR1,
                  'Generic Learning\nOutcome\n(30 Marks) B', sz=SZ_HDR)
        draw_cell(col_x(8), y, CW[8], span_h,
                  'Total\nOut\nof\n100\nA+B', sz=SZ_HDR)
        draw_cell(col_x(9),  y, CW[9],  span_h, 'Sign\nof\nTrainee', sz=SZ_HDR)
        draw_cell(col_x(10), y, CW[10], span_h, 'Sign\nof\nS.I',     sz=SZ_HDR)
        draw_cell(col_x(11), y, CW[11], span_h, 'Sign\nof\nF.I',     sz=SZ_HDR)
        draw_cell(col_x(12), y, CW[12], span_h, 'Sign\nof\nPrincipal', sz=SZ_HDR)
        y -= RH_HDR1

        # Row C: sub-headers for Attendance cols and GLO cols
        draw_cell(col_x(1), y, CW[1], RH_HDR2, 'Actual',   sz=SZ_HDR)
        draw_cell(col_x(2), y, CW[2], RH_HDR2, 'Possible', sz=SZ_HDR)
        draw_cell(col_x(3), y, CW[3], RH_HDR2, '%',        sz=SZ_HDR)
        draw_cell(col_x(5), y, CW[5], RH_HDR2, 'ES\n(10)', sz=SZ_HDR)
        draw_cell(col_x(6), y, CW[6], RH_HDR2, 'WS/C\n(10)', sz=SZ_HDR)
        draw_cell(col_x(7), y, CW[7], RH_HDR2, 'ED\n(10)', sz=SZ_HDR)
        y -= RH_HDR2
        return y

    # ── Single data row (H1/H2/H3/H4) ────────────────────────────────
    def draw_data_row(y_top, label, act, pos, pct, lo, es, wsc, ed, tot):
        def v(val): return str(val) if val != '' else ''
        draw_cell(col_x(0), y_top, CW[0],       RH_DATA, label, sz=SZ_DATA, bold=True)
        draw_cell(col_x(1), y_top, CW[1],       RH_DATA, v(act), sz=SZ_DATA)
        draw_cell(col_x(2), y_top, CW[2],       RH_DATA, v(pos), sz=SZ_DATA)
        draw_cell(col_x(3), y_top, CW[3],       RH_DATA, v(pct), sz=SZ_DATA)
        draw_cell(col_x(4), y_top, CW[4],       RH_DATA, v(lo),  sz=SZ_DATA)
        draw_cell(col_x(5), y_top, CW[5],       RH_DATA, v(es),  sz=SZ_DATA)
        draw_cell(col_x(6), y_top, CW[6],       RH_DATA, v(wsc), sz=SZ_DATA)
        draw_cell(col_x(7), y_top, CW[7],       RH_DATA, v(ed),  sz=SZ_DATA)
        draw_cell(col_x(8), y_top, CW[8],       RH_DATA, v(tot), sz=SZ_DATA, bold=True)
        draw_cell(col_x(9), y_top, CW[9],       RH_DATA, '', sz=SZ_DATA)
        draw_cell(col_x(10),y_top, CW[10],      RH_DATA, '', sz=SZ_DATA)
        draw_cell(col_x(11),y_top, CW[11],      RH_DATA, '', sz=SZ_DATA)
        draw_cell(col_x(12),y_top, CW[12],      RH_DATA, '', sz=SZ_DATA)
        return y_top - RH_DATA

    # ── Year summary block: header + data row ─────────────────────────
    def draw_year_summary(y_top, year_label):
        y = y_top
        # Header row — same 13 columns; FA Total header spans cols 4-8
        draw_cell(col_x(0),  y, CW[0],        RH_YRHDR, 'NO.\nof\nYear', sz=SZ_YR)
        draw_cell(col_x(1),  y, CW[1],        RH_YRHDR, 'Actual',        sz=SZ_YR)
        draw_cell(col_x(2),  y, CW[2],        RH_YRHDR, 'Possible',      sz=SZ_YR)
        draw_cell(col_x(3),  y, CW[3],        RH_YRHDR, '%',             sz=SZ_YR)
        draw_cell(col_x(4),  y, span_w(4,8),  RH_YRHDR,
                  'Formative Assessment\nTotal (Out of 200)',              sz=SZ_YR)
        draw_cell(col_x(9),  y, CW[9],        RH_YRHDR, 'Sign\nof\nTrainee', sz=SZ_YR)
        draw_cell(col_x(10), y, CW[10],       RH_YRHDR, 'Sign\nof\nS.I', sz=SZ_YR)
        draw_cell(col_x(11), y, CW[11],       RH_YRHDR, 'Sign\nof\nF.I', sz=SZ_YR)
        draw_cell(col_x(12), y, CW[12],       RH_YRHDR, 'Sign\nof\nPrincipal', sz=SZ_YR)
        y -= RH_YRHDR

        # Data row — FA Total value cell also spans cols 4-8 (merged)
        draw_cell(col_x(0),  y, CW[0],       RH_YRDAT, year_label, sz=SZ_DATA, bold=True)
        draw_cell(col_x(1),  y, CW[1],       RH_YRDAT, '', sz=SZ_DATA)
        draw_cell(col_x(2),  y, CW[2],       RH_YRDAT, '', sz=SZ_DATA)
        draw_cell(col_x(3),  y, CW[3],       RH_YRDAT, '', sz=SZ_DATA)
        draw_cell(col_x(4),  y, span_w(4,8), RH_YRDAT, '', sz=SZ_DATA)   # merged FA Total value
        draw_cell(col_x(9),  y, CW[9],       RH_YRDAT, '', sz=SZ_DATA)
        draw_cell(col_x(10), y, CW[10],      RH_YRDAT, '', sz=SZ_DATA)
        draw_cell(col_x(11), y, CW[11],      RH_YRDAT, '', sz=SZ_DATA)
        draw_cell(col_x(12), y, CW[12],      RH_YRDAT, '', sz=SZ_DATA)
        y -= RH_YRDAT
        return y

    # ─────────────────────────────────────────────────────────────────
    #  Per-trainee page loop
    # ─────────────────────────────────────────────────────────────────
    for t_idx, (_, tr) in enumerate(trainee_df.iterrows()):
        roll    = tr['rollno']
        nm      = ' '.join([str(tr.get(f,'') or '') for f in
                            ['Firstname','Fathername','Lastname']]).strip()
        adm_dt  = fmt_date(tr.get('DateofAdmission'))
        dob_dt  = fmt_date(tr.get('Birth_Date'))
        leav_dt = fmt_date(tr.get('DateofLeaving'))
        edu_col = next((col for col in tr.index if 'Edu' in str(col)), None)
        edu     = str(tr.get(edu_col,'') or '') if edu_col else ''
        es      = to_int(tr.get('ES',''))
        wsc     = to_int(tr.get('WC_SC',''))
        ed_val  = to_int(tr.get('ED',''))
        lo70    = to_int(tr.get('All_LO_average_base_on_70_SEM_I',''))
        wday    = to_int(tr.get('Working_days',''))
        aday    = to_int(tr.get('Attendeate_days',''))
        try:    att_pct = round(float(aday)/float(wday)*100)
        except: att_pct = ''
        try:    glo = int(es)+int(wsc)+int(ed_val)
        except: glo = ''
        try:    tot_100 = int(lo70)+int(glo)
        except: tot_100 = ''

        h1_act=aday    if sem_upper=='H1' else ''
        h1_pos=wday    if sem_upper=='H1' else ''
        h1_pct=att_pct if sem_upper=='H1' else ''
        h1_lo =lo70    if sem_upper=='H1' else ''
        h1_es =es      if sem_upper=='H1' else ''
        h1_wsc=wsc     if sem_upper=='H1' else ''
        h1_ed =ed_val  if sem_upper=='H1' else ''
        h1_tot=tot_100 if sem_upper=='H1' else ''
        h2_act=aday    if sem_upper=='H2' else ''
        h2_pos=wday    if sem_upper=='H2' else ''
        h2_pct=att_pct if sem_upper=='H2' else ''
        h2_lo =lo70    if sem_upper=='H2' else ''
        h2_es =es      if sem_upper=='H2' else ''
        h2_wsc=wsc     if sem_upper=='H2' else ''
        h2_ed =ed_val  if sem_upper=='H2' else ''
        h2_tot=tot_100 if sem_upper=='H2' else ''

        if t_idx > 0:
            c.showPage()

        # Start at top of card (centred on page for 1-year)
        y = PAGE_H - MT
        y_card_top = y

        # ── ROW 1 & 2: Institute name + PROGRESS CARD ─────────────────
        draw_cell(ML, y, TW, RH_TITLE,
                  f'INDUSTRIAL TRAINING INSTITUTE : {u_iti.upper()}',
                  sz=SZ_TITLE, bold=True)
        y -= RH_TITLE
        draw_cell(ML, y, TW, RH_TITLE,
                  'PROGRESS CARD', sz=SZ_TITLE, bold=True)
        y -= RH_TITLE

        # ── ROLL NO | roll | TRADE | trade name | BATCH | batch ───────
        draw_cell(col_x(0), y, span_w(0,1), RH_INFO,
                  'ROLL NO.', sz=SZ_INFO, bold=True)
        draw_cell(col_x(2), y, CW[2], RH_INFO,
                  str(roll), sz=SZ_INFO)
        draw_cell(col_x(3), y, span_w(3,4), RH_INFO,
                  'TRADE', sz=SZ_INFO, bold=True)
        draw_cell(col_x(5), y, span_w(5,10), RH_INFO,
                  u_trade, sz=SZ_INFO, align='left')
        draw_cell(col_x(11), y, CW[11], RH_INFO,
                  'BATCH', sz=SZ_INFO, bold=True)
        draw_cell(col_x(12), y, CW[12], RH_INFO,
                  str(u_batch), sz=SZ_INFO)
        y -= RH_INFO

        # ── NAME OF TRAINEE ───────────────────────────────────────────
        draw_cell(col_x(0), y, span_w(0,1), RH_INFO,
                  'NAME OF TRAINEE', sz=SZ_INFO, bold=True)
        draw_cell(col_x(2), y, span_w(2,12), RH_INFO,
                  nm, sz=SZ_INFO, align='left')
        y -= RH_INFO

        # ── DATE OF ADMISSION | DATE OF BIRTH ─────────────────────────
        draw_cell(col_x(0), y, span_w(0,2), RH_DATE,
                  'DATE OF ADMISSION', sz=SZ_DATE, bold=True)
        draw_cell(col_x(3), y, span_w(3,6), RH_DATE,
                  adm_dt, sz=SZ_DATE)
        draw_cell(col_x(7), y, span_w(7,9), RH_DATE,
                  'DATE OF BIRTH', sz=SZ_DATE, bold=True)
        draw_cell(col_x(10), y, span_w(10,12), RH_DATE,
                  dob_dt, sz=SZ_DATE)
        y -= RH_DATE

        # ── DATE OF LEAVING | EDU. QUA ────────────────────────────────
        draw_cell(col_x(0), y, span_w(0,2), RH_DATE,
                  'DATE OF LEAVING', sz=SZ_DATE, bold=True)
        draw_cell(col_x(3), y, span_w(3,6), RH_DATE,
                  leav_dt, sz=SZ_DATE)
        draw_cell(col_x(7), y, span_w(7,9), RH_DATE,
                  'EDU. QUA', sz=SZ_DATE, bold=True)
        draw_cell(col_x(10), y, span_w(10,12), RH_DATE,
                  edu, sz=SZ_DATE)
        y -= RH_DATE

        # ── 2-year: YEAR-1 ASSESSMENT label ───────────────────────────
        if is_2yr:
            draw_cell(ML, y, TW, RH_SEC,
                      'YEAR - 1 ASSESSMENT', sz=SZ_SEC, bold=True)
            y -= RH_SEC

        # ── Year-1 table: header + H1 + H2 + summary ──────────────────
        y = draw_table_header(y)
        y = draw_data_row(y, 'H1',
                          h1_act, h1_pos, h1_pct, h1_lo,
                          h1_es, h1_wsc, h1_ed, h1_tot)
        y = draw_data_row(y, 'H2',
                          h2_act, h2_pos, h2_pct, h2_lo,
                          h2_es, h2_wsc, h2_ed, h2_tot)
        y = draw_year_summary(y, 'YEAR-1')

        # ── 2-year: Year-2 block ───────────────────────────────────────
        if is_2yr:
            draw_cell(ML, y, TW, RH_SEC,
                      'YEAR - 2 ASSESSMENT', sz=SZ_SEC, bold=True)
            y -= RH_SEC
            y = draw_table_header(y)
            y = draw_data_row(y, 'H3', '', '', '', '', '', '', '', '')
            y = draw_data_row(y, 'H4', '', '', '', '', '', '', '', '')
            y = draw_year_summary(y, 'YEAR-2')

        # ── Outer border (thick) around entire card ────────────────────
        used_h = y_card_top - y
        c.setStrokeColor(BLACK)
        c.setLineWidth(1.0)
        c.rect(ML, y, TW, used_h, fill=0, stroke=1)

    c.save()
    buf.seek(0)
    return buf



def create_formative_pdf(trainee_df, lo_df, user_info=None, marks_cache=None):
    """
    Generate a print-ready A4 landscape PDF with one page per trainee.
    All trainees are combined into a single PDF file.
    Uses the same marks_cache as create_excel so numbers are identical.
    """
    if not _REPORTLAB_OK:
        raise RuntimeError("ReportLab is not installed. Run: pip install reportlab")

    ui = user_info or {}
    u_si_name            = ui.get('si_name', 'N/A')
    u_trade_name         = ui.get('trade_name', 'N/A')
    u_iti_name           = ui.get('iti_name', 'N/A')
    u_year_of_assessment = ui.get('year_of_assessment', '')
    u_assessment_location= ui.get('assessment_location', '')
    u_trade_duration     = ui.get('trade_duration', '')
    u_semester           = ui.get('semester', '')
    u_batch              = ui.get('batch', '')

    # ── Page geometry ─────────────────────────────────────────────────
    PAGE_W, PAGE_H = landscape(A4)   # 297 × 210 mm → pts: ~841.9 × 595.3
    L = R = 6 * mm
    T = B = 6 * mm
    USABLE_W = PAGE_W - L - R        # available content width in pts

    # ── Styles ────────────────────────────────────────────────────────
    NAVY  = colors.HexColor('#0B1D3A')
    BLUE  = colors.HexColor('#1B4F8A')
    LBLUE = colors.HexColor('#D6E4F0')
    WHITE = colors.white
    BLACK = colors.black
    LGRAY = colors.HexColor('#F8FAFC')

    def _ps(size=6, bold=False, color=BLACK, align=TA_CENTER):
        return ParagraphStyle('_', fontSize=size, leading=size+1.5,
                              fontName='Helvetica-Bold' if bold else 'Helvetica',
                              textColor=color, alignment=align,
                              wordWrap='LTR', spaceBefore=0, spaceAfter=0)

    # ── RotatedText: draws text rotated 90° inside a ReportLab table cell ──
    from reportlab.platypus import Flowable as _Flowable

    class RotatedText(_Flowable):
        """Draws text rotated 90° (bottom-to-top) centred in a fixed cell."""
        def __init__(self, text, cell_w, cell_h, font='Helvetica-Bold', size=6, color=None):
            _Flowable.__init__(self)
            self.text   = text
            self.width  = cell_w
            self.height = cell_h
            self._font  = font
            self._size  = size
            self._color = color or colors.black

        def draw(self):
            c = self.canv
            c.saveState()
            c.setFont(self._font, self._size)
            c.setFillColor(self._color)
            # Rotate around centre of cell
            cx = self.width / 2
            cy = self.height / 2
            c.translate(cx, cy)
            c.rotate(90)
            # After rotation: new x-axis is old y-axis.
            # Draw string centred on the rotated axis.
            lines = self.text.split('\n')
            line_h = self._size * 1.2
            total_h = len(lines) * line_h
            for i, line in enumerate(lines):
                tw = c.stringWidth(line, self._font, self._size)
                # y offset: start above centre, step down per line
                y_off = (total_h / 2) - (i + 0.5) * line_h
                c.drawString(-tw / 2, y_off, line)
            c.restoreState()

        def wrap(self, aW, aH):
            return self.width, self.height

    # ── Column definitions (41 data columns) ──────────────────────────
    # Match Excel: narrow sub-cols, wider LO/Pract/GrandTotal/Sig cols.
    # Proportions mirror the demo: A=1.6, B=1.1, subs=0.95, totals=1.2,
    #   Grand=1.3, SigT=1.6, SigSI=1.6
    raw_ws = [1.6, 1.1]   # A (LO No), B (Practical No)
    for _ in range(9):    # 9 categories × (3 sub + 1 total)
        raw_ws += [0.95, 0.95, 0.95, 1.2]
    raw_ws += [1.3, 1.6, 1.6]   # Grand Total, Sig Trainee, Sig SI
    total_raw = sum(raw_ws)
    COL_WS = [w / total_raw * USABLE_W for w in raw_ws]

    # Heights for the two header rows (matching Excel proportions)
    CAT_ROW_H  = 36   # category row: LO No / Practical No / Safety ... (rotated single-cols)
    SUB_ROW_H  = 110  # sub-column row: all cells rotated 90° (increased to fit 'Maintain Personal & Workplace Cleanliness')
    MAX_ROW_H  = 10   # max marks row

    # ── Build PDF story page-by-page ──────────────────────────────────
    buf = io.BytesIO()

    def _on_page(canvas, doc):
        """Draw thin outer border on every page."""
        canvas.saveState()
        canvas.setStrokeColor(BLUE)
        canvas.setLineWidth(0.8)
        canvas.rect(L - 2, B - 2, PAGE_W - L - R + 4, PAGE_H - T - B + 4)
        canvas.restoreState()

    doc = SimpleDocTemplate(
        buf,
        pagesize=landscape(A4),
        leftMargin=L, rightMargin=R,
        topMargin=T, bottomMargin=B,
    )

    story = []

    for t_idx, (_, trainee) in enumerate(trainee_df.iterrows()):
        roll = trainee['rollno']
        full_name = f"{trainee.get('Firstname','')} {trainee.get('Fathername','')} {trainee.get('Lastname','')}".strip()
        raw_70 = trainee.get('All_LO_average_base_on_70_SEM_I')
        if raw_70 is not None and pd.notna(raw_70):
            target_avg = round((float(raw_70) / 70) * 100)
        else:
            target_avg = None

        trainee_cache = (marks_cache or {}).get(roll, {})

        # ── Helper: Paragraph cell ────────────────────────────────────
        def P(txt, size=6, bold=False, color=BLACK, align=TA_CENTER):
            return Paragraph(str(txt), _ps(size, bold, color, align))

        def PL(txt, size=6, bold=False):
            return Paragraph(str(txt), _ps(size, bold, align=TA_LEFT))

        # ── Title row ─────────────────────────────────────────────────
        title_data = [[P('Internal Assessment', size=9, bold=True)]]
        title_tbl = Table(title_data, colWidths=[USABLE_W])
        title_tbl.setStyle(TableStyle([
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('BOX', (0,0), (-1,-1), 0.5, BLACK),
            ('TOPPADDING', (0,0), (-1,-1), 3),
            ('BOTTOMPADDING', (0,0), (-1,-1), 3),
        ]))
        story.append(title_tbl)

        # ── Info rows (4 rows × 3 segments) ───────────────────────────
        info_rows_data = [
            (f"Name of Trainee: {full_name}",       f"Roll No: {roll}",                          f"Year of Assessment: {u_year_of_assessment}"),
            (f"Name of ITI: {u_iti_name}",           "Date of Assessment: _____________",         f"Batch: {u_batch}"),
            (f"Name of Industry: {u_iti_name}",      f"Assessment Location: {u_assessment_location}", ""),
            (f"Trade Name: {u_trade_name}",           f"Duration of Trade: {u_trade_duration}",    f"S.I. Name: {u_si_name}"),
        ]
        # Three columns: ~40% / 25% / 35% of usable width
        seg_ws = [USABLE_W * 0.40, USABLE_W * 0.28, USABLE_W * 0.32]
        for seg1, seg2, seg3 in info_rows_data:
            info_tbl = Table(
                [[PL(seg1, size=6, bold=True), PL(seg2, size=6, bold=True), PL(seg3, size=6, bold=True)]],
                colWidths=seg_ws
            )
            info_tbl.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('BOX', (0,0), (-1,-1), 0.5, BLACK),
                ('INNERGRID', (0,0), (-1,-1), 0.3, BLACK),
                ('TOPPADDING', (0,0), (-1,-1), 1.5),
                ('BOTTOMPADDING', (0,0), (-1,-1), 1.5),
                ('LEFTPADDING', (0,0), (-1,-1), 2),
                ('RIGHTPADDING', (0,0), (-1,-1), 2),
            ]))
            story.append(info_tbl)

        # ── Category header row + Sub-column row + Max row ────────────
        # Row 0 (CAT): LO No [rot90], Practical No [rot90], Safety [wrap], ..., Grand Total [rot90], Sig [rot90]
        # Row 1 (SUB): all 41 cells rotated 90°
        # Row 2 (MAX): max marks, plain centred text

        def _rot(text, col_idx, row_h, bold=True):
            """RotatedText cell sized to match its column width."""
            fn = 'Helvetica-Bold' if bold else 'Helvetica'
            return RotatedText(text, COL_WS[col_idx], row_h, font=fn, size=5.5)

        def _wrap(text, size=5.5, bold=True):
            return P(text, size=size, bold=bold)

        # — Row 0: category headers —
        cat_row = []
        # col 0: LO No — rotated
        cat_row.append(_rot('LO No', 0, CAT_ROW_H))
        # col 1: Practical No — rotated
        cat_row.append(_rot('Practical No', 1, CAT_ROW_H))
        # cols 2-5: Safety span — wrap
        cat_row.append(_wrap('Safety\nConsciousness'))
        cat_row += ['', '', '']
        # cols 6-9: Hygiene span — wrap
        cat_row.append(_wrap('Workplace Hygiene\n& Economical Use\nof Materials'))
        cat_row += ['', '', '']
        # cols 10-13: Attendance — wrap
        cat_row.append(_wrap('Attendance/\nPunctuality'))
        cat_row += ['', '', '']
        # cols 14-17: Manuals — wrap
        cat_row.append(_wrap('Ability to Follow\nManuals/Written\nInstructions'))
        cat_row += ['', '', '']
        # cols 18-21: Knowledge — wrap
        cat_row.append(_wrap('Application of\nKnowledge'))
        cat_row += ['', '', '']
        # cols 22-25: Skills — wrap
        cat_row.append(_wrap('Skills to Handle\nTools & Equipment'))
        cat_row += ['', '', '']
        # cols 26-29: Speed — wrap
        cat_row.append(_wrap('Speed in\nDoing Work'))
        cat_row += ['', '', '']
        # cols 30-33: Quality — wrap
        cat_row.append(_wrap('Quality in\nWorkmanship'))
        cat_row += ['', '', '']
        # cols 34-37: VIVA — wrap
        cat_row.append(_wrap('VIVA'))
        cat_row += ['', '', '']
        # col 38: Grand Total — rotated
        cat_row.append(_rot('Grand Total', 38, CAT_ROW_H))
        # col 39: Signature Trainee — rotated
        cat_row.append(_rot('Signature Trainee', 39, CAT_ROW_H))
        # col 40: Signature SI — rotated
        cat_row.append(_rot('Signature SI', 40, CAT_ROW_H))

        # — Row 1: sub-column headers, all rotated 90° —
        sub_texts = [
            'LO No', 'Practical No',
            'Dress Code', 'Use PPE', 'Apply/ Practice Safety', 'Total',
            'Maintain Personal & Workplace Cleanliness', 'Dispose Scrap', 'Select Material', 'Total',
            'Initiative', 'Account- ability', 'Participative in Work', 'Total',
            'Select Right Manual', 'Search Topic', 'Read & Interpret', 'Total',
            'Plan the Work', 'Select Tools', 'Review Work', 'Total',
            'Handle & Use Tools', 'Maintain Safety', 'Care & Maintain', 'Total',
            'Properly Sequence', 'Use Approp. Technique', 'Review Execution', 'Total',
            'Achieve High Accuracy', 'Conform to Req.', 'Satisfy Purpose', 'Total',
            'Response with Clarity', 'Technical Understand.', 'Conscious towards Job Role', 'Total',
            '', '', '',
        ]
        sub_row = [
            RotatedText(t, COL_WS[i], SUB_ROW_H, font='Helvetica-Bold', size=5)
            if t else P('', 5)
            for i, t in enumerate(sub_texts)
        ]

        # — Row 2: max marks —
        max_labels = [
            '', '',
            2, 5, 8, 15, 3, 2, 5, 10, 3, 3, 4, 10,
            1, 2, 2, 5, 4, 3, 3, 10, 4, 3, 3, 10,
            3, 5, 2, 10, 7, 3, 5, 15, 7, 5, 3, 15,
            100, '', '',
        ]
        max_row = [P(str(m) if m else '', size=5.5, bold=True) for m in max_labels]

        hdr_data = [cat_row, sub_row, max_row]

        hdr_style = [
            ('GRID', (0,0), (-1,-1), 0.4, BLACK),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('TOPPADDING', (0,0), (-1,-1), 1),
            ('BOTTOMPADDING', (0,0), (-1,-1), 1),
            ('LEFTPADDING', (0,0), (-1,-1), 0),
            ('RIGHTPADDING', (0,0), (-1,-1), 0),
            ('BACKGROUND', (0,0), (-1,0), LBLUE),  # cat row
            ('BACKGROUND', (0,1), (-1,1), LBLUE),  # sub row
            # Row 0: category spans (each category header spans its 4 sub-cols)
            ('SPAN', (2,0),  (5,0)),   # Safety
            ('SPAN', (6,0),  (9,0)),   # Hygiene
            ('SPAN', (10,0), (13,0)),  # Attendance
            ('SPAN', (14,0), (17,0)),  # Manuals
            ('SPAN', (18,0), (21,0)),  # Knowledge
            ('SPAN', (22,0), (25,0)),  # Skills
            ('SPAN', (26,0), (29,0)),  # Speed
            ('SPAN', (30,0), (33,0)),  # Quality
            ('SPAN', (34,0), (37,0)),  # VIVA
            # LO No and Practical No span all 3 rows
            ('SPAN', (0,0), (0,2)),
            ('SPAN', (1,0), (1,2)),
            # Grand Total and Sig span rows 0-1
            ('SPAN', (38,0), (38,1)),
            ('SPAN', (39,0), (39,1)),
            ('SPAN', (40,0), (40,1)),
        ]

        hdr_tbl = Table(hdr_data, colWidths=COL_WS,
                        rowHeights=[CAT_ROW_H, SUB_ROW_H, MAX_ROW_H])
        hdr_tbl.setStyle(TableStyle(hdr_style))
        story.append(hdr_tbl)

        # ── Data rows ─────────────────────────────────────────────────
        all_lo_totals = []
        data_rows = []
        row_styles = []
        row_idx = 0

        for lo_num, lo_data in trainee_cache.items():
            pract_from  = lo_data['pract_from']
            pract_to    = lo_data['pract_to']
            lo_name     = lo_data['lo_name']
            pract_marks = lo_data['pract_marks']
            lo_grand_totals = []

            for pract_offset, pract_num in enumerate(range(pract_from, pract_to + 1)):
                m = pract_marks[pract_offset]
                lo_grand_totals.append(m['grand_total'])
                row_vals = [
                    f"LO-{lo_num}", pract_num,
                    m['dress'], m['ppe'], m['apply_safety'], m['safety_total'],
                    m['personal'], m['scrap'], m['material'], m['hygiene_total'],
                    m['initiative'], m['accountability'], m['participative'], m['attendance_total'],
                    m['select_manual'], m['search_topic'], m['read_manual'], m['manuals_total'],
                    m['plan_work'], m['select_tools'], m['review_work'], m['knowledge_total'],
                    m['handle_tools'], m['safety_handling'], m['care_maintain'], m['skills_total'],
                    m['sequence'], m['technique'], m['review_execution'], m['speed_total'],
                    m['accuracy'], m['conform'], m['satisfy'], m['quality_total'],
                    m['clarity'], m['technical'], m['conscious'], m['viva_total'],
                    m['grand_total'], "", "",
                ]
                data_rows.append([P(str(v), size=5.5) for v in row_vals])
                if row_idx % 2 == 0:
                    row_styles.append(('BACKGROUND', (0, row_idx), (-1, row_idx), LGRAY))
                row_idx += 1

            # LO summary row
            lo_avg = round(sum(lo_grand_totals) / len(lo_grand_totals))
            all_lo_totals.append(lo_avg)
            lo_summary = [P(f"{lo_name}    |    Average of LO{lo_num}: {lo_avg}", size=5.5, bold=True)] + [''] * 40
            data_rows.append(lo_summary)
            row_styles.append(('SPAN', (0, row_idx), (40, row_idx)))
            row_styles.append(('BACKGROUND', (0, row_idx), (-1, row_idx), LBLUE))
            row_idx += 1

        # Grand average row
        overall_avg = round(sum(all_lo_totals) / len(all_lo_totals)) if all_lo_totals else (target_avg or 0)
        grand_row = [P(f"Average of All LO:  {overall_avg}", size=6, bold=True)] + [''] * 40
        data_rows.append(grand_row)
        row_styles.append(('SPAN', (0, row_idx), (40, row_idx)))
        row_styles.append(('BACKGROUND', (0, row_idx), (-1, row_idx), colors.HexColor('#FFF2CC')))
        row_styles.append(('FONTNAME', (0, row_idx), (-1, row_idx), 'Helvetica-Bold'))

        data_style = [
            ('GRID', (0,0), (-1,-1), 0.3, BLACK),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('TOPPADDING', (0,0), (-1,-1), 1),
            ('BOTTOMPADDING', (0,0), (-1,-1), 1),
            ('LEFTPADDING', (0,0), (-1,-1), 1),
            ('RIGHTPADDING', (0,0), (-1,-1), 1),
        ] + row_styles

        if data_rows:
            data_tbl = Table(data_rows, colWidths=COL_WS, repeatRows=0)
            data_tbl.setStyle(TableStyle(data_style))
            story.append(data_tbl)

        # Page break between trainees (not after last)
        if t_idx < len(trainee_df) - 1:
            story.append(PageBreak())

    doc.build(story, onFirstPage=_on_page, onLaterPages=_on_page)
    buf.seek(0)
    return buf


@app.route('/generate-annexure-es-pdf', methods=['POST'])
def generate_annexure_es_pdf():
    return _serve_annexure_pdf('ES')

@app.route('/generate-annexure-wcs-pdf', methods=['POST'])
def generate_annexure_wcs_pdf():
    return _serve_annexure_pdf('WCS')

@app.route('/generate-annexure-ed-pdf', methods=['POST'])
def generate_annexure_ed_pdf():
    return _serve_annexure_pdf('ED')

def _serve_annexure_pdf(ann_type):
    if 'user_id' not in session:
        return jsonify({'error': 'Please login first'}), 401
    if not _REPORTLAB_OK:
        return jsonify({'error': 'ReportLab not installed on server. Run: pip install reportlab'}), 500
    try:
        uid = session.get('user_id')
        trainee_df = _get_df(uid, 'trainee', request.files.get('trainee'))
        if trainee_df is None:
            return jsonify({'error': 'No trainee file found. Please upload the Trainee Details file first.'}), 400
        user_info = {k: session.get(k,'') for k in
                     ['si_name','trade_name','iti_name','year_of_assessment',
                      'assessment_location','trade_duration','semester','batch']}
        output = create_annexure_pdf(ann_type, trainee_df, user_info)
        return send_file(output, mimetype='application/pdf',
                         as_attachment=True,
                         download_name=f'ANNEXURE_III_{ann_type}.pdf')
    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'detail': traceback.format_exc()}), 500


@app.route('/generate-lo-summary-pdf', methods=['POST'])
def generate_lo_summary_pdf():
    if 'user_id' not in session:
        return jsonify({'error': 'Please login first'}), 401
    if not _REPORTLAB_OK:
        return jsonify({'error': 'ReportLab not installed on server. Run: pip install reportlab'}), 500
    try:
        uid = session.get('user_id')
        trainee_df = _get_df(uid, 'trainee', request.files.get('trainee'))
        lo_df      = _get_df(uid, 'lo',      request.files.get('lo'))
        if trainee_df is None or lo_df is None:
            return jsonify({'error': 'Both files are required.'}), 400
        user_info = {k: session.get(k,'') for k in
                     ['si_name','trade_name','iti_name','year_of_assessment',
                      'assessment_location','near_trade','trade_duration','semester','batch']}
        t_hash = _file_hash(_load_user_file(uid, 'trainee'))
        l_hash = _file_hash(_load_user_file(uid, 'lo'))
        marks_cache = (_marks_cache_store.get(uid) or _load_marks_cache_disk(uid, t_hash, l_hash))
        if marks_cache is None:
            marks_cache = build_all_marks(trainee_df, lo_df)
            _save_marks_cache_disk(uid, marks_cache, t_hash, l_hash)
            _marks_cache_store[uid] = marks_cache
        output = create_lo_summary_pdf(trainee_df, lo_df, user_info, marks_cache=marks_cache)
        return send_file(output, mimetype='application/pdf',
                         as_attachment=True, download_name='LO_Wise_Summary.pdf')
    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'detail': traceback.format_exc()}), 500


@app.route('/generate-trainee-lo-pdf', methods=['POST'])
def generate_trainee_lo_pdf():
    if 'user_id' not in session:
        return jsonify({'error': 'Please login first'}), 401
    if not _REPORTLAB_OK:
        return jsonify({'error': 'ReportLab not installed on server. Run: pip install reportlab'}), 500
    try:
        uid = session.get('user_id')
        trainee_df = _get_df(uid, 'trainee', request.files.get('trainee'))
        lo_df      = _get_df(uid, 'lo',      request.files.get('lo'))
        if trainee_df is None or lo_df is None:
            return jsonify({'error': 'Both files are required.'}), 400
        user_info = {k: session.get(k,'') for k in
                     ['si_name','trade_name','iti_name','year_of_assessment',
                      'assessment_location','near_trade','trade_duration','semester','batch']}
        t_hash = _file_hash(_load_user_file(uid, 'trainee'))
        l_hash = _file_hash(_load_user_file(uid, 'lo'))
        marks_cache = (_marks_cache_store.get(uid) or _load_marks_cache_disk(uid, t_hash, l_hash))
        if marks_cache is None:
            marks_cache = build_all_marks(trainee_df, lo_df)
            _save_marks_cache_disk(uid, marks_cache, t_hash, l_hash)
            _marks_cache_store[uid] = marks_cache
        output = create_trainee_lo_pdf(trainee_df, lo_df, user_info, marks_cache=marks_cache)
        return send_file(output, mimetype='application/pdf',
                         as_attachment=True, download_name='Per_Trainee_Per_LO_Report.pdf')
    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'detail': traceback.format_exc()}), 500


@app.route('/generate-progress-card-pdf', methods=['POST'])
def generate_progress_card_pdf_route():
    if 'user_id' not in session:
        return jsonify({'error': 'Please login first'}), 401
    if not _REPORTLAB_OK:
        return jsonify({'error': 'ReportLab not installed on server. Run: pip install reportlab'}), 500
    try:
        uid = session.get('user_id')
        trainee_df = _get_df(uid, 'trainee', request.files.get('trainee'))
        if trainee_df is None:
            return jsonify({'error': 'No trainee file found. Please upload the Trainee Details file first.'}), 400
        user_info = {k: session.get(k,'') for k in
                     ['si_name','trade_name','iti_name','year_of_assessment',
                      'assessment_location','near_trade','trade_duration','semester','batch']}
        output = create_progress_card_pdf(trainee_df, user_info)
        return send_file(output, mimetype='application/pdf',
                         as_attachment=True, download_name='Progress_Cards.pdf')
    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'detail': traceback.format_exc()}), 500


@app.route('/generate-pdf', methods=['POST'])
def generate_pdf():
    """Generate a print-ready A4 landscape PDF (all trainees in one file)."""
    if 'user_id' not in session:
        return jsonify({'error': 'Please login first'}), 401
    if not _REPORTLAB_OK:
        return jsonify({'error': 'ReportLab not installed on server. Run: pip install reportlab'}), 500
    try:
        uid = session.get('user_id')
        trainee_df = _get_df(uid, 'trainee', request.files.get('trainee'))
        lo_df      = _get_df(uid, 'lo',      request.files.get('lo'))
        if trainee_df is None or lo_df is None:
            return jsonify({'error': 'Both files are required. Please upload Trainee Details and LO Details files.'}), 400
        user_info = {
            'si_name':            session.get('si_name', ''),
            'trade_name':         session.get('trade_name', ''),
            'iti_name':           session.get('iti_name', ''),
            'year_of_assessment': session.get('year_of_assessment', ''),
            'assessment_location':session.get('assessment_location', ''),
            'near_trade':         session.get('near_trade', ''),
            'trade_duration':     session.get('trade_duration', ''),
            'semester':           session.get('semester', ''),
            'batch':              session.get('batch', ''),
        }
        # Reuse cached marks if already generated (so numbers match Excel)
        t_hash = _file_hash(_load_user_file(uid, 'trainee'))
        l_hash = _file_hash(_load_user_file(uid, 'lo'))
        marks_cache = (_marks_cache_store.get(uid)
                       or _load_marks_cache_disk(uid, t_hash, l_hash))
        if marks_cache is None:
            marks_cache = build_all_marks(trainee_df, lo_df)
            _save_marks_cache_disk(uid, marks_cache, t_hash, l_hash)
        _marks_cache_store[uid]  = marks_cache
        _trainee_data_store[uid] = trainee_df
        output = create_formative_pdf(trainee_df, lo_df, user_info, marks_cache=marks_cache)
        return send_file(output, mimetype='application/pdf',
                         as_attachment=True, download_name='ITI_Assessment_Report.pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/generate', methods=['POST'])
def generate():
    if 'user_id' not in session:
        return jsonify({'error': 'Please login first'}), 401
    try:
        uid = session.get('user_id')
        trainee_df = _get_df(uid, 'trainee', request.files.get('trainee'))
        lo_df      = _get_df(uid, 'lo',      request.files.get('lo'))
        if trainee_df is None or lo_df is None:
            return jsonify({'error': 'Both files are required. Please upload Trainee Details and LO Details files.'}), 400
        user_info = {
            'si_name': session.get('si_name',''),
            'trade_name': session.get('trade_name',''),
            'iti_name': session.get('iti_name',''),
            'year_of_assessment': session.get('year_of_assessment',''),
            'assessment_location': session.get('assessment_location',''),
            'near_trade': session.get('near_trade',''),
            'trade_duration': session.get('trade_duration',''),
            'semester': session.get('semester',''),
            'batch': session.get('batch',''),
        }
        # ── Persistent marks cache: only regenerate if files changed ──
        t_hash = _file_hash(_load_user_file(uid, 'trainee'))
        l_hash = _file_hash(_load_user_file(uid, 'lo'))
        marks_cache = (_marks_cache_store.get(uid)
                       or _load_marks_cache_disk(uid, t_hash, l_hash))
        if marks_cache is None:
            marks_cache = build_all_marks(trainee_df, lo_df)
            _save_marks_cache_disk(uid, marks_cache, t_hash, l_hash)
        _marks_cache_store[uid]  = marks_cache
        _trainee_data_store[uid] = trainee_df
        output = create_excel(trainee_df, lo_df, user_info, marks_cache=marks_cache)
        return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                         as_attachment=True, download_name='ITI_Assessment_Report.xlsx')
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    print("="*50)
    print("ITI Assessment System Started!")
    print("Default Admin Password: admin123")
    print("URL: http://localhost:5001")
    print("="*50)
    app.run(host='0.0.0.0', port=5001, debug=False)